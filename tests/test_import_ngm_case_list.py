"""
Unit tests for scripts/import_ngm_case_list.py — the one-time import of
Nyari's master case-list Word doc, cross-checked against clients/matters
already in the database (this document is very likely the original source
for the clients already imported via earlier migrations).

Covers:
  - parse_case_list_docx(): robustness against the real document's field-
    label variants (backtick prefixes, case/spacing inconsistencies).
  - _combine_matter_description(): Re/Law combination, including the
    Re==Law collapse.
  - _matter_dedup_score() / _best_matter_dedup_match(): calibrated against
    real matter-text pairs (true positives: same matter, existing DB text
    has more detail appended; false-positive risks: generic Law-style Re
    values like "Trust"/"Estate" that recur across unrelated matters).
  - _build_plan(): the full cross-check pipeline against a FakeConnection —
    new client numbering, matched-client reuse, ambiguous-client review,
    matter dedup against both pre-existing DB matters and matters created
    earlier in the same run, client-only (no case detail) records, and
    unparseable rows.
"""

import asyncio
import io
import uuid
from datetime import datetime, timezone

import pytest
from docx import Document

from scripts.import_ngm_case_list import (
    FIRM_ID,
    _best_matter_dedup_match,
    _build_plan,
    _combine_matter_description,
    _matter_dedup_score,
    parse_case_list_docx,
)

FIRM_UUID = uuid.UUID(FIRM_ID)


# ── parse_case_list_docx ─────────────────────────────────────────────────

def _build_case_list_docx(records: list) -> str:
    """records: list of list-of-(label, value) pairs, one per table."""
    doc = Document()
    for fields in records:
        table = doc.add_table(rows=len(fields), cols=2)
        for i, (label, value) in enumerate(fields):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = value
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def test_parse_handles_clean_record():
    path = _build_case_list_docx([[
        ("File Name", "Ngm 1"),
        ("Name of client", "Munjeya Abigail"),
        ("Re", "Master: Estate Late Itai Sunguro"),
        ("Law", "Estate / Inheritance"),
        ("Action done", "Ngm appointed as Executor"),
        ("Next action", "Waiting for appointment of a curator"),
        ("Latest communication", "24/3/26"),
    ]])
    result = parse_case_list_docx(path)
    assert len(result["records"]) == 1
    r = result["records"][0]
    assert r["client_name"] == "Munjeya Abigail"
    assert r["re"] == "Master: Estate Late Itai Sunguro"
    assert r["law"] == "Estate / Inheritance"
    assert r["latest_communication"] == "24/3/26"
    assert result["unrecognized_field_warnings"] == []


def test_parse_handles_real_document_field_label_variants():
    """The real document uses inconsistent labels for the same field —
    backtick prefix, case variants, no-space, single-word abbreviations."""
    path = _build_case_list_docx([[
        ("`File Name", "Ngm 4"),
        ("Name of client", "Daniel Dube"),
        ("Re", "Transfer of farm"),
        ("Law", "Conveyancing"),
        ("Action done", "Communication with client"),
        ("Next action", "Client passed away"),
        ("Latest communication", ""),
    ], [
        ("File name", "Ngm 5"),                    # lowercase 'name'
        ("Name of client", "Austin Muroiwa"),
        ("Case number", "HC 1/26"),                 # lowercase 'number'
        ("Re", "Mary Pascalia Muroiwa"),
        ("Law", "Matrimonial / divorce"),
        ("Action done", "Curator appointed"),
        ("Next action latest", "Closed file"),      # odd variant seen once
        ("Latestcommunication", "June 2025"),       # no space
    ]])
    result = parse_case_list_docx(path)
    assert len(result["records"]) == 2
    assert result["records"][0]["file_name"] == "Ngm 4"
    r2 = result["records"][1]
    assert r2["file_name"] == "Ngm 5"
    assert r2["case_number"] == "HC 1/26"
    assert r2["next_action"] == "Closed file"
    assert r2["latest_communication"] == "June 2025"
    assert result["unrecognized_field_warnings"] == []


def test_parse_ignores_malformed_blank_label_rows():
    """A handful of real tables have a stray row with a blank label cell
    (['', '']) — a document artifact, not real data; must not crash or be
    reported as an unrecognized field."""
    doc = Document()
    table = doc.add_table(rows=4, cols=2)
    table.cell(0, 0).text = "Name of client"
    table.cell(0, 1).text = "Janga Christopher"
    table.cell(1, 0).text = ""
    table.cell(1, 1).text = ""
    table.cell(2, 0).text = "Re"
    table.cell(2, 1).text = "Estate Late Edna Mudzengi"
    table.cell(3, 0).text = "Law"
    table.cell(3, 1).text = "Estate"
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    result = parse_case_list_docx(buf)
    assert len(result["records"]) == 1
    assert result["records"][0]["client_name"] == "Janga Christopher"
    assert result["unrecognized_field_warnings"] == []


def test_parse_flags_genuinely_unrecognized_field_label():
    path = _build_case_list_docx([[
        ("Name of client", "Some Client"),
        ("Totally New Field", "unexpected"),
    ]])
    result = parse_case_list_docx(path)
    assert len(result["unrecognized_field_warnings"]) == 1
    assert result["unrecognized_field_warnings"][0]["unrecognized_label"] == "Totally New Field"


def test_parse_blank_client_name_field_is_none_not_empty_string():
    path = _build_case_list_docx([[
        ("Name of client", ""),
        ("Re", "Something"),
    ]])
    result = parse_case_list_docx(path)
    assert result["records"][0]["client_name"] is None


# ── _combine_matter_description ──────────────────────────────────────────

def test_combine_both_present_and_different():
    assert _combine_matter_description("Atlas", "Contract") == "Atlas — Contract"


def test_combine_re_equals_law_collapses_to_one():
    assert _combine_matter_description("Trust", "Trust") == "Trust"
    assert _combine_matter_description("Trust", "TRUST") == "Trust"  # case-insensitive


def test_combine_falls_back_to_whichever_is_present():
    assert _combine_matter_description("Atlas", "") == "Atlas"
    assert _combine_matter_description("", "Debt collection") == "Debt collection"


def test_combine_both_blank_is_none():
    assert _combine_matter_description("", "") is None
    assert _combine_matter_description(None, None) is None


# ── _matter_dedup_score / _best_matter_dedup_match ───────────────────────
# Calibrated against real matter text: the existing DB matter names (from
# the bulk-onboarding endpoint's own example data, itself very likely
# sourced from this document) have more detail appended than the source
# document's bare Re/Law fields, so exact-string matching won't work —
# these are the actual worked examples that set MATTER_DEDUP_THRESHOLD.

def test_dedup_score_true_positive_matches_score_high():
    pairs = [
        ("Mukweva and Paswa Civil",
         "Mukweva and Paswa Civil — Debt collection/fraud, replicated to special plea"),
        ("Mukweva", "Mukweva Criminal — Criminal fraud, following up constitutional application"),
        ("Kuvimba Mining Corporation",
         "Kuvimba Mining Corporation — Debt collection, settlement in place, awaiting monthly payment"),
        ("Atlas", "Atlas — Contract, letter of demand sent, discussing out-of-court settlement"),
        ("Tawanda Chamunorwa", "Tawanda Chamunorwa — Labour, response filed with Conciliator"),
    ]
    for doc_re, existing in pairs:
        score = _matter_dedup_score(doc_re, existing)
        assert score >= 0.75, f"{doc_re!r} vs {existing!r} scored {score}"


def test_dedup_score_different_matters_under_same_client_score_low():
    pairs = [
        ("Cupori and Lewis Tsvere Civil",
         "Mukweva and Paswa Civil — Debt collection/fraud, replicated to special plea"),
        ("Zhang Jian", "Mukweva Criminal — Criminal fraud, following up constitutional application"),
        ("Chakanetsa",
         "Kuvimba Mining Corporation — Debt collection, settlement in place, awaiting monthly payment"),
    ]
    for doc_re, existing in pairs:
        score = _matter_dedup_score(doc_re, existing)
        assert score < 0.75, f"{doc_re!r} vs {existing!r} scored {score} (false positive risk)"


def test_dedup_score_generic_re_terms_do_not_false_positive():
    """"Trust"/"Estate" alone are Law-style category words, not distinctive
    matter identifiers — recur across many genuinely different matters for
    the same client and must not trigger a containment-based false match."""
    pairs = [
        ("Trust", "Family Trust — Sent requirement of Trust to client"),
        ("Estate", "Estate Late Andipa Kashere — Registered the estate"),
    ]
    for doc_re, existing in pairs:
        score = _matter_dedup_score(doc_re, existing)
        assert score < 0.75, f"{doc_re!r} vs {existing!r} scored {score} (generic-term false positive)"


def test_best_matter_dedup_match_picks_highest_scorer_and_reports_source():
    candidates = [
        {"name": "Mukweva Criminal — Criminal fraud", "source": "existing_db"},
        {"name": "Mukweva and Paswa Civil — Debt collection/fraud, replicated to special plea",
         "source": "this_run"},
    ]
    match = _best_matter_dedup_match("Mukweva and Paswa Civil", candidates)
    assert match is not None
    assert match["source"] == "this_run"
    assert match["score"] >= 0.75


def test_best_matter_dedup_match_none_when_nothing_clears_threshold():
    candidates = [{"name": "Zhang Jian — Debt collection", "source": "existing_db"}]
    assert _best_matter_dedup_match("Mukweva and Paswa Civil", candidates) is None


def test_best_matter_dedup_match_none_for_blank_re():
    candidates = [{"name": "Mukweva and Paswa Civil — Debt collection", "source": "existing_db"}]
    assert _best_matter_dedup_match("", candidates) is None
    assert _best_matter_dedup_match(None, candidates) is None


# ── _build_plan (integration, FakeConnection) ────────────────────────────

class FakeConnection:
    def __init__(self, clients=None, matters=None):
        self.clients = clients if clients is not None else []
        self.matters = matters if matters is not None else []

    async def fetch(self, query, *args):
        q = " ".join(query.split())
        if q.startswith("SELECT id, full_name, client_number FROM clients WHERE firm_id=$1"):
            firm_id, = args
            return [dict(c) for c in self.clients if c["firm_id"] == firm_id]
        if q.startswith("SELECT name, matter_number FROM matters WHERE client_id=$1"):
            client_id, = args
            return [dict(m) for m in self.matters if m["client_id"] == client_id]
        raise NotImplementedError(f"FakeConnection.fetch: unhandled query: {q}")


def _client(name, client_number=None):
    return {"id": uuid.uuid4(), "firm_id": FIRM_UUID, "full_name": name, "client_number": client_number}


def _matter(client_id, name, matter_number=None):
    return {"client_id": client_id, "name": name, "matter_number": matter_number}


def _record(table_index, client_name, re=None, law=None, action_done=None,
            next_action=None, latest_communication=None, file_name=None, case_number=None):
    return {
        "table_index": table_index, "client_name": client_name, "re": re, "law": law,
        "action_done": action_done, "next_action": next_action,
        "latest_communication": latest_communication, "file_name": file_name, "case_number": case_number,
    }


def test_new_client_gets_next_available_number_continuing_from_existing():
    conn = FakeConnection(clients=[_client("Existing Client", client_number="NGM-012")])
    records = [_record(0, "Huang Li Qiang", re="Mukweva", law="Debt collection")]

    plan = asyncio.run(_build_plan(conn, "NGM", records))

    assert len(plan["clients"]["created"]) == 1
    assert plan["clients"]["created"][0]["client_number"] == "NGM-013"
    assert plan["matters"]["created"][0]["matter_number"] == "NGM-013-01"


def test_matched_existing_client_is_reused_not_duplicated():
    existing_id = uuid.uuid4()
    conn = FakeConnection(clients=[
        {"id": existing_id, "firm_id": FIRM_UUID, "full_name": "Huang Li Qiang", "client_number": "NGM-001"},
    ])
    records = [_record(0, "Huang Li Qiang", re="Zhang Jian", law="Debt collection")]

    plan = asyncio.run(_build_plan(conn, "NGM", records))

    assert plan["clients"]["created"] == []
    assert len(plan["clients"]["matched"]) == 1
    assert plan["clients"]["matched"][0]["matched_client_id"] == str(existing_id)
    assert plan["matters"]["created"][0]["client_id"] == str(existing_id)
    assert plan["matters"]["created"][0]["matter_number"] == "NGM-001-01"


def test_ambiguous_client_flagged_never_auto_merged_matter_still_created():
    conn = FakeConnection(clients=[
        _client("Kudzai Madzingira", client_number="NGM-005"),
        _client("Kudzai Ndanga", client_number="NGM-006"),
    ])
    records = [_record(0, "Ndanga Kudzai", re="Something", law="Divorce")]

    plan = asyncio.run(_build_plan(conn, "NGM", records))

    assert plan["clients"]["created"] == []
    assert plan["clients"]["matched"] == []
    assert len(plan["clients"]["review"]) == 1
    assert {c["full_name"] for c in plan["clients"]["review"][0]["candidates"]} == \
        {"Kudzai Madzingira", "Kudzai Ndanga"}
    # Matter still created, unlinked — nothing silently dropped.
    assert plan["matters"]["created"] == []
    assert len(plan["matters"]["created_unlinked_pending_review"]) == 1
    assert plan["matters"]["created_unlinked_pending_review"][0]["client_name"] == "Ndanga Kudzai"


def test_matter_skipped_as_duplicate_of_existing_db_matter():
    existing_id = uuid.uuid4()
    conn = FakeConnection(
        clients=[{"id": existing_id, "firm_id": FIRM_UUID, "full_name": "Huang Li Qiang", "client_number": "NGM-001"}],
        matters=[_matter(existing_id, "Mukweva and Paswa Civil — Debt collection/fraud, replicated to special plea",
                          "NGM-001-01")],
    )
    records = [_record(0, "Huang Li Qiang", re="Mukweva and Paswa Civil", law="Debt collection/ fraud")]

    plan = asyncio.run(_build_plan(conn, "NGM", records))

    assert plan["matters"]["created"] == []
    assert len(plan["matters"]["skipped_duplicate"]) == 1
    skip = plan["matters"]["skipped_duplicate"][0]
    assert skip["matched_source"] == "existing_db"
    assert skip["score"] >= 0.75


def test_genuinely_new_matter_for_existing_client_is_created_and_numbered():
    existing_id = uuid.uuid4()
    conn = FakeConnection(
        clients=[{"id": existing_id, "firm_id": FIRM_UUID, "full_name": "Huang Li Qiang", "client_number": "NGM-001"}],
        matters=[_matter(existing_id, "Mukweva and Paswa Civil — Debt collection/fraud", "NGM-001-01")],
    )
    records = [_record(0, "Huang Li Qiang", re="Kuvimba Mining Corporation", law="Debt collection mining")]

    plan = asyncio.run(_build_plan(conn, "NGM", records))

    assert len(plan["matters"]["created"]) == 1
    assert plan["matters"]["created"][0]["matter_name"] == "Kuvimba Mining Corporation — Debt collection mining"
    assert plan["matters"]["created"][0]["matter_number"] == "NGM-001-02"
    assert plan["matters"]["skipped_duplicate"] == []


def test_matter_skipped_as_duplicate_of_matter_created_earlier_in_same_run():
    """Two back-to-back document records for the same client that are
    really the same matter entered twice (as seen in the real document —
    Chitate Hebert's two "Trust"/"Trust" entries)."""
    conn = FakeConnection()
    records = [
        _record(0, "Chitate Hebert", re="Trust", law="Trust", action_done="Client gave instructions"),
        _record(1, "Chitate Hebert", re="Trust", law="Trust"),
    ]

    plan = asyncio.run(_build_plan(conn, "NGM", records))

    assert len(plan["clients"]["created"]) == 1  # one client, not two
    assert len(plan["matters"]["created"]) == 1
    assert len(plan["matters"]["skipped_duplicate"]) == 1
    assert plan["matters"]["skipped_duplicate"][0]["matched_source"] == "this_run"


def test_repeated_client_name_within_document_reuses_same_new_client():
    """A batch of records for a brand-new client (not yet in the DB) must
    all resolve to the SAME newly-created client, not one each."""
    conn = FakeConnection()
    records = [
        _record(0, "Vengesai Enterprises", re="Chinhende", law="Land and Conveyancing"),
        _record(1, "Vengsai Enterprices", re="Hatina wedu", law="Land and Conveyancing"),  # typo variant
        _record(2, "Vengesai Enterprices", re="Tototo", law="Land and Conveyancing"),      # another variant
    ]

    plan = asyncio.run(_build_plan(conn, "NGM", records))

    assert len(plan["clients"]["created"]) == 1
    assert len(plan["matters"]["created"]) == 3
    assert len({m["client_number"] for m in plan["matters"]["created"]}) == 1
    assert [m["matter_number"] for m in plan["matters"]["created"]] == ["NGM-001-01", "NGM-001-02", "NGM-001-03"]

    # A repeat mention with genuinely different text (typo variants) is
    # worth surfacing for sanity-checking, even though it's a single
    # confident match, not an ambiguous one.
    assert len(plan["clients"]["fuzzy_merged_within_run"]) == 2
    merged_names = {e["name"] for e in plan["clients"]["fuzzy_merged_within_run"]}
    assert merged_names == {"Vengsai Enterprices", "Vengesai Enterprices"}
    assert all(e["merged_into_name"] == "Vengesai Enterprises"
               for e in plan["clients"]["fuzzy_merged_within_run"])


def test_exact_repeat_client_name_is_not_reported_as_fuzzy_merge():
    """A trivial exact repeat (e.g. "Huang Li Qiang" named identically
    across many records) isn't a resolution event worth a report line —
    only when the text actually differs."""
    conn = FakeConnection()
    records = [
        _record(0, "Huang Li Qiang", re="Mukweva", law="Debt collection"),
        _record(1, "Huang Li Qiang", re="Zhang Jian", law="Debt collection"),
    ]

    plan = asyncio.run(_build_plan(conn, "NGM", records))

    assert plan["clients"]["fuzzy_merged_within_run"] == []
    assert plan["clients"]["matched"] == []  # not a DB match either — created once, reused once


def test_client_only_record_no_case_detail_creates_client_not_matter():
    conn = FakeConnection()
    records = [_record(0, "RBZ")]  # client name only, everything else blank

    plan = asyncio.run(_build_plan(conn, "NGM", records))

    assert len(plan["clients"]["created"]) == 1
    assert plan["matters"]["created"] == []
    assert plan["matters"]["skipped_duplicate"] == []
    assert len(plan["matters"]["client_only_no_matter"]) == 1
    assert plan["matters"]["client_only_no_matter"][0]["client_name"] == "RBZ"


def test_unparseable_row_no_client_name_is_skipped_and_listed():
    conn = FakeConnection()
    records = [_record(0, None, re="Something", law="Something else")]

    plan = asyncio.run(_build_plan(conn, "NGM", records))

    assert plan["clients"]["created"] == []
    assert plan["matters"]["created"] == []
    assert len(plan["unparseable_rows"]) == 1
    assert plan["unparseable_rows"][0]["table_index"] == 0


def test_no_matter_name_but_has_other_case_detail_still_creates_matter():
    """Re and Law both blank, but there's real narrative content (Action
    done etc.) — must not be silently dropped just because there's no
    name for it."""
    conn = FakeConnection()
    records = [_record(0, "Some Client", action_done="Did something", next_action="Follow up")]

    plan = asyncio.run(_build_plan(conn, "NGM", records))

    assert len(plan["matters"]["created"]) == 1
    assert plan["matters"]["created"][0]["matter_name"] == "Untitled matter"
