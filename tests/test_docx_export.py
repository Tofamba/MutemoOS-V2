"""
Unit tests for backend/docx_export.py and the /api/export-document-docx
endpoint.

Covers the actual bug reported: every input line — including blank ones —
was getting its own margined DOCX Paragraph, so a run of N blank lines
between two real paragraphs produced N+1 stacked, separately-spaced
paragraphs instead of one clean gap. The fix collapses a run of one or
more blank lines into a single paragraph boundary.
"""

import asyncio
import io

import pytest
from docx import Document as DocxReader

from backend.docx_export import (
    build_docx_bytes,
    paragraphs_from_html,
    paragraphs_from_plain_text,
)

SAMPLE_TEXT = (
    "IN THE HIGH COURT OF ZIMBABWE\n"
    "\n\n\n"  # three blank lines — must still be ONE paragraph boundary
    "CASE NO. HC 1234/26\n"
    "\n"
    "The plaintiff avers as follows:\n"
    "\n\n"  # two blank lines — must still be ONE paragraph boundary
    "WHEREFORE plaintiff prays for judgment."
)


# ── plain-text parsing ───────────────────────────────────────────────────────

def test_blank_line_runs_collapse_to_one_paragraph_boundary():
    blocks = paragraphs_from_plain_text(SAMPLE_TEXT)
    # 4 logical paragraphs in the sample, regardless of 1 vs 2 vs 3 blank lines between them
    assert len(blocks) == 4
    texts = [seg["text"] for segments, _ in blocks for seg in segments if "text" in seg]
    assert texts == [
        "IN THE HIGH COURT OF ZIMBABWE",
        "CASE NO. HC 1234/26",
        "The plaintiff avers as follows:",
        "WHEREFORE plaintiff prays for judgment.",
    ]


def test_more_blank_lines_does_not_add_more_paragraphs():
    """The actual regression check: 1, 3, and 7 BLANK LINES (i.e. 2, 4, 8
    newlines — a single \\n alone is just a line ending, not a blank line)
    between the same two paragraphs must all produce exactly 2 paragraphs,
    not 1 + (blank line count)."""
    for blank_line_count in (1, 3, 7):
        blank_run = "\n" * (blank_line_count + 1)
        text = f"First paragraph.{blank_run}Second paragraph."
        blocks = paragraphs_from_plain_text(text)
        assert len(blocks) == 2, f"expected 2 paragraphs for {blank_line_count} blank line(s), got {len(blocks)}"


def test_heading_lines_detected_and_centered():
    blocks = paragraphs_from_plain_text(SAMPLE_TEXT)
    heading_block = blocks[0]
    segments, alignment = heading_block
    assert alignment == "center"
    assert segments[0]["bold"] is True


def test_body_paragraph_is_justified_not_bold():
    blocks = paragraphs_from_plain_text(SAMPLE_TEXT)
    body_block = blocks[2]  # "The plaintiff avers as follows:"
    segments, alignment = body_block
    assert alignment == "justify"
    assert segments[0]["bold"] is False


def test_single_newline_within_a_block_becomes_a_break_not_a_new_paragraph():
    text = "123 Main Street\nHarare\nZimbabwe"
    blocks = paragraphs_from_plain_text(text)
    assert len(blocks) == 1
    segments, _ = blocks[0]
    break_count = sum(1 for s in segments if s.get("break"))
    assert break_count == 2  # 3 lines -> 2 internal breaks, still one paragraph


def test_empty_text_produces_no_paragraphs():
    assert paragraphs_from_plain_text("") == []
    assert paragraphs_from_plain_text("\n\n\n") == []


# ── HTML parsing ──────────────────────────────────────────────────────────────

def test_html_paragraphs_map_one_to_one_regardless_of_source_whitespace():
    """The editor's own HTML never has 'blank <p>' runs the way raw AI text
    can have blank lines — but confirm whitespace-only <p><br></p> blocks
    (which generateDocument() no longer emits, but old saved drafts might
    still contain) don't produce a spurious paragraph."""
    html = (
        '<p style="text-align:center"><strong>IN THE HIGH COURT</strong></p>'
        '<p><br></p>'
        '<p>The plaintiff avers as follows.</p>'
    )
    blocks = paragraphs_from_html(html)
    assert len(blocks) == 2  # the empty <p><br></p> block is dropped, not counted


def test_html_bold_italic_underline_detected():
    html = '<p>Normal <strong>bold</strong> and <em>italic</em> and <u>underlined</u> text.</p>'
    blocks = paragraphs_from_html(html)
    assert len(blocks) == 1
    segments, alignment = blocks[0]
    assert alignment == "justify"
    bold_segments = [s for s in segments if s.get("bold")]
    italic_segments = [s for s in segments if s.get("italic")]
    underline_segments = [s for s in segments if s.get("underline")]
    assert any(s["text"] == "bold" for s in bold_segments)
    assert any(s["text"] == "italic" for s in italic_segments)
    assert any(s["text"] == "underlined" for s in underline_segments)


def test_html_centered_paragraph_detected_via_style_attribute():
    html = '<p style="text-align:center">Centered caption</p>'
    blocks = paragraphs_from_html(html)
    assert blocks[0][1] == "center"


def test_html_br_within_paragraph_is_a_break_not_new_paragraph():
    html = '<p>Line one<br>Line two<br>Line three</p>'
    blocks = paragraphs_from_html(html)
    assert len(blocks) == 1
    segments, _ = blocks[0]
    break_count = sum(1 for s in segments if s.get("break"))
    assert break_count == 2


def test_empty_html_produces_no_paragraphs():
    assert paragraphs_from_html("") == []
    assert paragraphs_from_html("<p></p><p><br></p>") == []


# ── DOCX byte generation — the actual .docx paragraph count, read back ──────

def test_docx_paragraph_count_matches_logical_paragraphs_not_blank_lines():
    blocks = paragraphs_from_plain_text(SAMPLE_TEXT)
    docx_bytes = build_docx_bytes(blocks)
    assert len(docx_bytes) > 0

    doc = DocxReader(io.BytesIO(docx_bytes))
    non_empty_paragraphs = [p for p in doc.paragraphs if p.text.strip()]
    assert len(non_empty_paragraphs) == 4  # matches len(blocks), NOT the blank-line count in SAMPLE_TEXT


def test_docx_space_after_is_consistent_regardless_of_blank_line_count():
    """Every paragraph gets the same space_after (6pt) — spacing doesn't
    compound based on how many blank lines separated it from the next
    paragraph in the source text, since blank-line runs no longer produce
    their own paragraphs at all."""
    from docx.shared import Pt

    blocks = paragraphs_from_plain_text(SAMPLE_TEXT)
    docx_bytes = build_docx_bytes(blocks)
    doc = DocxReader(io.BytesIO(docx_bytes))
    non_empty_paragraphs = [p for p in doc.paragraphs if p.text.strip()]
    for p in non_empty_paragraphs:
        assert p.paragraph_format.space_after == Pt(6)


def test_build_docx_bytes_empty_input_still_produces_valid_docx():
    docx_bytes = build_docx_bytes([])
    assert len(docx_bytes) > 0
    doc = DocxReader(io.BytesIO(docx_bytes))  # must not raise
    assert doc.paragraphs == [] or all(not p.text.strip() for p in doc.paragraphs)


# ── /api/export-document-docx endpoint ──────────────────────────────────────
# Called directly rather than through FastAPI's TestClient, matching this
# repo's existing test convention (plain function imports, no HTTP layer) —
# AUTH_ENABLED is False by default with no OTP env vars configured, so
# get_current_user() never touches the `request` argument, and
# _check_permission() is a pure role check with no DB access.

def test_export_document_docx_endpoint_returns_valid_nonempty_docx():
    from backend.main import ExportDocumentDocxRequest, export_document_docx

    req = ExportDocumentDocxRequest(
        content_html=(
            '<p style="text-align:center"><strong>LAST WILL AND TESTAMENT</strong></p>'
            '<p>I, the undersigned, being of sound mind, declare this to be my will.</p>'
            '<p><strong>1.</strong> I revoke all previous wills.</p>'
        ),
        filename="Test Will",
    )
    response = asyncio.run(export_document_docx(req, None))

    assert response.status_code == 200
    assert response.media_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    assert 'filename="Test Will.docx"' in response.headers["content-disposition"]
    assert len(response.body) > 0

    doc = DocxReader(io.BytesIO(response.body))  # must parse as a real docx
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "LAST WILL AND TESTAMENT" in full_text
    assert "I revoke all previous wills." in full_text


def test_export_document_docx_rejects_empty_content():
    from fastapi import HTTPException

    from backend.main import ExportDocumentDocxRequest, export_document_docx

    req = ExportDocumentDocxRequest(content_html="<p></p><p><br></p>", filename="Empty")
    with pytest.raises(HTTPException) as exc_info:
        asyncio.run(export_document_docx(req, None))
    assert exc_info.value.status_code == 422
