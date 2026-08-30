"""
Mutemo Desk — Zimbabwe Legal Practice Operating System
FastAPI backend v2.0 — Production
Changes from v1.1:
  - PostgreSQL persistence (asyncpg) replaces JSON file store
  - firm_id on every data model — multi-tenancy foundation
  - Role-based access control: partner | associate | secretary | admin
  - Background OCR via FastAPI BackgroundTasks (all three upload endpoints)
  - OTP/session data migrated to DB (no more in-memory dicts)
  - Admin endpoints protected by X-Admin-Token header
"""

from fastapi import FastAPI, HTTPException, UploadFile, File, Form, Request, Response, BackgroundTasks, Depends
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, JSONResponse, StreamingResponse, Response
from pydantic import BaseModel
from typing import Optional, List
from contextlib import asynccontextmanager
import anthropic
import difflib
import subprocess
import tempfile
import os
import json
import uuid
import re
import asyncio
import secrets
import time
import hmac
from datetime import datetime, timedelta, date
from enum import Enum
from backend.grounding import compute_grounding, format_context, TEXTURE_RULES, apply_confidence_safeguard, display_label, FACT_EXTRACTION_RULES, LAWYER_JUDGMENT_RULES, STATUTORY_MECHANISM_PRECISION, IRAC_STRUCTURE_RULES, verify_citations, verify_inline_case_citations, enforce_confidence_consistency, run_legal_research_agent
from backend.client_migration import match_client_name
from backend.numbering import (
    generate_initials, disambiguate_initials, next_sequence,
    format_client_number, format_matter_number,
)
from backend.case_binder import provision_case_binder
from backend.practice_areas import PRACTICE_AREAS, classify_practice_area, extract_classification_text
from backend.conveyancing import CONVEYANCING_MILESTONES
from backend.matter_stages import resolve_stage_sequence, stage_storage_field
from backend.deadline_engine import try_compute_deadline
from backend.legal_taxonomy import classify_firm_document, classify_legal_update, classify_zlr_entry, authority_strength_for
from backend.authority_ranker import rerank
from backend.docx_export import paragraphs_from_plain_text, paragraphs_from_html, build_docx_bytes

# ── R2 / S3-compatible object storage ─────────────────────────────────────────
# R2 is optional — a deployment (e.g. a staging environment, or a future
# firm that doesn't need document storage in R2 yet) may have boto3
# installed but no R2_* vars set. boto3.client() raises ValueError
# immediately on an empty endpoint_url rather than failing lazily on first
# use, so the client must only be constructed once R2_ENABLED is already
# known true — constructing it unconditionally crashed the whole app at
# import time on any deployment missing R2 config.
try:
    import boto3
    from botocore.exceptions import ClientError
    R2_BUCKET = os.environ.get("R2_BUCKET", "mutemoos-documents")
    R2_ENABLED = bool(os.environ.get("R2_ENDPOINT") and os.environ.get("R2_ACCESS_KEY_ID"))
    if R2_ENABLED:
        _r2_client = boto3.client(
            "s3",
            endpoint_url=os.environ.get("R2_ENDPOINT", ""),
            aws_access_key_id=os.environ.get("R2_ACCESS_KEY_ID", ""),
            aws_secret_access_key=os.environ.get("R2_SECRET_ACCESS_KEY", ""),
            region_name="auto",
        )
        print(f"[r2] R2 storage enabled — bucket: {R2_BUCKET}")
    else:
        _r2_client = None
        print("[r2] R2 not configured — file storage disabled")
except ImportError:
    _r2_client = None
    R2_BUCKET = ""
    R2_ENABLED = False
    print("[r2] boto3 not installed — R2 storage disabled")

# ── Load .env file if present ─────────────────────────────────────────────────
def _load_dotenv():
    env_path = os.path.join(os.path.dirname(__file__), "..", ".env")
    if not os.path.exists(env_path):
        return
    with open(env_path) as f:
        for line in f:
            line = line.strip()
            if not line or line.startswith("#") or "=" not in line:
                continue
            key, _, value = line.partition("=")
            key, value = key.strip(), value.strip()
            if key and key not in os.environ:
                os.environ[key] = value

_load_dotenv()

# ── Database ──────────────────────────────────────────────────────────────────
import asyncpg

_db_pool: asyncpg.Pool = None

async def get_db() -> asyncpg.Pool:
    return _db_pool

async def init_db():
    """Create connection pool and run schema migrations."""
    global _db_pool
    database_url = os.environ.get("DATABASE_URL")
    if not database_url:
        raise RuntimeError("DATABASE_URL environment variable is not set. Add the Railway PostgreSQL plugin.")
    # asyncpg requires postgresql:// not postgres://
    if database_url.startswith("postgres://"):
        database_url = database_url.replace("postgres://", "postgresql://", 1)
    _db_pool = await asyncpg.create_pool(database_url, min_size=2, max_size=10)
    await run_migrations()
    print("[db] PostgreSQL connection pool ready")

async def run_migrations():
    """Idempotent schema creation — safe to run on every startup."""
    async with _db_pool.acquire() as conn:
        await conn.execute("""
        CREATE EXTENSION IF NOT EXISTS pgcrypto;

        CREATE TABLE IF NOT EXISTS firms (
            id          UUID PRIMARY KEY DEFAULT gen_random_uuid(),
            name        TEXT NOT NULL,
            short_name  TEXT,
            city        TEXT DEFAULT 'Harare',
            country     TEXT DEFAULT 'Zimbabwe',
            created_at  TIMESTAMPTZ NOT NULL DEFAULT NOW()
        );

        CREATE TABLE IF NOT EXISTS users (
            id              UUID PRIMARY KEY DEFAULT gen_random_uuid(),
            firm_id         UUID NOT NULL REFERENCES firms(id) ON DELETE CASCADE,
            phone           TEXT NOT NULL,
            email           TEXT,
            display_name    TEXT NOT NULL,
            role            TEXT NOT NULL CHECK (role IN ('partner','associate','secretary','admin')),
            is_active       BOOLEAN NOT NULL DEFAULT TRUE,
            created_at      TIMESTAMPTZ NOT NULL DEFAULT NOW(),
            UNIQUE (firm_id, phone)
        );
        ALTER TABLE users ADD COLUMN IF NOT EXISTS email TEXT;
        -- Prefix used in client_number/matter_number (backend/numbering.py),
        -- e.g. "NGM" for Nyaradzo Gilbertina Maphosa. Nullable — assigned
        -- lazily (auto-generated from display_name, with disambiguation on
        -- collision) the first time a user without one creates a client.
        ALTER TABLE users ADD COLUMN IF NOT EXISTS initials TEXT;
        CREATE UNIQUE INDEX IF NOT EXISTS idx_users_firm_initials
            ON users(firm_id, initials) WHERE initials IS NOT NULL;
        -- Fixed prefixes for the 5 partners, confirmed 2026-08-07 — set once,
        -- only if not already set, so this is safe to leave running on every
        -- startup. Anyone else (existing or new) gets initials auto-generated
        -- from their name on first use instead.
        UPDATE users SET initials='NGM' WHERE firm_id=(SELECT id FROM firms LIMIT 1) AND initials IS NULL AND lower(display_name)=lower('Nyaradzo Gilbertina Maphosa');
        UPDATE users SET initials='OM'  WHERE firm_id=(SELECT id FROM firms LIMIT 1) AND initials IS NULL AND lower(display_name)=lower('Ostern Mutero');
        UPDATE users SET initials='JRT' WHERE firm_id=(SELECT id FROM firms LIMIT 1) AND initials IS NULL AND lower(display_name)=lower('Jingini R. Tsivama');
        UPDATE users SET initials='HPM' WHERE firm_id=(SELECT id FROM firms LIMIT 1) AND initials IS NULL AND lower(display_name)=lower('Honour P Mkushi');
        UPDATE users SET initials='FS'  WHERE firm_id=(SELECT id FROM firms LIMIT 1) AND initials IS NULL AND lower(display_name)=lower('Farai Siyakurima');

        CREATE TABLE IF NOT EXISTS sessions (
            token       TEXT PRIMARY KEY,
            user_id     UUID NOT NULL REFERENCES users(id) ON DELETE CASCADE,
            firm_id     UUID NOT NULL REFERENCES firms(id) ON DELETE CASCADE,
            expires_at  TIMESTAMPTZ NOT NULL,
            created_at  TIMESTAMPTZ NOT NULL DEFAULT NOW()
        );
        CREATE INDEX IF NOT EXISTS idx_sessions_expires ON sessions(expires_at);
        -- Shared-device session hardening (2026-08-27) -- see
        -- SESSION_IDLE_TIMEOUT_SECONDS's own comment for why this exists
        -- alongside expires_at rather than replacing it. DEFAULT NOW()
        -- backfills every already-live session as "just active now" the
        -- moment this migration runs, rather than immediately idle-killing
        -- every existing logged-in user on deploy.
        ALTER TABLE sessions ADD COLUMN IF NOT EXISTS last_active TIMESTAMPTZ NOT NULL DEFAULT NOW();
        CREATE INDEX IF NOT EXISTS idx_sessions_last_active ON sessions(last_active);

        CREATE TABLE IF NOT EXISTS otp_store (
            phone       TEXT PRIMARY KEY,
            code        TEXT NOT NULL,
            attempts    INT NOT NULL DEFAULT 0,
            expires_at  TIMESTAMPTZ NOT NULL
        );
        -- Defense-in-depth (multi-tenancy hardening, Part 2) -- not
        -- currently exploitable: this deployment only ever has one firm
        -- (Option B: one deployment per firm), so a phone-number
        -- collision across firms can't happen in practice today. Added
        -- anyway so a query against this table is never the one place
        -- that silently assumes single-firm. phone stays the PRIMARY KEY
        -- (unchanged) -- this is an additive column, not a PK migration.
        -- Backfilled below (needs a bound FIRM_ID param -- this block has
        -- none) rather than here.
        ALTER TABLE otp_store ADD COLUMN IF NOT EXISTS firm_id UUID REFERENCES firms(id) ON DELETE CASCADE;

        CREATE TABLE IF NOT EXISTS matters (
            id              UUID PRIMARY KEY DEFAULT gen_random_uuid(),
            firm_id         UUID NOT NULL REFERENCES firms(id) ON DELETE CASCADE,
            name            TEXT NOT NULL,
            number          TEXT,
            internal_ref    TEXT,
            external_ref    TEXT,
            client_name     TEXT,
            matter_type     TEXT,
            status          TEXT NOT NULL DEFAULT 'Active',
            custom_status   TEXT,
            document_count  INT NOT NULL DEFAULT 0,
            last_activity   TIMESTAMPTZ,
            created_at      TIMESTAMPTZ NOT NULL DEFAULT NOW(),
            created_by      UUID REFERENCES users(id),
            next_deadline       DATE,
            next_deadline_note  TEXT
        );
        ALTER TABLE matters ADD COLUMN IF NOT EXISTS next_deadline DATE;
        ALTER TABLE matters ADD COLUMN IF NOT EXISTS next_deadline_note TEXT;

        -- Matter review safety net (2026-08-30): every matter gets a soft
        -- "please look at this" nudge date, distinct from next_deadline
        -- above (a hard court/filing deadline). Modeled on a real prior
        -- paper workflow -- writing a review date on a physical file
        -- folder -- so nothing goes untouched indefinitely without at
        -- least surfacing somewhere. See DEFAULT_REVIEW_INTERVAL_DAYS
        -- and _create_matter_row()/update_matter() for the actual
        -- defaulting behavior; this column alone has no default because
        -- it's set explicitly by application code on every create/update,
        -- not implicitly by the database.
        ALTER TABLE matters ADD COLUMN IF NOT EXISTS next_review_date DATE;
        -- Companion to next_review_date: when the matter was actually
        -- last looked at, stamped to NOW() by _resolve_review_dates()
        -- on every touch. Kept as a real, separate column rather than
        -- derived by subtracting DEFAULT_REVIEW_INTERVAL_DAYS from
        -- next_review_date, since that math breaks the moment a lawyer
        -- overrides next_review_date to something further out (e.g. a
        -- matter awaiting a court date, reviewed today but not due
        -- again for 6 months).
        ALTER TABLE matters ADD COLUMN IF NOT EXISTS last_reviewed_date DATE;
        CREATE INDEX IF NOT EXISTS idx_matters_firm ON matters(firm_id);
        CREATE INDEX IF NOT EXISTS idx_matters_status ON matters(firm_id, status);

        -- Sentinel matter for documents that aren't tied to a specific
        -- client matter (e.g. Rapid Precedent Capture's "General / Firm
        -- Precedents" option). documents.matter_id is NOT NULL with a hard
        -- FK -- rather than loosen that constraint (which would ripple
        -- into every join/document_count increment across the app), every
        -- firm gets one real matter row to attach general captures to.
        -- Identified by `number='GENERAL'`, not by name, so a firm
        -- renaming it doesn't break the lookup in the capture endpoint.
        --
        -- is_sentinel marks it as NOT a real client matter: list_matters(),
        -- the practice-area breakdown, and the inactivity digest all
        -- exclude it explicitly (see those queries) so it never inflates a
        -- matter count, shows up in an "active matters" list, or gets
        -- flagged as a stale/inactive matter in a firm-wide alert. The RBZ
        -- compliance export needs no such exclusion -- it's driven from
        -- clients LEFT JOIN matters, and the sentinel has no client_id, so
        -- it structurally cannot appear there regardless of this flag.
        ALTER TABLE matters ADD COLUMN IF NOT EXISTS is_sentinel BOOLEAN NOT NULL DEFAULT FALSE;
        INSERT INTO matters (firm_id, name, number, status, is_sentinel)
        SELECT id, 'General / Firm Precedents', 'GENERAL', 'Active', TRUE
        FROM firms
        WHERE NOT EXISTS (
            SELECT 1 FROM matters WHERE matters.firm_id = firms.id AND matters.number = 'GENERAL'
        );

        -- ── Clients ──────────────────────────────────────────────────────
        -- Separate from matters (one client can have many matters over time).
        -- client_id is nullable on matters until the standalone migration
        -- script (scripts/migrate_clients.py) backfills it for existing rows
        -- from the legacy client_name text column — see that script for why
        -- this can't just be done automatically at startup.
        CREATE TABLE IF NOT EXISTS clients (
            id                          UUID PRIMARY KEY DEFAULT gen_random_uuid(),
            firm_id                     UUID NOT NULL REFERENCES firms(id) ON DELETE CASCADE,
            full_name                   TEXT NOT NULL,
            email                       TEXT,
            phone                       TEXT,
            physical_address            TEXT,
            id_or_registration_number   TEXT,
            contact_person              TEXT,
            notes                       TEXT,
            created_at                  TIMESTAMPTZ NOT NULL DEFAULT NOW(),
            updated_at                  TIMESTAMPTZ NOT NULL DEFAULT NOW()
        );
        CREATE INDEX IF NOT EXISTS idx_clients_firm ON clients(firm_id);
        CREATE INDEX IF NOT EXISTS idx_clients_name_lower ON clients(firm_id, lower(full_name));
        -- Corporate/entity clients only — who to actually contact there.
        -- Left blank for individual clients (the client themselves is the
        -- contact).
        ALTER TABLE clients ADD COLUMN IF NOT EXISTS contact_person TEXT;
        -- Automatic numbering (backend/numbering.py): "{lawyer initials}-{seq:03d}",
        -- e.g. NGM-007. Assigned once at client creation from the creating/
        -- uploading lawyer's initials — immutable afterward. NULL for
        -- clients created before this feature, until backfilled (see
        -- scripts/backfill_client_matter_numbers.py).
        ALTER TABLE clients ADD COLUMN IF NOT EXISTS client_number TEXT;
        CREATE UNIQUE INDEX IF NOT EXISTS idx_clients_firm_number
            ON clients(firm_id, client_number) WHERE client_number IS NOT NULL;
        -- Atomic allocator backing _allocate_next_seq() (see DB helpers below) --
        -- one row per (firm_id, prefix), where prefix is either a lawyer's
        -- initials (client numbering) or a full client_number (matter
        -- numbering under that client) -- same generic prefix concept
        -- next_sequence() in backend/numbering.py already uses. Replaces the
        -- old MAX(existing)+1 scan, which raced under concurrent requests:
        -- two simultaneous creates could both read the same existing rows,
        -- compute the same next number, and the second INSERT would hit
        -- idx_clients_firm_number/idx_matters_firm_number with an unhandled
        -- UniqueViolationError instead of just getting a different number.
        CREATE TABLE IF NOT EXISTS numbering_counters (
            firm_id  UUID NOT NULL,
            prefix   TEXT NOT NULL,
            next_seq INT NOT NULL,
            PRIMARY KEY (firm_id, prefix)
        );
        -- Which lawyer created this client — the "My Clients" default-view
        -- filter's actual join key (client_number's initials prefix is a
        -- point-in-time formatted string, not reliably reversible to a
        -- specific users.id on collision/disambiguation, so it's not used
        -- for this). NULL for clients created before this column existed,
        -- or created while AUTH_ENABLED is False (get_current_user()
        -- returns a synthetic user with no real id in that case) — see
        -- scripts/backfill_client_ownership.py for best-effort backfill.
        ALTER TABLE clients ADD COLUMN IF NOT EXISTS created_by UUID REFERENCES users(id);

        -- client_id: who the firm has an actual client relationship with.
        -- case_parties: free text for everyone else named in the matter
        -- (co-respondents, companies, family members) who isn't a client —
        -- e.g. the client's own company or a relative named as a separate
        -- respondent alongside them.
        ALTER TABLE matters ADD COLUMN IF NOT EXISTS client_id UUID REFERENCES clients(id);
        ALTER TABLE matters ADD COLUMN IF NOT EXISTS case_parties TEXT;
        CREATE INDEX IF NOT EXISTS idx_matters_client ON matters(client_id);
        -- Automatic numbering (backend/numbering.py): "{client_number}-{seq:02d}",
        -- e.g. NGM-007-02 — sequential within that client. Assigned at
        -- matter creation only when the matter has a client_id whose client
        -- already has a client_number; NULL otherwise (unlinked matters,
        -- or a linked client not yet backfilled).
        ALTER TABLE matters ADD COLUMN IF NOT EXISTS matter_number TEXT;
        CREATE UNIQUE INDEX IF NOT EXISTS idx_matters_firm_number
            ON matters(firm_id, matter_number) WHERE matter_number IS NOT NULL;

        -- Fixed practice-area category (backend/practice_areas.py's
        -- PRACTICE_AREAS list) — nullable until backfilled (see
        -- scripts/backfill_practice_areas.py). No DB CHECK constraint,
        -- same convention as matters.status: enforced at the Pydantic/API
        -- layer (MatterCreate/MatterUpdate validators), not the schema.
        ALTER TABLE matters ADD COLUMN IF NOT EXISTS practice_area TEXT;
        CREATE INDEX IF NOT EXISTS idx_matters_practice_area ON matters(firm_id, practice_area);

        -- Fee tracking — the firm's own professional fees, manually entered.
        -- Explicitly NOT trust accounting: Sawyer & Mkushi's accounts
        -- department has its own separate system of record for client
        -- funds held in trust, and this must not become a second source of
        -- truth for that. Nullable, no default.
        ALTER TABLE matters ADD COLUMN IF NOT EXISTS amount_billed NUMERIC(14,2);
        ALTER TABLE matters ADD COLUMN IF NOT EXISTS amount_received NUMERIC(14,2);

        -- Conveyancing-specific fields (backend/conveyancing.py's
        -- CONVEYANCING_MILESTONES list) — only meaningful/shown when
        -- practice_area = 'Conveyancing/Property', but not hard-gated at
        -- write time (a matter's practice_area can be corrected later, and
        -- we don't want to reject or silently drop already-entered data).
        -- conveyancing_purchase_price is a reference fact for reporting
        -- only, not a balance or funds held.
        ALTER TABLE matters ADD COLUMN IF NOT EXISTS conveyancing_milestone TEXT;
        ALTER TABLE matters ADD COLUMN IF NOT EXISTS conveyancing_property_address TEXT;
        ALTER TABLE matters ADD COLUMN IF NOT EXISTS conveyancing_title_deed_number TEXT;
        ALTER TABLE matters ADD COLUMN IF NOT EXISTS conveyancing_purchase_price NUMERIC(14,2);
        ALTER TABLE matters ADD COLUMN IF NOT EXISTS conveyancing_other_conveyancer_contact TEXT;
        ALTER TABLE matters ADD COLUMN IF NOT EXISTS conveyancing_transfer_date DATE;
        ALTER TABLE matters ADD COLUMN IF NOT EXISTS conveyancing_rates_clearance_expiry DATE;
        ALTER TABLE matters ADD COLUMN IF NOT EXISTS conveyancing_bond_registration_deadline DATE;

        -- Matter Progress Tracker (visual stepper) -- generic current-stage
        -- column for matter_type values whose stage sequence isn't backed
        -- by an existing type-specific column. Conveyancing keeps using
        -- conveyancing_milestone above (see backend/matter_stages.py).
        -- stage_updated_at is shared across both storage fields -- set
        -- whenever either `stage` or `conveyancing_milestone` changes --
        -- so "days in this stage" is computed the same way regardless of
        -- which column actually holds the value.
        ALTER TABLE matters ADD COLUMN IF NOT EXISTS stage TEXT;
        ALTER TABLE matters ADD COLUMN IF NOT EXISTS stage_updated_at TIMESTAMPTZ;

        CREATE TABLE IF NOT EXISTS progress_notes (
            id          UUID PRIMARY KEY DEFAULT gen_random_uuid(),
            matter_id   UUID NOT NULL REFERENCES matters(id) ON DELETE CASCADE,
            firm_id     UUID NOT NULL REFERENCES firms(id) ON DELETE CASCADE,
            text        TEXT NOT NULL,
            author      TEXT NOT NULL DEFAULT 'Unknown',
            user_id     UUID REFERENCES users(id),
            created_at  TIMESTAMPTZ NOT NULL DEFAULT NOW()
        );
        CREATE INDEX IF NOT EXISTS idx_notes_matter ON progress_notes(matter_id);

        CREATE TABLE IF NOT EXISTS documents (
            id              UUID PRIMARY KEY DEFAULT gen_random_uuid(),
            matter_id       UUID NOT NULL REFERENCES matters(id) ON DELETE CASCADE,
            firm_id         UUID NOT NULL REFERENCES firms(id) ON DELETE CASCADE,
            filename        TEXT NOT NULL,
            document_type   TEXT,
            matter_type     TEXT,
            parties         TEXT,
            doc_date        DATE,
            court           TEXT,
            word_count      INT DEFAULT 0,
            page_count      INT DEFAULT 1,
            chunk_count     INT DEFAULT 0,
            ocr_used        BOOLEAN DEFAULT FALSE,
            ocr_confidence  FLOAT,
            needs_review    BOOLEAN DEFAULT FALSE,
            status          TEXT DEFAULT 'processing',
            error_message   TEXT,
            uploaded_at     TIMESTAMPTZ NOT NULL DEFAULT NOW(),
            uploaded_by     UUID REFERENCES users(id)
        );
        ALTER TABLE documents ADD COLUMN IF NOT EXISTS ocr_confidence FLOAT;
        ALTER TABLE documents ADD COLUMN IF NOT EXISTS needs_review BOOLEAN DEFAULT FALSE;
        -- Rapid Precedent Capture: distinguishes a phone/camera capture
        -- from a regular file upload, purely so the "Recently captured"
        -- list (GET /api/capture/recent) can filter to just these without
        -- guessing from filename or other heuristics.
        ALTER TABLE documents ADD COLUMN IF NOT EXISTS is_capture BOOLEAN DEFAULT FALSE;
        CREATE INDEX IF NOT EXISTS idx_documents_matter ON documents(matter_id);
        CREATE INDEX IF NOT EXISTS idx_documents_firm ON documents(firm_id);
        -- Legal source classification (backend/legal_taxonomy.py) — nullable,
        -- backfilled by /api/admin/backfill-legal-taxonomy, computed going
        -- forward wherever document_type is set.
        ALTER TABLE documents ADD COLUMN IF NOT EXISTS legal_source_type TEXT;
        ALTER TABLE documents ADD COLUMN IF NOT EXISTS authority_strength TEXT;
        CREATE INDEX IF NOT EXISTS idx_documents_legal_type ON documents(legal_source_type);
        CREATE INDEX IF NOT EXISTS idx_documents_capture ON documents(firm_id, is_capture, uploaded_at DESC);

        -- Document provenance metadata for client/matter (Vault) documents
        -- -- a genuinely separate concern from the two classification
        -- systems already on this table: document_type (AI-classified
        -- specific form, e.g. "affidavit"/"lease_agreement" -- feeds
        -- legal_source_type/authority_ranker.py, untouched here) and
        -- matter_type. Named provenance_document_type specifically to
        -- avoid colliding with the existing document_type column -- reusing
        -- that name would have meant either breaking
        -- classify_firm_document()'s FIRM_DOC_TYPE_MAP (new values like
        -- "Contract" aren't keys in it, so it would silently fall back to
        -- LegalSourceType.UNKNOWN) or overloading one column with two
        -- unrelated taxonomies.
        ALTER TABLE documents ADD COLUMN IF NOT EXISTS provenance_document_type TEXT
            CHECK (provenance_document_type IN (
                'Pleading', 'Contract', 'Correspondence', 'JudgmentOrder',
                'Evidence', 'Research', 'Precedent', 'General'
            ));
        -- document_status: a legal-document lifecycle stage, distinct from
        -- the existing `status` column (a processing-pipeline state:
        -- processing/complete). Defaults to Draft -- a document is a draft
        -- shell until someone actively moves it forward; this applies
        -- equally to a case-binder auto-provisioned starter document
        -- (definitionally untouched until a lawyer edits it) and a fresh
        -- manual upload with no status specified.
        ALTER TABLE documents ADD COLUMN IF NOT EXISTS document_status TEXT DEFAULT 'Draft'
            CHECK (document_status IN ('Draft', 'Review', 'Final', 'Executed', 'Superseded'));
        ALTER TABLE documents ADD COLUMN IF NOT EXISTS description TEXT;
        ALTER TABLE documents ADD COLUMN IF NOT EXISTS confidentiality TEXT DEFAULT 'Standard'
            CHECK (confidentiality IN ('Standard', 'Restricted', 'Privileged'));

        -- ── AML/KYC Client Compliance ───────────────────────────────────────
        -- Money Laundering and Proceeds of Crime Act [Chapter 9:24] --
        -- identity/ownership/PEP tracking only. Deliberately excludes
        -- trust/client-money accounting -- no ledger, no trust balance
        -- tracking; that is a separate concern this module does not touch.
        -- Every field below traces to a specific verified section, cited
        -- inline so the compliance logic stays auditable against its legal
        -- basis rather than free-floating "best practice" fields.
        --
        -- All additive and nullable at the schema level -- existing clients
        -- are not forced to backfill anything on this deploy.

        -- s17(a): identity particulars for a natural-person client -- date
        -- and place of birth required "to the extent not disclosed by the
        -- identity document" (hence nullable, not mandatory-on-write).
        ALTER TABLE clients ADD COLUMN IF NOT EXISTS client_type TEXT
            CHECK (client_type IN (
                'Individual', 'Company', 'Partnership', 'Trust', 'Estate',
                'NonProfit', 'Government', 'Other'
            ));
        ALTER TABLE clients ADD COLUMN IF NOT EXISTS date_of_birth DATE;
        ALTER TABLE clients ADD COLUMN IF NOT EXISTS place_of_birth TEXT;
        ALTER TABLE clients ADD COLUMN IF NOT EXISTS national_id_number TEXT;
        ALTER TABLE clients ADD COLUMN IF NOT EXISTS passport_number TEXT;
        ALTER TABLE clients ADD COLUMN IF NOT EXISTS id_expiry_date DATE;
        ALTER TABLE clients ADD COLUMN IF NOT EXISTS residential_address TEXT;
        ALTER TABLE clients ADD COLUMN IF NOT EXISTS occupation TEXT;
        ALTER TABLE clients ADD COLUMN IF NOT EXISTS employer_or_business TEXT;

        -- s17(a): identity particulars for a company/partnership/other legal
        -- person -- registration and controlling-document details, not just
        -- the free-text id_or_registration_number the client record already
        -- had (kept untouched for backward compatibility; these are the
        -- specific fields s17 actually names).
        ALTER TABLE clients ADD COLUMN IF NOT EXISTS registered_name TEXT;
        ALTER TABLE clients ADD COLUMN IF NOT EXISTS trading_name TEXT;
        ALTER TABLE clients ADD COLUMN IF NOT EXISTS registration_number TEXT;
        ALTER TABLE clients ADD COLUMN IF NOT EXISTS date_incorporated DATE;
        ALTER TABLE clients ADD COLUMN IF NOT EXISTS registered_office_address TEXT;
        ALTER TABLE clients ADD COLUMN IF NOT EXISTS principal_business_address TEXT;
        -- Document links, not file storage -- point at an existing Vault
        -- document (uploaded via /api/upload or /api/capture, both already
        -- built) rather than duplicating document handling here.
        ALTER TABLE clients ADD COLUMN IF NOT EXISTS proof_of_incorporation_document_id UUID REFERENCES documents(id);
        ALTER TABLE clients ADD COLUMN IF NOT EXISTS governing_document_id UUID REFERENCES documents(id);

        -- s17(c): "every trustee, settlor, and beneficiary" for a trust, or
        -- the equivalent for an estate -- a small, structured list per
        -- client rather than a join table, matching the existing
        -- calendar_events.attendees JSONB convention for this shape of data
        -- (see _row_to_client's handling below for the same
        -- str-or-already-decoded-list defensiveness that pattern uses).
        -- Each entry: {name, id_number, role}.
        ALTER TABLE clients ADD COLUMN IF NOT EXISTS trustees JSONB DEFAULT '[]'::jsonb;
        ALTER TABLE clients ADD COLUMN IF NOT EXISTS settlors JSONB DEFAULT '[]'::jsonb;
        ALTER TABLE clients ADD COLUMN IF NOT EXISTS beneficiaries JSONB DEFAULT '[]'::jsonb;

        -- s15 / s17(b): beneficial ownership -- "information necessary to
        -- understand ownership and control." No percentage threshold gate:
        -- ownership_percentage is nullable because not every basis of
        -- ownership/control is percentage-based (e.g. "de facto control via
        -- management agreement"). Multiple rows per client supported --
        -- ownership chains are not artificially capped at one owner.
        CREATE TABLE IF NOT EXISTS beneficial_owners (
            id                          UUID PRIMARY KEY DEFAULT gen_random_uuid(),
            client_id                   UUID NOT NULL REFERENCES clients(id) ON DELETE CASCADE,
            firm_id                     UUID NOT NULL REFERENCES firms(id) ON DELETE CASCADE,
            owner_name                  TEXT NOT NULL,
            date_of_birth               DATE,
            nationality                 TEXT,
            id_or_passport_number       TEXT,
            residential_address         TEXT,
            ownership_or_control_basis  TEXT,
            ownership_percentage        NUMERIC(5,2),
            verification_status         TEXT NOT NULL DEFAULT 'Unverified'
                CHECK (verification_status IN ('Unverified', 'Pending', 'Verified')),
            verified_date               DATE,
            verified_by                 UUID REFERENCES users(id),
            created_at                  TIMESTAMPTZ NOT NULL DEFAULT NOW()
        );
        CREATE INDEX IF NOT EXISTS idx_beneficial_owners_client ON beneficial_owners(client_id);

        -- s17(d): the person actually instructing the firm, when it isn't
        -- the client themselves -- their identity and the basis of their
        -- authority to act. authority_document is a Vault document link,
        -- same pattern as proof_of_incorporation/governing_document above.
        CREATE TABLE IF NOT EXISTS authorized_representatives (
            id                          UUID PRIMARY KEY DEFAULT gen_random_uuid(),
            client_id                   UUID NOT NULL REFERENCES clients(id) ON DELETE CASCADE,
            firm_id                     UUID NOT NULL REFERENCES firms(id) ON DELETE CASCADE,
            full_name                   TEXT NOT NULL,
            position_or_relationship    TEXT,
            id_or_passport_number       TEXT,
            contact_details             TEXT,
            authority_basis             TEXT
                CHECK (authority_basis IN (
                    'PowerOfAttorney', 'BoardResolution', 'Mandate',
                    'LetterOfAdministration', 'LetterOfExecutorship', 'Other'
                )),
            authority_document_id       UUID REFERENCES documents(id),
            verification_status         TEXT NOT NULL DEFAULT 'Unverified'
                CHECK (verification_status IN ('Unverified', 'Pending', 'Verified')),
            verified_date               DATE,
            created_at                  TIMESTAMPTZ NOT NULL DEFAULT NOW()
        );
        CREATE INDEX IF NOT EXISTS idx_authorized_representatives_client ON authorized_representatives(client_id);

        -- s20: politically exposed person screening, plus the overall
        -- compliance picture -- one row per client. is_pep is nullable
        -- (NULL = not yet assessed, distinct from FALSE = assessed and
        -- cleared) since compliance_status (see _compute_compliance_status)
        -- needs to tell "never checked" apart from "checked, not a PEP".
        --
        -- client_is_beneficial_owner backs Part 2's UI prompt ("is the
        -- client the beneficial owner? Yes/No/Unknown") -- lives here
        -- rather than on beneficial_owners itself since it's a client-level
        -- answer, not a per-owner record.
        --
        -- relationship_ended_date / retained_until back s24 (retention) --
        -- flag-only in this pass, no deletion/expiry enforcement.
        -- relationship_ended_date is auto-derived (set when every one of
        -- the client's matters reaches status='Closed', cleared if a
        -- matter reopens -- see _sync_client_relationship_ended below).
        -- retained_until is deliberately NOT auto-computed from
        -- relationship_ended_date: s24 sets a minimum retention period, and
        -- the exact figure was not available to verify at the time this was
        -- built -- left for a compliance officer to enter manually rather
        -- than risk baking in a wrong statutory duration.
        CREATE TABLE IF NOT EXISTS client_compliance (
            id                                      UUID PRIMARY KEY DEFAULT gen_random_uuid(),
            client_id                               UUID NOT NULL UNIQUE REFERENCES clients(id) ON DELETE CASCADE,
            firm_id                                  UUID NOT NULL REFERENCES firms(id) ON DELETE CASCADE,
            identity_verification_status            TEXT NOT NULL DEFAULT 'Unverified'
                CHECK (identity_verification_status IN ('Unverified', 'Pending', 'Verified')),
            client_is_beneficial_owner              TEXT
                CHECK (client_is_beneficial_owner IN ('Yes', 'No', 'Unknown')),
            is_pep                                   BOOLEAN,
            pep_basis                                TEXT
                CHECK (pep_basis IN ('Self', 'BeneficialOwner', 'CloseAssociate', 'NotApplicable')),
            pep_position                             TEXT,
            pep_country                              TEXT,
            senior_management_approval_required      BOOLEAN NOT NULL DEFAULT FALSE,
            senior_management_approved_by            UUID REFERENCES users(id),
            senior_management_approved_date          DATE,
            source_of_wealth                         TEXT,
            source_of_funds                          TEXT,
            enhanced_monitoring_required             BOOLEAN NOT NULL DEFAULT FALSE,
            risk_rating                              TEXT NOT NULL DEFAULT 'NotAssessed'
                CHECK (risk_rating IN ('Low', 'Medium', 'High', 'NotAssessed')),
            relationship_ended_date                  DATE,
            retained_until                           DATE,
            -- A real conflict-of-interest check DOES exist in this codebase
            -- (GET /api/matters/check-conflict, fuzzy name-similarity search
            -- against every existing matter — used today on the New Matter
            -- form). It's a live, on-demand similarity search, not a stored
            -- pass/fail outcome, so it has nothing to persist on its own —
            -- these two columns are what make "reviewed" a real, storable
            -- fact for compliance_status to require, reusing that existing
            -- endpoint (keyed off the client's own name) rather than
            -- building a second conflict-check mechanism.
            conflict_check_reviewed                  BOOLEAN NOT NULL DEFAULT FALSE,
            conflict_check_reviewed_by               UUID REFERENCES users(id),
            conflict_check_reviewed_date             DATE,
            created_at                               TIMESTAMPTZ NOT NULL DEFAULT NOW(),
            updated_at                               TIMESTAMPTZ NOT NULL DEFAULT NOW()
        );
        CREATE INDEX IF NOT EXISTS idx_client_compliance_client ON client_compliance(client_id);

        CREATE TABLE IF NOT EXISTS legal_updates (
            id              UUID PRIMARY KEY DEFAULT gen_random_uuid(),
            firm_id         UUID NOT NULL REFERENCES firms(id) ON DELETE CASCADE,
            filename        TEXT NOT NULL,
            source_type     TEXT,
            source_name     TEXT,
            reference       TEXT,
            document_type   TEXT,
            matter_type     TEXT,
            doc_date        DATE,
            court           TEXT,
            word_count      INT DEFAULT 0,
            chunk_count     INT DEFAULT 0,
            status          TEXT DEFAULT 'processing',
            ocr_used        BOOLEAN DEFAULT FALSE,
            error_message   TEXT,
            uploaded_at     TIMESTAMPTZ NOT NULL DEFAULT NOW(),
            source_url      TEXT,
            scraped_at      TIMESTAMPTZ,
            ocr_confidence  FLOAT,
            needs_review    BOOLEAN DEFAULT FALSE
        );
        ALTER TABLE legal_updates ADD COLUMN IF NOT EXISTS ocr_confidence FLOAT;
        ALTER TABLE legal_updates ADD COLUMN IF NOT EXISTS needs_review BOOLEAN DEFAULT FALSE;
        CREATE INDEX IF NOT EXISTS idx_legal_updates_firm ON legal_updates(firm_id);
        -- Migration: add source_url and scraped_at to existing deployments
        ALTER TABLE legal_updates ADD COLUMN IF NOT EXISTS source_url TEXT;
        ALTER TABLE legal_updates ADD COLUMN IF NOT EXISTS scraped_at TIMESTAMPTZ;
        -- Dedup cleanup: earlier pushes from the feed service could land as
        -- repeated rows for the same URL whenever the feed's own dedup state
        -- got reset (e.g. before a persistent volume was mounted). Clean up
        -- existing duplicates (keep earliest) before adding the constraint
        -- that prevents this going forward.
        DELETE FROM legal_updates a USING legal_updates b
        WHERE a.source_url IS NOT NULL
          AND a.source_url = b.source_url
          AND a.firm_id = b.firm_id
          AND a.uploaded_at > b.uploaded_at;
        CREATE UNIQUE INDEX IF NOT EXISTS idx_legal_updates_dedup
            ON legal_updates(firm_id, source_url) WHERE source_url IS NOT NULL;
        ALTER TABLE legal_updates ADD COLUMN IF NOT EXISTS legal_source_type TEXT;
        ALTER TABLE legal_updates ADD COLUMN IF NOT EXISTS authority_strength TEXT;
        CREATE INDEX IF NOT EXISTS idx_legal_updates_legal_type ON legal_updates(legal_source_type);

        CREATE TABLE IF NOT EXISTS zlr_entries (
            id                  UUID PRIMARY KEY DEFAULT gen_random_uuid(),
            firm_id             UUID NOT NULL REFERENCES firms(id) ON DELETE CASCADE,
            filename            TEXT,
            source              TEXT DEFAULT 'ZLR',
            jurisdiction        TEXT,
            authority_weight    TEXT,
            volume_year         TEXT,
            zimlii_url          TEXT,
            case_name           TEXT,
            citation            TEXT,
            judgment_number     TEXT,
            court               TEXT,
            judge               TEXT,
            case_type           TEXT,
            hearing_date        TEXT,
            judgment_date       TEXT,
            subject_chains      JSONB DEFAULT '[]',
            taxonomy_category   TEXT DEFAULT 'General',
            summary             TEXT,
            raw_text            TEXT,
            word_count          INT DEFAULT 0,
            chunk_count         INT DEFAULT 0,
            ocr_used            BOOLEAN DEFAULT FALSE,
            ocr_confidence      FLOAT,
            needs_review        BOOLEAN DEFAULT FALSE,
            uploaded_at         TIMESTAMPTZ NOT NULL DEFAULT NOW()
        );
        ALTER TABLE zlr_entries ADD COLUMN IF NOT EXISTS ocr_confidence FLOAT;
        ALTER TABLE zlr_entries ADD COLUMN IF NOT EXISTS needs_review BOOLEAN DEFAULT FALSE;
        CREATE INDEX IF NOT EXISTS idx_zlr_firm ON zlr_entries(firm_id);
        CREATE INDEX IF NOT EXISTS idx_zlr_category ON zlr_entries(firm_id, taxonomy_category);
        -- Same dedup cleanup as legal_updates, for the same reason.
        DELETE FROM zlr_entries a USING zlr_entries b
        WHERE a.zimlii_url IS NOT NULL
          AND a.zimlii_url = b.zimlii_url
          AND a.firm_id = b.firm_id
          AND a.uploaded_at > b.uploaded_at;
        CREATE UNIQUE INDEX IF NOT EXISTS idx_zlr_entries_dedup
            ON zlr_entries(firm_id, zimlii_url) WHERE zimlii_url IS NOT NULL;
        -- New column, distinct from the existing authority_weight (which is a
        -- coarse per-series Binding/Persuasive value tied to jurisdiction, and
        -- already displayed in the frontend) — left untouched for compatibility.
        ALTER TABLE zlr_entries ADD COLUMN IF NOT EXISTS legal_source_type TEXT;
        ALTER TABLE zlr_entries ADD COLUMN IF NOT EXISTS authority_strength TEXT;
        CREATE INDEX IF NOT EXISTS idx_zlr_entries_legal_type ON zlr_entries(legal_source_type);

        CREATE TABLE IF NOT EXISTS calendar_events (
            id          UUID PRIMARY KEY DEFAULT gen_random_uuid(),
            firm_id     UUID NOT NULL REFERENCES firms(id) ON DELETE CASCADE,
            matter_id   UUID REFERENCES matters(id) ON DELETE SET NULL,
            title       TEXT NOT NULL,
            date        DATE NOT NULL,
            time        TIME,
            event_type  TEXT DEFAULT 'other',
            court       TEXT,
            matter_name TEXT,
            notes       TEXT,
            source      TEXT DEFAULT 'manual',
            attendees   JSONB DEFAULT '[]'::jsonb,
            sequence    INTEGER NOT NULL DEFAULT 0,
            created_at  TIMESTAMPTZ NOT NULL DEFAULT NOW(),
            created_by  UUID REFERENCES users(id)
        );
        ALTER TABLE calendar_events ADD COLUMN IF NOT EXISTS attendees JSONB DEFAULT '[]'::jsonb;
        ALTER TABLE calendar_events ADD COLUMN IF NOT EXISTS sequence INTEGER NOT NULL DEFAULT 0;
        CREATE INDEX IF NOT EXISTS idx_calendar_firm ON calendar_events(firm_id);
        CREATE INDEX IF NOT EXISTS idx_calendar_date ON calendar_events(firm_id, date);

        CREATE TABLE IF NOT EXISTS reminder_settings (
            firm_id             UUID PRIMARY KEY REFERENCES firms(id) ON DELETE CASCADE,
            enabled             BOOLEAN NOT NULL DEFAULT FALSE,
            recipient_email     TEXT,
            send_hour_utc       INT NOT NULL DEFAULT 5,
            last_run_date       DATE,
            digest_enabled          BOOLEAN NOT NULL DEFAULT FALSE,
            digest_recipient_email  TEXT,
            digest_send_hour_utc    INT NOT NULL DEFAULT 6,
            digest_last_run_date    DATE
        );
        ALTER TABLE reminder_settings ADD COLUMN IF NOT EXISTS digest_enabled BOOLEAN NOT NULL DEFAULT FALSE;
        ALTER TABLE reminder_settings ADD COLUMN IF NOT EXISTS digest_recipient_email TEXT;
        ALTER TABLE reminder_settings ADD COLUMN IF NOT EXISTS digest_send_hour_utc INT NOT NULL DEFAULT 6;
        ALTER TABLE reminder_settings ADD COLUMN IF NOT EXISTS digest_last_run_date DATE;

        CREATE TABLE IF NOT EXISTS chunks (
            id              TEXT PRIMARY KEY,
            firm_id         UUID NOT NULL REFERENCES firms(id) ON DELETE CASCADE,
            document_id     UUID NOT NULL,
            matter_id       TEXT,
            chunk_source    TEXT NOT NULL,
            text            TEXT NOT NULL,
            chunk_index     INT NOT NULL DEFAULT 0,
            page_number     INT DEFAULT 1,
            zlr_item_id     TEXT,
            citation        TEXT,
            case_name       TEXT,
            taxonomy_category TEXT,
            source_type     TEXT,
            source_name     TEXT,
            reference       TEXT,
            created_at      TIMESTAMPTZ NOT NULL DEFAULT NOW()
        );
        CREATE INDEX IF NOT EXISTS idx_chunks_firm ON chunks(firm_id);
        CREATE INDEX IF NOT EXISTS idx_chunks_document ON chunks(document_id);
        CREATE INDEX IF NOT EXISTS idx_chunks_source ON chunks(firm_id, chunk_source);
        -- reconcile_chroma_index() used to compare ChromaDB against Postgres
        -- by COUNT(*) alone -- equal counts were treated as "in sync" even
        -- if the actual chunk contents differed (e.g. a stale/wrong vector
        -- under the right id, or two chunks silently swapped). content_hash
        -- lets reconciliation compare actual content per chunk_id instead.
        -- A GENERATED column rather than an application-computed value: it's
        -- derived automatically from `text` on every INSERT (all 5 existing
        -- INSERT INTO chunks call sites need no changes at all), and can
        -- never drift from what's actually stored. Postgres's built-in
        -- sha256() isn't marked IMMUTABLE (required for use in a generated
        -- column, even though it's fully deterministic), hence the small
        -- wrapper function.
        CREATE OR REPLACE FUNCTION mutemo_sha256_hex(input TEXT) RETURNS TEXT AS $$
            SELECT encode(sha256(convert_to(input, 'UTF8')), 'hex')
        $$ LANGUAGE SQL IMMUTABLE;
        ALTER TABLE chunks ADD COLUMN IF NOT EXISTS content_hash TEXT
            GENERATED ALWAYS AS (mutemo_sha256_hex(text)) STORED;
        CREATE TABLE IF NOT EXISTS invites (
            id          UUID PRIMARY KEY DEFAULT gen_random_uuid(),
            firm_id     UUID NOT NULL REFERENCES firms(id) ON DELETE CASCADE,
            email       TEXT NOT NULL,
            phone       TEXT,
            display_name TEXT NOT NULL,
            role        TEXT NOT NULL DEFAULT 'associate',
            invited_by  UUID REFERENCES users(id),
            cf_rule_id  TEXT,
            sent_at     TIMESTAMPTZ NOT NULL DEFAULT NOW(),
            accepted_at TIMESTAMPTZ,
            UNIQUE (firm_id, email)
        );
        CREATE INDEX IF NOT EXISTS idx_invites_firm ON invites(firm_id);
        ALTER TABLE invites ADD COLUMN IF NOT EXISTS phone TEXT;
        CREATE UNIQUE INDEX IF NOT EXISTS idx_invites_firm_phone ON invites(firm_id, phone) WHERE phone IS NOT NULL AND accepted_at IS NULL;

        -- ── Legal Corner spec-correct schema additions ─────────────────────────────────
        -- Branding and feature flags on firms
        ALTER TABLE firms ADD COLUMN IF NOT EXISTS firm_logo_url TEXT;
        ALTER TABLE documents ADD COLUMN IF NOT EXISTS r2_key TEXT;
        ALTER TABLE firms ADD COLUMN IF NOT EXISTS features JSONB DEFAULT '[]'::jsonb;

        -- Distinguishes a document a user actually uploaded (source IS
        -- NULL — the existing, default case) from one the system created
        -- on the matter's behalf, e.g. POST /api/onboarding/intake's Case
        -- Binder starter documents (source='auto_provisioned') — so the
        -- Vault/matter document list can eventually show that distinction
        -- rather than an auto-provisioned placeholder looking identical
        -- to something a lawyer actually uploaded.
        ALTER TABLE documents ADD COLUMN IF NOT EXISTS source TEXT;

        -- Idempotency guard for POST /api/onboarding/intake's commit=true
        -- path (a client-generated idempotency_key). Keyed on (firm_id,
        -- key) rather than key alone -- a key is only meaningful scoped to
        -- one firm, and this doubles as the lookup index for the guard's
        -- one query. response_body stores the exact response the first
        -- (real) commit produced, so a detected duplicate can return it
        -- byte-for-byte rather than reprocessing.
        CREATE TABLE IF NOT EXISTS intake_idempotency_keys (
            firm_id       UUID NOT NULL REFERENCES firms(id) ON DELETE CASCADE,
            key           TEXT NOT NULL,
            created_at    TIMESTAMPTZ NOT NULL DEFAULT NOW(),
            response_body JSONB NOT NULL,
            PRIMARY KEY (firm_id, key)
        );

        -- Organisation roles (ops_manager / panel_lawyer) — separate from firm-level role
        CREATE TABLE IF NOT EXISTS organisation_roles (
            id          UUID PRIMARY KEY DEFAULT gen_random_uuid(),
            firm_id     UUID NOT NULL REFERENCES firms(id) ON DELETE CASCADE,
            user_id     UUID NOT NULL REFERENCES users(id) ON DELETE CASCADE,
            role        TEXT NOT NULL CHECK (role IN ('ops_manager', 'panel_lawyer')),
            created_at  TIMESTAMPTZ NOT NULL DEFAULT now(),
            UNIQUE (firm_id, user_id)
        );
        CREATE INDEX IF NOT EXISTS idx_org_roles_firm ON organisation_roles(firm_id);
        CREATE INDEX IF NOT EXISTS idx_org_roles_user ON organisation_roles(user_id);

        -- Matter SLA and assignment fields for Legal Corner workflow
        ALTER TABLE matters ADD COLUMN IF NOT EXISTS assigned_lawyer_id UUID REFERENCES users(id);
        ALTER TABLE matters ADD COLUMN IF NOT EXISTS coverage_tier TEXT;
        ALTER TABLE matters ADD COLUMN IF NOT EXISTS sla_deadline TIMESTAMPTZ;
        ALTER TABLE matters ADD COLUMN IF NOT EXISTS assigned_by_id UUID REFERENCES users(id);
        ALTER TABLE matters ADD COLUMN IF NOT EXISTS service_type TEXT;
        CREATE UNIQUE INDEX IF NOT EXISTS idx_matters_external_ref
            ON matters(firm_id, external_ref) WHERE external_ref IS NOT NULL;

        -- Reassignment audit trail
        CREATE TABLE IF NOT EXISTS matter_reassignments (
            id                  UUID PRIMARY KEY DEFAULT gen_random_uuid(),
            matter_id           UUID NOT NULL REFERENCES matters(id) ON DELETE CASCADE,
            from_lawyer_id      UUID REFERENCES users(id),
            to_lawyer_id        UUID NOT NULL REFERENCES users(id),
            reassigned_by_id    UUID NOT NULL REFERENCES users(id),
            reason              TEXT,
            reassigned_at       TIMESTAMPTZ NOT NULL DEFAULT now()
        );
        CREATE INDEX IF NOT EXISTS idx_reassignments_matter ON matter_reassignments(matter_id);
        -- Defense-in-depth (multi-tenancy hardening, Part 2) -- not
        -- currently exploitable: the one existing query against this
        -- table (v_legal_corner_sla_status below) already correlates via
        -- matter_id to an already firm-scoped matters row. Added so a
        -- future direct query never has to re-derive firm scope through
        -- a join. Backfilled from matters.firm_id further down, after
        -- the matters table exists.
        ALTER TABLE matter_reassignments ADD COLUMN IF NOT EXISTS firm_id UUID REFERENCES firms(id) ON DELETE CASCADE;
        CREATE INDEX IF NOT EXISTS idx_reassignments_firm ON matter_reassignments(firm_id);
        -- No bound parameter needed here (unlike otp_store) -- matters
        -- already carries the right firm_id for every existing row.
        UPDATE matter_reassignments r SET firm_id = m.firm_id
        FROM matters m WHERE m.id = r.matter_id AND r.firm_id IS NULL;

        -- API keys for server-to-server auth (Legal Corner subscriber platform)
        CREATE TABLE IF NOT EXISTS firm_api_keys (
            id          UUID PRIMARY KEY DEFAULT gen_random_uuid(),
            firm_id     UUID NOT NULL REFERENCES firms(id) ON DELETE CASCADE,
            key_hash    TEXT NOT NULL,
            label       TEXT,
            created_at  TIMESTAMPTZ NOT NULL DEFAULT now(),
            revoked_at  TIMESTAMPTZ,
            UNIQUE (firm_id, label)
        );

        -- Read-optimised SLA status view for ops dashboard
        -- Note: u.full_name falls back to display_name since users table uses display_name
        CREATE OR REPLACE VIEW v_legal_corner_sla_status AS
        SELECT
            m.id AS matter_id,
            m.firm_id,
            m.name AS matter_name,
            m.client_name,
            m.assigned_lawyer_id,
            u.display_name AS lawyer_name,
            m.coverage_tier,
            m.service_type,
            m.sla_deadline,
            m.status,
            m.external_ref,
            CASE
                WHEN m.sla_deadline IS NULL THEN NULL
                WHEN m.status = 'complete' THEN false
                WHEN now() > m.sla_deadline THEN true
                ELSE false
            END AS is_overdue,
            (SELECT count(*) FROM matter_reassignments r WHERE r.matter_id = m.id AND r.firm_id = m.firm_id) AS reassignment_count
        FROM matters m
        LEFT JOIN users u ON u.id = m.assigned_lawyer_id;

        -- Migrate invites: add organisation_role column (spec-correct name, no FK)
        ALTER TABLE invites ADD COLUMN IF NOT EXISTS organisation_role TEXT;

        -- Search/audit trail (source-quality grounding rollout)
        CREATE TABLE IF NOT EXISTS audit_logs (
            id          UUID PRIMARY KEY DEFAULT gen_random_uuid(),
            firm_id     UUID NOT NULL REFERENCES firms(id),
            user_id     UUID REFERENCES users(id) ON DELETE SET NULL,
            actor_name  TEXT NOT NULL,
            actor_role  TEXT NOT NULL,
            action      TEXT NOT NULL,
            target_type TEXT NOT NULL,
            target_id   UUID,
            details     JSONB,
            created_at  TIMESTAMPTZ NOT NULL DEFAULT NOW()
        );
        CREATE INDEX IF NOT EXISTS idx_audit_logs_firm ON audit_logs(firm_id);
        CREATE INDEX IF NOT EXISTS idx_audit_logs_target ON audit_logs(target_type, target_id);

        -- RBZ compliance export audit trail — a dedicated table rather than
        -- another audit_logs row: report_history needs structured
        -- client_count/matter_count columns for the history table the
        -- frontend shows partners, not a generic JSONB blob to parse back
        -- out every time that's displayed.
        CREATE TABLE IF NOT EXISTS report_history (
            id                  UUID PRIMARY KEY DEFAULT gen_random_uuid(),
            firm_id             UUID NOT NULL REFERENCES firms(id) ON DELETE CASCADE,
            report_type         TEXT NOT NULL DEFAULT 'rbz_compliance_export',
            generated_by        UUID REFERENCES users(id) ON DELETE SET NULL,
            generated_by_name   TEXT NOT NULL,
            client_count        INT NOT NULL,
            matter_count        INT NOT NULL,
            generated_at        TIMESTAMPTZ NOT NULL DEFAULT NOW()
        );
        CREATE INDEX IF NOT EXISTS idx_report_history_firm ON report_history(firm_id, generated_at DESC);
        """)

        # Seed this deployment's firm row from its own env vars -- NOT
        # another firm's hardcoded details. Previously hardcoded
        # short_name/city/country to "S&M"/"Harare"/"Zimbabwe" (this
        # instance's own values) regardless of MUTEMO_FIRM_* env vars, and
        # silently ignored FIRM_CITY entirely -- confirmed via an actual
        # walkthrough of provisioning a fresh second firm, 2026-08-25.
        # MUTEMO_FIRM_NAME has no safe generic default: it flows straight
        # into every AI-generated document's system prompt via
        # get_firm_identity(). Refuse to start rather than silently seed
        # a blank or guessed firm name.
        if not FIRM_NAME:
            raise RuntimeError(
                "MUTEMO_FIRM_NAME is not set. This seeds the firms table "
                "and feeds every AI-generated document's system prompt via "
                "get_firm_identity() -- refusing to start with a blank "
                "firm name rather than silently seeding wrong data. Set "
                "MUTEMO_FIRM_NAME to this deployment's actual firm name "
                "and redeploy."
            )
        seed_short_name = FIRM_SHORT_NAME or _derive_firm_short_name(FIRM_NAME)
        await conn.execute("""
        INSERT INTO firms (id, name, short_name, city, country)
        VALUES ($1, $2, $3, $4, $5)
        ON CONFLICT (id) DO NOTHING
        """,
        FIRM_ID, FIRM_NAME, seed_short_name, FIRM_CITY, FIRM_COUNTRY)

        # Backfill otp_store.firm_id (Part 2, multi-tenancy hardening) --
        # needs a bound FIRM_ID param, so it can't live in the big
        # unparameterized migration block above.
        await conn.execute(
            "UPDATE otp_store SET firm_id=$1 WHERE firm_id IS NULL", FIRM_ID
        )

        await conn.execute("""
        INSERT INTO reminder_settings (firm_id)
        VALUES ($1) ON CONFLICT (firm_id) DO NOTHING
        """, FIRM_ID)

    print("[db] schema migrations complete")

# ── Firm identity ─────────────────────────────────────────────────────────────
# Multi-tenancy model: Option B (one deployment per firm) — full writeup in
# README.md's "Multi-Tenancy" section. Short version: FIRM_ID below is a
# fixed constant for the entire lifetime of this process, read once from
# MUTEMO_FIRM_ID at startup. There is exactly one firm's data in this
# deployment's Postgres/Chroma instance, ever. Many tables (and Chroma
# chunk metadata) also carry a firm_id column that most queries scope by
# — that's defense-in-depth / future-proofing, not evidence this instance
# actually serves more than one firm. It doesn't. Don't remove those
# columns/filters, but don't read their presence as "multi-tenancy is
# live" either. Second firm = its own Railway service + its own FIRM_ID,
# not a second row in this one.
FIRM_NAME = os.environ.get("MUTEMO_FIRM_NAME", "")
FIRM_CITY = os.environ.get("MUTEMO_FIRM_CITY", "Harare, Zimbabwe")
# Zimbabwe is a safe default across this product's whole domain (Zimbabwean
# case law/legislation, court structure) -- not another firm's specific
# detail leaking across tenants, unlike the short_name default below.
FIRM_COUNTRY = os.environ.get("MUTEMO_FIRM_COUNTRY", "Zimbabwe")
# No safe fixed default here (there's no such thing as a generically-right
# short name) -- left empty and derived from the real MUTEMO_FIRM_NAME at
# seed time via _derive_firm_short_name() below if not set explicitly.
FIRM_SHORT_NAME = os.environ.get("MUTEMO_FIRM_SHORT_NAME", "")
FIRM_ID_STR = os.environ.get("MUTEMO_FIRM_ID", "a1b2c3d4-0000-0000-0000-000000000001")
import uuid as _uuid_mod
FIRM_ID = _uuid_mod.UUID(FIRM_ID_STR)

def _derive_firm_short_name(name: str) -> str:
    """Fallback when MUTEMO_FIRM_SHORT_NAME isn't set: initials drawn from
    the firm's own real name (e.g. "Sawyer & Mkushi" -> "SM"), never
    another firm's hardcoded value. Falls back further to a plain
    truncation only if the name has no letter-starting words to draw
    initials from at all (e.g. a name that's just punctuation/numbers)."""
    words = re.findall(r"[A-Za-z][\w'-]*", name)
    initials = "".join(w[0].upper() for w in words)
    return initials or name[:8].upper()

async def get_firm_identity() -> dict:
    """
    Live DB lookup for this deployment's firm name/city — NOT the
    FIRM_NAME/FIRM_CITY env-derived constants above, which are frozen at
    process start. PATCH /api/settings can rename the firm (firms.name/
    city) without a restart; a prompt built from the stale constant would
    keep drafting documents under the old name until the next deploy.
    Used specifically for AI system prompts (contract review, legal
    research, affidavit, document drafting) where that staleness would
    land directly in generated legal documents — everywhere else in this
    file still reads the plain constants, unchanged.

    Falls back to the constants only if the firms row is somehow missing
    a name (defensive; run_migrations() always seeds one).
    """
    async with _db_pool.acquire() as conn:
        row = await conn.fetchrow("SELECT name, city FROM firms WHERE id=$1", FIRM_ID)
    name = (row["name"] if row else None) or FIRM_NAME
    city = (row["city"] if row else None) or FIRM_CITY
    return {"name": name, "city": city}

# ── RBAC helpers ──────────────────────────────────────────────────────────────
ROLE_WEIGHTS = {"admin": 4, "partner": 3, "associate": 2, "secretary": 1}

# Permission matrix
PERMISSIONS = {
    # Matters
    "matter:read":          {"admin", "partner", "associate", "secretary"},
    "matter:create":        {"admin", "partner", "associate", "secretary"},
    "matter:edit":          {"admin", "partner", "associate", "secretary"},
    "matter:delete":        {"admin", "partner"},
    # Clients
    "client:read":          {"admin", "partner", "associate", "secretary"},
    "client:create":        {"admin", "partner", "associate", "secretary"},
    "client:edit":          {"admin", "partner", "associate", "secretary"},
    # Documents
    "document:upload":      {"admin", "partner", "associate", "secretary"},
    "document:delete":      {"admin", "partner"},
    # Notes
    "note:create":          {"admin", "partner", "associate", "secretary"},
    "note:delete":          {"admin", "partner"},
    # Drafting (lawyer functions)
    "draft:affidavit":      {"admin", "partner", "associate"},
    "draft:document":       {"admin", "partner", "associate"},
    # Calendar
    "calendar:read":        {"admin", "partner", "associate", "secretary"},
    "calendar:create":      {"admin", "partner", "associate", "secretary"},
    "calendar:delete":      {"admin", "partner", "associate"},
    # Legal updates / ZLR
    "legal:upload":         {"admin", "partner", "associate"},
    "legal:delete":         {"admin", "partner"},
    # Search
    "search":               {"admin", "partner", "associate", "secretary"},
    # Admin
    "admin:settings":       {"admin", "partner"},
    "admin:users":          {"admin", "partner"},
    "admin:reindex":        {"admin"},
    # Reports — "Partner role only, 403 for Associates" in the spec, but
    # scoped {"admin", "partner"} like every other partner-tier permission
    # above rather than partner-excluding-admin: there's no existing
    # partner-only-excluding-admin precedent, and ROLE_WEIGHTS treats admin
    # as senior to partner, so locking admin out would be inconsistent with
    # every other permission in this table.
    "reports:rbz_compliance": {"admin", "partner"},
    "reports:practice_area_breakdown": {"admin", "partner"},
}

def _check_permission(user: dict, permission: str):
    """Raise 403 if user's role does not have the required permission."""
    if not user:
        raise HTTPException(status_code=401, detail="Authentication required")
    role = user.get("role", "secretary")
    allowed = PERMISSIONS.get(permission, set())
    if role not in allowed:
        raise HTTPException(
            status_code=403,
            detail=f"Your role ({role}) does not have permission for this action."
        )

# ── Lifespan ──────────────────────────────────────────────────────────────────
# lifespan() reaches `yield` (and Uvicorn starts accepting connections) as
# soon as warm_up() is *scheduled*, not when it finishes — reconcile_chroma_index()
# rebuilds the firm/legal/zlr collections sequentially and can take well
# over a minute after a redeploy with an out-of-sync volume. Confirmed in
# production: requests were served with legal=0 zlr=0 retrieval results at
# 18:14:38, while legal didn't finish re-indexing until 18:15:55 and zlr's
# rebuild hadn't even started — "[startup] semantic search ready" didn't
# fire until 18:16:05. _retrieval_ready gates the endpoints that actually
# read from these collections (see _require_retrieval_ready below) so a
# request landing in that window gets a clear 503 instead of a silent
# partial/empty result. Endpoints that don't touch ChromaDB (matters,
# calendar, generate-affidavit, legal-updates keyword search, etc.) are
# unaffected and keep serving traffic immediately, same as before.
_retrieval_ready = False

def _require_retrieval_ready():
    """
    Raise 503 if the post-deploy ChromaDB reconcile hasn't finished yet.
    Call this from any endpoint that reads the firm/legal/zlr collections
    (directly or via _semantic_search_firm/_semantic_search_legal/
    _zlr_semantic_search) — right after auth/permission checks, before any
    DB or ChromaDB access, so a not-ready request fails fast rather than
    running a search against a partially-rebuilt index.
    """
    if not _retrieval_ready:
        raise HTTPException(
            status_code=503,
            detail="Search is still initializing after a recent deploy. Please retry in a moment.",
            headers={"Retry-After": "20"},
        )

@asynccontextmanager
async def lifespan(app: FastAPI):
    await init_db()
    asyncio.create_task(reminder_scheduler_loop())
    async def warm_up():
        global _retrieval_ready
        try:
            await asyncio.to_thread(get_embedding_model)
            await asyncio.to_thread(get_chroma_collections)
            await reconcile_chroma_index()
            print("[startup] semantic search ready")
        except Exception as e:
            # Chroma/embeddings are unavailable entirely — a different,
            # already-handled degraded mode (keyword fallback), not the
            # race condition this flag guards against. Still release the
            # gate: the collections aren't "mid-rebuild", they're just not
            # going to be used, and _semantic_search_* already fall back
            # gracefully when ChromaDB isn't reachable.
            print(f"[startup] semantic search unavailable, will use keyword fallback: {e}")
        finally:
            _retrieval_ready = True
    asyncio.create_task(warm_up())
    yield
    if _db_pool:
        await _db_pool.close()
        print("[db] connection pool closed")

app = FastAPI(title="Mutemo Desk", version="2.1.0", lifespan=lifespan)

# ── AlertEngine instrumentation ───────────────────────────────────────────────
try:
    from fastapi_alertengine import instrument
    instrument(app)
    print("[startup] AlertEngine instrumentation active")
except Exception as e:
    print(f"[startup] AlertEngine instrumentation unavailable: {e}")

# ── AlertEngine health metrics ─────────────────────────────────────────────────
import time as _time
from collections import deque

_request_latencies: deque = deque(maxlen=200)
_request_errors: deque = deque(maxlen=200)

def _p95_latency() -> float:
    if not _request_latencies:
        return 0.0
    sorted_latencies = sorted(_request_latencies)
    idx = int(len(sorted_latencies) * 0.95)
    return round(sorted_latencies[min(idx, len(sorted_latencies) - 1)], 3)

def _error_rate() -> float:
    if not _request_errors:
        return 0.0
    return round(sum(_request_errors) / len(_request_errors), 3)

def _health_score() -> float:
    latency = _p95_latency()
    err = _error_rate()
    score = 100.0
    if latency > 2.0:
        score -= min(40, (latency - 2.0) * 10)
    if err > 0.01:
        score -= min(60, err * 200)
    return round(max(0.0, score), 1)

@app.get("/health/alerts")
async def health_alerts():
    """
    AlertEngine health endpoint — real-time API health metrics.

    WARNING: when fastapi_alertengine is actually installed (i.e. in
    production — it's absent locally, which is why this handler runs fine
    in dev/tests), instrument(app) above registers its OWN "/health/alerts"
    route first, at app-creation time. Starlette matches routes in
    registration order, so THIS handler never actually receives a request
    in production — confirmed live: the real response has a completely
    different shape (service_name/instance_id/metrics/health_score/alerts/
    adaptive_thresholds), not this one. retrieval_ready below is therefore
    dead code in production. Don't rely on this path for retrieval-
    readiness visibility — use /health/ready instead (unauthenticated,
    not shadowed by anything). Left this field here anyway since it's
    harmless and this handler + its tests remain meaningful wherever
    AlertEngine isn't loaded.

    "status" stays "ok" regardless of retrieval_ready below (on the rare
    chance this handler IS reached) — this reflects request-serving health
    (latency/error rate), not corpus readiness, and a post-deploy reconcile
    window is expected/normal, not an incident. This endpoint also isn't
    wired up as a Railway healthcheckPath (no railway.json in this repo
    configures one either way).
    """
    return {
        "status": "ok",
        "score": _health_score(),
        "p95_latency": _p95_latency(),
        "error_rate": _error_rate(),
        "retrieval_ready": _retrieval_ready,
        "timestamp": _time.time(),
    }

@app.get("/health/ready")
async def health_ready():
    """
    Unauthenticated, deliberately minimal readiness probe — not shadowed by
    AlertEngine (distinct path from /health/alerts, see the warning above),
    so this is the reliable way to check retrieval_ready from outside the
    process (e.g. during incident triage, or scripted polling after a
    deploy) without needing a logged-in session.
    """
    return {"retrieval_ready": _retrieval_ready}
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
    allow_credentials=True,
)

# ── Request size limit ────────────────────────────────────────────────────────
MAX_UPLOAD_SIZE = 50 * 1024 * 1024  # 50MB

@app.middleware("http")
async def metrics_middleware(request, call_next):
    start = _time.time()
    response = await call_next(request)
    duration = _time.time() - start
    _request_latencies.append(duration)
    _request_errors.append(1 if response.status_code >= 500 else 0)
    return response
@app.middleware("http")
async def size_limit_middleware(request, call_next):
    if request.method == "POST":
        content_length = request.headers.get("content-length")
        if content_length and int(content_length) > MAX_UPLOAD_SIZE:
            return JSONResponse(
                status_code=413,
                content={"detail": "File too large. Maximum upload size is 50MB."}
            )
    return await call_next(request)

# ── Admin token ───────────────────────────────────────────────────────────────
ADMIN_TOKEN = os.environ.get("MUTEMO_ADMIN_TOKEN")

def require_admin_token(request: Request):
    if ADMIN_TOKEN:
        token = request.headers.get("X-Admin-Token", "")
        if token != ADMIN_TOKEN:
            raise HTTPException(status_code=403, detail="Admin access required")

# ── Legal feed service token ─────────────────────────────────────────────────
# A separate, narrowly-scoped credential for mutemo-legal-feed's own
# machine-to-machine pushes — checked ONLY on FEED_UPLOAD_PATHS below, never
# treated as equivalent to ADMIN_TOKEN. Previously the feed reused
# MUTEMO_ADMIN_TOKEN, which also unlocks /api/admin/* (reindex,
# reset-chromadb, reclassify-zlr, etc.) — a leaked feed credential had far
# more reach than the feed itself ever needs. This token can only ever
# write legal-updates/zlr content, nothing else.
LEGAL_FEED_SERVICE_TOKEN = os.environ.get("LEGAL_FEED_SERVICE_TOKEN")
FEED_UPLOAD_PATHS = ("/api/legal-updates/upload", "/api/zlr/upload")

# ── OTP Authentication ────────────────────────────────────────────────────────
# Preferred channel is WhatsApp Business Cloud API (Meta) — reuses the same
# Meta WhatsApp Business setup already used for AlertEngine. This uses a
# dedicated AUTHENTICATION-category template, which (unlike regular
# utility/marketing templates) is NOT subject to the 24-hour conversation
# window restriction that affected AlertEngine's other WhatsApp delivery —
# authentication templates can be sent proactively any time, to anyone,
# once approved by Meta. Uses the "Copy Code" button type since this is a
# web app (one-tap/zero-tap autofill need native app integration).
# Twilio SMS is the middle fallback (see _send_sms_otp/_send_otp_code below)
# for firms without WhatsApp Business Verification yet; email via Resend
# remains the last-resort stopgap.
WHATSAPP_ACCESS_TOKEN     = os.environ.get("WHATSAPP_ACCESS_TOKEN")
WHATSAPP_PHONE_NUMBER_ID  = os.environ.get("WHATSAPP_PHONE_NUMBER_ID")
WHATSAPP_OTP_TEMPLATE     = os.environ.get("WHATSAPP_OTP_TEMPLATE_NAME", "mutemo_login_otp")
WHATSAPP_OTP_LANG         = os.environ.get("WHATSAPP_OTP_TEMPLATE_LANG", "en")
TWILIO_ACCOUNT_SID  = os.environ.get("TWILIO_ACCOUNT_SID")
TWILIO_AUTH_TOKEN   = os.environ.get("TWILIO_AUTH_TOKEN")
TWILIO_FROM_NUMBER  = os.environ.get("TWILIO_FROM_NUMBER")
_TWILIO_SMS_CONFIGURED = bool(TWILIO_ACCOUNT_SID and TWILIO_AUTH_TOKEN and TWILIO_FROM_NUMBER)
# Inlined rather than calling is_email_configured() (defined later in this
# file) — this line runs at import time, before that function exists yet.
_EMAIL_OTP_CONFIGURED = bool(os.environ.get("RESEND_API_KEY") or os.environ.get("SMTP_HOST"))
AUTH_ENABLED = bool(WHATSAPP_ACCESS_TOKEN and WHATSAPP_PHONE_NUMBER_ID) or _TWILIO_SMS_CONFIGURED or _EMAIL_OTP_CONFIGURED

# Any request landing while AUTH_ENABLED is False resolves to a synthetic
# dev user with no real identity (see get_current_user) — session auth is
# effectively off. Fine for a developer's own machine; not fine for
# anything reachable on the internet. RAILWAY_ENVIRONMENT_NAME is present
# on every Railway-hosted service and absent on a local machine, but it
# does NOT distinguish real production from staging here — MutemoOS-V2 and
# mutemoos-staging are two services inside the same Railway "production"
# environment, so RAILWAY_ENVIRONMENT_NAME reads "production" for both.
# Rather than trying to guess which service this is, any Railway
# deployment that genuinely wants the dev-auth fallback (staging today)
# must opt in explicitly via MUTEMO_ALLOW_DEV_AUTH=true — it's never
# silently available just because a service happens to be un-configured.
MUTEMO_ALLOW_DEV_AUTH = os.environ.get("MUTEMO_ALLOW_DEV_AUTH", "").strip().lower() == "true"
if os.environ.get("RAILWAY_ENVIRONMENT_NAME") and not AUTH_ENABLED and not MUTEMO_ALLOW_DEV_AUTH:
    raise RuntimeError(
        "AUTH_ENABLED is False on a Railway deployment "
        f"(RAILWAY_SERVICE_NAME={os.environ.get('RAILWAY_SERVICE_NAME')!r}) with no "
        "WhatsApp/Twilio/email OTP channel configured — every request would silently "
        "fall back to a synthetic dev user with no real authentication. If this service "
        "is intentionally running without real auth (e.g. staging), set "
        "MUTEMO_ALLOW_DEV_AUTH=true explicitly. Refusing to start otherwise."
    )

OTP_TTL_SECONDS     = 300
SESSION_TTL_SECONDS = 86400 * 7
MAX_OTP_ATTEMPTS    = 5
# Shared-device session hardening (2026-08-27): SESSION_TTL_SECONDS above is
# an absolute cap from login, not a safety net for someone who forgets to
# log out on a shared boardroom/library machine -- a session created at
# login stays fully live regardless of activity until that 7-day mark.
# This is a genuine sliding idle timeout layered on top, backed by
# sessions.last_active (touched on real request activity -- see
# _touch_session_last_active() below): a session is only valid while BOTH
# the absolute expires_at AND this idle window still hold. Configurable
# per deployment since "reasonable" depends on the firm/device -- default
# picked as a middle ground for someone reading a long document without
# triggering any request for a while, without leaving a shared machine
# live for hours after everyone's gone.
SESSION_IDLE_TIMEOUT_SECONDS = int(os.environ.get("MUTEMO_SESSION_IDLE_TIMEOUT_SECONDS", "2700"))  # 45 min

def _send_email_otp(email: str, code: str) -> bool:
    """
    Email OTP delivery via the same Resend integration already proven for
    calendar invites and the daily digest — a same-day stopgap while Meta
    Business Verification is pending for WhatsApp delivery (see
    _send_otp_code below, which prefers WhatsApp automatically once it's
    configured, no further code change needed to "switch over" later).
    """
    subject = "Your Mutemo Desk login code"
    html_body = (
        f"<p>Your Mutemo Desk login code is <strong>{code}</strong>. It expires in 5 minutes.</p>"
        f"<p style='color:#6b6b64;font-size:13px'>If you didn't request this, you can safely ignore this email.</p>"
    )
    text_body = f"Your Mutemo Desk login code is {code}. It expires in 5 minutes."
    try:
        _send_via_resend_sync(email, subject, html_body, text_body)
        return True
    except Exception as e:
        print(f"[otp] Email send failed: {e}")
        return False

def _send_otp_code(phone: str, email: Optional[str], code: str) -> Optional[str]:
    """
    Sends the OTP via whichever channel is actually configured. Prefers
    WhatsApp (the eventual target, once Meta Business Verification
    completes for the WhatsApp Business Account), falls back to Twilio SMS
    for firms without WhatsApp set up yet, and finally to email via Resend
    as a last-resort stopgap. Each channel activates automatically as soon
    as its own env vars are set on Railway — no further code change needed
    to "switch over" between them.

    Returns the channel actually used ("whatsapp" / "sms" / "email"), or
    None if nothing could be sent — this was previously just True/False,
    which is why the login screen kept saying "code sent to your phone"
    even when it had actually gone to email: the frontend had no way to
    know which channel was really used.
    """
    if WHATSAPP_ACCESS_TOKEN and WHATSAPP_PHONE_NUMBER_ID:
        if _send_whatsapp_otp(phone, code):
            return "whatsapp"
        print(f"[otp] WhatsApp send failed for {phone}, falling back")
    if _TWILIO_SMS_CONFIGURED:
        if _send_sms_otp(phone, code):
            return "sms"
        print(f"[otp] SMS send failed for {phone}, falling back")
    if email and _EMAIL_OTP_CONFIGURED:
        if _send_email_otp(email, code):
            return "email"
    print(f"[otp] No delivery channel available for {phone} (no email on file and neither WhatsApp nor SMS configured)")
    return None

def _send_whatsapp_otp(phone: str, code: str) -> bool:
    """
    Sends the OTP via an approved Meta WhatsApp AUTHENTICATION template.
    The template itself must already exist and be approved in Meta Business
    Manager (category=AUTHENTICATION, a COPY_CODE button) — that's a
    one-time setup step outside this code, not something this function can
    do on its own.
    """
    try:
        import httpx
        resp = httpx.post(
            f"https://graph.facebook.com/v20.0/{WHATSAPP_PHONE_NUMBER_ID}/messages",
            headers={"Authorization": f"Bearer {WHATSAPP_ACCESS_TOKEN}"},
            json={
                "messaging_product": "whatsapp",
                "to": phone,
                "type": "template",
                "template": {
                    "name": WHATSAPP_OTP_TEMPLATE,
                    "language": {"code": WHATSAPP_OTP_LANG},
                    "components": [
                        {"type": "body", "parameters": [{"type": "text", "text": code}]},
                        {
                            "type": "button", "sub_type": "url", "index": "0",
                            "parameters": [{"type": "text", "text": code}],
                        },
                    ],
                },
            },
            timeout=15,
        )
        if resp.status_code not in (200, 201):
            print(f"[otp] WhatsApp send failed {resp.status_code}: {resp.text[:300]}")
            return False
        return True
    except Exception as e:
        print(f"[otp] WhatsApp send failed: {e}")
        return False

def _send_sms_otp(phone: str, code: str) -> bool:
    """
    Sends the OTP as a plain SMS via Twilio's REST API — a direct httpx
    call rather than the Twilio SDK, matching _send_whatsapp_otp's style
    above rather than adding a second HTTP-client convention. Middle
    fallback: used when WhatsApp isn't configured (or its send failed) but
    Twilio is, ahead of the email stopgap.
    """
    try:
        import httpx
        resp = httpx.post(
            f"https://api.twilio.com/2010-04-01/Accounts/{TWILIO_ACCOUNT_SID}/Messages.json",
            auth=(TWILIO_ACCOUNT_SID, TWILIO_AUTH_TOKEN),
            data={
                "To": phone,
                "From": TWILIO_FROM_NUMBER,
                "Body": f"Your Mutemo Desk login code is {code}. It expires in 5 minutes.",
            },
            timeout=15,
        )
        if resp.status_code not in (200, 201):
            print(f"[otp] SMS send failed {resp.status_code}: {resp.text[:300]}")
            return False
        return True
    except Exception as e:
        print(f"[otp] SMS send failed: {e}")
        return False

class OTPRequestBody(BaseModel):
    phone: str

class OTPVerifyBody(BaseModel):
    phone: str
    code: str

@app.post("/api/auth/request-otp")
async def request_otp(req: OTPRequestBody):
    if not AUTH_ENABLED:
        raise HTTPException(status_code=503, detail="OTP login is not configured on this server.")
    phone = req.phone.strip()

    # Single source of truth for "who can log in" is now the invites/users
    # tables (managed entirely through the Invite New User screen) — not a
    # separate, manually-maintained env var list that's easy to forget to
    # update. Deliberately still return a generic {"sent": true} response
    # regardless of whether the number is actually recognized, so the
    # response itself doesn't leak which phone numbers are valid.
    async with _db_pool.acquire() as conn:
        known_row = await conn.fetchrow(
            """SELECT email FROM users WHERE firm_id=$1 AND phone=$2 AND is_active=TRUE
               UNION
               SELECT email FROM invites WHERE firm_id=$1 AND phone=$2 AND accepted_at IS NULL
               LIMIT 1""",
            FIRM_ID, phone
        )
    if not known_row:
        return {"sent": True, "message": "If this number is registered, a code has been sent."}
    known_email = known_row["email"]

    code = f"{secrets.randbelow(1000000):06d}"
    expires = datetime.utcnow() + timedelta(seconds=OTP_TTL_SECONDS)
    async with _db_pool.acquire() as conn:
        await conn.execute("""
            INSERT INTO otp_store (phone, code, attempts, expires_at, firm_id)
            VALUES ($1, $2, 0, $3, $4)
            ON CONFLICT (phone) DO UPDATE SET code=$2, attempts=0, expires_at=$3, firm_id=$4
        """, phone, code, expires, FIRM_ID)

    channel = await asyncio.to_thread(_send_otp_code, phone, known_email, code)
    if not channel:
        raise HTTPException(status_code=500, detail="Failed to send login code. Please try again.")
    # Channel is only returned for numbers we've already confirmed are known
    # (we returned early above for unknown ones with the same generic
    # message every time) — so this doesn't create a new way to probe
    # whether an arbitrary number is registered.
    return {
        "sent": True,
        "channel": channel,
        "message": f"A code has been sent via {channel}.",
    }

@app.post("/api/auth/verify-otp")
async def verify_otp(req: OTPVerifyBody, response: Response):
    if not AUTH_ENABLED:
        raise HTTPException(status_code=503, detail="OTP login is not configured on this server.")
    phone = req.phone.strip()

    async with _db_pool.acquire() as conn:
        # Clean expired entries
        await conn.execute("DELETE FROM otp_store WHERE expires_at < NOW()")
        # firm_id=$2 is defense-in-depth (Part 2, multi-tenancy hardening),
        # not a functional change under Option B -- phone remains the
        # actual PRIMARY KEY, so this can't yet change which row matches,
        # only make the intent explicit for whenever a second firm exists.
        entry = await conn.fetchrow("SELECT * FROM otp_store WHERE phone=$1 AND firm_id=$2", phone, FIRM_ID)

        if not entry:
            raise HTTPException(status_code=401, detail="No active code for this number. Request a new one.")

        new_attempts = entry["attempts"] + 1
        if new_attempts > MAX_OTP_ATTEMPTS:
            await conn.execute("DELETE FROM otp_store WHERE phone=$1 AND firm_id=$2", phone, FIRM_ID)
            raise HTTPException(status_code=429, detail="Too many attempts. Request a new code.")

        await conn.execute(
            "UPDATE otp_store SET attempts=$1 WHERE phone=$2 AND firm_id=$3", new_attempts, phone, FIRM_ID
        )

        if not hmac.compare_digest(entry["code"], req.code.strip()):
            raise HTTPException(status_code=401, detail="Incorrect code.")

        # Success — look up existing user, or provision from a matching invite
        await conn.execute("DELETE FROM otp_store WHERE phone=$1 AND firm_id=$2", phone, FIRM_ID)
        user = await conn.fetchrow(
            "SELECT * FROM users WHERE firm_id=$1 AND phone=$2 AND is_active=TRUE",
            FIRM_ID, phone
        )
        if not user:
            # Previously this auto-created a brand-new account for ANY phone
            # number that completed OTP, regardless of whether anyone
            # actually invited them — meaning the only real gate was
            # whichever email Cloudflare Access allows through, and anyone
            # who got that far could self-provision a MutemoOS account.
            # Now, account creation requires a matching, not-yet-accepted
            # invite for this exact phone number — set by an admin when
            # they invite someone, not just anyone who knows/receives a code.
            invite = await conn.fetchrow(
                "SELECT * FROM invites WHERE firm_id=$1 AND phone=$2 AND accepted_at IS NULL",
                FIRM_ID, phone
            )
            if not invite:
                raise HTTPException(
                    status_code=403,
                    detail="This phone number has not been invited to MutemoOS. Contact your administrator."
                )
            user = await conn.fetchrow("""
                INSERT INTO users (firm_id, phone, email, display_name, role)
                VALUES ($1, $2, $3, $4, $5)
                RETURNING *
            """, FIRM_ID, phone, invite["email"], invite["display_name"], invite["role"])
            await conn.execute("UPDATE invites SET accepted_at=NOW() WHERE id=$1", invite["id"])
            if invite["organisation_role"]:
                await conn.execute("""
                    INSERT INTO organisation_roles (firm_id, user_id, role)
                    VALUES ($1, $2, $3)
                    ON CONFLICT (firm_id, user_id) DO UPDATE SET role=$3
                """, FIRM_ID, user["id"], invite["organisation_role"])

        token = secrets.token_urlsafe(32)
        expires = datetime.utcnow() + timedelta(seconds=SESSION_TTL_SECONDS)
        await conn.execute("""
            INSERT INTO sessions (token, user_id, firm_id, expires_at)
            VALUES ($1, $2, $3, $4)
        """, token, user["id"], FIRM_ID, expires)

    response.set_cookie(
        key="mutemo_session", value=token,
        max_age=SESSION_TTL_SECONDS, httponly=True, secure=True, samesite="lax",
    )
    return {"verified": True, "phone": phone, "role": user["role"], "display_name": user["display_name"]}

async def _touch_session_last_active(conn, token: str) -> None:
    """
    Records real request activity for the idle timeout (see
    SESSION_IDLE_TIMEOUT_SECONDS). Throttled to at most once a minute per
    session -- without this, the 2-second job-status polling used by
    search/contract-review/document-generation while a job is running
    would turn into a sessions UPDATE on every single poll tick. A
    session genuinely idle for the full timeout window has had no real
    request at all in that window regardless of this throttle, so the
    idle timeout itself isn't weakened by it.
    """
    await conn.execute(
        "UPDATE sessions SET last_active=NOW() WHERE token=$1 AND last_active < NOW() - INTERVAL '60 seconds'",
        token,
    )

@app.post("/api/auth/logout")
async def logout(request: Request, response: Response):
    token = request.cookies.get("mutemo_session")
    user_email = None
    if token and _db_pool:
        async with _db_pool.acquire() as conn:
            # Looked up before deleting the row -- _revoke_cloudflare_access_session
            # below needs the user's email, not our own session token, and
            # there's nothing left to join against once the row is gone.
            row = await conn.fetchrow(
                "SELECT u.email FROM sessions s JOIN users u ON s.user_id = u.id WHERE s.token=$1",
                token,
            )
            user_email = row["email"] if row else None
            await conn.execute("DELETE FROM sessions WHERE token=$1", token)
    if user_email:
        await _revoke_cloudflare_access_session(user_email)
    response.delete_cookie("mutemo_session")
    return {"logged_out": True}

@app.get("/api/auth/status")
async def auth_status(request: Request):
    if not AUTH_ENABLED:
        return {"auth_enabled": False, "authenticated": True, "firm_name": FIRM_NAME}
    token = request.cookies.get("mutemo_session")
    if token and _db_pool:
        async with _db_pool.acquire() as conn:
            await conn.execute(
                "DELETE FROM sessions WHERE expires_at < NOW() OR last_active < NOW() - make_interval(secs => $1)",
                SESSION_IDLE_TIMEOUT_SECONDS,
            )
            row = await conn.fetchrow("""
                SELECT s.token, u.id, u.phone, u.role, u.display_name, u.initials
                FROM sessions s JOIN users u ON s.user_id = u.id
                WHERE s.token=$1 AND s.expires_at > NOW()
                  AND s.last_active > NOW() - make_interval(secs => $2)
            """, token, SESSION_IDLE_TIMEOUT_SECONDS)
            if row:
                # /api/auth/status is exempt from session_auth_middleware
                # (see open_paths below) so it never gets the middleware's
                # own touch -- called on every page load/reload, which is
                # itself a legitimate activity signal worth recording.
                await _touch_session_last_active(conn, token)
                return {
                    "auth_enabled": True, "authenticated": True,
                    "id": str(row["id"]), "phone": row["phone"], "role": row["role"],
                    "display_name": row["display_name"], "initials": row["initials"],
                    "firm_name": FIRM_NAME,
                }
    return {"auth_enabled": True, "authenticated": False, "firm_name": FIRM_NAME}

async def get_current_user(request: Request) -> Optional[dict]:
    """Return the current user dict, or None if not authenticated."""
    if not AUTH_ENABLED:
        # Return a synthetic partner user when auth is disabled (dev/demo
        # mode). Deliberately a generic label, not a person's name — this
        # shows up as the author on every note/report/action taken while
        # running without real auth (staging demo included), so it must
        # never resemble anyone real.
        return {"id": None, "firm_id": FIRM_ID, "phone": None, "email": None, "role": "partner", "display_name": "Demo User"}
    token = request.cookies.get("mutemo_session")
    if not token or not _db_pool:
        return None
    async with _db_pool.acquire() as conn:
        # Idle-window check duplicated here rather than relying solely on
        # session_auth_middleware's own gate -- this function is also
        # called directly (by tests, and any future non-HTTP code path)
        # bypassing the middleware entirely, same reasoning this file
        # already applies to expires_at being checked independently in
        # both places rather than just once.
        row = await conn.fetchrow("""
            SELECT u.id, u.firm_id, u.phone, u.email, u.role, u.display_name
            FROM sessions s JOIN users u ON s.user_id = u.id
            WHERE s.token=$1 AND s.expires_at > NOW()
              AND s.last_active > NOW() - make_interval(secs => $2)
        """, token, SESSION_IDLE_TIMEOUT_SECONDS)
        if row:
            return dict(row)
    return None

@app.middleware("http")
async def session_auth_middleware(request, call_next):
    if not AUTH_ENABLED:
        return await call_next(request)
    open_paths = (
        "/api/health", "/api/auth/request-otp", "/api/auth/verify-otp",
        "/api/auth/status", "/api/matters/template", "/api/matters/template-excel"
    )
    if request.url.path in open_paths or not request.url.path.startswith("/api/"):
        return await call_next(request)

    token = request.cookies.get("mutemo_session")
    if token and _db_pool:
        async with _db_pool.acquire() as conn:
            row = await conn.fetchrow(
                "SELECT token FROM sessions WHERE token=$1 AND expires_at > NOW() "
                "AND last_active > NOW() - make_interval(secs => $2)",
                token, SESSION_IDLE_TIMEOUT_SECONDS,
            )
            if row:
                # The one universal touch point -- every authenticated
                # /api/ request except the open_paths above passes through
                # here before reaching its route handler.
                await _touch_session_last_active(conn, token)
                return await call_next(request)

    # A valid X-Admin-Token is also sufficient for general admin tooling
    # (curl + MUTEMO_ADMIN_TOKEN against /api/admin/* etc. — see
    # require_admin_token() below). This was a real, live regression once:
    # when AUTH_ENABLED actually started being enforced (following the
    # invite-gating/OTP fixes), this middleware had no token exemption at
    # all, meaning any token-based caller started getting a silent 401 the
    # moment auth was properly turned on.
    admin_token_header = request.headers.get("X-Admin-Token", "")
    if ADMIN_TOKEN and admin_token_header == ADMIN_TOKEN:
        return await call_next(request)

    # mutemo-legal-feed's pusher authenticates its machine-to-machine calls
    # with its own dedicated LEGAL_FEED_SERVICE_TOKEN (X-Feed-Service-Token
    # header) — deliberately NOT the general admin token, and deliberately
    # only exempted for its two actual upload paths, so a leaked feed
    # credential can never reach /api/admin/* or anything else.
    if request.url.path in FEED_UPLOAD_PATHS and LEGAL_FEED_SERVICE_TOKEN:
        feed_token_header = request.headers.get("X-Feed-Service-Token", "")
        if feed_token_header == LEGAL_FEED_SERVICE_TOKEN:
            return await call_next(request)

    return JSONResponse(status_code=401, content={"detail": "Authentication required"})

# ── User management endpoints ─────────────────────────────────────────────────

@app.get("/api/users")
async def list_users(request: Request):
    user = await get_current_user(request)
    _check_permission(user, "admin:users")
    async with _db_pool.acquire() as conn:
        rows = await conn.fetch(
            "SELECT id, phone, display_name, role, is_active, created_at FROM users WHERE firm_id=$1 ORDER BY created_at",
            FIRM_ID
        )
    return [dict(r) for r in rows]

@app.patch("/api/users/{user_id}")
async def update_user(user_id: str, body: dict, request: Request):
    """Admin only: update a user's role, display_name, email, or phone."""
    user = await get_current_user(request)
    _check_permission(user, "admin:users")
    allowed_fields = {"role", "display_name", "is_active", "email", "phone"}
    updates = {k: v for k, v in body.items() if k in allowed_fields}
    if not updates:
        raise HTTPException(status_code=400, detail="No valid fields to update")
    if "role" in updates and updates["role"] not in ("partner", "associate", "secretary", "admin"):
        raise HTTPException(status_code=400, detail="Invalid role")
    set_clauses = ", ".join(f"{k}=${i+2}" for i, k in enumerate(updates.keys()))
    values = list(updates.values())
    async with _db_pool.acquire() as conn:
        row = await conn.fetchrow(
            f"UPDATE users SET {set_clauses} WHERE id=$1 AND firm_id=${len(values)+2} RETURNING *",
            _uuid_mod.UUID(user_id), *values, FIRM_ID
        )
    if not row:
        raise HTTPException(status_code=404, detail="User not found")
    return dict(row)

# ── Invites ───────────────────────────────────────────────────────────────────

def _get_cf_vars():
    return (
        os.environ.get("CLOUDFLARE_API_TOKEN"),
        os.environ.get("CLOUDFLARE_ACCOUNT_ID"),
        os.environ.get("CLOUDFLARE_ACCESS_APP_ID"),
    )

async def _add_email_to_cloudflare_access(email: str) -> Optional[str]:
    """Add an email to the Cloudflare Access policy. Returns the rule ID or None."""
    CLOUDFLARE_API_TOKEN, CLOUDFLARE_ACCOUNT_ID, CLOUDFLARE_ACCESS_APP_ID = _get_cf_vars()
    if not all([CLOUDFLARE_API_TOKEN, CLOUDFLARE_ACCOUNT_ID, CLOUDFLARE_ACCESS_APP_ID]):
        print("[invite] Cloudflare vars not set — skipping CF Access update")
        return None
    try:
        import httpx
        async with httpx.AsyncClient(timeout=15) as http:
            resp = await http.get(
                f"https://api.cloudflare.com/client/v4/accounts/{CLOUDFLARE_ACCOUNT_ID}/access/apps/{CLOUDFLARE_ACCESS_APP_ID}/policies",
                headers={"Authorization": f"Bearer {CLOUDFLARE_API_TOKEN}"},
            )
            resp.raise_for_status()
            policies = resp.json().get("result", [])

            allow_policy = next((p for p in policies if p.get("decision") == "allow"), None)
            if not allow_policy:
                print("[invite] No Allow policy found in Cloudflare Access app")
                return None

            policy_id = allow_policy["id"]
            existing_include = allow_policy.get("include", [])

            already_there = any(
                r.get("email", {}).get("email") == email
                for r in existing_include
            )
            if already_there:
                print(f"[invite] {email} already in Cloudflare Access policy")
                return policy_id

            new_include = existing_include + [{"email": {"email": email}}]

            update_resp = await http.put(
                f"https://api.cloudflare.com/client/v4/accounts/{CLOUDFLARE_ACCOUNT_ID}/access/apps/{CLOUDFLARE_ACCESS_APP_ID}/policies/{policy_id}",
                headers={"Authorization": f"Bearer {CLOUDFLARE_API_TOKEN}", "Content-Type": "application/json"},
                json={
                    "name": allow_policy["name"],
                    "decision": "allow",
                    "include": new_include,
                    "exclude": allow_policy.get("exclude", []),
                    "require": allow_policy.get("require", []),
                },
            )
            update_resp.raise_for_status()
            print(f"[invite] Added {email} to Cloudflare Access policy")
            return policy_id

    except Exception as e:
        print(f"[invite] Cloudflare Access update failed: {e}")
        return None

async def _revoke_cloudflare_access_session(email: str) -> bool:
    """
    Shared-device session hardening (2026-08-27): ends the user's
    Cloudflare Access session (the CF_Authorization cookie / login to the
    Access application gating this app), not just our own app-level
    session. Before this, /api/auth/logout only ever touched our own
    `sessions` table -- confirmed by grepping this whole codebase for any
    CF_Authorization reference, there was none. That meant on a shared
    device, app logout alone did not fully sign the browser out: a second
    person could sit down and, while Cloudflare's own session was still
    live, never see a fresh Cloudflare Access login prompt at all.

    Real, documented endpoint (verified against Cloudflare's own API docs
    directly, not assumed): POST /accounts/{account_id}/access/
    organizations/revoke_user, keyed by email. Two real caveats, accepted
    deliberately rather than silently:
      - This revokes the user's Access session ORG-WIDE, across every
        Cloudflare Access application in this Cloudflare account, not
        scoped to just MutemoOS -- Cloudflare's API has no per-app
        variant of this call. For a law firm's shared-device threat model
        (don't leave client data exposed to whoever sits down next),
        erring toward "signed out of everything" is the right direction,
        not a bug -- but worth knowing if this account ever fronts
        another internal tool with its own Access app.
      - Requires the "Access: Organizations, Identity Providers, and
        Groups Write" API token scope -- broader than what
        CLOUDFLARE_API_TOKEN was originally scoped for (the existing
        _add_email_to_cloudflare_access() above only ever needed "Access:
        Apps and Policies Edit"). If the token lacks this scope,
        Cloudflare returns 403 and this silently no-ops -- logged, never
        raised, never blocking the app-level logout that already
        happened by the time this runs. Whether the current token
        actually has this scope has NOT been verified end-to-end locally
        (Railway CLI access to read the real token was unavailable this
        session) -- the real, authoritative answer is whatever the
        production logs show after a real logout; see them before
        assuming this silently does nothing.
    """
    CLOUDFLARE_API_TOKEN, CLOUDFLARE_ACCOUNT_ID, _ = _get_cf_vars()
    if not CLOUDFLARE_API_TOKEN or not CLOUDFLARE_ACCOUNT_ID:
        print("[logout] Cloudflare vars not set — skipping Access session revoke")
        return False
    try:
        import httpx
        async with httpx.AsyncClient(timeout=10) as http:
            resp = await http.post(
                f"https://api.cloudflare.com/client/v4/accounts/{CLOUDFLARE_ACCOUNT_ID}/access/organizations/revoke_user",
                headers={"Authorization": f"Bearer {CLOUDFLARE_API_TOKEN}", "Content-Type": "application/json"},
                json={"email": email},
            )
        if resp.status_code == 200 and resp.json().get("success"):
            print(f"[logout] Cloudflare Access session revoked for {email}")
            return True
        print(f"[logout] Cloudflare Access revoke failed (non-fatal): {resp.status_code} {resp.text[:300]}")
        return False
    except Exception as e:
        print(f"[logout] Cloudflare Access revoke error (non-fatal): {e}")
        return False

@app.post("/api/admin/cf-set-session-duration")
async def cf_set_session_duration(request: Request, duration: str = "6h"):
    """TEMPORARY, one-time endpoint (2026-08-30) -- not a standing feature.

    Changes the Cloudflare Access application's session_duration, per an
    explicit, informed decision: automating this was originally judged too
    risky (see docs/shared-device-logout.md) because it was never
    confirmed whether PUT /accounts/{id}/access/apps/{id} is a partial
    update or a full-document replace, and a blind partial body risked
    wiping the live application's domain/allowed_idps/policies -- real
    production impact for a real firm (Sawyer & Mkushi) currently using
    this exact application to log in.

    Mitigation used here: GET the complete, current object immediately
    before the PUT (not a stale/earlier copy), strip only the fields that
    are unambiguously read-only response metadata (id, uid, created_at,
    updated_at, aud -- none of these are configuration a client could
    sensibly set), and echo every other field back exactly as-is with
    only session_duration changed. If the PUT endpoint does turn out to
    be a full replace, this is the closest thing to a safe one, since
    nothing is actually omitted except read-only metadata. Returns both
    the pre-change and post-change object so the result is verifiable,
    not assumed -- and so the pre-change state is recoverable by hand via
    the dashboard if anything looks wrong.

    Read-only from the caller's perspective except for the one intended
    write; X-Admin-Token gated, same pattern as tonight's other temporary
    endpoints. Safe to delete once this change is verified.
    """
    admin_token_header = request.headers.get("X-Admin-Token", "")
    if not ADMIN_TOKEN or admin_token_header != ADMIN_TOKEN:
        raise HTTPException(status_code=403, detail="Invalid admin token")

    CLOUDFLARE_API_TOKEN, CLOUDFLARE_ACCOUNT_ID, CLOUDFLARE_ACCESS_APP_ID = _get_cf_vars()
    if not all([CLOUDFLARE_API_TOKEN, CLOUDFLARE_ACCOUNT_ID, CLOUDFLARE_ACCESS_APP_ID]):
        return {"error": "Cloudflare vars not fully configured on this deployment"}

    url = f"https://api.cloudflare.com/client/v4/accounts/{CLOUDFLARE_ACCOUNT_ID}/access/apps/{CLOUDFLARE_ACCESS_APP_ID}"
    headers = {"Authorization": f"Bearer {CLOUDFLARE_API_TOKEN}", "Content-Type": "application/json"}

    try:
        import httpx
        async with httpx.AsyncClient(timeout=15) as http:
            get_resp = await http.get(url, headers=headers)
            if get_resp.status_code != 200 or not get_resp.json().get("success"):
                return {"error": "Fresh GET failed, refusing to PUT blind", "status_code": get_resp.status_code, "body": get_resp.text[:500]}

            current = get_resp.json()["result"]
            before_session_duration = current.get("session_duration")

            READ_ONLY_FIELDS = {"id", "uid", "created_at", "updated_at", "aud"}
            put_body = {k: v for k, v in current.items() if k not in READ_ONLY_FIELDS}
            put_body["session_duration"] = duration

            put_resp = await http.put(url, headers=headers, json=put_body)
            put_json = put_resp.json() if put_resp.headers.get("content-type", "").startswith("application/json") else {"raw_text": put_resp.text[:1000]}

            verify_resp = await http.get(url, headers=headers)
            verify_json = verify_resp.json() if verify_resp.status_code == 200 else None
            after = verify_json["result"] if verify_json and verify_json.get("success") else None
    except Exception as e:
        return {"error": f"Request to Cloudflare API failed: {e}"}

    return {
        "put_status_code": put_resp.status_code,
        "put_response": put_json,
        "before_session_duration": before_session_duration,
        "after_session_duration": after.get("session_duration") if after else None,
        "before_domain": current.get("domain"),
        "after_domain": after.get("domain") if after else None,
        "before_allowed_idps": current.get("allowed_idps"),
        "after_allowed_idps": after.get("allowed_idps") if after else None,
        "before_policy_count": len(current.get("policies") or []),
        "after_policy_count": len(after.get("policies") or []) if after else None,
        "before_policy_emails": [
            e.get("email", {}).get("email")
            for p in (current.get("policies") or [])
            for e in (p.get("include") or [])
            if "email" in e
        ],
        "after_policy_emails": [
            e.get("email", {}).get("email")
            for p in ((after or {}).get("policies") or [])
            for e in (p.get("include") or [])
            if "email" in e
        ] if after else None,
    }

@app.get("/api/admin/session-check")
async def session_check(request: Request, user_id: str):
    """TEMPORARY diagnostic endpoint (2026-08-30) -- not a standing feature.

    Purely a read: a plain SELECT against the sessions table, with no
    touch of last_active and no interaction with get_current_user() or
    the session middleware at all. Built specifically so the real 90-
    minute idle-timeout verification could check a session's state at
    multiple points in time without each check itself resetting the
    idle clock -- /api/auth/status (the normal way to check) *does*
    touch last_active by design (that's how real traffic keeps a session
    alive), which would have silently invalidated the test.

    Read-only, X-Admin-Token gated, same pattern as tonight's other
    temporary endpoints. Safe to delete once this verification is done.
    """
    admin_token_header = request.headers.get("X-Admin-Token", "")
    if not ADMIN_TOKEN or admin_token_header != ADMIN_TOKEN:
        raise HTTPException(status_code=403, detail="Invalid admin token")

    async with _db_pool.acquire() as conn:
        row = await conn.fetchrow(
            """
            SELECT token, created_at, last_active, expires_at,
                   EXTRACT(EPOCH FROM (NOW() - last_active)) AS idle_seconds,
                   EXTRACT(EPOCH FROM (NOW() - created_at)) AS session_age_seconds,
                   (expires_at > NOW()) AS within_absolute_cap,
                   (last_active > NOW() - make_interval(secs => $2)) AS within_idle_window
            FROM sessions
            WHERE user_id = $1
            ORDER BY created_at DESC
            LIMIT 1
            """,
            _uuid_mod.UUID(user_id), SESSION_IDLE_TIMEOUT_SECONDS,
        )
    if not row:
        return {"error": "No session found for this user_id"}
    return {
        "last_active": row["last_active"].isoformat(),
        "created_at": row["created_at"].isoformat(),
        "expires_at": row["expires_at"].isoformat(),
        "idle_seconds": row["idle_seconds"],
        "idle_minutes": round(row["idle_seconds"] / 60, 1),
        "session_age_minutes": round(row["session_age_seconds"] / 60, 1),
        "within_absolute_cap": row["within_absolute_cap"],
        "within_idle_window": row["within_idle_window"],
        "would_currently_be_authenticated": row["within_absolute_cap"] and row["within_idle_window"],
        "configured_idle_timeout_seconds": SESSION_IDLE_TIMEOUT_SECONDS,
        "configured_idle_timeout_minutes": SESSION_IDLE_TIMEOUT_SECONDS / 60,
    }

async def _send_invite_email(email: str, display_name: str, invited_by_name: str) -> bool:
    """Send welcome invite email via Resend."""
    try:
        html = f"""
        <div style="font-family:Georgia,serif;max-width:560px;margin:0 auto">
            <div style="background:#1b4d2e;color:white;padding:16px 20px;border-radius:6px 6px 0 0">
                <strong style="font-size:18px">&#9878; Mutemo Desk</strong><br/>
                <span style="font-size:13px;opacity:0.8">You have been invited</span>
            </div>
            <div style="padding:24px 20px;border:1px solid #d8d3c8;border-top:none;border-radius:0 0 6px 6px">
                <p>Hi {display_name},</p>
                <p>{invited_by_name} has invited you to access <strong>Mutemo Desk</strong> — the legal practice management system for {FIRM_NAME}.</p>
                <p style="margin:24px 0">
                    <a href="https://mutemo.tofamba.com"
                       style="background:#1b4d2e;color:white;padding:12px 24px;border-radius:6px;text-decoration:none;font-weight:bold">
                        Access Mutemo Desk
                    </a>
                </p>
                <p style="font-size:13px;color:#666">
                    When prompted, enter your email address <strong>{email}</strong> to receive a one-time login code.
                </p>
                <p style="font-size:13px;color:#666">
                    If you have any issues accessing the system, contact {invited_by_name}.
                </p>
            </div>
        </div>
        """
        text = f"Hi {display_name},\n\n{invited_by_name} has invited you to Mutemo Desk.\n\nAccess it at: https://mutemo.tofamba.com\n\nUse your email {email} to log in.\n\n— Mutemo Desk"
        await asyncio.to_thread(
            _send_via_resend_sync,
            email,
            f"You've been invited to Mutemo Desk — {FIRM_NAME}",
            html,
            text,
        )
        return True
    except Exception as e:
        print(f"[invite] email send failed: {e}")
        return False

class InviteRequest(BaseModel):
    email: str
    phone: str
    display_name: str
    role: str = "associate"
    organisation_role: Optional[str] = None  # ops_manager | panel_lawyer

@app.post("/api/admin/invite")
async def invite_user(req: InviteRequest, request: Request):
    user = await get_current_user(request)
    _check_permission(user, "admin:users")

    if req.role not in ("partner", "associate", "secretary", "admin"):
        raise HTTPException(status_code=400, detail="Invalid role")

    if req.organisation_role and req.organisation_role not in ("ops_manager", "panel_lawyer"):
        raise HTTPException(status_code=400, detail="Invalid organisation_role")

    phone = req.phone.strip()
    if not phone:
        raise HTTPException(status_code=400, detail="Phone number is required — this is what gates account creation.")

    try:
        cf_rule_id = await _add_email_to_cloudflare_access(req.email)
    except Exception as e:
        print(f"[invite] CF access error (non-fatal): {e}")
        cf_rule_id = None

    async with _db_pool.acquire() as conn:
        try:
            row = await conn.fetchrow("""
                INSERT INTO invites
                    (firm_id, email, phone, display_name, role, invited_by, cf_rule_id,
                     organisation_role)
                VALUES ($1, $2, $3, $4, $5, $6, $7, $8)
                ON CONFLICT (firm_id, email) DO UPDATE SET
                    phone=$3, display_name=$4, role=$5, sent_at=NOW(), cf_rule_id=$7,
                    organisation_role=$8, accepted_at=NULL
                RETURNING *
            """,
            FIRM_ID, req.email, phone, req.display_name, req.role,
            _uuid_mod.UUID(str(user["id"])) if user.get("id") else None,
            str(cf_rule_id) if cf_rule_id else None,
            req.organisation_role
            )
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Could not create invite: {e}")

    invited_by_name = user.get("display_name") or "Your administrator"
    email_sent = await _send_invite_email(req.email, req.display_name, invited_by_name)

    return {
        "invited": True,
        "email": req.email,
        "display_name": req.display_name,
        "role": req.role,
        "organisation_role": req.organisation_role,
        "cloudflare_updated": cf_rule_id is not None,
        "email_sent": email_sent,
    }

class BootstrapAdminRequest(BaseModel):
    phone: str
    email: str
    display_name: str

@app.post("/api/admin/bootstrap")
async def bootstrap_admin(req: BootstrapAdminRequest, request: Request):
    """
    One-time, self-disabling admin provisioning for a freshly deployed
    instance. Closes a real gap found this week: a brand-new instance has
    no supported way to create its first admin — the only path was
    manual SQL directly against Postgres (which is exactly how Lenard had
    to bootstrap his own account on this very instance). This is not an
    authentication feature — the OTP/invite system works fine once a user
    exists. It's a bootstrap *lifecycle* problem: how does a freshly
    deployed instance go from "deployed software" to "owned and usable
    firm system."

    Deliberately stricter than require_admin_token() elsewhere in this
    file, which silently allows access if MUTEMO_ADMIN_TOKEN isn't
    configured — too permissive for an operation this sensitive. Here,
    a missing/unset token means bootstrap is refused outright, not
    silently allowed.

    Self-disabling by construction: refuses unconditionally the moment
    any admin already exists for this firm, regardless of whether the
    token is valid — a provisioning credential for a one-time gap, not a
    standing backdoor.
    """
    if not ADMIN_TOKEN:
        raise HTTPException(status_code=503, detail="Bootstrap is not available - MUTEMO_ADMIN_TOKEN is not configured on this server.")
    token = request.headers.get("X-Admin-Token", "")
    if token != ADMIN_TOKEN:
        raise HTTPException(status_code=403, detail="Invalid admin token")

    async with _db_pool.acquire() as conn:
        existing_admin = await conn.fetchval(
            "SELECT EXISTS(SELECT 1 FROM users WHERE firm_id=$1 AND role='admin' AND is_active=TRUE)",
            FIRM_ID
        )
    if existing_admin:
        raise HTTPException(
            status_code=403,
            detail="An admin already exists for this firm — bootstrap is no longer available. Use the normal invite flow instead."
        )

    phone = req.phone.strip()
    if not phone:
        raise HTTPException(status_code=400, detail="Phone number is required — this is what gates account creation.")

    # Reuses the exact same invites-table + normal phone/OTP login path as
    # every other user, rather than creating a parallel account-creation
    # mechanism — the first admin ends up going through the same tested,
    # audited path as everyone who comes after them.
    try:
        cf_rule_id = await _add_email_to_cloudflare_access(req.email)
    except Exception as e:
        print(f"[bootstrap] CF access error (non-fatal): {e}")
        cf_rule_id = None

    async with _db_pool.acquire() as conn:
        try:
            await conn.execute("""
                INSERT INTO invites (firm_id, email, phone, display_name, role, cf_rule_id, sent_at)
                VALUES ($1, $2, $3, $4, 'admin', $5, NOW())
                ON CONFLICT (firm_id, email) DO UPDATE SET
                    phone=$3, display_name=$4, role='admin', cf_rule_id=$5, sent_at=NOW(), accepted_at=NULL
            """, FIRM_ID, req.email, phone, req.display_name, str(cf_rule_id) if cf_rule_id else None)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Could not create bootstrap admin invite: {e}")

    email_sent = await _send_invite_email(req.email, req.display_name, "Mutemo Desk setup")

    return {
        "bootstrapped": True,
        "message": f"Admin invite created for {phone}. Log in normally via the phone/OTP login screen to complete setup.",
        "cloudflare_updated": cf_rule_id is not None,
        "email_sent": email_sent,
    }

@app.get("/api/admin/export-legal-corpus")
async def export_legal_corpus(request: Request, source: str = "legal", limit: int = 50, offset: int = 0):
    """
    Exports this firm's already-scraped/uploaded legal content — legislation
    and news (source='legal') or ZLR judgments (source='zlr') — in a form
    ready to be replayed into a brand-new firm's instance via the normal
    /api/legal-updates/upload or /api/zlr/upload endpoints. Built so a new
    firm doesn't have to start Search Vault from zero and wait weeks for
    the daily scraper to slowly catch up.

    Deliberately only ever touches legal_updates/zlr_entries — the genuinely
    shared, non-firm-specific legal corpus. Never exports anything from
    documents/matters, which stay private to this firm regardless.
    """
    require_admin_token(request)
    if source not in ("legal", "zlr"):
        raise HTTPException(status_code=400, detail="source must be 'legal' or 'zlr'")

    if source == "legal":
        async with _db_pool.acquire() as conn:
            rows = await conn.fetch("""
                SELECT * FROM legal_updates WHERE firm_id=$1 AND status='complete'
                ORDER BY uploaded_at ASC LIMIT $2 OFFSET $3
            """, FIRM_ID, limit, offset)
            total = await conn.fetchval(
                "SELECT COUNT(*) FROM legal_updates WHERE firm_id=$1 AND status='complete'", FIRM_ID
            )

        items = []
        for r in rows:
            async with _db_pool.acquire() as conn:
                chunk_rows = await conn.fetch(
                    "SELECT text FROM chunks WHERE document_id=$1 AND chunk_source='legal' ORDER BY chunk_index",
                    r["id"]
                )
            full_text = "\n\n".join(c["text"] for c in chunk_rows)
            if not full_text.strip():
                continue  # nothing to replay for this item — skip rather than push an empty document
            items.append({
                "title": r["reference"] or r["filename"],
                "source_type": r["source_type"],
                "source_name": r["source_name"],
                "reference": r["reference"],
                "document_type": r["document_type"],
                "matter_type": r["matter_type"],
                "doc_date": str(r["doc_date"]) if r["doc_date"] else None,
                "court": r["court"],
                "source_url": r["source_url"],
                "text": full_text,
            })
        return {"source": "legal", "total": total, "limit": limit, "offset": offset, "items": items}

    else:  # zlr
        async with _db_pool.acquire() as conn:
            rows = await conn.fetch("""
                SELECT * FROM zlr_entries
                WHERE firm_id=$1 AND raw_text IS NOT NULL AND raw_text != ''
                ORDER BY uploaded_at ASC LIMIT $2 OFFSET $3
            """, FIRM_ID, limit, offset)
            total = await conn.fetchval(
                "SELECT COUNT(*) FROM zlr_entries WHERE firm_id=$1 AND raw_text IS NOT NULL AND raw_text != ''",
                FIRM_ID
            )

        items = [{
            "case_name": r["case_name"],
            "citation": r["citation"],
            "judgment_number": r["judgment_number"],
            "court": r["court"],
            "judge": r["judge"],
            "case_type": r["case_type"],
            "hearing_date": r["hearing_date"],
            "judgment_date": r["judgment_date"],
            "taxonomy_category": r["taxonomy_category"],
            "summary": r["summary"],
            "jurisdiction": r["jurisdiction"],
            "authority_weight": r["authority_weight"],
            "zimlii_url": r["zimlii_url"],
            "text": r["raw_text"],
        } for r in rows]
        return {"source": "zlr", "total": total, "limit": limit, "offset": offset, "items": items}

@app.get("/api/admin/invites")
async def list_invites(request: Request):
    user = await get_current_user(request)
    _check_permission(user, "admin:users")
    async with _db_pool.acquire() as conn:
        rows = await conn.fetch(
            "SELECT * FROM invites WHERE firm_id=$1 ORDER BY sent_at DESC",
            FIRM_ID
        )
    result = []
    for r in rows:
        d = dict(r)
        d["id"] = str(d["id"])
        d["firm_id"] = str(d["firm_id"])
        if d.get("invited_by"):
            d["invited_by"] = str(d["invited_by"])
        if d.get("sent_at"):
            d["sent_at"] = d["sent_at"].isoformat()
        if d.get("accepted_at"):
            d["accepted_at"] = d["accepted_at"].isoformat()
        result.append(d)
    return result

@app.delete("/api/admin/invites/{invite_id}")
async def cancel_invite(invite_id: str, request: Request):
    user = await get_current_user(request)
    _check_permission(user, "admin:users")
    async with _db_pool.acquire() as conn:
        result = await conn.execute(
            "DELETE FROM invites WHERE id=$1 AND firm_id=$2",
            _uuid_mod.UUID(invite_id), FIRM_ID
        )
    if result == "DELETE 0":
        raise HTTPException(status_code=404, detail="Invite not found")
    return {"deleted": True}

@app.patch("/api/invites/{invite_id}/accept")
async def accept_invite(invite_id: str, request: Request):
    """
    Mark an invite as accepted and create the user account.
    Called by the onboarding flow after OTP verification.
    The authenticated user's session must already exist (they logged in via OTP).
    """
    user = await get_current_user(request)

    try:
        invite_uuid = _uuid_mod.UUID(invite_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid invite_id")

    async with _db_pool.acquire() as conn:
        invite = await conn.fetchrow(
            "SELECT * FROM invites WHERE id=$1 AND firm_id=$2",
            invite_uuid, FIRM_ID
        )
        if not invite:
            raise HTTPException(status_code=404, detail="Invite not found")
        if invite["accepted_at"]:
            return {"already_accepted": True, "accepted_at": invite["accepted_at"].isoformat()}

        # Mark invite as accepted
        await conn.execute(
            "UPDATE invites SET accepted_at=NOW() WHERE id=$1",
            invite_uuid
        )

        # If invite has an org role, insert into organisation_roles (spec-correct)
        if invite.get("organisation_role"):
            user_id = _uuid_mod.UUID(str(user["id"])) if user.get("id") else None
            if user_id:
                await conn.execute("""
                    INSERT INTO organisation_roles (firm_id, user_id, role)
                    VALUES ($1, $2, $3)
                    ON CONFLICT (firm_id, user_id) DO UPDATE SET role = $3
                """,
                FIRM_ID, user_id, invite["organisation_role"]
                )

    return {
        "accepted": True,
        "invite_id": invite_id,
        "role": invite["role"],
        "organisation_role": invite.get("organisation_role"),
    }

class UpdateInviteRequest(BaseModel):
    organisation_role: Optional[str] = None
    role: Optional[str] = None  # firm-level role

@app.patch("/api/admin/invites/{invite_id}")
async def update_invite(invite_id: str, req: UpdateInviteRequest, request: Request):
    """
    Update a pending invite's organisation_role or firm role before it is accepted.
    Requires admin:users permission.
    """
    user = await get_current_user(request)
    _check_permission(user, "admin:users")

    try:
        invite_uuid = _uuid_mod.UUID(invite_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid invite_id")

    if req.organisation_role and req.organisation_role not in ("ops_manager", "panel_lawyer"):
        raise HTTPException(status_code=400, detail="Invalid organisation_role")
    if req.role and req.role not in ("partner", "associate", "secretary", "admin"):
        raise HTTPException(status_code=400, detail="Invalid role")

    async with _db_pool.acquire() as conn:
        invite = await conn.fetchrow(
            "SELECT * FROM invites WHERE id=$1 AND firm_id=$2",
            invite_uuid, FIRM_ID
        )
        if not invite:
            raise HTTPException(status_code=404, detail="Invite not found")
        if invite["accepted_at"]:
            raise HTTPException(status_code=409, detail="Invite already accepted — cannot modify")

        updates = []
        params = []
        param_idx = 1
        if req.organisation_role is not None:
            updates.append(f"organisation_role=${param_idx}")
            params.append(req.organisation_role)
            param_idx += 1
        if req.role is not None:
            updates.append(f"role=${param_idx}")
            params.append(req.role)
            param_idx += 1

        if not updates:
            raise HTTPException(status_code=400, detail="No fields to update")

        params.append(invite_uuid)
        await conn.execute(
            f"UPDATE invites SET {', '.join(updates)} WHERE id=${param_idx}",
            *params
        )

        updated = await conn.fetchrow("SELECT * FROM invites WHERE id=$1", invite_uuid)

    d = dict(updated)
    d["id"] = str(d["id"])
    d["firm_id"] = str(d["firm_id"])
    if d.get("invited_by"):
        d["invited_by"] = str(d["invited_by"])
    if d.get("sent_at"):
        d["sent_at"] = d["sent_at"].isoformat()
    return d

import hashlib
# ── Legal Corner — spec-correct endpoints ─────────────────────────────────────────────────

from datetime import timezone

# ── API key auth helper (Bearer token, firm_api_keys table) ────────────────────
async def verify_firm_api_key(request: Request) -> str:
    """
    Validate Authorization: Bearer <key> against firm_api_keys table.
    Returns the firm_id string or raises 401.
    """
    auth_header = request.headers.get("Authorization", "")
    if not auth_header.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Missing or malformed API key")
    raw_key = auth_header.removeprefix("Bearer ").strip()
    key_hash = hashlib.sha256(raw_key.encode()).hexdigest()
    async with _db_pool.acquire() as conn:
        row = await conn.fetchrow(
            "SELECT firm_id FROM firm_api_keys WHERE key_hash=$1 AND revoked_at IS NULL",
            key_hash
        )
    if not row:
        raise HTTPException(status_code=401, detail="Invalid or revoked API key")
    return str(row["firm_id"])

# ── Org role permission helper ───────────────────────────────────────────────
async def _check_org_role(user: dict, firm_id, required_role: str):
    """
    Verify the authenticated user holds the required organisation role.
    Raises 403 if not.
    """
    user_id = _uuid_mod.UUID(str(user["id"])) if user.get("id") else None
    if not user_id:
        raise HTTPException(status_code=403, detail="Org role required")
    if isinstance(firm_id, str):
        firm_id = _uuid_mod.UUID(firm_id)
    async with _db_pool.acquire() as conn:
        row = await conn.fetchrow(
            "SELECT role FROM organisation_roles WHERE firm_id=$1 AND user_id=$2",
            firm_id, user_id
        )
    if not row or row["role"] != required_role:
        raise HTTPException(status_code=403, detail=f"{required_role} role required")

# ── SLA deadline calculation ────────────────────────────────────────────────────────
# PLACEHOLDER: tier-to-hours mapping needs confirming with Legal Corner before go-live
SLA_HOURS_BY_TIER = {"tier_1": 24, "tier_2": 48, "tier_3": 72, "tier_4": 168}

def calculate_sla_deadline(created_at: datetime, coverage_tier: str) -> datetime:
    hours = SLA_HOURS_BY_TIER.get(coverage_tier, 48)
    return created_at + timedelta(hours=hours)

# ── Pydantic models for spec-correct Legal Corner endpoints ─────────────────────
class AutoCreateMatterRequest(BaseModel):
    external_ref: str
    client_name: str
    assigned_lawyer_id: str
    coverage_tier: str
    service_type: str
    description: Optional[str] = None
    client_id: Optional[str] = None
    case_parties: Optional[str] = None

class ReassignRequest(BaseModel):
    to_lawyer_id: str
    reason: Optional[str] = None

class FirmApiKeyRequest(BaseModel):
    label: str = "default"

# ── POST /api/matters/auto-create — server-to-server, API key auth ──────────────
@app.post("/api/matters/auto-create", status_code=201)
async def auto_create_matter(req: AutoCreateMatterRequest, request: Request):
    """
    Auto-create a matter from Legal Corner's subscriber platform.
    Authenticates via Authorization: Bearer <api_key>.
    Idempotent on external_ref.
    """
    firm_id_str = await verify_firm_api_key(request)
    firm_uuid = _uuid_mod.UUID(firm_id_str)

    async with _db_pool.acquire() as conn:
        existing = await conn.fetchrow(
            "SELECT * FROM matters WHERE firm_id=$1 AND external_ref=$2",
            firm_uuid, req.external_ref
        )
        if existing:
            return {**_row_to_doc(existing), "created": False,
                    "message": "Matter already exists for this external_ref"}

        # Verify assigned lawyer is a panel_lawyer in this firm
        try:
            lawyer_uuid = _uuid_mod.UUID(req.assigned_lawyer_id)
        except ValueError:
            raise HTTPException(status_code=422, detail="Invalid assigned_lawyer_id")

        lawyer_check = await conn.fetchrow(
            "SELECT role FROM organisation_roles WHERE firm_id=$1 AND user_id=$2",
            firm_uuid, lawyer_uuid
        )
        if not lawyer_check or lawyer_check["role"] != "panel_lawyer":
            raise HTTPException(status_code=422,
                detail="assigned_lawyer_id is not a panel_lawyer in this firm")

        client_id_uuid = None
        client_name = req.client_name
        if req.client_id:
            try:
                client_id_uuid = _uuid_mod.UUID(req.client_id)
            except ValueError:
                raise HTTPException(status_code=422, detail="client_id must be a valid UUID")
            client_row = await conn.fetchrow(
                "SELECT full_name FROM clients WHERE id=$1 AND firm_id=$2", client_id_uuid, firm_uuid
            )
            if not client_row:
                raise HTTPException(status_code=404, detail="Client not found")
            # Same sync-from-client-record behavior as create_matter/update_matter.
            # Also used below for the derived matter name, since (unlike
            # create_matter) this endpoint doesn't take an explicit `name` —
            # it's built from client_name, so a resolved client should win there too.
            client_name = client_row["full_name"]

        created_at = datetime.now(timezone.utc)
        sla_deadline = calculate_sla_deadline(created_at, req.coverage_tier)

        # last_activity intentionally not passed — unchanged from before
        # this consolidation, stays NULL for this path (see
        # _create_matter_row()'s docstring; a deliberate non-change, not
        # an oversight being carried forward silently).
        row = await _create_matter_row(
            conn, firm_uuid, f"{client_name} — {req.service_type}",
            external_ref=req.external_ref,
            client_name=client_name, client_id=client_id_uuid, case_parties=req.case_parties,
            status="Active",
            assigned_lawyer_id=lawyer_uuid, coverage_tier=req.coverage_tier,
            service_type=req.service_type, sla_deadline=sla_deadline,
            created_at=created_at,
        )

    return {**_row_to_doc(row), "created": True}

# ── POST /api/matters/{matter_id}/reassign — ops_manager only ──────────────────
@app.post("/api/matters/{matter_id}/reassign")
async def reassign_matter_spec(matter_id: str, req: ReassignRequest, request: Request):
    """
    Reassign a matter to a different panel lawyer.
    Requires ops_manager organisation role.
    Records full audit trail in matter_reassignments.
    """
    user = await get_current_user(request)

    try:
        matter_uuid = _uuid_mod.UUID(matter_id)
        to_lawyer_uuid = _uuid_mod.UUID(req.to_lawyer_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid UUID")

    async with _db_pool.acquire() as conn:
        matter = await conn.fetchrow(
            "SELECT * FROM matters WHERE id=$1 AND firm_id=$2",
            matter_uuid, FIRM_ID
        )
        if not matter:
            raise HTTPException(status_code=404, detail="Matter not found")

        await _check_org_role(user, matter["firm_id"], "ops_manager")

        user_id = _uuid_mod.UUID(str(user["id"]))

        # Record audit trail
        await conn.execute("""
            INSERT INTO matter_reassignments
                (matter_id, from_lawyer_id, to_lawyer_id, reassigned_by_id, reason, firm_id)
            VALUES ($1, $2, $3, $4, $5, $6)
        """,
        matter_uuid, matter.get("assigned_lawyer_id"),
        to_lawyer_uuid, user_id, req.reason, matter["firm_id"]
        )

        # Update matter
        await conn.execute(
            "UPDATE matters SET assigned_lawyer_id=$1, assigned_by_id=$2 WHERE id=$3",
            to_lawyer_uuid, user_id, matter_uuid
        )

    return {
        "status": "reassigned",
        "matter_id": matter_id,
        "to_lawyer_id": req.to_lawyer_id,
    }

# ── GET /api/organisations/{firm_id}/lawyers — list panel lawyers ───────────────
@app.get("/api/organisations/{firm_id}/lawyers")
async def list_org_lawyers(firm_id: str, request: Request):
    """
    List panel lawyers for a firm.
    Accepts session auth (ops_manager) or API key auth (matching firm_id).
    """
    try:
        firm_uuid = _uuid_mod.UUID(firm_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid firm_id")

    # Try session auth first, then API key
    try:
        user = await get_current_user(request)
        await _check_org_role(user, firm_uuid, "ops_manager")
    except HTTPException:
        api_firm_id = await verify_firm_api_key(request)
        if api_firm_id != firm_id:
            raise HTTPException(status_code=403, detail="API key does not match firm_id")

    async with _db_pool.acquire() as conn:
        rows = await conn.fetch("""
            SELECT u.id, u.display_name, u.phone
            FROM users u
            JOIN organisation_roles o ON o.user_id = u.id
            WHERE o.firm_id = $1 AND o.role = 'panel_lawyer'
            ORDER BY u.display_name
        """, firm_uuid)

    return {"lawyers": [{"id": str(r["id"]), "display_name": r["display_name"],
                         "phone": r["phone"]} for r in rows]}

# ── GET /api/organisations/{firm_id}/matters — ops manager SLA view ────────────
@app.get("/api/organisations/{firm_id}/matters")
async def list_org_matters(firm_id: str, request: Request):
    """List all matters with SLA status. Requires ops_manager role."""
    try:
        firm_uuid = _uuid_mod.UUID(firm_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid firm_id")

    user = await get_current_user(request)
    await _check_org_role(user, firm_uuid, "ops_manager")

    async with _db_pool.acquire() as conn:
        rows = await conn.fetch("""
            SELECT * FROM v_legal_corner_sla_status
            WHERE firm_id = $1
            ORDER BY sla_deadline ASC NULLS LAST
        """, firm_uuid)

    result = []
    for r in rows:
        d = dict(r)
        for k, v in d.items():
            if hasattr(v, 'isoformat'):
                d[k] = v.isoformat()
            elif isinstance(v, _uuid_mod.UUID):
                d[k] = str(v)
        result.append(d)
    return {"matters": result}

# ── GET /api/organisations/{firm_id}/dashboard — ops dashboard stats ───────────
@app.get("/api/organisations/{firm_id}/dashboard")
async def org_dashboard_spec(firm_id: str, request: Request):
    """Return SLA dashboard stats. Requires ops_manager role."""
    try:
        firm_uuid = _uuid_mod.UUID(firm_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid firm_id")

    user = await get_current_user(request)
    await _check_org_role(user, firm_uuid, "ops_manager")

    async with _db_pool.acquire() as conn:
        stats = await conn.fetchrow("""
            SELECT
                count(*) FILTER (WHERE status != 'complete') AS active_matters,
                count(*) FILTER (WHERE is_overdue = true) AS overdue_matters,
                count(DISTINCT assigned_lawyer_id) AS active_lawyers,
                count(*) FILTER (WHERE reassignment_count > 0) AS reassigned_matters
            FROM v_legal_corner_sla_status
            WHERE firm_id = $1
        """, firm_uuid)

    return {
        "firm_id": firm_id,
        "active_matters": stats["active_matters"],
        "overdue_matters": stats["overdue_matters"],
        "active_lawyers": stats["active_lawyers"],
        "reassigned_matters": stats["reassigned_matters"],
    }

# ── POST /api/admin/firm-api-keys — generate a firm API key ─────────────────────
@app.post("/api/admin/firm-api-keys", status_code=201)
async def create_firm_api_key(req: FirmApiKeyRequest, request: Request):
    """Generate a new API key for server-to-server auth. Requires admin:users permission."""
    user = await get_current_user(request)
    _check_permission(user, "admin:users")

    raw_key = f"fk_{secrets.token_urlsafe(32)}"
    key_hash = hashlib.sha256(raw_key.encode()).hexdigest()

    async with _db_pool.acquire() as conn:
        try:
            row = await conn.fetchrow("""
                INSERT INTO firm_api_keys (firm_id, key_hash, label)
                VALUES ($1, $2, $3) RETURNING id, label, created_at
            """, FIRM_ID, key_hash, req.label)
        except asyncpg.UniqueViolationError:
            raise HTTPException(status_code=409, detail="A key with this label already exists")

    return {
        "id": str(row["id"]),
        "label": row["label"],
        "api_key": raw_key,  # shown once — store securely
        "created_at": row["created_at"].isoformat(),
        "note": "Store this key securely. It will not be shown again.",
    }

# ── GET /api/auth/profile — return user profile including org role ───────────────
@app.get("/api/auth/profile")
async def get_user_profile(request: Request):
    """
    Return the current user's profile including firm role and organisation role.
    Used by the frontend on login to determine which UI features to show.
    """
    user = await get_current_user(request)
    user_id = _uuid_mod.UUID(str(user["id"])) if user.get("id") else None

    org_role = None
    firm_features = []
    firm_logo_url = None

    async with _db_pool.acquire() as conn:
        if user_id:
            org_row = await conn.fetchrow(
                "SELECT role FROM organisation_roles WHERE firm_id=$1 AND user_id=$2",
                FIRM_ID, user_id
            )
            org_role = org_row["role"] if org_row else None

        firm_row = await conn.fetchrow(
            "SELECT features, firm_logo_url FROM firms WHERE id=$1", FIRM_ID
        )
        if firm_row:
            firm_features = firm_row["features"] or []
            firm_logo_url = firm_row["firm_logo_url"]

    return {
        "id": str(user["id"]),
        "display_name": user.get("display_name"),
        "role": user.get("role"),
        "org_role": org_role,
        "firm_id": str(FIRM_ID),
        "firm_name": FIRM_NAME,
        "firm_logo_url": firm_logo_url,
        "features": firm_features,
    }

# ── Frontend static files ─────────────────────────────────────────────────────
frontend_path = os.path.join(os.path.dirname(__file__), "../frontend")
assets_path = os.path.join(frontend_path, "assets")
if os.path.exists(assets_path):
    app.mount("/assets", StaticFiles(directory=assets_path), name="assets")

client = anthropic.Anthropic()

# ── Semantic Search: embeddings + ChromaDB ────────────────────────────────────
EMBEDDING_MODEL_NAME = "all-MiniLM-L6-v2"
_embedding_model = None
_chroma_client = None
_firm_collection = None
_legal_collection = None
_zlr_collection = None

def get_embedding_model():
    global _embedding_model
    if _embedding_model is None:
        from sentence_transformers import SentenceTransformer
        print(f"[embeddings] loading model '{EMBEDDING_MODEL_NAME}'...")
        _embedding_model = SentenceTransformer(EMBEDDING_MODEL_NAME)
        print("[embeddings] model loaded")
    return _embedding_model

def embed_texts(texts: list) -> list:
    import numpy as np
    model = get_embedding_model()
    vectors = model.encode(texts, convert_to_numpy=True, show_progress_bar=False)
    vectors = np.array(vectors)
    if vectors.ndim == 1:
        vectors = vectors.reshape(1, -1)
    elif vectors.ndim > 2:
        vectors = vectors.reshape(len(texts), -1)
    return [v.tolist() for v in vectors]

def get_chroma_collections():
    global _chroma_client, _firm_collection, _legal_collection, _zlr_collection
    if _chroma_client is None:
        import chromadb
        # CHROMA_DATA_DIR should point at a mounted persistent volume in
        # production (e.g. /data/chroma) — without it, the vector index
        # lives on the container's ephemeral filesystem and is wiped on
        # every redeploy, even though Postgres (the real source of truth
        # for chunk text) survives fine. Falls back to the old relative
        # path for local/dev use where persistence across restarts doesn't
        # matter.
        chroma_path = os.environ.get(
            "CHROMA_DATA_DIR",
            os.path.join(os.path.dirname(__file__), "..", "data", "chroma")
        )
        os.makedirs(chroma_path, exist_ok=True)
        _chroma_client = chromadb.PersistentClient(path=chroma_path)
        _firm_collection = _chroma_client.get_or_create_collection(
            "firm_precedents", metadata={"hnsw:space": "cosine"}
        )
        _legal_collection = _chroma_client.get_or_create_collection(
            "legal_updates", metadata={"hnsw:space": "cosine"}
        )
        _zlr_collection = _chroma_client.get_or_create_collection(
            "zlr_index", metadata={"hnsw:space": "cosine"}
        )
        print("[vector_store] ChromaDB initialized")
    return _firm_collection, _legal_collection, _zlr_collection

async def reconcile_chroma_index():
    """
    Self-healing check, run once at startup: compare each ChromaDB
    collection against Postgres per chunk_id, using content_hash rather
    than a bare count. Postgres is the real source of truth (chunk text is
    already fully extracted and chunked there) — ChromaDB is a derived
    index that can go out of sync if it's ever reset (volume detached,
    first deploy after adding a volume, manual intervention, etc.).

    A count comparison alone can't catch every drift: equal counts don't
    prove equal content (e.g. a wrong/stale vector sitting under the right
    chunk_id after a partial rebuild). Comparing content_hash per chunk_id
    catches that. Three outcomes per chunk_id:
      - in Postgres, missing from Chroma      -> index it
      - in both, but hashes differ            -> re-index it (upsert)
      - in Chroma only (stray, e.g. an orphan
        left by a deleted document)           -> remove it
    Already-matching chunk_ids are never touched — this fetches ids+hashes
    (cheap: two short strings each) for every chunk on every boot, but only
    ever re-embeds ones that actually need it, same as the count-only
    version's stated goal, just with real detection instead of a count that
    can silently miss drift.
    """
    if not _db_pool:
        return
    try:
        firm_col, legal_col, zlr_col = get_chroma_collections()
    except Exception as e:
        print(f"[reconcile] ChromaDB unavailable, skipping: {e}")
        return

    async with _db_pool.acquire() as conn:
        pg_rows_by_source = {}
        for source in ("firm", "legal", "zlr"):
            rows = await conn.fetch(
                "SELECT id, content_hash FROM chunks WHERE firm_id=$1 AND chunk_source=$2",
                FIRM_ID, source
            )
            pg_rows_by_source[source] = {r["id"]: r["content_hash"] for r in rows}

    collections = {"firm": firm_col, "legal": legal_col, "zlr": zlr_col}
    for source, col in collections.items():
        pg_hashes = pg_rows_by_source.get(source, {})
        if not pg_hashes:
            continue  # nothing to index for this source

        try:
            chroma_data = col.get(include=["metadatas"])
            chroma_hashes = {
                cid: (meta or {}).get("content_hash")
                for cid, meta in zip(chroma_data["ids"], chroma_data["metadatas"])
            }
        except Exception as e:
            print(f"[reconcile] {source}: failed to read ChromaDB collection, skipping: {e}")
            continue

        missing = pg_hashes.keys() - chroma_hashes.keys()
        mismatched = {
            cid for cid in (pg_hashes.keys() & chroma_hashes.keys())
            if pg_hashes[cid] != chroma_hashes[cid]
        }
        stray = chroma_hashes.keys() - pg_hashes.keys()
        to_reindex = missing | mismatched

        if not to_reindex and not stray:
            continue  # already in sync, nothing to do

        print(f"[reconcile] {source}: {len(missing)} missing, {len(mismatched)} content-mismatched, "
              f"{len(stray)} stray in ChromaDB — repairing")

        if to_reindex:
            async with _db_pool.acquire() as conn:
                rows = await conn.fetch(
                    "SELECT * FROM chunks WHERE firm_id=$1 AND chunk_source=$2 AND id = ANY($3)",
                    FIRM_ID, source, list(to_reindex)
                )
            chunks_to_index = [{
                "id": r["id"],
                "text": r["text"],
                "document_id": str(r["document_id"]),
                "matter_id": r["matter_id"],
                "chunk_index": r["chunk_index"],
                "page_number": r["page_number"],
                "content_hash": r["content_hash"],
            } for r in rows]
            if chunks_to_index:
                await asyncio.to_thread(index_chunks_in_chroma, chunks_to_index, source)
                print(f"[reconcile] {source}: re-indexed {len(chunks_to_index)} chunk(s) "
                      f"({len(missing)} missing, {len(mismatched)} mismatched)")

        if stray:
            await asyncio.to_thread(remove_chunks_from_chroma, list(stray), source)
            print(f"[reconcile] {source}: removed {len(stray)} stray chunk(s) with no Postgres row")

# ── Pydantic Models ───────────────────────────────────────────────────────────

MATTER_STATUSES = ["Active", "Awaiting Client", "Awaiting Court", "On Hold", "Closed"]

# matter_type enum for POST /api/onboarding/intake specifically — the
# general create_matter/update_matter matter_type field stays free-text
# (Optional[str], validated nowhere) as it already is today; this pass
# doesn't touch that. Mirrors the existing frontend New Matter dropdown's
# option values, plus "litigation_general" — a generic-litigation catch-all
# the existing list doesn't have one of (it only has specific litigation
# sub-types: eviction, matrimonial, criminal, etc.) — added because
# config/case_binder_templates.yml seeds a starter checklist for it.
# Deliberately a superset of case_binder.known_matter_types(): most of
# these values don't have a case-binder entry yet, and that's expected —
# provision_case_binder() returns no starter documents for those, not an
# error (see backend/case_binder.py's own docstring).
INTAKE_MATTER_TYPES = [
    "eviction", "estate", "trust", "employment", "commercial_property",
    "commercial_contract", "conveyancing", "customary_law", "matrimonial",
    "family_law", "company_law", "criminal", "constitutional",
    "debt_collection", "mining", "litigation_general", "other",
]

# Document provenance metadata for client/matter (Vault) documents -- see
# the documents.provenance_document_type/document_status/confidentiality
# migration comment (run_migrations()) for why provenance_document_type
# is a deliberately separate concept/column from the existing
# document_type (AI-classified specific form, feeds legal_source_type).
# Kept as real Python lists (matching PRACTICE_AREAS's convention) rather
# than deriving from the DB CHECK constraints, so a mismatch between the
# two would show up immediately as a real bug, not silently resolve to
# whatever Postgres happens to accept.
PROVENANCE_DOCUMENT_TYPES = [
    "Pleading", "Contract", "Correspondence", "JudgmentOrder",
    "Evidence", "Research", "Precedent", "General",
]
DOCUMENT_STATUSES = ["Draft", "Review", "Final", "Executed", "Superseded"]
DOCUMENT_CONFIDENTIALITY_LEVELS = ["Standard", "Restricted", "Privileged"]

# AML/KYC Client Compliance (Money Laundering and Proceeds of Crime Act
# [Chapter 9:24]) -- kept as real Python lists mirroring the DB CHECK
# constraints, same convention/reasoning as PROVENANCE_DOCUMENT_TYPES above.
CLIENT_TYPES = [
    "Individual", "Company", "Partnership", "Trust", "Estate",
    "NonProfit", "Government", "Other",
]
# Legal-person types for which Part 2 (beneficial ownership, s15/s17(b))
# applies -- an Individual is inherently their own beneficial owner.
LEGAL_PERSON_CLIENT_TYPES = [t for t in CLIENT_TYPES if t != "Individual"]
VERIFICATION_STATUSES = ["Unverified", "Pending", "Verified"]
AUTHORITY_BASIS_TYPES = [
    "PowerOfAttorney", "BoardResolution", "Mandate",
    "LetterOfAdministration", "LetterOfExecutorship", "Other",
]
PEP_BASIS_TYPES = ["Self", "BeneficialOwner", "CloseAssociate", "NotApplicable"]
RISK_RATINGS = ["Low", "Medium", "High", "NotAssessed"]
CLIENT_IS_BENEFICIAL_OWNER_VALUES = ["Yes", "No", "Unknown"]

class MatterCreate(BaseModel):
    name: str
    number: Optional[str] = None
    internal_ref: Optional[str] = None
    external_ref: Optional[str] = None
    matter_type: Optional[str] = None
    practice_area: Optional[str] = None  # validated against PRACTICE_AREAS in create_matter/update_matter
    status: Optional[str] = "Active"
    client_name: Optional[str] = None
    client_id: Optional[str] = None
    case_parties: Optional[str] = None
    custom_status: Optional[str] = None
    next_deadline: Optional[str] = None
    next_deadline_note: Optional[str] = None

class MatterUpdate(BaseModel):
    name: Optional[str] = None
    internal_ref: Optional[str] = None
    external_ref: Optional[str] = None
    matter_type: Optional[str] = None
    practice_area: Optional[str] = None
    status: Optional[str] = None
    client_name: Optional[str] = None
    client_id: Optional[str] = None
    case_parties: Optional[str] = None
    custom_status: Optional[str] = None
    next_deadline: Optional[str] = None       # ISO date string, e.g. "2026-08-15"
    next_deadline_note: Optional[str] = None
    # Matter review safety net — a soft "please look at this" nudge date,
    # distinct from next_deadline above. If omitted from a PATCH entirely,
    # update_matter() re-defaults it to today + DEFAULT_REVIEW_INTERVAL_DAYS
    # rather than leaving it untouched — see that function for why.
    next_review_date: Optional[str] = None    # ISO date string, e.g. "2026-08-15"
    # Fee tracking — the firm's own professional fees only, manually
    # entered. Not trust accounting; see the schema comment in
    # run_migrations() for why this must not become a second source of
    # truth for client funds held in trust.
    amount_billed: Optional[float] = None
    amount_received: Optional[float] = None
    # Conveyancing-specific (backend/conveyancing.py's CONVEYANCING_MILESTONES)
    # — accepted regardless of the matter's current practice_area; the
    # practice_area === 'Conveyancing/Property' check that gates whether
    # this section is shown is a frontend display condition only.
    conveyancing_milestone: Optional[str] = None
    conveyancing_property_address: Optional[str] = None
    conveyancing_title_deed_number: Optional[str] = None
    conveyancing_purchase_price: Optional[float] = None  # reference fact only — not a balance or funds held
    conveyancing_other_conveyancer_contact: Optional[str] = None
    conveyancing_transfer_date: Optional[str] = None                # ISO date string
    conveyancing_rates_clearance_expiry: Optional[str] = None       # ISO date string
    conveyancing_bond_registration_deadline: Optional[str] = None   # ISO date string
    # Matter Progress Tracker (visual stepper) — generic current-stage
    # field for matter_type values without a type-specific column
    # (backend/matter_stages.py). Validated against that matter's
    # resolved stage sequence in update_matter(), not against a fixed
    # list here — which sequence applies depends on the matter's own
    # matter_type/practice_area.
    stage: Optional[str] = None

class ClientCreate(BaseModel):
    full_name: str
    email: Optional[str] = None
    phone: Optional[str] = None
    physical_address: Optional[str] = None
    id_or_registration_number: Optional[str] = None
    contact_person: Optional[str] = None  # corporate/entity clients only — blank for individuals
    notes: Optional[str] = None

class TrustParty(BaseModel):
    """s17(c): one trustee/settlor/beneficiary entry."""
    name: str
    id_number: Optional[str] = None
    role: Optional[str] = None

class ClientUpdate(BaseModel):
    full_name: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    physical_address: Optional[str] = None
    id_or_registration_number: Optional[str] = None
    contact_person: Optional[str] = None
    notes: Optional[str] = None
    # ── AML/KYC (s17) — validated against CLIENT_TYPES etc. in update_client() ──
    client_type: Optional[str] = None
    # s17(a): individual
    date_of_birth: Optional[str] = None
    place_of_birth: Optional[str] = None
    national_id_number: Optional[str] = None
    passport_number: Optional[str] = None
    id_expiry_date: Optional[str] = None
    residential_address: Optional[str] = None
    occupation: Optional[str] = None
    employer_or_business: Optional[str] = None
    # s17(a): company/partnership/other legal person
    registered_name: Optional[str] = None
    trading_name: Optional[str] = None
    registration_number: Optional[str] = None
    date_incorporated: Optional[str] = None
    registered_office_address: Optional[str] = None
    principal_business_address: Optional[str] = None
    proof_of_incorporation_document_id: Optional[str] = None
    governing_document_id: Optional[str] = None
    # s17(c): trust/estate
    trustees: Optional[List[TrustParty]] = None
    settlors: Optional[List[TrustParty]] = None
    beneficiaries: Optional[List[TrustParty]] = None

# ── AML/KYC: beneficial ownership (s15/s17(b)) ──────────────────────────────

class BeneficialOwnerCreate(BaseModel):
    owner_name: str
    date_of_birth: Optional[str] = None
    nationality: Optional[str] = None
    id_or_passport_number: Optional[str] = None
    residential_address: Optional[str] = None
    ownership_or_control_basis: Optional[str] = None
    ownership_percentage: Optional[float] = None  # not required — basis need not be percentage-based

class BeneficialOwnerUpdate(BaseModel):
    owner_name: Optional[str] = None
    date_of_birth: Optional[str] = None
    nationality: Optional[str] = None
    id_or_passport_number: Optional[str] = None
    residential_address: Optional[str] = None
    ownership_or_control_basis: Optional[str] = None
    ownership_percentage: Optional[float] = None
    verification_status: Optional[str] = None
    verified_date: Optional[str] = None

# ── AML/KYC: authorized representative (s17(d)) ─────────────────────────────

class AuthorizedRepresentativeCreate(BaseModel):
    full_name: str
    position_or_relationship: Optional[str] = None
    id_or_passport_number: Optional[str] = None
    contact_details: Optional[str] = None
    authority_basis: Optional[str] = None
    authority_document_id: Optional[str] = None

class AuthorizedRepresentativeUpdate(BaseModel):
    full_name: Optional[str] = None
    position_or_relationship: Optional[str] = None
    id_or_passport_number: Optional[str] = None
    contact_details: Optional[str] = None
    authority_basis: Optional[str] = None
    authority_document_id: Optional[str] = None
    verification_status: Optional[str] = None
    verified_date: Optional[str] = None

# ── AML/KYC: compliance / PEP (s20) ──────────────────────────────────────────

class ClientComplianceUpdate(BaseModel):
    identity_verification_status: Optional[str] = None
    client_is_beneficial_owner: Optional[str] = None
    is_pep: Optional[bool] = None
    pep_basis: Optional[str] = None
    pep_position: Optional[str] = None
    pep_country: Optional[str] = None
    senior_management_approved_by: Optional[str] = None
    senior_management_approved_date: Optional[str] = None
    source_of_wealth: Optional[str] = None
    source_of_funds: Optional[str] = None
    enhanced_monitoring_required: Optional[bool] = None
    risk_rating: Optional[str] = None
    retained_until: Optional[str] = None
    # Reuses the real, existing GET /api/matters/check-conflict —
    # conflict_check_reviewed is set true by the frontend once a lawyer
    # has actually run that check for this client and confirmed it (see
    # runClientConflictCheck() / markConflictReviewed() in index.html).
    conflict_check_reviewed: Optional[bool] = None

class ProgressNote(BaseModel):
    text: str
    author: Optional[str] = None
    # Matter review safety net — adding a note is the canonical "I just
    # worked on/reviewed this matter" action, so it stamps the review
    # clock the same way a PATCH does (see add_progress_note() and
    # _resolve_review_dates()). Optional: leave unset for the default
    # today + DEFAULT_REVIEW_INTERVAL_DAYS.
    next_review_date: Optional[str] = None    # ISO date string, e.g. "2026-08-15"

class AffidavitRequest(BaseModel):
    matter_type: Optional[str] = None
    court: Optional[str] = "High Court of Zimbabwe"
    deponent_name: Optional[str] = None
    deponent_id: Optional[str] = None
    deponent_capacity: Optional[str] = None
    parties: Optional[str] = None
    matter_summary: str
    key_facts: Optional[str] = None
    relief: Optional[str] = None
    precedent_context: Optional[dict] = None

class DocumentRequest(BaseModel):
    doc_type: str
    plaintiff: Optional[str] = None
    defendant: Optional[str] = None
    court: Optional[str] = None
    case_number: Optional[str] = None
    facts: str
    instructions: Optional[str] = None
    precedent_context: Optional[dict] = None

class SearchRequest(BaseModel):
    query: str
    matter_type: Optional[str] = None
    document_type: Optional[str] = None
    matter_id: Optional[str] = None
    limit: int = 8
    include_legal_updates: bool = True

class ExportRequest(BaseModel):
    affidavit_text: str
    deponent_name: Optional[str] = "Deponent"
    document_id: Optional[str] = "DOC"

class ExportDocumentDocxRequest(BaseModel):
    content_html: str
    filename: Optional[str] = "Document"

class Attendee(BaseModel):
    email: str
    name: Optional[str] = None

class CalendarInviteRequest(BaseModel):
    attendees: List[Attendee]
    invite_message: Optional[str] = None

class CalendarInviteResponseRequest(BaseModel):
    status: str  # "accepted" or "declined"

class CalendarEventUpdate(BaseModel):
    title: Optional[str] = None
    date: Optional[str] = None
    time: Optional[str] = None
    court: Optional[str] = None
    notes: Optional[str] = None
    event_type: Optional[str] = None
    update_message: Optional[str] = None  # note to attendees explaining the change

class CalendarEvent(BaseModel):
    title: str
    matter_id: Optional[str] = None
    matter_name: Optional[str] = None
    event_type: str
    date: str
    time: Optional[str] = None
    court: Optional[str] = None
    notes: Optional[str] = None
    attendees: Optional[List[Attendee]] = None
    invite_message: Optional[str] = None

class LegalUpdateSearchRequest(BaseModel):
    query: str
    source_type: Optional[str] = None
    limit: int = 8

class ReminderSettings(BaseModel):
    enabled: bool
    recipient_email: str
    send_hour_utc: int = 5

class DigestSettings(BaseModel):
    enabled: bool
    recipient_email: str
    send_hour_utc: int = 6

# ── DB helpers ────────────────────────────────────────────────────────────────

async def _resolve_user_initials(conn, firm_id, user_id, display_name, persist: bool = True) -> str:
    """
    Returns a user's numbering-prefix initials (backend/numbering.py),
    generating and disambiguating them against the firm's other users if
    not already set. `user_id` may be None — the synthetic user returned by
    get_current_user() when AUTH_ENABLED is False has no real users row —
    in which case this only computes a value, never persists.

    `persist=False` also computes without writing, for preview-only paths
    (e.g. the bulk-upload endpoint's commit=False mode) that must not touch
    the DB even for an existing user whose initials aren't set yet.
    """
    if user_id is not None:
        row = await conn.fetchrow("SELECT initials FROM users WHERE id=$1", user_id)
        if row and row["initials"]:
            return row["initials"]
    existing = await conn.fetch(
        "SELECT initials FROM users WHERE firm_id=$1 AND initials IS NOT NULL", firm_id
    )
    initials = disambiguate_initials(generate_initials(display_name), {r["initials"] for r in existing})
    if persist and user_id is not None:
        await conn.execute("UPDATE users SET initials=$1 WHERE id=$2", initials, user_id)
    return initials

async def _allocate_next_seq(conn, firm_id, prefix: str, compute_seed) -> int:
    """
    Atomically allocates the next sequence number for (firm_id, prefix) via
    numbering_counters, replacing the old MAX(existing)+1 scan that raced
    under concurrent requests.

    `compute_seed` is an async callable, invoked ONLY the first time this
    prefix is ever allocated (existing counter rows skip straight to the
    atomic UPDATE below, no scan at all) — it should return the correct
    starting value by scanning the real table, exactly like the old logic
    did, so a prefix with pre-existing numbers (e.g. NGM-001..NGM-071)
    seeds itself correctly the first time it's touched under this scheme,
    without any separate backfill migration.

    The UPDATE...RETURNING is the atomic step: Postgres takes a row lock
    for its duration, so two concurrent callers serialize on it and always
    get distinct, consecutive values — never the same one. The seed INSERT
    uses ON CONFLICT DO NOTHING, so even a true race on the very first-ever
    call for a brand new prefix (two concurrent callers both finding no
    counter row yet) is safe: at most one INSERT wins, the other silently
    no-ops, and both then proceed through the same atomic UPDATE.
    """
    existing = await conn.fetchval(
        "SELECT 1 FROM numbering_counters WHERE firm_id=$1 AND prefix=$2", firm_id, prefix
    )
    if existing is None:
        seed = await compute_seed()
        await conn.execute(
            """
            INSERT INTO numbering_counters (firm_id, prefix, next_seq)
            VALUES ($1, $2, $3)
            ON CONFLICT (firm_id, prefix) DO NOTHING
            """,
            firm_id, prefix, seed,
        )
    row = await conn.fetchrow(
        """
        UPDATE numbering_counters
        SET next_seq = next_seq + 1
        WHERE firm_id=$1 AND prefix=$2
        RETURNING next_seq - 1 AS allocated
        """,
        firm_id, prefix,
    )
    return row["allocated"]

async def _next_client_number(conn, firm_id, initials: str) -> str:
    async def compute_seed():
        rows = await conn.fetch(
            "SELECT client_number FROM clients WHERE firm_id=$1 AND client_number LIKE $2",
            firm_id, f"{initials}-%",
        )
        return next_sequence([r["client_number"] for r in rows], initials)
    seq = await _allocate_next_seq(conn, firm_id, initials, compute_seed)
    return format_client_number(initials, seq)

async def _next_matter_number(conn, firm_id, client_number: str) -> str:
    async def compute_seed():
        rows = await conn.fetch(
            "SELECT matter_number FROM matters WHERE firm_id=$1 AND matter_number LIKE $2",
            firm_id, f"{client_number}-%",
        )
        return next_sequence([r["matter_number"] for r in rows], client_number)
    seq = await _allocate_next_seq(conn, firm_id, client_number, compute_seed)
    return format_matter_number(client_number, seq)


# ── Shared client/matter row creation ──────────────────────────────────────
# Phase 1a consolidation (2026-08-26): create_client(), create_matter(),
# auto_create_matter(), bulk_import_matters(), bulk_onboard_from_excel(),
# and client_intake() each used to hand-write their own INSERT INTO
# clients/matters, with no shared function -- a real, live drift risk
# (any change to numbering/defaults had to be made correctly in up to
# five places by hand). Diffed all five field-by-field before writing
# this; see git history for the full comparison. Two real behavioral
# differences were found and are fixed here, not just consolidated:
#   - bulk_onboard_from_excel's commit-time client/matter numbering used
#     to be a manual in-Python counter seeded from a SELECT...LIKE scan,
#     never touching numbering_counters -- not concurrency-safe, unlike
#     every other path. Now goes through the same atomic
#     _next_client_number()/_next_matter_number() (real row lock on
#     numbering_counters) as everyone else. Preview mode (commit=False)
#     is untouched -- it must never burn a real sequence number for a
#     plan that might not be applied.
#   - bulk_onboard_from_excel's client rows never got created_by set at
#     all. Now set to the onboarded/matched lawyer's id, same as the
#     matter rows it creates in the same request.
# Every other field-by-field difference found (client_intake's Case
# Binder provisioning and audit log, auto_create_matter's SLA fields,
# bulk_import_matters' progress-notes side effect and lack of a client
# link, differing permission checks, auto_create_matter's NULL
# last_activity, etc.) is preserved exactly as each caller had it --
# these two functions are pure row-write primitives with no validation
# and no policy of their own; each caller decides what to pass.


async def _create_client_row(
    conn, firm_id, initials: str, full_name: str, *,
    email: Optional[str] = None,
    phone: Optional[str] = None,
    physical_address: Optional[str] = None,
    id_or_registration_number: Optional[str] = None,
    contact_person: Optional[str] = None,
    notes: Optional[str] = None,
    created_by: Optional[_uuid_mod.UUID] = None,
    client_type: Optional[str] = None,
    client_is_beneficial_owner: Optional[str] = None,
) -> "asyncpg.Record":
    """
    Shared INSERT for a new clients row. Atomically allocates the next
    client_number under `initials` via _next_client_number() -- always a
    real write; never call this from a preview branch (preview's
    "would-be number" stays a separate, read-only next_sequence()
    estimate in each caller, unchanged).

    Takes `initials` (already resolved), not a user to resolve them
    from: callers differ on whose initials to number under (the
    submitting user for create_client; the assigned/onboarded lawyer for
    client_intake and bulk_onboard_from_excel), and bulk_onboard_from_excel
    specifically needs to resolve initials once per upload and reuse them
    across every client it creates rather than re-resolving per row.
    Call _resolve_user_initials() yourself first.

    Also creates the client's client_compliance row, in the same default
    "not yet assessed" state _DEFAULT_CLIENT_COMPLIANCE encodes (every
    other column left to its own DB DEFAULT: identity_verification_status
    'Unverified', risk_rating 'NotAssessed', the boolean flags FALSE) --
    closes a real gap where bulk-migrated clients (via
    bulk_onboard_from_excel/client_intake) previously got no
    client_compliance row at creation at all, unlike clients created via
    the normal single-client flow -- both got none before this, but
    "none" is now uniformly replaced with an explicit not-yet-assessed
    row for every path. `client_type` (a clients-table column) and
    `client_is_beneficial_owner` (a client_compliance-table column) are
    optional overrides for the two compliance fields the bulk-onboard
    Excel template can populate directly when a lawyer/secretary already
    knows them off-hand; every other compliance field stays at its
    default until someone opens the compliance modal. create_client()
    and client_intake() don't collect either field today and pass
    neither, unchanged from their prior behavior.
    """
    cid = _uuid_mod.uuid4()
    now = datetime.utcnow()
    client_number = await _next_client_number(conn, firm_id, initials)
    row = await conn.fetchrow("""
        INSERT INTO clients (id, firm_id, full_name, email, phone,
                             physical_address, id_or_registration_number, contact_person, notes,
                             client_number, client_type, created_by, created_at, updated_at)
        VALUES ($1,$2,$3,$4,$5,$6,$7,$8,$9,$10,$11,$12,$13,$13)
        RETURNING *
    """,
    cid, firm_id, full_name, email, phone,
    physical_address, id_or_registration_number, contact_person, notes,
    client_number, client_type, created_by, now,
    )
    await conn.execute("""
        INSERT INTO client_compliance (client_id, firm_id, client_is_beneficial_owner)
        VALUES ($1, $2, $3)
    """, cid, firm_id, client_is_beneficial_owner)
    return row


# Matter review safety net (2026-08-30): confirmed by the firm at 30 days.
# Named constant, not a magic number, so this stays easy to change later
# without hunting through _create_matter_row()/update_matter(). See the
# next_review_date column comment in run_migrations() for the concept.
DEFAULT_REVIEW_INTERVAL_DAYS = 30
# How far ahead (in days) a matter's upcoming next_review_date has to be
# before it's worth surfacing in the daily digest as "approaching" —
# overdue ones (past their date already) always show regardless of this.
REVIEW_DIGEST_LOOKAHEAD_DAYS = 7


def _resolve_review_dates(explicit_next_review_date: Optional[str]) -> tuple:
    """
    Returns (next_review_date, last_reviewed_date) for a matter being
    touched right now -- shared by every place a matter can be "reviewed"
    (update_matter()'s PATCH, add_progress_note()'s note-adding flow, and
    any future one), so the defaulting logic lives in exactly one place.

    last_reviewed_date is always today -- any touch that reaches this
    function IS the review happening right now, whether or not a real
    next_review_date was also given. This is deliberately a real, honest
    timestamp of when the matter was actually looked at, not derived by
    subtracting DEFAULT_REVIEW_INTERVAL_DAYS from next_review_date --
    that would misrepresent history the moment a lawyer ever overrides
    next_review_date to something further out (e.g. 6 months for a
    matter awaiting a court date), which is a real, expected case, not
    an edge case to ignore.

    next_review_date: the explicit value if given (parsed and respected
    verbatim, restarting the clock from today); otherwise
    today + DEFAULT_REVIEW_INTERVAL_DAYS.
    """
    today = date.today()
    if explicit_next_review_date:
        try:
            next_review = date.fromisoformat(explicit_next_review_date)
        except ValueError:
            raise HTTPException(status_code=400, detail="next_review_date must be in YYYY-MM-DD format")
    else:
        next_review = today + timedelta(days=DEFAULT_REVIEW_INTERVAL_DAYS)
    return next_review, today


async def _create_matter_row(
    conn, firm_id, name: str, *,
    number: Optional[str] = None,
    internal_ref: Optional[str] = None,
    external_ref: Optional[str] = None,
    client_name: Optional[str] = None,
    client_id: Optional[_uuid_mod.UUID] = None,
    case_parties: Optional[str] = None,
    matter_type: Optional[str] = None,
    practice_area: Optional[str] = None,
    status: str = "Active",
    custom_status: Optional[str] = None,
    next_deadline: Optional[date] = None,
    next_deadline_note: Optional[str] = None,
    assigned_lawyer_id: Optional[_uuid_mod.UUID] = None,
    coverage_tier: Optional[str] = None,
    service_type: Optional[str] = None,
    sla_deadline: Optional[datetime] = None,
    numbering_client_number: Optional[str] = None,
    created_by: Optional[_uuid_mod.UUID] = None,
    created_at: Optional[datetime] = None,
    last_activity: Optional[datetime] = None,
    next_review_date: Optional[date] = None,
    last_reviewed_date: Optional[date] = None,
) -> "asyncpg.Record":
    """
    Shared INSERT for a new matters row. Atomically allocates
    matter_number (via _next_matter_number()) iff `numbering_client_number`
    is passed -- the caller decides whether numbering applies at all, by
    passing the linked client's client_number or leaving it out. This
    reproduces every existing caller's numbering rule unchanged:
    create_matter already only numbers when the linked client has a
    client_number, which falls out naturally here since it's the
    caller's job to pass (or not pass) that value; auto_create_matter
    and bulk_import_matters never request numbering at all, by never
    passing this.

    `created_at`/`last_activity` both default to None, meaning "leave
    unset" -- last_activity has no DB default and NULL is a normal,
    sorted-around state elsewhere (list_matters' `ORDER BY last_activity
    DESC NULLS LAST`). Pass both explicitly to get "now"; deliberately
    NOT defaulted to now() here so auto_create_matter's existing
    NULL-last_activity behavior (unchanged, not a bug being fixed as
    part of this consolidation) falls out by simply not passing it,
    rather than needing a special case.

    `next_review_date` defaults here (unconditionally, unless a caller
    ever passes one explicitly -- none currently do, since MatterCreate
    has no such field) rather than at each of the 5 call sites, so every
    matter-creation path gets the review safety net uniformly for free.
    """
    mid = _uuid_mod.uuid4()
    matter_number = (
        await _next_matter_number(conn, firm_id, numbering_client_number)
        if numbering_client_number else None
    )
    if next_review_date is None:
        next_review_date = date.today() + timedelta(days=DEFAULT_REVIEW_INTERVAL_DAYS)
    if last_reviewed_date is None:
        last_reviewed_date = date.today()
    return await conn.fetchrow("""
        INSERT INTO matters (
            id, firm_id, name, number, internal_ref, external_ref,
            client_name, client_id, case_parties, matter_type, practice_area,
            status, custom_status, next_deadline, next_deadline_note,
            assigned_lawyer_id, coverage_tier, service_type, sla_deadline,
            matter_number, created_by, created_at, last_activity,
            next_review_date, last_reviewed_date
        )
        VALUES ($1,$2,$3,$4,$5,$6,$7,$8,$9,$10,$11,$12,$13,$14,$15,$16,$17,$18,$19,$20,$21,$22,$23,$24,$25)
        RETURNING *
    """,
    mid, firm_id, name, number, internal_ref, external_ref,
    client_name, client_id, case_parties, matter_type, practice_area,
    status, custom_status, next_deadline, next_deadline_note,
    assigned_lawyer_id, coverage_tier, service_type, sla_deadline,
    matter_number, created_by, created_at, last_activity,
    next_review_date, last_reviewed_date,
    )


_CONVEYANCING_DATE_EVENT_TITLES = {
    "conveyancing_transfer_date": "Conveyancing: Transfer Date",
    "conveyancing_rates_clearance_expiry": "Conveyancing: Rates Clearance Expiry",
    "conveyancing_bond_registration_deadline": "Conveyancing: Bond Registration Deadline",
}

async def _sync_conveyancing_calendar_events(conn, matter, touched_fields):
    """
    Keeps calendar_events in sync with a matter's conveyancing key dates —
    this is what "feeding into the existing deadline/calendar system"
    means concretely: a set date shows up in GET /api/calendar, the daily
    reminder digest, and the client detail view's upcoming-deadlines list
    for free, with no separate conveyancing-aware code needed there.

    Upserts one calendar_events row per touched field, keyed on
    (matter_id, source='conveyancing_sync', title) so a re-save updates
    rather than duplicates. `matter` is the raw asyncpg Record from the
    UPDATE ... RETURNING * in update_matter() (real UUID/date types, not
    yet through _row_to_matter()'s JSON-safe conversion).

    The delete-when-cleared branch below is defensive/idempotent rather
    than reachable today: update_matter()'s PATCH, like every other
    Optional field on MatterUpdate, filters out None values before this
    is ever called (matching this API's existing convention — there is no
    "explicitly clear a field" path for anything else here either), so in
    practice `matter[field]` is never None for a field that made it into
    touched_fields. Kept correct in case that ever changes.
    """
    for field in touched_fields:
        title = _CONVEYANCING_DATE_EVENT_TITLES[field]
        value = matter[field]
        existing = await conn.fetchrow(
            "SELECT id FROM calendar_events WHERE matter_id=$1 AND source='conveyancing_sync' AND title=$2",
            matter["id"], title,
        )
        if value is None:
            if existing:
                await conn.execute("DELETE FROM calendar_events WHERE id=$1", existing["id"])
            continue
        if existing:
            await conn.execute("UPDATE calendar_events SET date=$1 WHERE id=$2", value, existing["id"])
        else:
            await conn.execute(
                """INSERT INTO calendar_events
                       (id, firm_id, matter_id, title, date, event_type, matter_name, source)
                   VALUES ($1,$2,$3,$4,$5,$6,$7,$8)""",
                _uuid_mod.uuid4(), matter["firm_id"], matter["id"], title, value,
                "deadline", matter["name"], "conveyancing_sync",
            )

async def _sync_client_relationship_ended(conn, client_id) -> None:
    """
    s24 retention flag (Part 5): "end of business relationship" is not a
    concept that exists on the client record itself today (confirmed by
    inspection before building this module — clients has no status
    field). Derived instead: relationship_ended_date is set the moment
    every one of a client's matters reaches status='Closed', and cleared
    again the moment that stops being true (a matter reopening, a new
    matter being created). Called from update_matter() whenever a
    matter's status changes and it's linked to a client.

    Deliberately does not compute retained_until here — s24 sets a
    minimum retention period, and the exact statutory duration was not
    available to verify when this was built. A compliance officer enters
    that deadline manually (PATCH .../compliance); this only tracks the
    trigger date a future feature could measure that period from.
    """
    matter_rows = await conn.fetch(
        "SELECT status FROM matters WHERE client_id=$1 AND firm_id=$2",
        client_id, FIRM_ID
    )
    all_closed = bool(matter_rows) and all(m["status"] == "Closed" for m in matter_rows)
    ended_date = date.today() if all_closed else None

    existing = await conn.fetchrow(
        "SELECT id, relationship_ended_date FROM client_compliance WHERE client_id=$1 AND firm_id=$2",
        client_id, FIRM_ID
    )
    if existing:
        if (existing["relationship_ended_date"] is not None) == all_closed:
            return  # already in the right state — avoid an unnecessary write
        await conn.execute(
            "UPDATE client_compliance SET relationship_ended_date=$1, updated_at=NOW() WHERE id=$2",
            ended_date, existing["id"]
        )
    elif all_closed:
        await conn.execute(
            "INSERT INTO client_compliance (client_id, firm_id, relationship_ended_date) VALUES ($1,$2,$3)",
            client_id, FIRM_ID, ended_date
        )

def _row_to_matter(row) -> dict:
    d = dict(row)
    for k in ("id", "firm_id", "created_by", "client_id"):
        if d.get(k):
            d[k] = str(d[k])
    for k in ("created_at", "last_activity"):
        if d.get(k):
            d[k] = d[k].isoformat()
    if d.get("next_deadline"):
        d["next_deadline"] = d["next_deadline"].isoformat()
    if d.get("next_review_date"):
        d["next_review_date"] = d["next_review_date"].isoformat()
    if d.get("last_reviewed_date"):
        d["last_reviewed_date"] = d["last_reviewed_date"].isoformat()
    for k in ("conveyancing_transfer_date", "conveyancing_rates_clearance_expiry",
              "conveyancing_bond_registration_deadline"):
        if d.get(k):
            d[k] = d[k].isoformat()
    # asyncpg returns NUMERIC as Decimal — cast explicitly rather than
    # relying on the default JSON encoder's handling of it.
    for k in ("amount_billed", "amount_received", "conveyancing_purchase_price"):
        if d.get(k) is not None:
            d[k] = float(d[k])
    # Computed here (not stored) so there's exactly one place this
    # arithmetic lives, rather than duplicating "billed minus received" in
    # frontend JS too. None when neither figure has been entered yet.
    if d.get("amount_billed") is not None or d.get("amount_received") is not None:
        d["fee_balance"] = (d.get("amount_billed") or 0.0) - (d.get("amount_received") or 0.0)
    else:
        d["fee_balance"] = None
    # Matter Progress Tracker (visual stepper) — stage_info is None (not
    # an empty tracker) when this matter_type/practice_area has no
    # defined sequence, so the frontend falls back to the plain-text
    # status chip cleanly. days_in_stage computed from the raw
    # stage_updated_at before it gets stringified below — normalized to
    # UTC-aware first since it's written as naive (datetime.utcnow(),
    # matching last_activity's convention elsewhere) but may round-trip
    # back from Postgres as either naive or tz-aware depending on
    # asyncpg/driver behavior; assuming one or the other crashed this on
    # the naive path (can't subtract offset-naive and offset-aware).
    sequence = resolve_stage_sequence(d.get("matter_type"), d.get("practice_area"))
    if sequence:
        field = stage_storage_field(d.get("matter_type"), d.get("practice_area"))
        current_stage = d.get(field)
        days_in_stage = None
        if d.get("stage_updated_at"):
            stage_updated = d["stage_updated_at"]
            if stage_updated.tzinfo is None:
                stage_updated = stage_updated.replace(tzinfo=timezone.utc)
            days_in_stage = (datetime.now(timezone.utc) - stage_updated).days
        d["stage_info"] = {
            "sequence": sequence,
            "current_stage": current_stage,
            "current_index": sequence.index(current_stage) if current_stage in sequence else None,
            "days_in_stage": days_in_stage,
        }
    else:
        d["stage_info"] = None
    if d.get("stage_updated_at"):
        d["stage_updated_at"] = d["stage_updated_at"].isoformat()
    return d

def _row_to_client(row) -> dict:
    d = dict(row)
    for k in ("id", "firm_id", "created_by", "proof_of_incorporation_document_id", "governing_document_id"):
        if d.get(k):
            d[k] = str(d[k])
    for k in ("created_at", "updated_at"):
        if d.get(k):
            d[k] = d[k].isoformat()
    for k in ("date_of_birth", "id_expiry_date", "date_incorporated"):
        if d.get(k):
            d[k] = str(d[k])
    # asyncpg may return jsonb as a raw string depending on codec config —
    # same defensive handling as calendar_events.attendees.
    for k in ("trustees", "settlors", "beneficiaries"):
        if isinstance(d.get(k), str):
            try:
                d[k] = json.loads(d[k])
            except (ValueError, TypeError):
                d[k] = []
        elif d.get(k) is None:
            d[k] = []
    return d

def _row_to_beneficial_owner(row) -> dict:
    d = dict(row)
    for k in ("id", "client_id", "firm_id", "verified_by"):
        if d.get(k):
            d[k] = str(d[k])
    if d.get("created_at"):
        d["created_at"] = d["created_at"].isoformat()
    if d.get("date_of_birth"):
        d["date_of_birth"] = str(d["date_of_birth"])
    if d.get("verified_date"):
        d["verified_date"] = str(d["verified_date"])
    if d.get("ownership_percentage") is not None:
        d["ownership_percentage"] = float(d["ownership_percentage"])
    return d

def _row_to_authorized_representative(row) -> dict:
    d = dict(row)
    for k in ("id", "client_id", "firm_id", "authority_document_id"):
        if d.get(k):
            d[k] = str(d[k])
    if d.get("created_at"):
        d["created_at"] = d["created_at"].isoformat()
    if d.get("verified_date"):
        d["verified_date"] = str(d["verified_date"])
    return d

def _row_to_client_compliance(row) -> dict:
    d = dict(row)
    for k in ("id", "client_id", "firm_id", "senior_management_approved_by", "conflict_check_reviewed_by"):
        if d.get(k):
            d[k] = str(d[k])
    for k in ("created_at", "updated_at"):
        if d.get(k):
            d[k] = d[k].isoformat()
    for k in ("senior_management_approved_date", "relationship_ended_date", "retained_until", "conflict_check_reviewed_date"):
        if d.get(k):
            d[k] = str(d[k])
    return d

_DEFAULT_CLIENT_COMPLIANCE = {
    "identity_verification_status": "Unverified",
    "client_is_beneficial_owner": None,
    "is_pep": None,
    "pep_basis": None,
    "pep_position": None,
    "pep_country": None,
    "senior_management_approval_required": False,
    "senior_management_approved_by": None,
    "senior_management_approved_date": None,
    "source_of_wealth": None,
    "source_of_funds": None,
    "enhanced_monitoring_required": False,
    "risk_rating": "NotAssessed",
    "relationship_ended_date": None,
    "retained_until": None,
    "conflict_check_reviewed": False,
    "conflict_check_reviewed_by": None,
    "conflict_check_reviewed_date": None,
}

def _compute_compliance_status(client: dict, compliance: Optional[dict], beneficial_owners: list) -> dict:
    """
    "Cleared" requires every one of: identity verified, beneficial
    ownership resolved (verified, or not applicable), PEP screening done
    (and senior management approval if PEP), and the conflict check
    reviewed.

    Conflict check: reuses the real, existing GET /api/matters/check-conflict
    (fuzzy name-similarity search against every matter, already live on the
    New Matter form) — not a stub. That endpoint is a live, on-demand
    search with no stored outcome of its own, so conflict_check_reviewed
    on client_compliance is what "Cleared" actually gates on: a lawyer
    has run the check for this client and confirmed there's no conflict
    (see runClientConflictCheck()/markConflictReviewed() in index.html).
    """
    compliance = compliance or _DEFAULT_CLIENT_COMPLIANCE
    missing = []

    client_type = client.get("client_type")
    if not client_type:
        return {"compliance_status": "Action Required", "missing": ["Client type not recorded"]}

    if compliance.get("identity_verification_status") != "Verified":
        missing.append("Identity not verified")

    if client_type in LEGAL_PERSON_CLIENT_TYPES:
        is_bo = compliance.get("client_is_beneficial_owner")
        if is_bo == "No":
            if not any(o.get("verification_status") == "Verified" for o in beneficial_owners):
                missing.append("Beneficial ownership not verified")
        elif is_bo in (None, "Unknown"):
            missing.append("Beneficial ownership not assessed")
        # is_bo == "Yes" -> client itself is the beneficial owner, satisfied

    is_pep = compliance.get("is_pep")
    if is_pep is None:
        missing.append("PEP screening not completed")
    elif is_pep is True and not compliance.get("senior_management_approved_by"):
        missing.append("Senior management approval required (PEP)")

    if not compliance.get("conflict_check_reviewed"):
        missing.append("Conflict check not reviewed")

    return {
        "compliance_status": "Cleared" if not missing else "Action Required",
        "missing": missing,
    }

def _row_to_doc(row) -> dict:
    d = dict(row)
    for k in ("id", "matter_id", "firm_id", "uploaded_by"):
        if d.get(k):
            d[k] = str(d[k])
    for k in ("uploaded_at",):
        if d.get(k):
            d[k] = d[k].isoformat()
    if d.get("doc_date"):
        d["doc_date"] = str(d["doc_date"])
    return d

def _row_to_note(row) -> dict:
    d = dict(row)
    for k in ("id", "matter_id", "firm_id", "user_id"):
        if d.get(k):
            d[k] = str(d[k])
    if d.get("created_at"):
        d["created_at"] = d["created_at"].isoformat()
    return d

def _row_to_event(row) -> dict:
    d = dict(row)
    for k in ("id", "firm_id", "matter_id", "created_by"):
        if d.get(k):
            d[k] = str(d[k])
    if d.get("created_at"):
        d["created_at"] = d["created_at"].isoformat()
    if d.get("date"):
        d["date"] = str(d["date"])
    if d.get("time"):
        d["time"] = str(d["time"])[:5]  # HH:MM
    if d.get("attendees") is not None:
        # asyncpg may return jsonb as a raw string depending on codec config —
        # handle both an already-decoded list and a raw JSON string safely.
        if isinstance(d["attendees"], str):
            try:
                d["attendees"] = json.loads(d["attendees"])
            except (ValueError, TypeError):
                d["attendees"] = []
    else:
        d["attendees"] = []
    return d

# ── Health ────────────────────────────────────────────────────────────────────

@app.get("/api/health")
async def health():
    import shutil
    embeddings_ok = _embedding_model is not None
    if not embeddings_ok:
        try:
            import sentence_transformers
            embeddings_ok = True
        except Exception:
            pass

    db_ok = False
    try:
        if _db_pool:
            async with _db_pool.acquire() as conn:
                await conn.fetchval("SELECT 1")
            db_ok = True
    except Exception:
        pass

    deps = {
        "anthropic_key": bool(os.environ.get("ANTHROPIC_API_KEY")),
        "database": db_ok,
        "tesseract": shutil.which("tesseract") is not None,
        "pdftoppm": shutil.which("pdftoppm") is not None,
        "node": shutil.which("node") is not None,
        "smtp_configured": is_email_configured(),
        "semantic_search": embeddings_ok,
    }
    status = "ok" if (deps["anthropic_key"] and deps["database"]) else "degraded"
    return {
        "status": status,
        "version": app.version,
        "service": "Mutemo Desk",
        "dependencies": deps,
    }

@app.post("/api/admin/reindex")
async def reindex_semantic_search(request: Request):
    require_admin_token(request)
    async with _db_pool.acquire() as conn:
        chunk_rows = await conn.fetch(
            "SELECT * FROM chunks WHERE firm_id=$1", FIRM_ID
        )
    chunks = [dict(r) for r in chunk_rows]
    firm_chunks = [c for c in chunks if c["chunk_source"] == "firm"]
    legal_chunks = [c for c in chunks if c["chunk_source"] == "legal"]
    zlr_chunks_list = [c for c in chunks if c["chunk_source"] == "zlr"]
    if firm_chunks:
        await asyncio.to_thread(index_chunks_in_chroma, firm_chunks, "firm")
    if legal_chunks:
        await asyncio.to_thread(index_chunks_in_chroma, legal_chunks, "legal")
    if zlr_chunks_list:
        await asyncio.to_thread(index_chunks_in_chroma, zlr_chunks_list, "zlr")
    return {"reindexed": len(chunks), "firm": len(firm_chunks), "legal": len(legal_chunks), "zlr": len(zlr_chunks_list)}

@app.post("/api/admin/reset-chromadb")
async def admin_reset_chromadb(request: Request):
    """
    Drop and recreate all ChromaDB collections (firm/legal/zlr), then
    rebuild them from Postgres via reconcile_chroma_index(). For use when
    the vector index is suspected corrupt or out of sync in a way the
    normal startup reconcile check won't catch on its own.
    """
    require_admin_token(request)
    global _chroma_client, _firm_collection, _legal_collection, _zlr_collection

    get_chroma_collections()  # ensure _chroma_client is initialized

    collection_names = ["firm_precedents", "legal_updates", "zlr_index"]
    deleted = []
    for name in collection_names:
        _chroma_client.delete_collection(name)
        deleted.append(name)

    _firm_collection = _chroma_client.get_or_create_collection(
        "firm_precedents", metadata={"hnsw:space": "cosine"}
    )
    _legal_collection = _chroma_client.get_or_create_collection(
        "legal_updates", metadata={"hnsw:space": "cosine"}
    )
    _zlr_collection = _chroma_client.get_or_create_collection(
        "zlr_index", metadata={"hnsw:space": "cosine"}
    )

    await reconcile_chroma_index()

    return {"status": "success", "deleted_collections": deleted}

@app.post("/api/admin/reindex-from-db")
async def reindex_from_db(request: Request):
    """
    Rebuild ChromaDB vectors from raw_text stored in PostgreSQL.
    Use after migration — populates chunks table and ChromaDB from existing DB records.
    """
    require_admin_token(request)
    indexed_zlr = 0
    all_chunks = []

    async with _db_pool.acquire() as conn:
        zlr_rows = await conn.fetch(
            "SELECT id, raw_text, case_name, citation, taxonomy_category FROM zlr_entries WHERE firm_id=$1 AND raw_text IS NOT NULL",
            FIRM_ID
        )

    for row in zlr_rows:
        item_id = str(row["id"])
        new_chunks = chunk_text(row["raw_text"], 1, item_id, "zlr")
        for c in new_chunks:
            c["chunk_source"] = "zlr"
            c["zlr_item_id"] = item_id
            c["citation"] = row.get("citation")
            c["case_name"] = row.get("case_name")
            c["taxonomy_category"] = row.get("taxonomy_category")
        all_chunks.extend(new_chunks)
        indexed_zlr += 1

    if all_chunks:
        await asyncio.to_thread(index_chunks_in_chroma, all_chunks, "zlr")
        async with _db_pool.acquire() as conn:
            for c in all_chunks:
                await conn.execute("""
                    INSERT INTO chunks (id, firm_id, document_id, matter_id, chunk_source,
                                       text, chunk_index, page_number, zlr_item_id, citation,
                                       case_name, taxonomy_category, created_at)
                    VALUES ($1,$2,$3,'zlr','zlr',$4,$5,$6,$7,$8,$9,$10,NOW())
                    ON CONFLICT (id) DO NOTHING
                """,
                c["id"], FIRM_ID, _uuid_mod.UUID(c["document_id"]),
                c["text"], c["chunk_index"], c.get("page_number", 1),
                c.get("zlr_item_id"), c.get("citation"),
                c.get("case_name"), c.get("taxonomy_category")
                )

    return {
        "zlr_entries_processed": indexed_zlr,
        "chunks_created": len(all_chunks),
    }

@app.post("/api/admin/reclassify-zlr")
async def reclassify_zlr(request: Request):
    require_admin_token(request)
    async with _db_pool.acquire() as conn:
        rows = await conn.fetch("SELECT id, raw_text, filename FROM zlr_entries WHERE firm_id=$1", FIRM_ID)
    updated = 0
    for row in rows:
        if not row["raw_text"]:
            continue
        ai_meta = await asyncio.to_thread(classify_case_with_ai, row["raw_text"], row["filename"] or "")
        if ai_meta and ai_meta.get("taxonomy_category") and ai_meta["taxonomy_category"] != "General":
            async with _db_pool.acquire() as conn:
                await conn.execute(
                    "UPDATE zlr_entries SET taxonomy_category=$1, summary=$2 WHERE id=$3",
                    ai_meta["taxonomy_category"], ai_meta.get("summary"), row["id"]
                )
            updated += 1
    return {"reclassified": updated, "total": len(rows)}

@app.post("/api/admin/backfill-legal-taxonomy")
async def backfill_legal_taxonomy(request: Request):
    """
    One-time backfill of legal_source_type/authority_strength for rows
    that predate this classification (backend/legal_taxonomy.py) — pure,
    deterministic, no AI call, safe to re-run any time (idempotent: only
    rows still missing legal_source_type are touched).
    """
    require_admin_token(request)
    counts = {"documents": 0, "legal_updates": 0, "zlr_entries": 0}

    async with _db_pool.acquire() as conn:
        rows = await conn.fetch(
            "SELECT id, document_type FROM documents WHERE firm_id=$1 AND legal_source_type IS NULL",
            FIRM_ID
        )
    for row in rows:
        source_type = classify_firm_document(row["document_type"])
        strength = authority_strength_for(source_type)
        async with _db_pool.acquire() as conn:
            await conn.execute(
                "UPDATE documents SET legal_source_type=$1, authority_strength=$2 WHERE id=$3",
                source_type.value, strength.value, row["id"]
            )
        counts["documents"] += 1

    async with _db_pool.acquire() as conn:
        rows = await conn.fetch(
            "SELECT id, source_type, reference FROM legal_updates WHERE firm_id=$1 AND legal_source_type IS NULL",
            FIRM_ID
        )
    for row in rows:
        source_type = classify_legal_update(row["source_type"], row["reference"])
        strength = authority_strength_for(source_type)
        async with _db_pool.acquire() as conn:
            await conn.execute(
                "UPDATE legal_updates SET legal_source_type=$1, authority_strength=$2 WHERE id=$3",
                source_type.value, strength.value, row["id"]
            )
        counts["legal_updates"] += 1

    async with _db_pool.acquire() as conn:
        rows = await conn.fetch(
            "SELECT id, court FROM zlr_entries WHERE firm_id=$1 AND legal_source_type IS NULL",
            FIRM_ID
        )
    for row in rows:
        source_type = classify_zlr_entry(row["court"])
        strength = authority_strength_for(source_type)
        async with _db_pool.acquire() as conn:
            await conn.execute(
                "UPDATE zlr_entries SET legal_source_type=$1, authority_strength=$2 WHERE id=$3",
                source_type.value, strength.value, row["id"]
            )
        counts["zlr_entries"] += 1

    return {"backfilled": counts}

# TEMPORARY — one-time production backfill for the content_hash reconciliation
# fix. Runs scripts/backfill_chunk_content_hash.py's real build_plan()/
# apply_plan() in-process (against this app's own live _db_pool and
# already-initialized Chroma client, rather than opening a second
# PersistentClient against the same on-disk data from within the same
# process) — `railway run` executes on the operator's machine, not inside
# the container, so it can't reach the private DB host or the volume the
# standalone script needs. Remove this endpoint once the one-time backfill
# has been confirmed applied to production.
@app.post("/api/admin/backfill-chunk-hashes")
async def admin_backfill_chunk_hashes(request: Request):
    require_admin_token(request)
    from scripts.backfill_chunk_content_hash import build_plan, apply_plan

    firm_col, legal_col, zlr_col = get_chroma_collections()
    collections = {"firm": firm_col, "legal": legal_col, "zlr": zlr_col}

    async with _db_pool.acquire() as conn:
        plan = await build_plan(conn, lambda source: collections[source], FIRM_ID)
        summary = await apply_plan(plan, lambda source: collections[source])

    return summary

# TEMPORARY — read-only companion to the backfill endpoint above, for
# verifying its result against real per-chunk data rather than trusting its
# own summary count. Returns actual Postgres/Chroma content_hash values
# side by side for a sample of real chunk_ids per source. Remove alongside
# the backfill endpoint once no longer needed.
@app.get("/api/admin/verify-chunk-hashes")
async def admin_verify_chunk_hashes(request: Request, sample: int = 5):
    require_admin_token(request)
    firm_col, legal_col, zlr_col = get_chroma_collections()
    collections = {"firm": firm_col, "legal": legal_col, "zlr": zlr_col}

    results = {}
    async with _db_pool.acquire() as conn:
        for source, col in collections.items():
            rows = await conn.fetch(
                "SELECT id, content_hash FROM chunks WHERE firm_id=$1 AND chunk_source=$2 LIMIT $3",
                FIRM_ID, source, sample,
            )
            if not rows:
                results[source] = []
                continue
            ids = [r["id"] for r in rows]
            chroma_data = col.get(ids=ids, include=["metadatas"])
            chroma_meta = dict(zip(chroma_data["ids"], chroma_data["metadatas"]))
            entries = []
            for r in rows:
                cid = r["id"]
                pg_hash = r["content_hash"]
                meta = chroma_meta.get(cid)
                chroma_hash = (meta or {}).get("content_hash") if meta else None
                entries.append({
                    "chunk_id": cid,
                    "postgres_hash": pg_hash,
                    "chroma_hash": chroma_hash,
                    "present_in_chroma": cid in chroma_meta,
                    "match": pg_hash == chroma_hash,
                })
            results[source] = entries
    return results

# TEMPORARY — one-time backfill for Multi-tenancy hardening (Part 3): puts
# firm_id into every pre-existing firm_precedents chunk's Chroma metadata
# so _semantic_search_firm()'s new where={"firm_id": ...} filter doesn't
# make already-indexed documents invisible. Same in-process reuse pattern
# as /api/admin/backfill-chunk-hashes above and for the same reason:
# `railway run` executes on the operator's machine, not inside the
# container, so it can't reach the private DB host or the volume the
# standalone script needs. Remove this endpoint once the one-time backfill
# has been confirmed applied to production.
@app.post("/api/admin/backfill-chroma-firm-id")
async def admin_backfill_chroma_firm_id(request: Request):
    require_admin_token(request)
    from scripts.backfill_chroma_firm_id import build_plan, apply_plan

    firm_col, _, _ = get_chroma_collections()

    async with _db_pool.acquire() as conn:
        plan = await build_plan(conn, lambda: firm_col, FIRM_ID)
        summary = await apply_plan(plan, lambda: firm_col)

    return summary

# ── Matters ───────────────────────────────────────────────────────────────────

@app.get("/api/matters")
async def list_matters(request: Request):
    user = await get_current_user(request)
    _check_permission(user, "matter:read")

    user_id = _uuid_mod.UUID(str(user["id"])) if user.get("id") else None

    async with _db_pool.acquire() as conn:
        org_role = None
        if user_id:
            org_role_row = await conn.fetchrow(
                "SELECT role FROM organisation_roles WHERE firm_id=$1 AND user_id=$2",
                FIRM_ID, user_id
            )
            org_role = org_role_row["role"] if org_role_row else None

        if org_role == "panel_lawyer":
            # Panel lawyers only see matters assigned to them, not the whole firm's docket.
            rows = await conn.fetch(
                "SELECT * FROM matters WHERE firm_id=$1 AND assigned_lawyer_id=$2 AND NOT is_sentinel "
                "ORDER BY last_activity DESC NULLS LAST, created_at DESC",
                FIRM_ID, user_id
            )
        else:
            # Excludes the sentinel "General / Firm Precedents" matter
            # (Rapid Precedent Capture) -- it's a system bucket, not a real
            # client matter, and this list feeds the frontend's global
            # `matters` array that dashboard stats and the Matters tab both
            # read from directly.
            rows = await conn.fetch(
                "SELECT * FROM matters WHERE firm_id=$1 AND NOT is_sentinel "
                "ORDER BY last_activity DESC NULLS LAST, created_at DESC",
                FIRM_ID
            )
    matters = []
    for row in rows:
        m = _row_to_matter(row)
        # Attach progress notes
        async with _db_pool.acquire() as conn:
            note_rows = await conn.fetch(
                "SELECT * FROM progress_notes WHERE matter_id=$1 ORDER BY created_at ASC",
                row["id"]
            )
        m["progress_notes"] = [_row_to_note(n) for n in note_rows]
        matters.append(m)
    return matters

def _name_tokens(s: str) -> set:
    return {w for w in re.split(r'[\s,.]+', (s or "").lower()) if len(w) >= 3}

def _name_similarity(a: str, b: str) -> float:
    score = difflib.SequenceMatcher(None, (a or "").lower().strip(), (b or "").lower().strip()).ratio()
    # A shared significant word (surname, distinctive party/company name) is a
    # stronger, more reliable signal than raw character similarity for this
    # purpose — "D. Ngorima" vs "Deliwe Ngorima" share "ngorima" outright,
    # which matters more than their overall string similarity ratio.
    if _name_tokens(a) & _name_tokens(b):
        score = max(score, 0.8)
    return score

@app.get("/api/matters/check-conflict")
async def check_matter_conflict(name: str = "", client_name: str = "", request: Request = None):
    """
    Checks a proposed new matter's name/client against every existing matter
    — including closed ones, since a conflict doesn't expire just because a
    matter did — using fuzzy + shared-word name matching, so near-variations
    ("D. Ngorima" vs "Deliwe Ngorima") or a shared opposing party aren't
    missed by a plain exact-match search. This surfaces a warning; it does
    not block matter creation — the actual judgment call belongs to whoever
    is opening the matter, this just makes sure they don't miss something
    worth checking first.
    """
    user = await get_current_user(request)
    _check_permission(user, "matter:read")

    query_terms = [t.strip() for t in [name, client_name] if t and t.strip()]
    if not query_terms:
        return {"matches": []}

    async with _db_pool.acquire() as conn:
        rows = await conn.fetch(
            "SELECT id, name, client_name, number, status FROM matters WHERE firm_id=$1", FIRM_ID
        )

    matches = {}
    for row in rows:
        candidates = [c for c in [row["name"], row["client_name"]] if c]
        best_score = 0.0
        for qt in query_terms:
            for ec in candidates:
                best_score = max(best_score, _name_similarity(qt, ec))
        if best_score >= 0.55:
            mid = str(row["id"])
            if mid not in matches or matches[mid]["score"] < best_score:
                matches[mid] = {
                    "id": mid, "name": row["name"], "client_name": row["client_name"],
                    "number": row["number"], "status": row["status"],
                    "score": round(best_score, 2),
                }

    ranked = sorted(matches.values(), key=lambda m: m["score"], reverse=True)
    return {"matches": ranked[:8]}

@app.post("/api/matters")
async def create_matter(matter: MatterCreate, request: Request):
    user = await get_current_user(request)
    _check_permission(user, "matter:create")
    if matter.practice_area and matter.practice_area not in PRACTICE_AREAS:
        raise HTTPException(status_code=422, detail=f"practice_area must be one of: {', '.join(PRACTICE_AREAS)}")
    now = datetime.utcnow()
    parsed_deadline = None
    if matter.next_deadline:
        try:
            parsed_deadline = date.fromisoformat(matter.next_deadline)
        except ValueError:
            raise HTTPException(status_code=400, detail="next_deadline must be in YYYY-MM-DD format")

    client_id = None
    client_name = matter.client_name
    numbering_client_number = None
    async with _db_pool.acquire() as conn:
        if matter.client_id:
            try:
                client_id = _uuid_mod.UUID(matter.client_id)
            except ValueError:
                raise HTTPException(status_code=400, detail="client_id must be a valid UUID")
            client_row = await conn.fetchrow(
                "SELECT full_name, client_number FROM clients WHERE id=$1 AND firm_id=$2", client_id, FIRM_ID
            )
            if not client_row:
                raise HTTPException(status_code=404, detail="Client not found")
            # client_name is a display fallback/audit trail once client_id is
            # set — keep it in sync with the actual client record rather than
            # trusting a possibly-stale value passed alongside client_id.
            client_name = client_row["full_name"]
            # Only clients that already have a client_number (post-numbering,
            # or backfilled — see scripts/backfill_client_matter_numbers.py)
            # get their matters numbered; NULL otherwise. _create_matter_row()
            # only allocates a matter_number when numbering_client_number is
            # passed, so this falls out naturally.
            numbering_client_number = client_row["client_number"]

        row = await _create_matter_row(
            conn, FIRM_ID, matter.name,
            number=matter.number or matter.internal_ref,
            internal_ref=matter.internal_ref, external_ref=matter.external_ref,
            client_name=client_name, client_id=client_id, case_parties=matter.case_parties,
            matter_type=matter.matter_type, practice_area=matter.practice_area,
            status=matter.status or "Active", custom_status=matter.custom_status,
            next_deadline=parsed_deadline, next_deadline_note=matter.next_deadline_note,
            numbering_client_number=numbering_client_number,
            created_by=_uuid_mod.UUID(str(user["id"])) if user.get("id") else None,
            created_at=now, last_activity=now,
        )
    m = _row_to_matter(row)
    m["progress_notes"] = []
    return m

@app.post("/api/clients", status_code=201)
async def create_client(client: ClientCreate, request: Request):
    user = await get_current_user(request)
    _check_permission(user, "client:create")
    async with _db_pool.acquire() as conn:
        initials = await _resolve_user_initials(
            conn, FIRM_ID, user.get("id"), user.get("display_name") or "Client"
        )
        row = await _create_client_row(
            conn, FIRM_ID, initials, client.full_name,
            email=client.email, phone=client.phone,
            physical_address=client.physical_address,
            id_or_registration_number=client.id_or_registration_number,
            contact_person=client.contact_person, notes=client.notes,
            created_by=_uuid_mod.UUID(str(user["id"])) if user.get("id") else None,
        )
    return _row_to_client(row)

@app.get("/api/clients")
async def list_clients(request: Request):
    user = await get_current_user(request)
    _check_permission(user, "client:read")
    async with _db_pool.acquire() as conn:
        rows = await conn.fetch(
            "SELECT * FROM clients WHERE firm_id=$1 ORDER BY full_name ASC", FIRM_ID
        )
    return [_row_to_client(r) for r in rows]

@app.get("/api/clients/{client_id}")
async def get_client(client_id: str, request: Request):
    user = await get_current_user(request)
    _check_permission(user, "client:read")
    try:
        cid = _uuid_mod.UUID(client_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="client_id must be a valid UUID")
    async with _db_pool.acquire() as conn:
        row = await conn.fetchrow("SELECT * FROM clients WHERE id=$1 AND firm_id=$2", cid, FIRM_ID)
        if not row:
            raise HTTPException(status_code=404, detail="Client not found")
        matter_rows = await conn.fetch(
            "SELECT * FROM matters WHERE client_id=$1 AND firm_id=$2 "
            "ORDER BY last_activity DESC NULLS LAST, created_at DESC",
            cid, FIRM_ID
        )
        matter_ids = [m["id"] for m in matter_rows]

        # Batched (not N+1) across every matter this client has — progress
        # notes and documents feed the client detail view's "recent
        # activity" feed, calendar events feed its "upcoming deadlines"
        # list alongside each matter's own next_deadline (already embedded
        # via _row_to_matter below, no extra query needed for that part).
        notes_by_matter, docs_by_matter, calendar_events = {}, {}, []
        if matter_ids:
            note_rows = await conn.fetch(
                "SELECT * FROM progress_notes WHERE matter_id = ANY($1) ORDER BY created_at DESC",
                matter_ids
            )
            for n in note_rows:
                notes_by_matter.setdefault(str(n["matter_id"]), []).append(_row_to_note(n))

            doc_rows = await conn.fetch(
                "SELECT * FROM documents WHERE matter_id = ANY($1) AND status='complete' "
                "ORDER BY uploaded_at DESC",
                matter_ids
            )
            for doc in doc_rows:
                docs_by_matter.setdefault(str(doc["matter_id"]), []).append(_row_to_doc(doc))

            event_rows = await conn.fetch(
                "SELECT * FROM calendar_events WHERE matter_id = ANY($1) AND date >= CURRENT_DATE "
                "ORDER BY date ASC",
                matter_ids
            )
            calendar_events = [_row_to_event(e) for e in event_rows]

        compliance_row = await conn.fetchrow(
            "SELECT * FROM client_compliance WHERE client_id=$1 AND firm_id=$2", cid, FIRM_ID
        )
        owner_rows = await conn.fetch(
            "SELECT verification_status FROM beneficial_owners WHERE client_id=$1 AND firm_id=$2", cid, FIRM_ID
        )

    c = _row_to_client(row)
    c["matters"] = []
    for m in matter_rows:
        md = _row_to_matter(m)
        mid = str(m["id"])
        md["progress_notes"] = notes_by_matter.get(mid, [])
        md["documents"] = docs_by_matter.get(mid, [])
        c["matters"].append(md)
    c["calendar_events"] = calendar_events
    # Compliance badge — same shape/spirit as the document_status chips
    # already built this session, computed from whatever exists so far
    # rather than requiring a second round-trip to /compliance.
    compliance = _row_to_client_compliance(compliance_row) if compliance_row else dict(_DEFAULT_CLIENT_COMPLIANCE)
    status = _compute_compliance_status(dict(row), compliance, [dict(r) for r in owner_rows])
    c["compliance_status"] = status["compliance_status"]
    c["compliance_missing"] = status["missing"]
    return c

@app.patch("/api/clients/{client_id}")
async def update_client(client_id: str, update: ClientUpdate, request: Request):
    user = await get_current_user(request)
    _check_permission(user, "client:edit")
    try:
        cid = _uuid_mod.UUID(client_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="client_id must be a valid UUID")
    fields = {k: v for k, v in update.dict().items() if v is not None}
    if not fields:
        raise HTTPException(status_code=400, detail="No fields to update")

    if "client_type" in fields and fields["client_type"] not in CLIENT_TYPES:
        raise HTTPException(status_code=422, detail=f"client_type must be one of: {', '.join(CLIENT_TYPES)}")

    date_fields = ("date_of_birth", "id_expiry_date", "date_incorporated")
    for k in date_fields:
        if k in fields:
            try:
                fields[k] = date.fromisoformat(fields[k])
            except ValueError:
                raise HTTPException(status_code=400, detail=f"{k} must be in YYYY-MM-DD format")

    for k in ("proof_of_incorporation_document_id", "governing_document_id"):
        if k in fields:
            try:
                fields[k] = _uuid_mod.UUID(fields[k])
            except ValueError:
                raise HTTPException(status_code=400, detail=f"{k} must be a valid UUID")

    # s17(c) trust/estate party lists — JSONB columns need an explicit
    # ::jsonb cast + json.dumps, same convention as calendar_events.attendees
    # (see update_calendar_event) and firms.features.
    jsonb_fields = ("trustees", "settlors", "beneficiaries")
    for k in jsonb_fields:
        if k in fields:
            fields[k] = json.dumps([p.dict() if hasattr(p, "dict") else p for p in fields[k]])

    fields["updated_at"] = datetime.utcnow()
    set_clauses = ", ".join(
        f"{k}=${i+2}::jsonb" if k in jsonb_fields else f"{k}=${i+2}"
        for i, k in enumerate(fields.keys())
    )
    values = list(fields.values())
    async with _db_pool.acquire() as conn:
        row = await conn.fetchrow(
            f"UPDATE clients SET {set_clauses} WHERE id=$1 AND firm_id=${len(values)+2} RETURNING *",
            cid, *values, FIRM_ID
        )
    if not row:
        raise HTTPException(status_code=404, detail="Client not found")
    return _row_to_client(row)

async def _get_client_or_404(conn, cid) -> dict:
    row = await conn.fetchrow("SELECT * FROM clients WHERE id=$1 AND firm_id=$2", cid, FIRM_ID)
    if not row:
        raise HTTPException(status_code=404, detail="Client not found")
    return dict(row)

def _parse_client_id(client_id: str):
    try:
        return _uuid_mod.UUID(client_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="client_id must be a valid UUID")

# ── AML/KYC: beneficial owners (s15/s17(b)) ─────────────────────────────────

@app.get("/api/clients/{client_id}/beneficial-owners")
async def list_beneficial_owners(client_id: str, request: Request):
    user = await get_current_user(request)
    _check_permission(user, "client:read")
    cid = _parse_client_id(client_id)
    async with _db_pool.acquire() as conn:
        await _get_client_or_404(conn, cid)
        rows = await conn.fetch(
            "SELECT * FROM beneficial_owners WHERE client_id=$1 AND firm_id=$2 ORDER BY created_at ASC",
            cid, FIRM_ID
        )
    return [_row_to_beneficial_owner(r) for r in rows]

@app.post("/api/clients/{client_id}/beneficial-owners", status_code=201)
async def create_beneficial_owner(client_id: str, owner: BeneficialOwnerCreate, request: Request):
    user = await get_current_user(request)
    _check_permission(user, "client:edit")
    cid = _parse_client_id(client_id)
    async with _db_pool.acquire() as conn:
        await _get_client_or_404(conn, cid)
        row = await conn.fetchrow("""
            INSERT INTO beneficial_owners
                (client_id, firm_id, owner_name, date_of_birth, nationality,
                 id_or_passport_number, residential_address, ownership_or_control_basis,
                 ownership_percentage)
            VALUES ($1,$2,$3,$4,$5,$6,$7,$8,$9) RETURNING *
        """,
        cid, FIRM_ID, owner.owner_name,
        date.fromisoformat(owner.date_of_birth) if owner.date_of_birth else None,
        owner.nationality, owner.id_or_passport_number, owner.residential_address,
        owner.ownership_or_control_basis, owner.ownership_percentage,
        )
    return _row_to_beneficial_owner(row)

@app.patch("/api/clients/{client_id}/beneficial-owners/{owner_id}")
async def update_beneficial_owner(client_id: str, owner_id: str, update: BeneficialOwnerUpdate, request: Request):
    user = await get_current_user(request)
    _check_permission(user, "client:edit")
    cid = _parse_client_id(client_id)
    try:
        oid = _uuid_mod.UUID(owner_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="owner_id must be a valid UUID")
    fields = {k: v for k, v in update.dict().items() if v is not None}
    if not fields:
        raise HTTPException(status_code=400, detail="No fields to update")
    if "verification_status" in fields and fields["verification_status"] not in VERIFICATION_STATUSES:
        raise HTTPException(status_code=422, detail=f"verification_status must be one of: {', '.join(VERIFICATION_STATUSES)}")
    for k in ("date_of_birth", "verified_date"):
        if k in fields:
            try:
                fields[k] = date.fromisoformat(fields[k])
            except ValueError:
                raise HTTPException(status_code=400, detail=f"{k} must be in YYYY-MM-DD format")
    if fields.get("verification_status") == "Verified":
        fields["verified_by"] = _uuid_mod.UUID(str(user["id"])) if user.get("id") else None
        fields.setdefault("verified_date", datetime.utcnow().date())
    set_clauses = ", ".join(f"{k}=${i+3}" for i, k in enumerate(fields.keys()))
    values = list(fields.values())
    async with _db_pool.acquire() as conn:
        row = await conn.fetchrow(
            f"UPDATE beneficial_owners SET {set_clauses} WHERE id=$1 AND client_id=$2 AND firm_id=${len(values)+3} RETURNING *",
            oid, cid, *values, FIRM_ID
        )
    if not row:
        raise HTTPException(status_code=404, detail="Beneficial owner not found")
    return _row_to_beneficial_owner(row)

@app.delete("/api/clients/{client_id}/beneficial-owners/{owner_id}", status_code=204)
async def delete_beneficial_owner(client_id: str, owner_id: str, request: Request):
    user = await get_current_user(request)
    _check_permission(user, "client:edit")
    cid = _parse_client_id(client_id)
    try:
        oid = _uuid_mod.UUID(owner_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="owner_id must be a valid UUID")
    async with _db_pool.acquire() as conn:
        result = await conn.execute(
            "DELETE FROM beneficial_owners WHERE id=$1 AND client_id=$2 AND firm_id=$3",
            oid, cid, FIRM_ID
        )
    if result == "DELETE 0":
        raise HTTPException(status_code=404, detail="Beneficial owner not found")

# ── AML/KYC: authorized representatives (s17(d)) ────────────────────────────

@app.get("/api/clients/{client_id}/authorized-representatives")
async def list_authorized_representatives(client_id: str, request: Request):
    user = await get_current_user(request)
    _check_permission(user, "client:read")
    cid = _parse_client_id(client_id)
    async with _db_pool.acquire() as conn:
        await _get_client_or_404(conn, cid)
        rows = await conn.fetch(
            "SELECT * FROM authorized_representatives WHERE client_id=$1 AND firm_id=$2 ORDER BY created_at ASC",
            cid, FIRM_ID
        )
    return [_row_to_authorized_representative(r) for r in rows]

@app.post("/api/clients/{client_id}/authorized-representatives", status_code=201)
async def create_authorized_representative(client_id: str, rep: AuthorizedRepresentativeCreate, request: Request):
    user = await get_current_user(request)
    _check_permission(user, "client:edit")
    cid = _parse_client_id(client_id)
    if rep.authority_basis is not None and rep.authority_basis not in AUTHORITY_BASIS_TYPES:
        raise HTTPException(status_code=422, detail=f"authority_basis must be one of: {', '.join(AUTHORITY_BASIS_TYPES)}")
    async with _db_pool.acquire() as conn:
        await _get_client_or_404(conn, cid)
        row = await conn.fetchrow("""
            INSERT INTO authorized_representatives
                (client_id, firm_id, full_name, position_or_relationship,
                 id_or_passport_number, contact_details, authority_basis, authority_document_id)
            VALUES ($1,$2,$3,$4,$5,$6,$7,$8) RETURNING *
        """,
        cid, FIRM_ID, rep.full_name, rep.position_or_relationship,
        rep.id_or_passport_number, rep.contact_details, rep.authority_basis,
        _uuid_mod.UUID(rep.authority_document_id) if rep.authority_document_id else None,
        )
    return _row_to_authorized_representative(row)

@app.patch("/api/clients/{client_id}/authorized-representatives/{rep_id}")
async def update_authorized_representative(
    client_id: str, rep_id: str, update: AuthorizedRepresentativeUpdate, request: Request
):
    user = await get_current_user(request)
    _check_permission(user, "client:edit")
    cid = _parse_client_id(client_id)
    try:
        rid = _uuid_mod.UUID(rep_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="rep_id must be a valid UUID")
    fields = {k: v for k, v in update.dict().items() if v is not None}
    if not fields:
        raise HTTPException(status_code=400, detail="No fields to update")
    if "authority_basis" in fields and fields["authority_basis"] not in AUTHORITY_BASIS_TYPES:
        raise HTTPException(status_code=422, detail=f"authority_basis must be one of: {', '.join(AUTHORITY_BASIS_TYPES)}")
    if "verification_status" in fields and fields["verification_status"] not in VERIFICATION_STATUSES:
        raise HTTPException(status_code=422, detail=f"verification_status must be one of: {', '.join(VERIFICATION_STATUSES)}")
    if "authority_document_id" in fields:
        try:
            fields["authority_document_id"] = _uuid_mod.UUID(fields["authority_document_id"])
        except ValueError:
            raise HTTPException(status_code=400, detail="authority_document_id must be a valid UUID")
    if "verified_date" in fields:
        try:
            fields["verified_date"] = date.fromisoformat(fields["verified_date"])
        except ValueError:
            raise HTTPException(status_code=400, detail="verified_date must be in YYYY-MM-DD format")
    set_clauses = ", ".join(f"{k}=${i+3}" for i, k in enumerate(fields.keys()))
    values = list(fields.values())
    async with _db_pool.acquire() as conn:
        row = await conn.fetchrow(
            f"UPDATE authorized_representatives SET {set_clauses} "
            f"WHERE id=$1 AND client_id=$2 AND firm_id=${len(values)+3} RETURNING *",
            rid, cid, *values, FIRM_ID
        )
    if not row:
        raise HTTPException(status_code=404, detail="Authorized representative not found")
    return _row_to_authorized_representative(row)

@app.delete("/api/clients/{client_id}/authorized-representatives/{rep_id}", status_code=204)
async def delete_authorized_representative(client_id: str, rep_id: str, request: Request):
    user = await get_current_user(request)
    _check_permission(user, "client:edit")
    cid = _parse_client_id(client_id)
    try:
        rid = _uuid_mod.UUID(rep_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="rep_id must be a valid UUID")
    async with _db_pool.acquire() as conn:
        result = await conn.execute(
            "DELETE FROM authorized_representatives WHERE id=$1 AND client_id=$2 AND firm_id=$3",
            rid, cid, FIRM_ID
        )
    if result == "DELETE 0":
        raise HTTPException(status_code=404, detail="Authorized representative not found")

# ── AML/KYC: compliance / PEP / compliance_status (s20, s24) ───────────────

@app.get("/api/clients/{client_id}/compliance")
async def get_client_compliance(client_id: str, request: Request):
    user = await get_current_user(request)
    _check_permission(user, "client:read")
    cid = _parse_client_id(client_id)
    async with _db_pool.acquire() as conn:
        client_row = await _get_client_or_404(conn, cid)
        compliance_row = await conn.fetchrow(
            "SELECT * FROM client_compliance WHERE client_id=$1 AND firm_id=$2", cid, FIRM_ID
        )
        owner_rows = await conn.fetch(
            "SELECT verification_status FROM beneficial_owners WHERE client_id=$1 AND firm_id=$2", cid, FIRM_ID
        )
    compliance = _row_to_client_compliance(compliance_row) if compliance_row else dict(_DEFAULT_CLIENT_COMPLIANCE)
    status = _compute_compliance_status(client_row, compliance, [dict(r) for r in owner_rows])
    return {**compliance, **status}

@app.patch("/api/clients/{client_id}/compliance")
async def update_client_compliance(client_id: str, update: ClientComplianceUpdate, request: Request):
    """
    Upsert — client_compliance has no row for a client until the first
    write here, rather than every client getting one at creation time
    (most fields start meaningless — Unverified/NotAssessed/null — so
    there's nothing to gain from a row existing before anyone has actually
    looked at compliance for this client).
    """
    user = await get_current_user(request)
    _check_permission(user, "client:edit")
    cid = _parse_client_id(client_id)
    fields = {k: v for k, v in update.dict().items() if v is not None}
    if not fields:
        raise HTTPException(status_code=400, detail="No fields to update")

    if "pep_basis" in fields and fields["pep_basis"] not in PEP_BASIS_TYPES:
        raise HTTPException(status_code=422, detail=f"pep_basis must be one of: {', '.join(PEP_BASIS_TYPES)}")
    if "risk_rating" in fields and fields["risk_rating"] not in RISK_RATINGS:
        raise HTTPException(status_code=422, detail=f"risk_rating must be one of: {', '.join(RISK_RATINGS)}")
    if "identity_verification_status" in fields and fields["identity_verification_status"] not in VERIFICATION_STATUSES:
        raise HTTPException(status_code=422, detail=f"identity_verification_status must be one of: {', '.join(VERIFICATION_STATUSES)}")
    if "client_is_beneficial_owner" in fields and fields["client_is_beneficial_owner"] not in CLIENT_IS_BENEFICIAL_OWNER_VALUES:
        raise HTTPException(status_code=422, detail=f"client_is_beneficial_owner must be one of: {', '.join(CLIENT_IS_BENEFICIAL_OWNER_VALUES)}")
    for k in ("senior_management_approved_date", "retained_until"):
        if k in fields:
            try:
                fields[k] = date.fromisoformat(fields[k])
            except ValueError:
                raise HTTPException(status_code=400, detail=f"{k} must be in YYYY-MM-DD format")
    if "senior_management_approved_by" in fields:
        try:
            fields["senior_management_approved_by"] = _uuid_mod.UUID(fields["senior_management_approved_by"])
        except ValueError:
            raise HTTPException(status_code=400, detail="senior_management_approved_by must be a valid UUID")

    # s20: senior management approval is mandatory once a client is flagged
    # PEP — forced true here (not merely defaulted), so it can't be
    # silently left false by omission when is_pep flips to true.
    if fields.get("is_pep") is True:
        fields["senior_management_approval_required"] = True

    # conflict_check_reviewed_by/date are set server-side from whoever is
    # actually marking it reviewed, not client-supplied — same pattern as
    # beneficial_owners.verified_by/verified_date.
    if fields.get("conflict_check_reviewed") is True:
        fields["conflict_check_reviewed_by"] = _uuid_mod.UUID(str(user["id"])) if user.get("id") else None
        fields["conflict_check_reviewed_date"] = datetime.utcnow().date()
    elif fields.get("conflict_check_reviewed") is False:
        fields["conflict_check_reviewed_by"] = None
        fields["conflict_check_reviewed_date"] = None

    fields["updated_at"] = datetime.utcnow()

    async with _db_pool.acquire() as conn:
        await _get_client_or_404(conn, cid)
        existing = await conn.fetchrow(
            "SELECT id FROM client_compliance WHERE client_id=$1 AND firm_id=$2", cid, FIRM_ID
        )
        try:
            if existing:
                set_clauses = ", ".join(f"{k}=${i+3}" for i, k in enumerate(fields.keys()))
                values = list(fields.values())
                row = await conn.fetchrow(
                    f"UPDATE client_compliance SET {set_clauses} WHERE client_id=$1 AND firm_id=$2 "
                    f"RETURNING *",
                    cid, FIRM_ID, *values
                )
            else:
                cols = ["client_id", "firm_id"] + list(fields.keys())
                placeholders = ", ".join(f"${i+1}" for i in range(len(cols)))
                row = await conn.fetchrow(
                    f"INSERT INTO client_compliance ({', '.join(cols)}) VALUES ({placeholders}) RETURNING *",
                    cid, FIRM_ID, *fields.values()
                )
        except asyncpg.exceptions.ForeignKeyViolationError as e:
            # senior_management_approved_by / conflict_check_reviewed_by both
            # reference users(id) — a nonexistent id previously surfaced as a
            # raw 500 (the FK constraint was correctly rejecting it, just
            # not caught into a clean validation error). Constraint names
            # are Postgres's default "{table}_{column}_fkey" — used to name
            # which field actually failed rather than a generic message.
            constraint = (e.constraint_name or "")
            if "senior_management_approved_by" in constraint:
                field = "senior_management_approved_by"
            elif "conflict_check_reviewed_by" in constraint:
                field = "conflict_check_reviewed_by"
            else:
                field = "a referenced field"
            raise HTTPException(status_code=422, detail=f"{field} must be a valid user ID")
        client_row = await _get_client_or_404(conn, cid)
        owner_rows = await conn.fetch(
            "SELECT verification_status FROM beneficial_owners WHERE client_id=$1 AND firm_id=$2", cid, FIRM_ID
        )
    compliance = _row_to_client_compliance(row)
    status = _compute_compliance_status(client_row, compliance, [dict(r) for r in owner_rows])
    return {**compliance, **status}

@app.get("/api/matters/template")
async def download_matter_template():
    tpl = os.path.join(frontend_path, "MutemoDesk_Matter_Import_Template.docx")
    if os.path.exists(tpl):
        return FileResponse(tpl, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            filename="MutemoDesk_Matter_Import_Template.docx")
    raise HTTPException(status_code=404, detail="Template not found")

@app.get("/api/matters/template-excel")
async def download_matter_template_excel():
    tpl = os.path.join(frontend_path, "MutemoDesk_Matter_Import_Template.xlsx")
    if os.path.exists(tpl):
        return FileResponse(tpl, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            filename="MutemoDesk_Matter_Import_Template.xlsx")
    raise HTTPException(status_code=404, detail="Template not found")

@app.patch("/api/matters/{matter_id}")
async def update_matter(matter_id: str, update: MatterUpdate, request: Request):
    user = await get_current_user(request)
    _check_permission(user, "matter:edit")
    fields = {k: v for k, v in update.dict().items() if v is not None}
    if not fields:
        raise HTTPException(status_code=400, detail="No fields to update")
    if "practice_area" in fields and fields["practice_area"] not in PRACTICE_AREAS:
        raise HTTPException(status_code=422, detail=f"practice_area must be one of: {', '.join(PRACTICE_AREAS)}")
    if "next_deadline" in fields:
        # asyncpg is strict about type matching for DATE columns — a raw
        # Python str isn't reliably accepted the way it might be with a
        # text-based driver. Parse explicitly rather than risk an
        # intermittent DataError depending on asyncpg/codec version.
        try:
            fields["next_deadline"] = date.fromisoformat(fields["next_deadline"])
        except ValueError:
            raise HTTPException(status_code=400, detail="next_deadline must be in YYYY-MM-DD format")
    if "conveyancing_milestone" in fields and fields["conveyancing_milestone"] not in CONVEYANCING_MILESTONES:
        raise HTTPException(
            status_code=422,
            detail=f"conveyancing_milestone must be one of: {', '.join(CONVEYANCING_MILESTONES)}"
        )
    if "stage" in fields:
        # Which sequence applies depends on the matter's own matter_type/
        # practice_area — read from this same PATCH if being set together,
        # otherwise from the matter's existing stored values.
        mt, pa = fields.get("matter_type"), fields.get("practice_area")
        if mt is None or pa is None:
            async with _db_pool.acquire() as conn:
                existing = await conn.fetchrow(
                    "SELECT matter_type, practice_area FROM matters WHERE id=$1 AND firm_id=$2",
                    _uuid_mod.UUID(matter_id), FIRM_ID
                )
            if not existing:
                raise HTTPException(status_code=404, detail="Matter not found")
            mt = mt if mt is not None else existing["matter_type"]
            pa = pa if pa is not None else existing["practice_area"]
        sequence = resolve_stage_sequence(mt, pa)
        if not sequence:
            raise HTTPException(status_code=422, detail="This matter type has no defined stage sequence")
        if fields["stage"] not in sequence:
            raise HTTPException(status_code=422, detail=f"stage must be one of: {', '.join(sequence)}")
    if "stage" in fields or "conveyancing_milestone" in fields:
        fields["stage_updated_at"] = datetime.utcnow()
    conveyancing_date_fields = (
        "conveyancing_transfer_date", "conveyancing_rates_clearance_expiry",
        "conveyancing_bond_registration_deadline",
    )
    for k in conveyancing_date_fields:
        if k in fields:
            try:
                fields[k] = date.fromisoformat(fields[k])
            except ValueError:
                raise HTTPException(status_code=400, detail=f"{k} must be in YYYY-MM-DD format")
    fields["last_activity"] = datetime.utcnow()
    # Matter review safety net: every update stamps last_reviewed_date to
    # today (this touch IS a review happening now) and either respects an
    # explicit next_review_date given in this same PATCH, or re-defaults
    # it to today + DEFAULT_REVIEW_INTERVAL_DAYS. Runs unconditionally on
    # EVERY update, matching last_activity's own unconditional set just
    # above — the whole point is that a matter genuinely being worked on
    # (touched by any update, not just a dedicated "review" action) keeps
    # pushing its own review date forward and never surfaces in the
    # digest, while one that goes untouched keeps its stale date and
    # eventually does.
    fields["next_review_date"], fields["last_reviewed_date"] = _resolve_review_dates(fields.get("next_review_date"))

    async with _db_pool.acquire() as conn:
        if "client_id" in fields:
            # Same UUID-typing caveat as next_deadline above, plus a firm-
            # ownership check so a matter can't be pointed at another firm's
            # client record. Resolved before the SET clause is built so the
            # values list only ever needs one construction pass.
            try:
                client_uuid = _uuid_mod.UUID(fields["client_id"])
            except ValueError:
                raise HTTPException(status_code=400, detail="client_id must be a valid UUID")
            client_row = await conn.fetchrow(
                "SELECT full_name FROM clients WHERE id=$1 AND firm_id=$2", client_uuid, FIRM_ID
            )
            if not client_row:
                raise HTTPException(status_code=404, detail="Client not found")
            fields["client_id"] = client_uuid
            if "client_name" not in fields:
                fields["client_name"] = client_row["full_name"]

        set_clauses = ", ".join(f"{k}=${i+2}" for i, k in enumerate(fields.keys()))
        values = list(fields.values())
        row = await conn.fetchrow(
            f"UPDATE matters SET {set_clauses} WHERE id=$1 AND firm_id=${len(values)+2} RETURNING *",
            _uuid_mod.UUID(matter_id), *values, FIRM_ID
        )
    if not row:
        raise HTTPException(status_code=404, detail="Matter not found")
    touched_date_fields = [k for k in conveyancing_date_fields if k in fields]
    if touched_date_fields:
        async with _db_pool.acquire() as conn:
            await _sync_conveyancing_calendar_events(conn, row, touched_date_fields)
    if "status" in fields and row.get("client_id"):
        async with _db_pool.acquire() as conn:
            await _sync_client_relationship_ended(conn, row["client_id"])
    m = _row_to_matter(row)
    async with _db_pool.acquire() as conn:
        note_rows = await conn.fetch(
            "SELECT * FROM progress_notes WHERE matter_id=$1 ORDER BY created_at ASC",
            _uuid_mod.UUID(matter_id)
        )
    m["progress_notes"] = [_row_to_note(n) for n in note_rows]
    return m

@app.post("/api/matters/{matter_id}/notes")
async def add_progress_note(matter_id: str, note: ProgressNote, request: Request):
    user = await get_current_user(request)
    _check_permission(user, "note:create")
    async with _db_pool.acquire() as conn:
        matter = await conn.fetchrow(
            "SELECT id, name, internal_ref FROM matters WHERE id=$1 AND firm_id=$2",
            _uuid_mod.UUID(matter_id), FIRM_ID
        )
    if not matter:
        raise HTTPException(status_code=404, detail="Matter not found")

    author = note.author or (user.get("display_name") if user else None) or "Unknown"
    now = datetime.utcnow()
    nid = _uuid_mod.uuid4()

    async with _db_pool.acquire() as conn:
        note_row = await conn.fetchrow("""
            INSERT INTO progress_notes (id, matter_id, firm_id, text, author, user_id, created_at)
            VALUES ($1,$2,$3,$4,$5,$6,$7) RETURNING *
        """,
        nid, _uuid_mod.UUID(matter_id), FIRM_ID, note.text, author,
        _uuid_mod.UUID(str(user["id"])) if user.get("id") else None, now
        )
        # Matter review safety net: adding a note is a real "reviewed
        # this matter" action, same as a PATCH — see _resolve_review_dates().
        next_review, last_reviewed = _resolve_review_dates(note.next_review_date)
        await conn.execute(
            "UPDATE matters SET last_activity=$1, next_review_date=$2, last_reviewed_date=$3 WHERE id=$4",
            now, next_review, last_reviewed, _uuid_mod.UUID(matter_id)
        )

    entry = _row_to_note(note_row)

    # Quietly scan the note for actionable dates
    detected_dates = []
    try:
        today = datetime.utcnow().date().isoformat()
        matter_name = matter["name"]
        internal_ref = matter["internal_ref"] or ""

        def scan_note_sync():
            msg = client.messages.create(
                model="claude-sonnet-4-5",
                max_tokens=400,
                messages=[{"role": "user", "content": f"""Scan this legal progress note for any specific dates, deadlines, or appointments mentioned.
Today is {today}.

Return ONLY valid JSON — no other text:
{{
  "dates": [
    {{
      "title": "brief description of the action",
      "date": "YYYY-MM-DD",
      "time": "HH:MM or null",
      "event_type": "deadline|hearing|meeting|filing|other"
    }}
  ]
}}

If no specific dates are mentioned, return {{"dates": []}}.
Only include dates with a specific day — ignore vague references like "next week" or "soon".

Note text: {note.text}

JSON:"""}]
            )
            raw = msg.content[0].text.strip()
            raw = re.sub(r'^```json\s*|\s*```$', '', raw, flags=re.MULTILINE).strip()
            parsed = json.loads(raw)
            return parsed.get("dates", [])

        detected_dates = await asyncio.to_thread(scan_note_sync)
        for d in detected_dates:
            d["matter_id"] = matter_id
            d["matter_name"] = matter_name
            d["internal_ref"] = internal_ref
            d["source"] = "progress_note"
    except Exception as e:
        print(f"[notes] date scan failed: {e}")
        detected_dates = []

    # next_review_date/last_reviewed_date included so the frontend can
    # merge them into its local matter cache without a full re-fetch —
    # renderMatterPanel() reads from that cache, not a fresh GET, same
    # reasoning as updateMatterField()'s "merge full server response"
    # discipline elsewhere in the frontend.
    return {
        **entry, "detected_dates": detected_dates,
        "next_review_date": next_review.isoformat(),
        "last_reviewed_date": last_reviewed.isoformat(),
    }

@app.delete("/api/matters/{matter_id}/notes/{note_id}")
async def delete_progress_note(matter_id: str, note_id: str, request: Request):
    user = await get_current_user(request)
    _check_permission(user, "note:delete")
    async with _db_pool.acquire() as conn:
        result = await conn.execute(
            "DELETE FROM progress_notes WHERE id=$1 AND matter_id=$2 AND firm_id=$3",
            _uuid_mod.UUID(note_id), _uuid_mod.UUID(matter_id), FIRM_ID
        )
    if result == "DELETE 0":
        raise HTTPException(status_code=404, detail="Note not found")
    return {"deleted": True}

@app.delete("/api/matters/{matter_id}")
async def delete_matter(matter_id: str, request: Request):
    user = await get_current_user(request)
    _check_permission(user, "matter:delete")
    async with _db_pool.acquire() as conn:
        # Get chunk IDs for ChromaDB cleanup
        chunk_rows = await conn.fetch(
            "SELECT id FROM chunks WHERE matter_id=$1 AND firm_id=$2",
            matter_id, FIRM_ID
        )
        chunk_ids = [r["id"] for r in chunk_rows]
        result = await conn.execute(
            "DELETE FROM matters WHERE id=$1 AND firm_id=$2",
            _uuid_mod.UUID(matter_id), FIRM_ID
        )
    if result == "DELETE 0":
        raise HTTPException(status_code=404, detail="Matter not found")
    if chunk_ids:
        await asyncio.to_thread(remove_chunks_from_chroma, chunk_ids, "firm")
    return {"deleted": True}

# ── Bulk Matter Import ─────────────────────────────────────────────────────────

@app.post("/api/matters/bulk-import")
async def bulk_import_matters(file: UploadFile = File(...), request: Request = None):
    if request:
        user = await get_current_user(request)
        _check_permission(user, "matter:create")
    content = await file.read()
    ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
    if ext not in ("docx", "doc", "xlsx", "xlsm"):
        raise HTTPException(status_code=422, detail="Only .docx, .doc, .xlsx or .xlsm files supported")

    VALID_STATUSES = {"Active", "Awaiting Client", "Awaiting Court", "On Hold", "Closed"}
    LAW_TYPE_MAP = {
        "matrimonial": "matrimonial", "divorce": "matrimonial",
        "estate": "estate", "inheritance": "estate",
        "trust": "trust",
        "conveyancing": "conveyancing", "transfer": "conveyancing",
        "eviction": "eviction",
        "labour": "employment", "employment": "employment",
        "criminal": "criminal",
        "debt": "debt_collection", "debt collection": "debt_collection",
        "mining": "mining",
        "company": "company_law", "commercial": "commercial_contract",
        "property": "commercial_property", "land": "commercial_property",
        "family": "family_law", "custody": "family_law", "guardianship": "family_law",
        "lease": "eviction", "constitutional": "constitutional",
    }

    def detect_matter_type(law_text: str) -> str:
        if not law_text:
            return "other"
        law_lower = law_text.lower()
        for key, val in LAW_TYPE_MAP.items():
            if key in law_lower:
                return val
        return "other"

    def detect_status(next_action: str, action_done: str) -> str:
        combined = (f"{next_action} {action_done}").lower()
        if any(w in combined for w in ["n/a", "file closed", "closed file", "client passed", "passed away", "deceased"]):
            return "Closed"
        if any(w in combined for w in ["awaiting client", "awaiting further instructions", "awaiting instructions"]):
            return "Awaiting Client"
        if any(w in combined for w in ["awaiting set down", "awaiting court", "awaiting hearing", "awaiting order", "awaiting judgment"]):
            return "Awaiting Court"
        if any(w in combined for w in ["on hold", "sleeping dogs", "in abeyance"]):
            return "On Hold"
        return "Active"

    def build_matter_dict(internal_ref, client_name, subject, law_text, external_ref,
                          action_done, next_action, raw_status, latest_comm, opposing=""):
        if not client_name and not internal_ref:
            return None, "No client name or internal ref"
        if client_name and subject:
            matter_name = f"{client_name} — {subject}"
        elif client_name:
            matter_name = client_name
        elif subject:
            matter_name = subject
        else:
            matter_name = internal_ref
        status = raw_status if raw_status in VALID_STATUSES else detect_status(next_action or "", action_done or "")
        matter_type = detect_matter_type(law_text or "")
        now = datetime.utcnow()
        mid = _uuid_mod.uuid4()
        notes = []
        if action_done and str(action_done).lower() not in ("", "n/a", "-"):
            notes.append({"text": f"Action done: {action_done}", "author": "Import"})
        if next_action and str(next_action).lower() not in ("", "n/a", "-"):
            notes.append({"text": f"Next action: {next_action}", "author": "Import"})
        if latest_comm and str(latest_comm).strip():
            notes.append({"text": f"Latest communication: {latest_comm}", "author": "Import"})
        return {
            "id": mid, "name": matter_name, "number": internal_ref,
            "internal_ref": internal_ref or "", "external_ref": external_ref or "",
            "client_name": client_name or "", "matter_type": matter_type,
            "status": status, "custom_status": "",
            "created_at": now, "last_activity": now,
            "document_count": 0, "notes": notes,
            # No per-row client_id here — the import template has no client
            # picker (it's a bulk file upload, one row per matter, no way to
            # search/select an existing client mid-import). client_name is
            # still recorded as free text, same as every other matter created
            # before Client existed; scripts/migrate_clients.py's grouping
            # logic is the intended follow-up step to link these to Client
            # records afterward, same as any pre-existing matter.
            "case_parties": opposing or None,
        }, None

    created = []
    skipped = []
    matters_to_insert = []

    if ext in ("xlsx", "xlsm"):
        import openpyxl, io as _io
        wb = openpyxl.load_workbook(_io.BytesIO(content), data_only=True, read_only=True)
        ws = wb.active
        header_row = None
        header_map = {}
        COL_ALIASES = {
            "internal ref": "internal_ref", "file name": "internal_ref",
            "client name": "client_name", "client": "client_name",
            "matter description": "subject", "matter": "subject", "re": "subject",
            "opposing party": "opposing", "opposing party / re": "subject",
            "area of law": "law_type", "law": "law_type",
            "external ref": "external_ref", "case number": "external_ref",
            "status": "status", "action done": "action_done",
            "next action": "next_action",
            "latest communication": "latest_comm", "latest": "latest_comm",
        }
        for r_idx, row in enumerate(ws.iter_rows(values_only=True), start=1):
            row_vals = [str(c).lower().strip().rstrip("*").strip() if c else "" for c in row]
            if any(v in COL_ALIASES for v in row_vals):
                header_row = r_idx
                for c_idx, val in enumerate(row_vals):
                    canonical = COL_ALIASES.get(val)
                    if canonical:
                        header_map[c_idx] = canonical
                break
        if not header_row:
            raise HTTPException(status_code=422, detail="Could not find a header row in the Excel file.")
        for row in ws.iter_rows(min_row=header_row + 2, values_only=True):
            if not any(row):
                continue
            def g(field):
                for c_idx, f in header_map.items():
                    if f == field and c_idx < len(row):
                        v = row[c_idx]
                        return str(v).strip() if v is not None else ""
                return ""
            if g("internal_ref").upper().startswith("EXAMPLE"):
                continue
            matter, err = build_matter_dict(
                g("internal_ref"), g("client_name"), g("subject") or g("opposing"),
                g("law_type"), g("external_ref"), g("action_done"),
                g("next_action"), g("status"), g("latest_comm"), g("opposing")
            )
            if matter:
                matters_to_insert.append(matter)
            else:
                skipped.append({"reason": err, "row": str(row)[:100]})
        wb.close()
    else:
        import docx as docx_lib, io as _io
        try:
            doc = docx_lib.Document(_io.BytesIO(content))
        except Exception as e:
            raise HTTPException(status_code=422, detail=f"Could not read document: {e}")
        FIELD_MAP = {
            "file name": "internal_ref", "file name / internal ref": "internal_ref",
            "internal ref": "internal_ref",
            "name of client": "client_name", "client": "client_name",
            "re": "subject", "re (opposing party / subject)": "subject",
            "area of law": "law_type", "law": "law_type",
            "external reference": "external_ref", "case number": "external_ref",
            "action done": "action_done", "next action": "next_action",
            "status": "status",
            "latest communication": "latest_communication", "latest": "latest_communication",
        }
        for table in doc.tables:
            fields = {}
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if len(cells) >= 2:
                    label = cells[0].lower().strip().rstrip(":")
                    value = "\n".join(cells[1:]).strip()
                    canonical = FIELD_MAP.get(label)
                    if canonical:
                        fields[canonical] = value
            matter, err = build_matter_dict(
                fields.get("internal_ref", ""), fields.get("client_name", ""),
                fields.get("subject", ""), fields.get("law_type", ""),
                fields.get("external_ref", ""), fields.get("action_done", ""),
                fields.get("next_action", ""), fields.get("status", ""),
                fields.get("latest_communication", ""),
            )
            if matter:
                matters_to_insert.append(matter)
            else:
                skipped.append({"reason": err, "fields": {k: v[:50] for k, v in fields.items()}})

    # Bulk insert into PostgreSQL
    async with _db_pool.acquire() as conn:
        for m in matters_to_insert:
            row = await _create_matter_row(
                conn, FIRM_ID, m["name"],
                number=m["number"], internal_ref=m["internal_ref"], external_ref=m["external_ref"],
                client_name=m["client_name"], case_parties=m["case_parties"],
                matter_type=m["matter_type"], status=m["status"], custom_status=m["custom_status"],
                created_at=m["created_at"], last_activity=m["last_activity"],
            )
            for note in m.get("notes", []):
                await conn.execute("""
                    INSERT INTO progress_notes (matter_id, firm_id, text, author, created_at)
                    VALUES ($1,$2,$3,$4,$5)
                """, row["id"], FIRM_ID, note["text"], note["author"], m["created_at"])
            created.append({
                "id": str(row["id"]), "name": m["name"],
                "internal_ref": m["internal_ref"], "client_name": m["client_name"],
                "status": m["status"], "matter_type": m["matter_type"]
            })

    return {"created": len(created), "skipped": len(skipped), "matters": created, "skipped_details": skipped}

# ── Lawyer/Client/Matter Onboarding (bulk Excel upload) ─────────────────────
# Fixed layout, matching Sawyer_Mkushi_Client_Database_Form.xlsx exactly —
# unlike bulk_import_matters above, this is deliberately NOT a flexible
# header-alias parser. The form's cell positions are the source of truth:
#   B3/B4/B5/B6  — lawyer Name / Phone / Email / Role ("Partner"/"Associate")
#   Row 11       — column headers (not parsed — position-based, see below)
#   Row 12+      — one row per matter. Column A (+B/C/D/F/G) filled only
#                  on a client's first row, blank on subsequent matter
#                  rows for that same client. Column D is Contact Person
#                  — companies/entities only, blank for individuals.
#                  Column E is one matter's free-text "Reference/Case
#                  No. — description", stored as-is. Columns F (Client
#                  Type) and G (Is the client itself the beneficial
#                  owner?) are OPTIONAL compliance fields added
#                  2026-08-26 — see _normalize_client_type_cell()/
#                  _normalize_beneficial_owner_cell() below; a blank
#                  cell in either is normal, not an error, and leaves
#                  the client's compliance record in its default
#                  "not yet assessed" state.

def _cell_str(ws, coord: str) -> str:
    v = ws[coord].value
    return str(v).strip() if v is not None else ""


def _normalize_client_type_cell(raw: str) -> Optional[str]:
    """
    Optional Client Type column (2026-08-26 compliance-gap fix) — blank is
    the normal, expected case, not an error. Case-insensitive match
    against CLIENT_TYPES; anything that doesn't match one of those exact
    values (a typo, an unrelated note) is treated the same as blank
    rather than blocking the upload — same "opportunistic, best-effort,
    never blocks" discipline this endpoint already applies to matter
    practice_area classification just below.
    """
    raw = (raw or "").strip()
    if not raw:
        return None
    for ct in CLIENT_TYPES:
        if ct.lower() == raw.lower():
            return ct
    return None


def _normalize_beneficial_owner_cell(raw: str) -> Optional[str]:
    """
    Optional "is the client itself the beneficial owner?" column
    (2026-08-26 compliance-gap fix) — Yes/No/blank, case-insensitive.
    Deliberately narrower than client_compliance.client_is_beneficial_owner's
    full Yes/No/Unknown range: a bulk-migration spreadsheet cell is either
    answered or left blank (-> not yet assessed, same as every other
    compliance field at creation), never explicitly "Unknown" -- that
    value only makes sense as something a lawyer sets deliberately via
    the compliance modal.
    """
    raw = (raw or "").strip().lower()
    if raw == "yes":
        return "Yes"
    if raw == "no":
        return "No"
    return None


@app.post("/api/onboarding/bulk-upload")
async def bulk_onboard_from_excel(
    request: Request,
    file: UploadFile = File(...),
    commit: bool = Form(False),
):
    """
    Onboards a lawyer plus their existing client base from a filled copy of
    the firm's Client Database Excel form in one upload.

    Preview-first, matching this project's migrate_clients.py report/--yes
    convention: commit=False (the default) parses, matches, and returns the
    exact plan without writing anything. Re-submit the SAME file with
    commit=True to actually apply it. This uses a request parameter rather
    than a two-step job-store flow (parse once server-side, commit later)
    — simpler, and parsing is idempotent given the same file, so there's no
    need to persist in-progress state between the preview and commit calls.
    KNOWN LIMITATION: unlike migrate_clients.py (idempotent via a
    client_id IS NULL filter), this endpoint has no built-in guard against
    being committed twice with the same file — that would create duplicate
    matters. Review the preview, then commit once.

    Client name matching reuses match_client_name() (backend/client_migration.py)
    — the same fuzzy-matching primitives migrate_clients.py's own grouping
    uses. A name that's ambiguous against existing clients (or against
    another client already resolved earlier in this same upload) is NEVER
    auto-merged or guessed: its matters are still created (client_id left
    NULL, client_name set to the raw uploaded name — the same "legacy"
    state every pre-Client-entity matter is already in) and the name is
    listed in the response's review list, so a human can resolve it
    afterward via the existing Clients tab / matter detail relink UI —
    nothing is silently dropped.

    Newly created clients/matters get client_number/matter_number assigned
    under the uploading lawyer's initials (backend/numbering.py), same as
    the single-record /api/clients and /api/matters endpoints. Preview mode
    computes but never persists the numbers (or the lawyer's initials, if
    they weren't set yet) shown in the response.
    """
    user = await get_current_user(request)
    _check_permission(user, "admin:users")

    filename = file.filename or "upload.xlsx"
    if not filename.lower().endswith(".xlsx"):
        raise HTTPException(status_code=422, detail="Only .xlsx files are supported for this form.")

    content = await file.read()
    import openpyxl, io as _io
    try:
        wb = openpyxl.load_workbook(_io.BytesIO(content), data_only=True)
    except Exception as e:
        raise HTTPException(status_code=422, detail=f"Could not read Excel file: {e}")
    ws = wb.active

    # ── Header block: lawyer info ────────────────────────────────────────
    lawyer_name = _cell_str(ws, "B3")
    lawyer_phone = _cell_str(ws, "B4")
    lawyer_email = _cell_str(ws, "B5") or None
    lawyer_role_raw = _cell_str(ws, "B6")
    lawyer_role = {"partner": "partner", "associate": "associate"}.get(lawyer_role_raw.lower())

    header_errors = []
    if not lawyer_name:
        header_errors.append("Missing lawyer name (cell B3).")
    if not lawyer_phone:
        header_errors.append("Missing lawyer phone number (cell B4).")
    if not lawyer_role:
        header_errors.append(f'Lawyer role (cell B6) must be "Partner" or "Associate" — got {lawyer_role_raw!r}.')
    if header_errors:
        raise HTTPException(status_code=422, detail="; ".join(header_errors))

    # ── Client/matter rows ────────────────────────────────────────────────
    # One block per client: {"row", "name", "phone", "email", "contact_person", "matters": [text, ...]}
    # F (Client Type) and G (Is the client itself the beneficial owner?)
    # are optional compliance columns (2026-08-26) -- a blank cell in
    # either is normal, expected input, not an error; see
    # _normalize_client_type_cell()/_normalize_beneficial_owner_cell()
    # below for exactly how each is parsed.
    blocks = []
    row_errors = []
    current = None
    for row_idx in range(12, ws.max_row + 1):
        name = _cell_str(ws, f"A{row_idx}")
        phone = _cell_str(ws, f"B{row_idx}") or None
        email = _cell_str(ws, f"C{row_idx}") or None
        contact_person = _cell_str(ws, f"D{row_idx}") or None
        matter_text = _cell_str(ws, f"E{row_idx}")
        client_type_cell = _cell_str(ws, f"F{row_idx}")
        beneficial_owner_cell = _cell_str(ws, f"G{row_idx}")

        if name:
            current = {"row": row_idx, "name": name, "phone": phone, "email": email,
                       "contact_person": contact_person, "matters": [],
                       "client_type": _normalize_client_type_cell(client_type_cell),
                       "client_is_beneficial_owner": _normalize_beneficial_owner_cell(beneficial_owner_cell)}
            blocks.append(current)
        if matter_text:
            if current is None:
                row_errors.append(f"Row {row_idx}: matter text with no client established yet — skipped.")
                continue
            current["matters"].append(matter_text)

    if not blocks:
        raise HTTPException(status_code=422, detail="No client rows found starting at row 12.")

    async with _db_pool.acquire() as conn:
        async with conn.transaction():
            # ── Lawyer: match existing by phone, or create ─────────────────
            existing_lawyer = await conn.fetchrow(
                "SELECT * FROM users WHERE firm_id=$1 AND phone=$2", FIRM_ID, lawyer_phone
            )
            if existing_lawyer:
                lawyer_id = existing_lawyer["id"]
                lawyer_result = {
                    "action": "matched", "user_id": str(lawyer_id),
                    "display_name": existing_lawyer["display_name"], "phone": lawyer_phone,
                    "role": existing_lawyer["role"],
                }
            else:
                lawyer_id = _uuid_mod.uuid4()
                lawyer_result = {
                    "action": "created" if commit else "would_create", "user_id": str(lawyer_id),
                    "display_name": lawyer_name, "phone": lawyer_phone, "role": lawyer_role,
                }
                if commit:
                    await conn.execute(
                        """INSERT INTO users (id, firm_id, phone, email, display_name, role, is_active, created_at)
                           VALUES ($1,$2,$3,$4,$5,$6,$7,$8)""",
                        lawyer_id, FIRM_ID, lawyer_phone, lawyer_email, lawyer_name, lawyer_role, True, datetime.utcnow(),
                    )

            # Initials for client numbering. Only user_id if commit actually
            # wrote (or matched) a real users row — a not-yet-created lawyer
            # in preview mode has no row to read/persist against.
            # persist=commit: preview must not write anything, including a
            # matched-existing-lawyer's not-yet-set initials.
            lawyer_result["initials"] = lawyer_initials = await _resolve_user_initials(
                conn, FIRM_ID,
                existing_lawyer["id"] if existing_lawyer else (lawyer_id if commit else None),
                existing_lawyer["display_name"] if existing_lawyer else lawyer_name,
                persist=commit,
            )
            # Numbering, commit vs preview. Commit uses the same atomic
            # _next_client_number()/_next_matter_number() (via
            # _create_client_row()/_create_matter_row() below) as every
            # other creation path -- a real row lock on numbering_counters,
            # so two concurrent writers (e.g. this upload and someone
            # hitting POST /api/clients at the same moment) can never be
            # handed the same number. This used to be a manual in-Python
            # counter seeded from a SELECT...LIKE scan that never touched
            # numbering_counters at all -- not concurrency-safe, fixed as
            # part of the Phase 1a consolidation (see the note above
            # _create_client_row()). Preview must NOT call the atomic
            # functions -- that would burn a real sequence number for a
            # plan that might never be applied -- so it keeps the original
            # read-only running-estimate logic, unchanged.
            next_client_seq = None
            matter_seq_cache = {}  # preview only: client_number -> next matter sequence estimate
            if not commit:
                existing_client_numbers = await conn.fetch(
                    "SELECT client_number FROM clients WHERE firm_id=$1 AND client_number LIKE $2",
                    FIRM_ID, f"{lawyer_initials}-%",
                )
                next_client_seq = next_sequence(
                    [r["client_number"] for r in existing_client_numbers], lawyer_initials
                )

            async def _next_matter_seq_estimate(client_number):
                """Preview-only read-only estimate — never allocated for real."""
                if client_number not in matter_seq_cache:
                    mrows = await conn.fetch(
                        "SELECT matter_number FROM matters WHERE firm_id=$1 AND matter_number LIKE $2",
                        FIRM_ID, f"{client_number}-%",
                    )
                    matter_seq_cache[client_number] = next_sequence(
                        [r["matter_number"] for r in mrows], client_number
                    )
                seq = matter_seq_cache[client_number]
                matter_seq_cache[client_number] = seq + 1
                return seq

            # ── Candidate pool for client-name matching: existing DB
            # clients, growing as new clients are resolved within this
            # same upload (so a repeated near-identical name later in the
            # same form matches the one just created, not a fresh one).
            # client_number rides along for matched candidates so their
            # matters can be numbered too — match_client_name() only reads
            # "full_name", so the extra key is harmless.
            existing_client_rows = await conn.fetch(
                "SELECT id, full_name, client_number FROM clients WHERE firm_id=$1", FIRM_ID
            )
            pool = [{"id": str(r["id"]), "full_name": r["full_name"], "client_number": r["client_number"]}
                    for r in existing_client_rows]

            clients_created, clients_matched, clients_review = [], [], []
            matters_created, matters_unlinked = [], []
            now = datetime.utcnow()

            for block in blocks:
                match = match_client_name(block["name"], pool)
                client_number_for_block = None

                if match["status"] == "matched":
                    client_id = _uuid_mod.UUID(match["candidate"]["id"])
                    client_number_for_block = match["candidate"].get("client_number")
                    clients_matched.append({
                        "row": block["row"], "name": block["name"],
                        "matched_client_id": match["candidate"]["id"],
                        "matched_client_name": match["candidate"]["full_name"],
                        "matched_client_number": client_number_for_block,
                    })
                elif match["status"] == "no_match":
                    if commit:
                        client_row = await _create_client_row(
                            conn, FIRM_ID, lawyer_initials, block["name"],
                            email=block["email"], phone=block["phone"],
                            contact_person=block["contact_person"],
                            created_by=lawyer_id,  # FIX (Phase 1a): previously omitted entirely
                            client_type=block["client_type"],
                            client_is_beneficial_owner=block["client_is_beneficial_owner"],
                        )
                        new_id = client_row["id"]
                        client_number_for_block = client_row["client_number"]
                    else:
                        new_id = _uuid_mod.uuid4()
                        client_number_for_block = format_client_number(lawyer_initials, next_client_seq)
                        next_client_seq += 1
                    client_id = new_id
                    clients_created.append({
                        "row": block["row"], "name": block["name"], "client_id": str(new_id),
                        "contact_person": block["contact_person"], "client_number": client_number_for_block,
                    })
                    pool.append({"id": str(new_id), "full_name": block["name"], "client_number": client_number_for_block})
                else:  # ambiguous — never guess; matters below stay unlinked
                    client_id = None
                    clients_review.append({
                        "row": block["row"], "name": block["name"], "candidates": match["candidates"],
                    })

                for matter_text in block["matters"]:
                    # Opportunistic, best-effort — never blocks the upload.
                    # Only a confident single-category keyword match is
                    # applied; ambiguous/no-match rows are simply left
                    # NULL, same as any matter created without a
                    # practice_area today (covered by the backfill script,
                    # not guessed here).
                    classification = classify_practice_area(extract_classification_text(matter_text))
                    practice_area = classification.get("practice_area") if classification["status"] == "matched" else None

                    if commit:
                        matter_row = await _create_matter_row(
                            conn, FIRM_ID, matter_text,
                            client_name=block["name"], client_id=client_id,
                            status="Active", practice_area=practice_area,
                            numbering_client_number=client_number_for_block,
                            created_by=lawyer_id, created_at=now, last_activity=now,
                        )
                        matter_id = matter_row["id"]
                        matter_number = matter_row["matter_number"]
                    else:
                        matter_id = _uuid_mod.uuid4()
                        matter_number = (
                            format_matter_number(client_number_for_block, await _next_matter_seq_estimate(client_number_for_block))
                            if client_number_for_block else None
                        )
                    entry = {"row": block["row"], "client_name": block["name"],
                             "name": matter_text, "matter_id": str(matter_id), "matter_number": matter_number,
                             "practice_area": practice_area}
                    (matters_created if client_id is not None else matters_unlinked).append(entry)

    return {
        "committed": commit,
        "lawyer": lawyer_result,
        "clients": {
            "created": clients_created,
            "matched": clients_matched,
            "review": clients_review,
        },
        "matters": {
            "created": matters_created,
            "created_unlinked_pending_review": matters_unlinked,
        },
        "errors": row_errors,
    }

# ── Single-client intake (Case Binder provisioning) ─────────────────────────
# Preview-first (a `commit` field, default False), same convention as
# bulk_onboard_from_excel just above. Client-identity resolution reuses
# match_client_name() (backend/client_migration.py) the same way:
# zero matches -> create, one match -> reuse, 2+ -> flag for manual review
# and never guess.
#
# Numbering (backend/numbering.py) is only actually *allocated* on commit,
# via the same atomic _next_client_number()/_next_matter_number()
# create_client/create_matter already use — those functions atomically
# increment numbering_counters the moment they're called, so calling them
# in preview mode would burn a real sequence number nothing ends up using
# (the next real commit would then skip a number). Preview instead computes
# a *would-be* number via numbering.py's next_sequence() directly against
# current rows — a read-only estimate, not a reservation, same technique
# bulk_onboard_from_excel's own preview already relies on for exactly this
# reason.

class ClientIntakeRequest(BaseModel):
    client_name: str
    phone: str
    email: Optional[str] = None
    contact_person: Optional[str] = None
    matter_type: str  # validated against INTAKE_MATTER_TYPES below
    matter_description: str
    assigned_lawyer_id: str
    existing_client_id: Optional[str] = None
    commit: bool = False
    # Client-generated (a UUID is recommended, but any string is accepted)
    # — guards commit=true against an accidental duplicate submit (double-
    # click, a network retry, a client-side retry-on-timeout that actually
    # succeeded server-side). Optional and additive: omitting it reproduces
    # today's exact behaviour with no duplicate-submit protection at all —
    # see the docstring below for what that means in practice.
    idempotency_key: Optional[str] = None

@app.post("/api/onboarding/intake")
async def client_intake(req: ClientIntakeRequest, request: Request):
    """
    Single-client version of bulk_onboard_from_excel above: takes one new
    (or existing) client plus one new matter, and auto-provisions that
    matter's Case Binder starter documents (backend/case_binder.py) —
    intended for the normal one-at-a-time "a new client just walked in"
    intake, not a batch upload.

    commit=False (the default) returns the exact plan — client match/create
    decision, matter number, and the case-binder document list — without
    writing anything. Re-submit with commit=True to actually apply it.

    Duplicate-submit protection on commit=true is opt-in via
    idempotency_key: pass the same key on a retry of the same logical
    request, and a detected duplicate short-circuits to the ORIGINAL
    response — the exact one the first, real commit produced — without
    creating a second client, matter, or set of case-binder documents, and
    without a second audit_logs entry. Checked before anything else in the
    request happens, keyed per-firm. KNOWN LIMITATION: if idempotency_key
    is omitted (it's optional, not required), there is no protection at
    all — a genuine double-click or retry with no key will create a
    duplicate client/matter/documents exactly as it would today. Preview
    mode (commit=false) needs no such guard — it never writes anything, so
    repeating a preview call is already safe on its own.

    An ambiguous client-name match (2+ existing clients score similarly)
    blocks the ENTIRE request rather than proceeding with anything — unlike
    bulk_onboard_from_excel, there is no batch of other rows that still need
    processing here, so there's nothing to gain by partially creating a
    matter with no resolved client. Resolve the ambiguity (e.g. by
    re-submitting with existing_client_id set to the correct candidate)
    and try again. (Not idempotency-guarded — an ambiguous match writes
    nothing on any attempt, so retrying it is already naturally safe.)
    """
    user = await get_current_user(request)
    _check_permission(user, "client:create")

    if req.matter_type not in INTAKE_MATTER_TYPES:
        raise HTTPException(
            status_code=422,
            detail=f"matter_type must be one of: {', '.join(INTAKE_MATTER_TYPES)}",
        )

    try:
        lawyer_uuid = _uuid_mod.UUID(req.assigned_lawyer_id)
    except ValueError:
        raise HTTPException(status_code=422, detail="assigned_lawyer_id must be a valid UUID")

    async with _db_pool.acquire() as conn:
        if req.commit and req.idempotency_key:
            # Checked first, before any lookup/validation work below —
            # a genuine duplicate returns the ORIGINAL response as-is,
            # not a freshly-recomputed one.
            existing_key_row = await conn.fetchrow(
                "SELECT response_body FROM intake_idempotency_keys WHERE firm_id=$1 AND key=$2",
                FIRM_ID, req.idempotency_key,
            )
            if existing_key_row:
                return json.loads(existing_key_row["response_body"])

        lawyer_row = await conn.fetchrow(
            "SELECT id, display_name FROM users WHERE firm_id=$1 AND id=$2", FIRM_ID, lawyer_uuid
        )
        if not lawyer_row:
            raise HTTPException(status_code=404, detail="assigned_lawyer_id is not a user in this firm")

        # ── Client resolution (read-only — safe regardless of commit) ──────
        client_id = None
        client_number = None
        client_full_name = req.client_name
        client_result = {}

        if req.existing_client_id:
            try:
                existing_uuid = _uuid_mod.UUID(req.existing_client_id)
            except ValueError:
                raise HTTPException(status_code=422, detail="existing_client_id must be a valid UUID")
            existing = await conn.fetchrow(
                "SELECT id, full_name, client_number FROM clients WHERE firm_id=$1 AND id=$2",
                FIRM_ID, existing_uuid,
            )
            if not existing:
                raise HTTPException(status_code=404, detail="existing_client_id not found")
            client_id = existing["id"]
            client_number = existing["client_number"]
            client_full_name = existing["full_name"]
            client_result = {
                "action": "matched_explicit", "client_id": str(client_id),
                "full_name": client_full_name, "client_number": client_number,
            }
        else:
            existing_client_rows = await conn.fetch(
                "SELECT id, full_name, client_number FROM clients WHERE firm_id=$1", FIRM_ID
            )
            candidate_pool = [
                {"id": str(r["id"]), "full_name": r["full_name"], "client_number": r["client_number"]}
                for r in existing_client_rows
            ]
            match = match_client_name(req.client_name, candidate_pool)

            if match["status"] == "matched":
                client_id = _uuid_mod.UUID(match["candidate"]["id"])
                client_number = match["candidate"].get("client_number")
                client_full_name = match["candidate"]["full_name"]
                client_result = {
                    "action": "matched", "client_id": str(client_id),
                    "full_name": client_full_name, "client_number": client_number,
                }
            elif match["status"] == "ambiguous":
                return {
                    "committed": False,
                    "client": {"action": "review_required", "candidates": match["candidates"]},
                    "matter": None,
                    "case_binder": [],
                    "errors": [
                        "Client name is ambiguous against existing clients — resolve by "
                        "re-submitting with existing_client_id set to the correct candidate, "
                        "or use a more specific name."
                    ],
                }
            else:  # no_match
                client_full_name = req.client_name
                if req.commit:
                    # Actual id/client_number allocation is deferred to the
                    # transaction below (via _create_client_row(), called
                    # alongside the matter/case-binder/audit-log writes) --
                    # bundling atomic numbering with the insert inside that
                    # same transaction means a failure anywhere in it rolls
                    # back the number allocation too, not just the row.
                    # (Previously numbering was its own committed statement
                    # before the transaction even opened, so a burned
                    # sequence number could survive a rolled-back commit --
                    # fixed as an incidental consequence of the Phase 1a
                    # consolidation, not a separate deliberate change.)
                    client_result = {
                        "action": "created", "client_id": None,
                        "full_name": client_full_name, "client_number": None,
                    }
                else:
                    initials = await _resolve_user_initials(
                        conn, FIRM_ID, lawyer_row["id"], lawyer_row["display_name"], persist=False
                    )
                    # Would-be number only — see module docstring above for
                    # why this can't call _next_client_number() in preview.
                    existing_numbers = await conn.fetch(
                        "SELECT client_number FROM clients WHERE firm_id=$1 AND client_number LIKE $2",
                        FIRM_ID, f"{initials}-%",
                    )
                    would_be_seq = next_sequence([r["client_number"] for r in existing_numbers], initials)
                    client_number = format_client_number(initials, would_be_seq)
                    client_result = {
                        "action": "would_create", "client_id": None,
                        "full_name": client_full_name, "client_number": client_number,
                    }

        # ── Matter + Case Binder ────────────────────────────────────────────
        if req.commit:
            async with conn.transaction():
                now = datetime.utcnow()
                if client_result["action"] == "created":
                    initials = await _resolve_user_initials(
                        conn, FIRM_ID, lawyer_row["id"], lawyer_row["display_name"], persist=True
                    )
                    client_row = await _create_client_row(
                        conn, FIRM_ID, initials, client_full_name,
                        email=req.email, phone=req.phone, contact_person=req.contact_person,
                        created_by=lawyer_row["id"],
                    )
                    client_id = client_row["id"]
                    client_number = client_row["client_number"]
                    client_result["client_id"] = str(client_id)
                    client_result["client_number"] = client_number

                matter_row = await _create_matter_row(
                    conn, FIRM_ID, req.matter_description,
                    client_name=client_full_name, client_id=client_id,
                    status="Active", matter_type=req.matter_type,
                    assigned_lawyer_id=lawyer_row["id"],
                    numbering_client_number=client_number,
                    created_by=lawyer_row["id"], created_at=now, last_activity=now,
                )
                matter_id = matter_row["id"]
                matter_number = matter_row["matter_number"]
                matter_result = {
                    "action": "created", "matter_id": str(matter_id), "matter_number": matter_number,
                    "name": req.matter_description, "matter_type": req.matter_type,
                }

                binder_items = provision_case_binder(
                    {"matter_number": matter_number}, req.matter_type, {"full_name": client_full_name}
                )
                case_binder_result = []
                for item in binder_items:
                    doc_id = _uuid_mod.uuid4()
                    content_bytes = item["content"].encode("utf-8")
                    r2_key = None
                    # R2 upload failure is best-effort, same resilience as
                    # the normal single-document upload path — a starter
                    # document still gets its row (so it shows up on the
                    # matter and can be re-uploaded later) even if the
                    # object store round-trip fails.
                    if R2_ENABLED and _r2_client:
                        try:
                            r2_key = f"{FIRM_ID}/{matter_id}/{doc_id}/{item['name']}.txt"
                            await asyncio.to_thread(
                                _r2_client.put_object, Bucket=R2_BUCKET, Key=r2_key,
                                Body=content_bytes, ContentType="text/plain",
                            )
                        except Exception as e:
                            print(f"[case-binder] R2 upload failed for {item['name']!r}: {e}")
                            r2_key = None
                    await conn.execute(
                        """INSERT INTO documents (id, matter_id, firm_id, filename, source, status,
                                                  word_count, page_count, r2_key, uploaded_at, uploaded_by,
                                                  document_status, provenance_document_type)
                           VALUES ($1,$2,$3,$4,'auto_provisioned','complete',$5,1,$6,$7,$8,'Draft',$9)""",
                        doc_id, matter_id, FIRM_ID, item["name"], len(item["content"].split()),
                        r2_key, now, lawyer_row["id"], item.get("provenance_document_type") or "General",
                    )
                    case_binder_result.append({
                        "action": "created", "document_id": str(doc_id), "name": item["name"],
                        "template_source": item["template_source"], "r2_uploaded": r2_key is not None,
                    })

                await conn.execute(
                    """INSERT INTO audit_logs (firm_id, user_id, actor_name, actor_role, action,
                                               target_type, target_id, details)
                       VALUES ($1, $2, $3, $4, 'CLIENT_INTAKE', 'MATTER', $5, $6)""",
                    FIRM_ID,
                    _uuid_mod.UUID(str(user["id"])) if user.get("id") else None,
                    user.get("display_name", "Unknown"), user.get("role", "unknown"),
                    matter_id,
                    json.dumps({
                        "client_action": client_result["action"],
                        "client_id": client_result.get("client_id"),
                        "matter_number": matter_number,
                        "matter_type": req.matter_type,
                        "assigned_lawyer_id": str(lawyer_row["id"]),
                        "case_binder_documents_created": len(case_binder_result),
                    }),
                )

                final_response = {
                    "committed": True,
                    "client": client_result,
                    "matter": matter_result,
                    "case_binder": case_binder_result,
                    "errors": [],
                }
                if req.idempotency_key:
                    # Stored in this same transaction as the writes above —
                    # a crash between "processed" and "key stored" can't
                    # happen; either both land or neither does. ON CONFLICT
                    # DO NOTHING as a defensive backstop against a
                    # genuinely simultaneous duplicate racing the initial
                    # SELECT above (this guard is a check-then-act, not a
                    # row lock — adequate for the double-click/retry/
                    # accidental-refire case this exists for, not a
                    # guarantee against millisecond-scale true concurrency).
                    await conn.execute(
                        """INSERT INTO intake_idempotency_keys (firm_id, key, response_body)
                           VALUES ($1, $2, $3)
                           ON CONFLICT (firm_id, key) DO NOTHING""",
                        FIRM_ID, req.idempotency_key, json.dumps(final_response),
                    )
                return final_response
        else:
            # Preview: matter number is a would-be estimate (see module
            # docstring); nothing below this point writes anything.
            if client_number:
                existing_matter_numbers = await conn.fetch(
                    "SELECT matter_number FROM matters WHERE firm_id=$1 AND matter_number LIKE $2",
                    FIRM_ID, f"{client_number}-%",
                )
                would_be_matter_seq = next_sequence(
                    [r["matter_number"] for r in existing_matter_numbers], client_number
                )
                matter_number = format_matter_number(client_number, would_be_matter_seq)
            else:
                matter_number = None
            matter_result = {
                "action": "would_create", "matter_id": None, "matter_number": matter_number,
                "name": req.matter_description, "matter_type": req.matter_type,
            }
            binder_items = provision_case_binder(
                {"matter_number": matter_number}, req.matter_type, {"full_name": client_full_name}
            )
            case_binder_result = [
                {"action": "would_create", "document_id": None, "name": item["name"],
                 "template_source": item["template_source"]}
                for item in binder_items
            ]

    return {
        "committed": req.commit,
        "client": client_result,
        "matter": matter_result,
        "case_binder": case_binder_result,
        "errors": [],
    }

# ── Reports (RBZ compliance export) ─────────────────────────────────────────
# Partner-tier only (see PERMISSIONS["reports:rbz_compliance"]). Firm-wide —
# deliberately ignores the panel_lawyer scoping list_matters applies, since
# a compliance export exists precisely to show everything, not one lawyer's
# subset. Every generation is logged to report_history (who/when/counts),
# including a zero-client run, so the firm can show a habit of regular
# generation rather than a one-off produced right before an inspection.

async def _fetch_rbz_compliance_rows(conn) -> list:
    """Firm-wide clients + their linked matters, one row per matter (or one
    placeholder row for a client with none) — the shared dataset behind
    both export formats below."""
    return await conn.fetch("""
        SELECT c.id AS client_id, c.client_number, c.full_name AS client_name, c.contact_person,
               c.phone, c.email,
               m.matter_number, m.status AS matter_status
        FROM clients c
        LEFT JOIN matters m ON m.client_id = c.id
        WHERE c.firm_id=$1
        ORDER BY c.full_name ASC, m.created_at ASC
    """, FIRM_ID)

def _rbz_compliance_counts(rows: list) -> tuple:
    client_count = len({r["client_id"] for r in rows})
    # matter_status (not matter_number) signals a real joined matter row vs.
    # the LEFT JOIN's all-NULL placeholder for a client with no matters —
    # matter_number can legitimately be NULL on a real, not-yet-backfilled
    # matter, so it isn't a safe "row exists" check.
    matter_count = len([r for r in rows if r["matter_status"] is not None])
    return client_count, matter_count

async def _log_rbz_compliance_report(conn, user: dict, report_type: str, client_count: int, matter_count: int) -> None:
    await conn.execute("""
        INSERT INTO report_history (firm_id, report_type, generated_by, generated_by_name,
                                    client_count, matter_count)
        VALUES ($1, $2, $3, $4, $5, $6)
    """,
    FIRM_ID, report_type,
    _uuid_mod.UUID(str(user["id"])) if user.get("id") else None,
    user.get("display_name") or "Unknown", client_count, matter_count)

@app.get("/api/reports/rbz-compliance-export")
async def export_rbz_compliance_report(request: Request):
    user = await get_current_user(request)
    _check_permission(user, "reports:rbz_compliance")

    async with _db_pool.acquire() as conn:
        rows = await _fetch_rbz_compliance_rows(conn)
        client_count, matter_count = _rbz_compliance_counts(rows)
        await _log_rbz_compliance_report(conn, user, "rbz_compliance_export", client_count, matter_count)

    import csv, io as _io
    buf = _io.StringIO()
    writer = csv.writer(buf)
    writer.writerow(["Client Number", "Client Name", "Contact Person", "Phone", "Email",
                      "Matter Number", "Matter Status"])
    for r in rows:
        writer.writerow([
            r["client_number"] or "", r["client_name"] or "", r["contact_person"] or "",
            r["phone"] or "", r["email"] or "",
            r["matter_number"] or "", r["matter_status"] or "",
        ])

    filename = f"rbz_compliance_export_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.csv"
    return Response(
        content=buf.getvalue(),
        media_type="text/csv",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )

@app.get("/api/reports/rbz-compliance-export-pdf")
async def export_rbz_compliance_report_pdf(request: Request):
    """
    Same data and access restriction as the CSV export above — a second
    output format, not a separate feature. Deliberately plain: a readable
    table, not a designed document. fpdf2 (pure Python, no system
    dependencies like weasyprint/wkhtmltopdf would need) was the only PDF
    library added for this — nothing else in the codebase generates PDFs
    (pdfplumber, already a dependency, only reads them).
    """
    user = await get_current_user(request)
    _check_permission(user, "reports:rbz_compliance")

    async with _db_pool.acquire() as conn:
        rows = await _fetch_rbz_compliance_rows(conn)
        client_count, matter_count = _rbz_compliance_counts(rows)
        await _log_rbz_compliance_report(conn, user, "rbz_compliance_export_pdf", client_count, matter_count)

    pdf_bytes = _build_rbz_compliance_pdf(rows)

    filename = f"rbz_compliance_export_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.pdf"
    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )

def _build_rbz_compliance_pdf(rows: list) -> bytes:
    """
    Groups the flat client+matter rows (already ordered by client, then by
    matter created_at — see _fetch_rbz_compliance_rows) into one block per
    client: a header line with client_number/name/contact person/phone/
    email, then that client's matters listed underneath. Deliberately
    simple — one readable table, no per-page styling beyond a title.
    """
    from fpdf import FPDF

    pdf = FPDF(orientation="P", unit="mm", format="A4")
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(0, 8, _pdf_safe(FIRM_NAME), new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 10)
    pdf.cell(0, 6, "RBZ Compliance Export", new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 6, f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M')} UTC", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)

    if not rows:
        pdf.set_font("Helvetica", "I", 10)
        pdf.cell(0, 6, "No clients on file.", new_x="LMARGIN", new_y="NEXT")
        return bytes(pdf.output())

    current_client_id = None
    for r in rows:
        if r["client_id"] != current_client_id:
            current_client_id = r["client_id"]
            pdf.ln(3)
            pdf.set_font("Helvetica", "B", 11)
            header = f"{r['client_number']} - {r['client_name']}" if r["client_number"] else r["client_name"]
            pdf.cell(0, 6, _pdf_safe(header), new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Helvetica", "", 9)
            contact_bits = [b for b in (r["contact_person"], r["phone"], r["email"]) if b]
            if contact_bits:
                pdf.cell(0, 5, _pdf_safe("  " + " | ".join(contact_bits)), new_x="LMARGIN", new_y="NEXT")

        if r["matter_status"] is not None:
            pdf.set_font("Helvetica", "", 9)
            matter_label = r["matter_number"] or "(unnumbered)"
            pdf.cell(0, 5, _pdf_safe(f"    - {matter_label}: {r['matter_status']}"), new_x="LMARGIN", new_y="NEXT")

    return bytes(pdf.output())

def _pdf_safe(text: str) -> str:
    """
    fpdf2's core Helvetica font is latin-1 only — a bundled Unicode font
    would be needed to render arbitrary characters, which is more weight
    than this deliberately simple export calls for. Common punctuation
    that free-text matter descriptions/names use (em/en dashes, curly
    quotes — the onboarding template's own examples use em-dashes) gets a
    plain-ASCII substitute; anything else outside latin-1 (e.g. an
    accented name) degrades to "?" via latin-1's replace mode rather than
    crashing the whole export over one character.
    """
    if not text:
        return ""
    substitutions = {
        "—": "-", "–": "-",       # em dash, en dash
        "‘": "'", "’": "'",       # curly single quotes
        "“": '"', "”": '"',       # curly double quotes
        "…": "...",                    # ellipsis
    }
    for char, replacement in substitutions.items():
        text = text.replace(char, replacement)
    return text.encode("latin-1", errors="replace").decode("latin-1")

@app.get("/api/reports/history")
async def list_report_history(request: Request):
    user = await get_current_user(request)
    _check_permission(user, "reports:rbz_compliance")
    async with _db_pool.acquire() as conn:
        rows = await conn.fetch(
            "SELECT * FROM report_history WHERE firm_id=$1 ORDER BY generated_at DESC LIMIT 100",
            FIRM_ID
        )
    result = []
    for r in rows:
        d = dict(r)
        for k in ("id", "firm_id", "generated_by"):
            if d.get(k):
                d[k] = str(d[k])
        if d.get("generated_at"):
            d["generated_at"] = d["generated_at"].isoformat()
        result.append(d)
    return result

@app.get("/api/reports/practice-area-breakdown")
async def practice_area_breakdown(request: Request):
    """
    Partner-facing breakdown: matter count per practice_area, firm-wide.
    Matters without a practice_area yet (not backfilled — see
    scripts/backfill_practice_areas.py — or created before this feature)
    are grouped under "Uncategorized" rather than silently excluded from
    the total.
    """
    user = await get_current_user(request)
    _check_permission(user, "reports:practice_area_breakdown")
    async with _db_pool.acquire() as conn:
        rows = await conn.fetch(
            "SELECT practice_area, COUNT(*) AS matter_count FROM matters "
            "WHERE firm_id=$1 AND NOT is_sentinel GROUP BY practice_area ORDER BY matter_count DESC",
            FIRM_ID
        )
    return [
        {"practice_area": r["practice_area"] or "Uncategorized", "matter_count": r["matter_count"]}
        for r in rows
    ]

# ── Documents ─────────────────────────────────────────────────────────────────

@app.get("/api/matters/{matter_id}/documents")
async def list_documents(matter_id: str, request: Request):
    user = await get_current_user(request)
    _check_permission(user, "matter:read")
    async with _db_pool.acquire() as conn:
        rows = await conn.fetch(
            "SELECT * FROM documents WHERE matter_id=$1 AND firm_id=$2 ORDER BY uploaded_at DESC",
            _uuid_mod.UUID(matter_id), FIRM_ID
        )
    return [_row_to_doc(r) for r in rows]

@app.get("/api/documents/recent")
async def list_recent_documents(request: Request, limit: int = 10):
    """
    Firm-wide recent completed uploads, joined with their matter for display
    context — backs the Dashboard tab's activity feed (the document-upload
    half; progress notes are already attached per-matter by GET /api/matters
    and don't need a separate fetch). Same matter:read gate as the
    per-matter document list above; no extra role scoping beyond that.
    """
    user = await get_current_user(request)
    _check_permission(user, "matter:read")
    async with _db_pool.acquire() as conn:
        rows = await conn.fetch("""
            SELECT d.*, m.name AS matter_name, m.client_name AS matter_client_name
            FROM documents d
            JOIN matters m ON m.id = d.matter_id
            WHERE d.firm_id=$1 AND d.status='complete'
            ORDER BY d.uploaded_at DESC
            LIMIT $2
        """, FIRM_ID, limit)
    result = []
    for r in rows:
        d = _row_to_doc(r)
        d["matter_name"] = r["matter_name"]
        d["matter_client_name"] = r["matter_client_name"]
        result.append(d)
    return result

async def _process_document_background(doc_id: str, matter_id: str, content: bytes, filename: str, ext: str):
    """
    Background task: extract text, classify, chunk, and index a document.
    Updates the document record in PostgreSQL when complete.
    This runs after the upload endpoint has already returned 202 to the client.
    """
    text = ""
    word_count = 0
    page_count = 1
    ocr_used = False
    ocr_confidence = None

    try:
        if ext == "pdf":
            text, page_count, ocr_used, ocr_confidence = extract_pdf_text(content)
        elif ext in ("docx", "doc"):
            text = extract_docx_text(content)
        elif ext in ("xlsx", "xlsm"):
            text = extract_xlsx_text(content)
        else:
            text = content.decode("utf-8", errors="replace")
        word_count = len(text.split())
    except Exception as e:
        print(f"[upload] text extraction failed for {filename}: {e}")

    metadata = {}
    if text:
        try:
            metadata = await asyncio.to_thread(classify_document_sync, text[:2000])
        except Exception:
            metadata = {}

    legal_source_type = classify_firm_document(metadata.get("document_type"))
    authority_strength = authority_strength_for(legal_source_type)

    chunk_count = 0
    if text:
        new_chunks = chunk_text(text, page_count, doc_id, matter_id)
        for c in new_chunks:
            c["chunk_source"] = "firm"
        if new_chunks:
            await asyncio.to_thread(index_chunks_in_chroma, new_chunks, "firm")
            # Persist chunks to PostgreSQL
            async with _db_pool.acquire() as conn:
                for c in new_chunks:
                    await conn.execute("""
                        INSERT INTO chunks (id, firm_id, document_id, matter_id, chunk_source,
                                           text, chunk_index, page_number, created_at)
                        VALUES ($1,$2,$3,$4,$5,$6,$7,$8,NOW())
                        ON CONFLICT (id) DO NOTHING
                    """,
                    c["id"], FIRM_ID, _uuid_mod.UUID(doc_id),
                    matter_id, "firm", c["text"], c["chunk_index"], c.get("page_number", 1)
                    )
            chunk_count = len(new_chunks)

    # Parse doc_date safely
    raw_date = metadata.get("doc_date")
    doc_date = None
    if raw_date:
        try:
            doc_date = datetime.strptime(raw_date, "%Y-%m-%d").date()
        except Exception:
            doc_date = None

    # Upload original file to R2 for view/download
    r2_key = None
    if R2_ENABLED and _r2_client and content:
        try:
            r2_key = f"{FIRM_ID}/{matter_id}/{doc_id}/{filename}"
            content_type = (
                "application/pdf" if ext == "pdf" else
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document" if ext in ("docx","doc") else
                "application/octet-stream"
            )
            await asyncio.to_thread(
                _r2_client.put_object,
                Bucket=R2_BUCKET,
                Key=r2_key,
                Body=content,
                ContentType=content_type,
            )
            print(f"[r2] uploaded {filename} → {r2_key}")
        except Exception as e:
            print(f"[r2] upload failed for {filename}: {e}")
            r2_key = None

    async with _db_pool.acquire() as conn:
        needs_review = ocr_used and (ocr_confidence is not None) and (ocr_confidence < 80)
        await conn.execute("""
            UPDATE documents SET
                document_type=$1, matter_type=$2, parties=$3,
                doc_date=$4, court=$5, word_count=$6, page_count=$7,
                chunk_count=$8, ocr_used=$9, status='complete', r2_key=$12,
                ocr_confidence=$13, needs_review=$14,
                legal_source_type=$15, authority_strength=$16
            WHERE id=$10 AND firm_id=$11
        """,
        metadata.get("document_type"), metadata.get("matter_type"),
        str(metadata.get("parties", "")) if metadata.get("parties") else None,
        doc_date, metadata.get("court"),
        word_count, page_count, chunk_count, ocr_used,
        _uuid_mod.UUID(doc_id), FIRM_ID, r2_key,
        ocr_confidence, needs_review,
        legal_source_type.value, authority_strength.value
        )
        await conn.execute(
            "UPDATE matters SET document_count = document_count + 1, last_activity=NOW() WHERE id=$1 AND firm_id=$2",
            _uuid_mod.UUID(matter_id), FIRM_ID
        )

    if needs_review:
        print(f"[upload] ⚠ {filename}: OCR confidence {ocr_confidence}% (below 80%) — flagged for manual review")
    print(f"[upload] processed {filename}: {word_count} words, {chunk_count} chunks, ocr={ocr_used}")

@app.post("/api/upload", status_code=202)
async def upload_document(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    matter_id: str = Form(...),
    provenance_document_type: Optional[str] = Form(None),
    document_status: Optional[str] = Form(None),
    description: Optional[str] = Form(None),
    confidentiality: Optional[str] = Form(None),
    request: Request = None,
):
    if request:
        user = await get_current_user(request)
        _check_permission(user, "document:upload")
    else:
        user = None

    if provenance_document_type is not None and provenance_document_type not in PROVENANCE_DOCUMENT_TYPES:
        raise HTTPException(status_code=422, detail=f"provenance_document_type must be one of {PROVENANCE_DOCUMENT_TYPES}")
    if document_status is not None and document_status not in DOCUMENT_STATUSES:
        raise HTTPException(status_code=422, detail=f"document_status must be one of {DOCUMENT_STATUSES}")
    if confidentiality is not None and confidentiality not in DOCUMENT_CONFIDENTIALITY_LEVELS:
        raise HTTPException(status_code=422, detail=f"confidentiality must be one of {DOCUMENT_CONFIDENTIALITY_LEVELS}")
    document_status = document_status or "Draft"
    confidentiality = confidentiality or "Standard"

    async with _db_pool.acquire() as conn:
        matter = await conn.fetchrow(
            "SELECT id FROM matters WHERE id=$1 AND firm_id=$2",
            _uuid_mod.UUID(matter_id), FIRM_ID
        )
    if not matter:
        raise HTTPException(status_code=404, detail="Matter not found")

    content = await file.read()
    filename = file.filename or "document"
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else "bin"
    doc_id = str(_uuid_mod.uuid4())

    # Insert document record immediately with status='processing'
    async with _db_pool.acquire() as conn:
        row = await conn.fetchrow("""
            INSERT INTO documents (id, matter_id, firm_id, filename, status, uploaded_at, uploaded_by,
                                    provenance_document_type, document_status, description, confidentiality)
            VALUES ($1,$2,$3,$4,'processing',NOW(),$5,$6,$7,$8,$9) RETURNING *
        """,
        _uuid_mod.UUID(doc_id), _uuid_mod.UUID(matter_id), FIRM_ID, filename,
        _uuid_mod.UUID(str(user["id"])) if user and user.get("id") else None,
        provenance_document_type, document_status, description, confidentiality,
        )

    # Schedule heavy processing in the background — returns immediately to client
    background_tasks.add_task(
        _process_document_background, doc_id, matter_id, content, filename, ext
    )

    return {**_row_to_doc(row), "processing": True,
            "message": "Document received. Text extraction and indexing are running in the background."}

# ── Rapid Precedent Capture ───────────────────────────────────────────────────
# Fast, low-friction phone/camera capture path, distinct from /api/upload
# above: minimal required input at capture time (a matter, or nothing —
# defaults to the firm's sentinel "General / Firm Precedents" matter), no
# manual tagging gate, and support for either (a) one or more photographed
# pages of a single physical document, combined into one document, or (b)
# a single already-digital file picked from the device (e.g. an email
# attachment just downloaded to the phone) -- see capture_documents().
#
# document_status defaults to 'Final' here, not 'Draft' -- deliberately
# distinct from /api/upload's Draft default (line ~4771 above). A lawyer
# photographing a document is virtually always capturing something already
# finished (a signed judgment, an executed contract, a filed pleading),
# not a work-in-progress draft; case-binder shells and general uploads
# have the opposite default because that's overwhelmingly what's happening
# there instead. Not shared logic, so not factored into one constant --
# these are two independently-meaningful defaults that happen to differ.

def _combine_capture_pages_to_pdf(pages: list) -> bytes:
    """
    pages: list of (content: bytes, ext: str) for photographed image pages,
    in capture order. Combines them into a single multi-page PDF so a
    multi-photo capture session becomes one coherent, downloadable
    document -- matching every other document in the Vault -- rather than
    N loose page images with nowhere to live under documents.r2_key's
    single-file-per-document shape. Pillow can write a multi-page PDF
    directly (save_all=True, append_images=...); already a dependency for
    OCR image handling, so this needs nothing new.
    """
    from PIL import Image
    import io
    images = [Image.open(io.BytesIO(content)).convert("RGB") for content, _ext in pages]
    buf = io.BytesIO()
    if len(images) == 1:
        images[0].save(buf, format="PDF")
    else:
        images[0].save(buf, format="PDF", save_all=True, append_images=images[1:])
    return buf.getvalue()


async def _process_capture_background(doc_id: str, matter_id: str, files: list):
    """
    files: list of {"content": bytes, "filename": str}, in capture order.
    Mirrors _process_document_background's pipeline (extract -> AI
    classify -> chunk -> embed -> legal_source_type/authority_strength)
    but assembles potentially several captured pages into one document's
    text and storage, reusing _extract_attached_document_text -- the same
    per-file extraction dispatch (PDF/DOCX/image-OCR) already used by the
    multi-file Search Vault attach path -- rather than duplicating it.
    """
    combined_text_parts = []
    ocr_confidences = []
    any_ocr_used = False
    image_pages = []        # (content, ext) -- photographed pages, combined into one PDF
    non_image_files = []    # (content, filename, ext) -- already-digital files (e.g. an emailed PDF)

    for f in files:
        filename = f["filename"]
        ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
        try:
            page_text, page_confidence = _extract_attached_document_text(f["content"], filename)
        except ValueError as e:
            print(f"[capture:{doc_id}] page extraction failed for {filename}: {e}")
            continue
        combined_text_parts.append(page_text)
        if ext in ("jpg", "jpeg", "png", "webp"):
            any_ocr_used = True
            if page_confidence is not None:
                ocr_confidences.append(page_confidence)
            image_pages.append((f["content"], ext))
        else:
            non_image_files.append((f["content"], filename, ext))

    text = "\n\n".join(combined_text_parts)
    word_count = len(text.split())
    page_count = len(files)
    ocr_confidence = round(sum(ocr_confidences) / len(ocr_confidences), 1) if ocr_confidences else None

    metadata = {}
    if text:
        try:
            metadata = await asyncio.to_thread(classify_document_sync, text[:2000])
        except Exception:
            metadata = {}

    legal_source_type = classify_firm_document(metadata.get("document_type"))
    authority_strength = authority_strength_for(legal_source_type)

    chunk_count = 0
    if text:
        new_chunks = chunk_text(text, page_count, doc_id, matter_id)
        for c in new_chunks:
            c["chunk_source"] = "firm"
        if new_chunks:
            await asyncio.to_thread(index_chunks_in_chroma, new_chunks, "firm")
            async with _db_pool.acquire() as conn:
                for c in new_chunks:
                    await conn.execute("""
                        INSERT INTO chunks (id, firm_id, document_id, matter_id, chunk_source,
                                           text, chunk_index, page_number, created_at)
                        VALUES ($1,$2,$3,$4,$5,$6,$7,$8,NOW())
                        ON CONFLICT (id) DO NOTHING
                    """,
                    c["id"], FIRM_ID, _uuid_mod.UUID(doc_id),
                    matter_id, "firm", c["text"], c["chunk_index"], c.get("page_number", 1)
                    )
            chunk_count = len(new_chunks)

    raw_date = metadata.get("doc_date")
    doc_date = None
    if raw_date:
        try:
            doc_date = datetime.strptime(raw_date, "%Y-%m-%d").date()
        except Exception:
            doc_date = None

    # Storage: a single already-digital file (email attachment picked from
    # the device) is stored as-is. One or more photographed pages are
    # combined into a single PDF. A batch mixing both is a rare edge case
    # (real usage is "photograph N pages" OR "pick one existing file", not
    # both in the same session) -- the digital file wins as the stored
    # document in that case; the photographed pages' text still made it
    # into the combined text above, just not into the stored file itself.
    r2_key = None
    if R2_ENABLED and _r2_client:
        try:
            if non_image_files:
                store_content, store_filename, store_ext = non_image_files[0]
                content_type = (
                    "application/pdf" if store_ext == "pdf" else
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    if store_ext in ("docx", "doc") else "application/octet-stream"
                )
            elif image_pages:
                store_content = await asyncio.to_thread(_combine_capture_pages_to_pdf, image_pages)
                store_filename = "Captured Document.pdf"
                content_type = "application/pdf"
            else:
                store_content = None
            if store_content:
                r2_key = f"{FIRM_ID}/{matter_id}/{doc_id}/{store_filename}"
                await asyncio.to_thread(
                    _r2_client.put_object, Bucket=R2_BUCKET, Key=r2_key,
                    Body=store_content, ContentType=content_type,
                )
        except Exception as e:
            print(f"[capture] R2 upload failed for {doc_id}: {e}")
            r2_key = None

    async with _db_pool.acquire() as conn:
        needs_review = any_ocr_used and (ocr_confidence is not None) and (ocr_confidence < 80)
        await conn.execute("""
            UPDATE documents SET
                document_type=$1, matter_type=$2, parties=$3,
                doc_date=$4, court=$5, word_count=$6, page_count=$7,
                chunk_count=$8, ocr_used=$9, status='complete', r2_key=$12,
                ocr_confidence=$13, needs_review=$14,
                legal_source_type=$15, authority_strength=$16
            WHERE id=$10 AND firm_id=$11
        """,
        metadata.get("document_type"), metadata.get("matter_type"),
        str(metadata.get("parties", "")) if metadata.get("parties") else None,
        doc_date, metadata.get("court"),
        word_count, page_count, chunk_count, any_ocr_used,
        _uuid_mod.UUID(doc_id), FIRM_ID, r2_key,
        ocr_confidence, needs_review,
        legal_source_type.value, authority_strength.value
        )
        await conn.execute(
            "UPDATE matters SET document_count = document_count + 1, last_activity=NOW() WHERE id=$1 AND firm_id=$2",
            _uuid_mod.UUID(matter_id), FIRM_ID
        )

    if needs_review:
        print(f"[capture] ⚠ {doc_id}: OCR confidence {ocr_confidence}% (below 80%) — flagged for manual review")
    print(f"[capture] processed {doc_id}: {word_count} words, {chunk_count} chunks across {page_count} page(s)")


@app.post("/api/capture", status_code=202)
async def capture_documents(
    background_tasks: BackgroundTasks,
    files: List[UploadFile] = File(...),
    matter_id: Optional[str] = Form(None),
    request: Request = None,
):
    if request:
        user = await get_current_user(request)
        _check_permission(user, "document:upload")
    else:
        user = None

    if not files:
        raise HTTPException(status_code=422, detail="At least one photo or file is required")

    async with _db_pool.acquire() as conn:
        if matter_id:
            matter = await conn.fetchrow(
                "SELECT id FROM matters WHERE id=$1 AND firm_id=$2",
                _uuid_mod.UUID(matter_id), FIRM_ID
            )
            if not matter:
                raise HTTPException(status_code=404, detail="Matter not found")
        else:
            # No matter specified -- the fast/low-friction default path.
            matter = await conn.fetchrow(
                "SELECT id FROM matters WHERE firm_id=$1 AND number='GENERAL'", FIRM_ID
            )
            if not matter:
                raise HTTPException(status_code=500, detail="General/Firm Precedents matter not provisioned")
        resolved_matter_id = str(matter["id"])

    file_payloads = []
    for f in files:
        content = await f.read()
        if content:
            file_payloads.append({"content": content, "filename": f.filename or "capture.jpg"})
    if not file_payloads:
        raise HTTPException(status_code=422, detail="At least one photo or file is required")

    doc_id = str(_uuid_mod.uuid4())
    filename = (
        file_payloads[0]["filename"] if len(file_payloads) == 1
        else f"Captured Document ({len(file_payloads)} pages)"
    )

    async with _db_pool.acquire() as conn:
        row = await conn.fetchrow("""
            INSERT INTO documents (id, matter_id, firm_id, filename, status, uploaded_at, uploaded_by,
                                    document_status, is_capture)
            VALUES ($1,$2,$3,$4,'processing',NOW(),$5,'Final',TRUE) RETURNING *
        """,
        _uuid_mod.UUID(doc_id), _uuid_mod.UUID(resolved_matter_id), FIRM_ID, filename,
        _uuid_mod.UUID(str(user["id"])) if user and user.get("id") else None,
        )

    background_tasks.add_task(_process_capture_background, doc_id, resolved_matter_id, file_payloads)

    return {**_row_to_doc(row), "processing": True,
            "message": f"{len(file_payloads)} page(s) received. Processing in the background."}


@app.get("/api/capture/recent")
async def recent_captures(request: Request):
    user = await get_current_user(request)
    _check_permission(user, "matter:read")
    async with _db_pool.acquire() as conn:
        rows = await conn.fetch("""
            SELECT d.*, m.name AS matter_name FROM documents d
            JOIN matters m ON m.id = d.matter_id
            WHERE d.firm_id=$1 AND d.is_capture=TRUE
            ORDER BY d.uploaded_at DESC LIMIT 50
        """, FIRM_ID)
    return {"documents": [_row_to_doc(r) for r in rows]}


# ── Legal Updates ─────────────────────────────────────────────────────────────

@app.get("/api/legal-updates")
async def list_legal_updates(source_type: Optional[str] = None, request: Request = None):
    if request:
        user = await get_current_user(request)
        _check_permission(user, "matter:read")
    async with _db_pool.acquire() as conn:
        if source_type:
            rows = await conn.fetch(
                "SELECT * FROM legal_updates WHERE firm_id=$1 AND source_type=$2 ORDER BY uploaded_at DESC",
                FIRM_ID, source_type
            )
        else:
            rows = await conn.fetch(
                "SELECT * FROM legal_updates WHERE firm_id=$1 ORDER BY uploaded_at DESC",
                FIRM_ID
            )
    return [_row_to_doc(r) for r in rows]

async def _process_legal_update_background(item_id: str, content: bytes, filename: str, ext: str,
                                            source_type: str, source_name: str, reference: str,
                                            summary: str = ""):
    """Background task: extract, classify, chunk, and index a legal update document."""
    text = ""
    word_count = 0
    page_count = 1
    ocr_used = False
    ocr_confidence = None

    if content:
        try:
            if ext == "pdf":
                text, page_count, ocr_used, ocr_confidence = extract_pdf_text(content)
            elif ext in ("docx", "doc"):
                text = extract_docx_text(content)
            else:
                text = content.decode("utf-8", errors="replace")
            word_count = len(text.split())
        except Exception as e:
            print(f"[legal-update] text extraction failed for {filename}: {e}")
    elif summary:
        # No file attached (e.g. a scraped news article) — fall back to the
        # scraper-provided summary so the item still lands as usable content
        # instead of an empty 'error' row.
        text = summary
        word_count = len(text.split())

    metadata = {}
    if text:
        try:
            metadata = await asyncio.to_thread(classify_document_sync, text[:2000])
        except Exception:
            metadata = {}

    chunk_count = 0
    if text:
        new_chunks = chunk_text(text, page_count, item_id, "legal_updates")
        for c in new_chunks:
            c["chunk_source"] = "legal"
            c["source_type"] = source_type
            c["source_name"] = source_name
            c["reference"] = reference
        if new_chunks:
            await asyncio.to_thread(index_chunks_in_chroma, new_chunks, "legal")
            async with _db_pool.acquire() as conn:
                for c in new_chunks:
                    await conn.execute("""
                        INSERT INTO chunks (id, firm_id, document_id, matter_id, chunk_source,
                                           text, chunk_index, page_number, source_type, source_name, reference, created_at)
                        VALUES ($1,$2,$3,'legal_updates','legal',$4,$5,$6,$7,$8,$9,NOW())
                        ON CONFLICT (id) DO NOTHING
                    """,
                    c["id"], FIRM_ID, _uuid_mod.UUID(item_id),
                    c["text"], c["chunk_index"], c.get("page_number", 1),
                    source_type, source_name, reference
                    )
            chunk_count = len(new_chunks)

    raw_date = metadata.get("doc_date")
    doc_date = None
    if raw_date:
        try:
            doc_date = datetime.strptime(raw_date, "%Y-%m-%d").date()
        except Exception:
            doc_date = None

    async with _db_pool.acquire() as conn:
        needs_review = ocr_used and (ocr_confidence is not None) and (ocr_confidence < 80)
        await conn.execute("""
            UPDATE legal_updates SET
                document_type=$1, matter_type=$2, doc_date=$3, court=$4,
                word_count=$5, chunk_count=$6, ocr_used=$7,
                status=CASE WHEN $5 > 0 THEN 'complete' ELSE 'error' END,
                error_message=CASE WHEN $5 = 0 THEN 'Could not extract text' ELSE NULL END,
                ocr_confidence=$10, needs_review=$11
            WHERE id=$8 AND firm_id=$9
        """,
        metadata.get("document_type"), metadata.get("matter_type"), doc_date,
        metadata.get("court"), word_count, chunk_count, ocr_used,
        _uuid_mod.UUID(item_id), FIRM_ID, ocr_confidence, needs_review
        )
        if needs_review:
            print(f"[legal-update] ⚠ {filename}: OCR confidence {ocr_confidence}% (below 80%) — flagged for manual review")

@app.post("/api/legal-updates/upload", status_code=202)
async def upload_legal_update(
    background_tasks: BackgroundTasks,
    file: Optional[UploadFile] = File(None),
    source_type: str = Form(...),
    source_name: str = Form(""),
    reference: str = Form(""),
    source_url: str = Form(""),
    scraped_at: str = Form(""),
    summary: str = Form(""),
    title: str = Form(""),
    request: Request = None,
):
    if request:
        feed_token_header = request.headers.get("X-Feed-Service-Token", "")
        if not (LEGAL_FEED_SERVICE_TOKEN and feed_token_header == LEGAL_FEED_SERVICE_TOKEN):
            # Not a valid feed-token call (the pusher's normal path) — fall
            # back to requiring a genuine logged-in user, same as before.
            user = await get_current_user(request)
            _check_permission(user, "legal:upload")

    # Items without a PDF (e.g. scraped news articles) arrive with no file —
    # `file` is optional precisely for that case. Prefer the real article/
    # item title as the display name (there's no dedicated `title` column,
    # `filename` is what the UI renders as the heading) — only fall back to
    # the generic "<source>.txt" placeholder if no title was actually sent.
    if file is not None:
        content = await file.read()
        filename = file.filename or "document"
    else:
        content = b""
        filename = title.strip() or f"{source_name or source_type or 'item'}.txt"
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else "bin"
    item_id = str(_uuid_mod.uuid4())

    # Parse scraped_at timestamp if provided by the feed service
    scraped_at_ts = None
    if scraped_at:
        try:
            scraped_at_ts = datetime.fromisoformat(scraped_at.replace("Z", "+00:00"))
        except ValueError:
            scraped_at_ts = None

    legal_source_type = classify_legal_update(source_type, reference)
    authority_strength = authority_strength_for(legal_source_type)

    async with _db_pool.acquire() as conn:
        row = await conn.fetchrow("""
            INSERT INTO legal_updates
                (id, firm_id, filename, source_type, source_name, reference,
                 source_url, scraped_at, status, uploaded_at,
                 legal_source_type, authority_strength)
            VALUES ($1,$2,$3,$4,$5,$6,$7,$8,'processing',NOW(),$9,$10)
            ON CONFLICT (firm_id, source_url) WHERE source_url IS NOT NULL DO NOTHING
            RETURNING *
        """,
        _uuid_mod.UUID(item_id), FIRM_ID, filename, source_type, source_name, reference,
        source_url or None, scraped_at_ts,
        legal_source_type.value, authority_strength.value
        )

    if not row:
        # Already have this URL for this firm — the feed service pushed
        # something it thought was new (its own dedup state may have reset),
        # but we already have it. No need to process or index it again.
        return {"status": "duplicate", "source_url": source_url}

    background_tasks.add_task(
        _process_legal_update_background, item_id, content, filename, ext,
        source_type, source_name, reference, summary
    )

    return {**_row_to_doc(row), "processing": True,
            "message": "Document received. Indexing is running in the background."}

@app.delete("/api/legal-updates/{item_id}")
async def delete_legal_update(item_id: str, request: Request):
    user = await get_current_user(request)
    _check_permission(user, "legal:delete")
    async with _db_pool.acquire() as conn:
        chunk_rows = await conn.fetch(
            "SELECT id FROM chunks WHERE document_id=$1 AND firm_id=$2",
            _uuid_mod.UUID(item_id), FIRM_ID
        )
        chunk_ids = [r["id"] for r in chunk_rows]
        result = await conn.execute(
            "DELETE FROM legal_updates WHERE id=$1 AND firm_id=$2",
            _uuid_mod.UUID(item_id), FIRM_ID
        )
    if result == "DELETE 0":
        raise HTTPException(status_code=404, detail="Not found")
    if chunk_ids:
        await asyncio.to_thread(remove_chunks_from_chroma, chunk_ids, "legal")
    return {"deleted": True}

@app.post("/api/legal-updates/search")
async def search_legal_updates(req: LegalUpdateSearchRequest, request: Request):
    user = await get_current_user(request)
    _check_permission(user, "search")
    async with _db_pool.acquire() as conn:
        chunk_rows = await conn.fetch(
            "SELECT * FROM chunks WHERE firm_id=$1 AND chunk_source='legal'",
            FIRM_ID
        )
    chunks = [dict(r) for r in chunk_rows]
    if not chunks:
        return {"answer": None, "results": [], "message": "No legislation or case law indexed yet."}

    query_words = set(req.query.lower().split())
    scored = []
    async with _db_pool.acquire() as conn:
        items_rows = await conn.fetch("SELECT * FROM legal_updates WHERE firm_id=$1", FIRM_ID)
    items_map = {str(r["id"]): dict(r) for r in items_rows}

    for chunk in chunks:
        if req.source_type and chunk.get("source_type") != req.source_type:
            continue
        item = items_map.get(str(chunk["document_id"]), {})
        chunk_words = set(chunk["text"].lower().split())
        overlap = len(query_words & chunk_words)
        total = len(query_words | chunk_words)
        score = overlap / total if total > 0 else 0
        if req.query.lower() in chunk["text"].lower():
            score += 0.3
        if score > 0:
            scored.append((score, chunk, item))

    scored.sort(key=lambda x: x[0], reverse=True)
    top = scored[:req.limit]
    if not top:
        return {"answer": None, "results": [], "message": f'No relevant results for: "{req.query}"'}

    results = []
    for score, chunk, item in top:
        results.append({
            "chunk_id": chunk["id"], "text": chunk["text"],
            "similarity": round(score, 3),
            "document_id": str(chunk["document_id"]),
            "filename": item.get("filename", "Unknown"),
            "source_type": item.get("source_type"),
            "source_name": item.get("source_name"),
            "reference": item.get("reference"),
            "document_type": item.get("document_type"),
            "doc_date": str(item["doc_date"]) if item.get("doc_date") else None,
            "court": item.get("court"),
            "page_number": chunk.get("page_number"),
            "chunk_index": chunk.get("chunk_index"),
        })
    return {"answer": None, "results": results}

# ── Text extraction helpers ───────────────────────────────────────────────────

def extract_pdf_text(content: bytes):
    """Extract text from PDF. Falls back to OCR for scanned/image-only pages.
    Returns (text, page_count, ocr_used, ocr_confidence).
    ocr_confidence is the average tesseract word-confidence (0-100) across
    any OCR'd pages, or None if no OCR was needed."""
    try:
        import pdfplumber, io
        pages = []
        needs_ocr_pages = []
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for i, page in enumerate(pdf.pages):
                t = page.extract_text()
                if t and t.strip():
                    pages.append(t)
                else:
                    pages.append(None)
                    needs_ocr_pages.append(i)
        total_pages = len(pages)
        ocr_used = bool(needs_ocr_pages)
        page_confidences = []
        if needs_ocr_pages:
            ocr_results = ocr_pdf_pages(content, needs_ocr_pages)
            for i, (ocr_text, ocr_conf) in ocr_results.items():
                pages[i] = ocr_text
                if ocr_conf is not None:
                    page_confidences.append(ocr_conf)
        avg_confidence = round(sum(page_confidences) / len(page_confidences), 1) if page_confidences else None
        final_pages = [p for p in pages if p and p.strip()]
        return "\n\n".join(final_pages), max(1, total_pages), ocr_used, avg_confidence
    except Exception:
        # Previously fell back to a raw UTF-8 decode of the undecodable
        # bytes. For content that isn't actually a valid PDF -- confirmed
        # cause of RTF markup ("\par \pard\plain...") ending up indexed as
        # judgment text for a couple of legacy ZimLII entries that were
        # RTF content saved/served under a .pdf name -- that silently
        # stored binary/markup garbage as if it were the real text.
        # Returning empty text lets the caller's normal "no text
        # extracted" handling deal with it honestly instead.
        return "", 1, False, None

def extract_rtf_text(content: bytes) -> str:
    """Real RTF-to-text extraction, not a plain byte decode. Genuine RTF
    files were previously routed through the same handling as .txt (a raw
    UTF-8 decode), which leaves every control word ("\\pard", "\\rtf1",
    etc.) in the output verbatim -- confirmed as the source of RTF markup
    contaminating a couple of legacy ZimLII entries' indexed text."""
    try:
        from striprtf.striprtf import rtf_to_text
        return rtf_to_text(content.decode("utf-8", errors="replace"))
    except Exception:
        return ""

def ocr_pdf_pages(content: bytes, page_indices: list) -> dict:
    """Returns {page_index: (text, confidence)} for each successfully-OCR'd page."""
    import subprocess as sp
    results = {}
    if not page_indices:
        return results
    with tempfile.TemporaryDirectory() as tmpdir:
        pdf_path = os.path.join(tmpdir, "input.pdf")
        with open(pdf_path, "wb") as f:
            f.write(content)
        for idx in page_indices:
            page_num = idx + 1
            try:
                img_prefix = os.path.join(tmpdir, f"page_{page_num}")
                sp.run(["pdftoppm", "-png", "-r", "200", "-f", str(page_num), "-l", str(page_num), pdf_path, img_prefix],
                       capture_output=True, timeout=60, check=False)
                candidates = [f"{img_prefix}-{page_num}.png", f"{img_prefix}.png", f"{img_prefix}-1.png"]
                img_path = next((c for c in candidates if os.path.exists(c)), None)
                if not img_path:
                    for fn in os.listdir(tmpdir):
                        if fn.startswith(f"page_{page_num}") and fn.endswith(".png"):
                            img_path = os.path.join(tmpdir, fn)
                            break
                if not img_path:
                    continue
                ocr_result = sp.run(["tesseract", img_path, "stdout", "-l", "eng"],
                                    capture_output=True, text=True, timeout=60, check=False)
                text = ocr_result.stdout.strip()
                if text:
                    confidence = _tesseract_confidence(img_path)
                    results[idx] = (text, confidence)
            except Exception:
                continue
    return results

def extract_xlsx_text(content: bytes):
    try:
        import openpyxl, io
        wb = openpyxl.load_workbook(io.BytesIO(content), data_only=True, read_only=True)
        lines = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            lines.append(f"=== Sheet: {sheet_name} ===")
            for row in ws.iter_rows(values_only=True):
                if all(cell is None for cell in row):
                    continue
                row_text = " | ".join(str(cell) if cell is not None else "" for cell in row).strip(" |")
                if row_text:
                    lines.append(row_text)
        wb.close()
        return "\n".join(lines)
    except Exception as e:
        print(f"[extract_xlsx_text] failed: {e}")
        return ""

def _tesseract_confidence(image_path: str) -> Optional[float]:
    """
    Run tesseract in TSV mode to get per-word confidence scores and return
    the average (0-100). Separate from the plain-text extraction call
    (which stays unchanged, to avoid touching working text-reconstruction
    logic) — this is purely for quality signal so low-confidence OCR can be
    flagged for manual review before it ends up misquoted in something like
    a court filing (e.g. "US$120" misread as "US$12O").
    """
    import subprocess as sp
    try:
        result = sp.run(["tesseract", image_path, "stdout", "-l", "eng", "tsv"],
                        capture_output=True, text=True, timeout=60, check=False)
        lines = result.stdout.strip().split("\n")
        if len(lines) < 2:
            return None
        header = lines[0].split("\t")
        try:
            conf_idx = header.index("conf")
            text_idx = header.index("text")
        except ValueError:
            return None
        confidences = []
        for line in lines[1:]:
            fields = line.split("\t")
            if len(fields) <= max(conf_idx, text_idx):
                continue
            try:
                conf = float(fields[conf_idx])
            except ValueError:
                continue
            # -1 marks structural rows (page/block/para/line), not actual
            # recognized words — only real word-level confidence counts.
            if conf < 0:
                continue
            if not fields[text_idx].strip():
                continue
            confidences.append(conf)
        if not confidences:
            return None
        return round(sum(confidences) / len(confidences), 1)
    except Exception:
        return None

def ocr_image_bytes(content: bytes, ext: str) -> tuple:
    """
    OCR a photographed document (jpg/png/webp/etc), returning (text, confidence).
    Same tesseract-based approach already used for ZLR image uploads —
    factored out here so it's used consistently everywhere a raw image gets
    OCR'd, with a confidence score attached so low-quality reads can be
    flagged rather than silently trusted. Real-world use in a Zimbabwean
    practice means people will very often attach a phone photo of a paper
    document, not a clean PDF/docx — without this, those uploads silently
    produce garbage text (a raw UTF-8 decode of binary image bytes) instead
    of an actual error or actual content.
    """
    import shutil
    import subprocess as sp
    if not shutil.which("tesseract"):
        return "", None
    with tempfile.NamedTemporaryFile(suffix=f".{ext}", delete=False) as tmp:
        tmp.write(content)
        tmp_path = tmp.name
    try:
        ocr_result = sp.run(["tesseract", tmp_path, "stdout", "-l", "eng"],
                            capture_output=True, text=True, timeout=60, check=False)
        text = ocr_result.stdout.strip()
        confidence = _tesseract_confidence(tmp_path)
        return text, confidence
    finally:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)

def extract_docx_text(content: bytes):
    if content[:4] == b'\xd0\xcf\x11\xe0':
        try:
            import subprocess, tempfile
            with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
                tmp.write(content)
                tmp_path = tmp.name
            try:
                result = subprocess.run(["antiword", tmp_path], capture_output=True, text=True, timeout=30)
                if result.returncode == 0 and result.stdout.strip():
                    return result.stdout.strip()
            finally:
                if os.path.exists(tmp_path):
                    os.remove(tmp_path)
        except Exception as e:
            print(f"[extract_docx_text] antiword failed: {e}")
        return ""
    try:
        import docx, io
        doc = docx.Document(io.BytesIO(content))
        return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    except Exception:
        return ""

def chunk_text(text: str, page_count: int, doc_id: str, matter_id: str) -> list:
    CHUNK_WORDS = 500
    OVERLAP_WORDS = 50
    words = text.split()
    if not words:
        return []
    chunks = []
    start = 0
    idx = 0
    while start < len(words):
        end = min(start + CHUNK_WORDS, len(words))
        chunk_str = " ".join(words[start:end])
        progress = start / len(words)
        page_num = max(1, round(progress * page_count))
        chunks.append({
            "id": str(_uuid_mod.uuid4()),
            "document_id": doc_id,
            "matter_id": matter_id,
            "text": chunk_str,
            "chunk_index": idx,
            "page_number": page_num,
            # Matches chunks.content_hash's GENERATED column expression
            # exactly (encode(sha256(convert_to(text,'UTF8')),'hex')) —
            # computed here too so it's available immediately for Chroma's
            # metadata without a round-trip read back from Postgres.
            "content_hash": hashlib.sha256(chunk_str.encode("utf-8")).hexdigest(),
        })
        idx += 1
        start = end - OVERLAP_WORDS
        if start >= len(words) - OVERLAP_WORDS:
            break
    return chunks

def index_chunks_in_chroma(chunks: list, collection_type: str = "firm"):
    if not chunks:
        return
    try:
        firm_col, legal_col, zlr_col = get_chroma_collections()
        collection = {"firm": firm_col, "legal": legal_col, "zlr": zlr_col}.get(collection_type, firm_col)
        texts = [c["text"] for c in chunks]
        ids = [c["id"] for c in chunks]
        embeddings = embed_texts(texts)
        metadatas = [{
            "document_id": c["document_id"],
            "matter_id": c.get("matter_id") or "zlr",
            "chunk_index": c["chunk_index"],
            "page_number": c.get("page_number") or 0,
            # Backs reconcile_chroma_index()'s per-chunk drift detection —
            # must match chunks.content_hash exactly (both are sha256 hex of
            # the same UTF-8 text). Computed by the caller (chunk_text()) or
            # passed straight from a Postgres row on reconciliation repairs.
            "content_hash": c.get("content_hash") or "",
        } for c in chunks]
        if collection_type == "firm":
            # Multi-tenancy hardening (Part 3) -- firm_precedents is the one
            # Chroma collection holding actual firm client/matter documents;
            # legal_updates/zlr_index are shared corpora deliberately visible
            # to every firm and stay unscoped. Backfilled onto pre-existing
            # chunks by scripts/backfill_chroma_firm_id.py (this write path
            # alone only covers chunks indexed from here on). Consumed by
            # _semantic_search_firm()'s explicit where filter below.
            for meta in metadatas:
                meta["firm_id"] = str(FIRM_ID)
        # upsert, not add: add() raises/silently fails on an id that's
        # already present (e.g. reconciliation repairing a chunk whose
        # Chroma entry exists but has drifted) — upsert() correctly
        # overwrites it instead.
        collection.upsert(ids=ids, embeddings=embeddings, documents=texts, metadatas=metadatas)
    except Exception as e:
        print(f"[vector_store] failed to index chunks ({collection_type}): {e}")

def remove_chunks_from_chroma(chunk_ids: list, collection_type: str = "firm"):
    if not chunk_ids:
        return
    try:
        firm_col, legal_col, zlr_col = get_chroma_collections()
        collection = {"firm": firm_col, "legal": legal_col, "zlr": zlr_col}.get(collection_type, firm_col)
        collection.delete(ids=chunk_ids)
    except Exception as e:
        print(f"[vector_store] failed to remove chunks ({collection_type}): {e}")

def classify_document_sync(text_preview: str) -> dict:
    try:
        msg = client.messages.create(
            model="claude-haiku-4-5",
            max_tokens=512,
            messages=[{"role": "user", "content": f"""Zimbabwean law firm document classifier.
Return ONLY valid JSON with keys:
document_type, parties, matter_type, doc_date (YYYY-MM-DD or null), court (or null)

document_type options: affidavit, founding_affidavit, opposing_affidavit, replying_affidavit, lease_agreement, heads_of_argument, correspondence, court_order, summons, declaration, plea, notice_of_motion, deed_of_settlement, power_of_attorney, will_and_testament, contract, opinion, other

matter_type options: eviction, estate, employment, commercial_property, commercial_contract, customary_law, matrimonial, company_law, criminal, constitutional, other

Excerpt:
{text_preview[:2000]}

JSON only:"""}]
        )
        raw = msg.content[0].text
        m = re.search(r'\{[\s\S]*\}', raw)
        return json.loads(m.group(0)) if m else {}
    except Exception:
        return {}

# ── Zimbabwe Law Reports Index ─────────────────────────────────────────────────

JURISDICTION_MAP = {
    "ZLR": "Zimbabwe", "ZimLII": "Zimbabwe", "SC": "Zimbabwe",
    "Laws.Africa": "Zimbabwe",
    "SADC": "SADC", "ECOWAS": "ECOWAS",
    "UKSC": "United Kingdom", "UKHL": "United Kingdom",
    "NZCA": "New Zealand", "NZSC": "New Zealand",
    "HCA": "Australia", "FCAFC": "Australia",
    "SCA": "South Africa", "ZACC": "South Africa", "ZASCA": "South Africa",
}
AUTHORITY_WEIGHT = {
    "Zimbabwe": "Binding", "SADC": "Persuasive", "ECOWAS": "Persuasive",
    "United Kingdom": "Persuasive", "New Zealand": "Persuasive",
    "Australia": "Persuasive", "South Africa": "Persuasive", "Other": "Persuasive",
}

def get_jurisdiction(source: str) -> str:
    return JURISDICTION_MAP.get(source, "Other")

def get_authority_weight(source: str) -> str:
    return AUTHORITY_WEIGHT.get(get_jurisdiction(source), "Persuasive")

ZLR_SUBJECT_TAXONOMY = {
    "constitutional": "Constitutional Law",
    "administrative": "Administrative Law & Review",
    "civil procedure": "Civil Procedure",
    "appeal": "Appeals & Review",
    "contract": "Contract Law",
    "property": "Property Law",
    "family": "Family Law & Matrimonial",
    "matrimonial": "Family Law & Matrimonial",
    "customary": "Customary Law & Succession",
    "succession": "Customary Law & Succession",
    "company": "Company & Commercial Law",
    "commercial": "Company & Commercial Law",
    "employment": "Employment & Labour Law",
    "labour": "Employment & Labour Law",
    "delict": "Delict",
    "criminal": "Criminal Law & Procedure",
    "revenue": "Revenue & Tax Law",
    "tax": "Revenue & Tax Law",
    "insolvency": "Insolvency & Sequestration",
    "liquidation": "Insolvency & Sequestration",
    "intellectual property": "Intellectual Property",
    "mining": "Environmental & Mining Law",
    "environmental": "Environmental & Mining Law",
    "human rights": "Human Rights",
    "stock exchange": "Company & Commercial Law",
    "banking": "Company & Commercial Law",
    "land": "Property Law",
    "evidence": "Civil Procedure",
    "prescription": "Civil Procedure",
    "costs": "Civil Procedure",
    "interdict": "Civil Procedure",
    "urgent": "Civil Procedure",
}

def classify_zlr_subject(subject_chains: list) -> str:
    text = " ".join(subject_chains).lower()
    for keyword, category in ZLR_SUBJECT_TAXONOMY.items():
        if keyword in text:
            return category
    return "General"

def parse_zlr_headnote(text: str) -> dict:
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    result = {
        "citation": None, "case_name": None, "court": None,
        "judgment_number": None, "judge": None, "case_type": None,
        "hearing_date": None, "judgment_date": None,
        "subject_chains": [], "taxonomy_category": None,
        "summary": None, "zimlii_url": None,
    }
    # Old-style printed-report citation ("YYYY (N) ZLR NNN") -- this is
    # the judgment's OWN citation only if it appears in the header; the
    # same pattern also matches whenever the judgment quotes a precedent
    # case elsewhere in its reasoning (e.g. "...as held in PTC v Mahachi
    # 1999 (1) ZLR 176 (H)..."), which a full-body scan can't distinguish
    # from the judgment's own citation. Restricted to the header zone
    # (matching how case_name below only looks at lines[:5]) and captures
    # only the matched span, not the whole line -- previously stored the
    # entire line, which produced full sentence fragments as "citation"
    # whenever the header itself ran past word-wrap onto one long line.
    citation_pattern = re.compile(r'\d{4}\s*\(\d+\)\s*ZLR\s*\d+')
    for line in lines[:10]:
        m = citation_pattern.search(line)
        if m:
            result["citation"] = m.group(0).strip()
            break
    # Same false-positive risk as the citation regex above, confirmed live
    # during a citation backfill: a full-body scan can lock onto a
    # different, cross-referenced/related matter's judgment number quoted
    # within the reasoning (e.g. a consolidated or appeal-linked case)
    # instead of this judgment's own. Restricted to the same header zone.
    judgment_number_pattern = re.compile(r'(?:Judgment No\.?\s*)?((?:HH|SC|CCZ|LC|HB|HM|HMT)[-\s]?\d+[-/]\d+)', re.IGNORECASE)
    for line in lines[:10]:
        m = judgment_number_pattern.search(line)
        if m:
            result["judgment_number"] = m.group(1).strip()
            break
    courts = ["High Court, Harare", "High Court, Bulawayo", "High Court, Masvingo",
              "High Court, Mutare", "Supreme Court", "Constitutional Court",
              "Labour Court", "Administrative Court", "Magistrates Court"]
    for line in lines:
        for court in courts:
            if court.lower() in line.lower():
                result["court"] = court
                break
    for line in lines[:5]:
        if re.search(r'\bv\b', line, re.IGNORECASE) and len(line) > 10:
            if not re.search(r'\d{4}.*ZLR', line):
                result["case_name"] = line.strip()
                break
    for line in lines:
        if re.search(r'\b(J|JA|CJ|DCJ|AJA|JP|AJ)\b$', line.strip()):
            result["judge"] = line.strip()
            break
    case_types = ["Chamber application", "Urgent application", "Appeal", "Review",
                  "Action", "Application", "Trial", "Motion"]
    for line in lines:
        for ct in case_types:
            if ct.lower() == line.lower().strip():
                result["case_type"] = ct
                break
    for line in lines:
        if "Date of Judgment" in line or "Judgment date" in line.lower():
            result["judgment_date"] = re.sub(r'Date of Judgment:?\s*', '', line).strip()
        elif re.search(r'\d+\s+(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}', line):
            if not result["hearing_date"]:
                result["hearing_date"] = line.strip()
    chains = []
    for line in lines:
        if (' – ' in line or ' — ' in line or ' - ' in line) and not re.search(r'\d{4}.*ZLR', line):
            chains.append(line.strip())
    for line in lines:
        if re.search(r'[A-Z][a-z]+ (law|procedure|Act|rights) —', line):
            if line not in chains:
                chains.append(line.strip())
    result["subject_chains"] = chains
    result["taxonomy_category"] = classify_zlr_subject(chains)
    for line in lines:
        if (len(line) > 50 and ' – ' not in line and ' — ' not in line
                and not re.search(r'\d{4}.*ZLR', line)
                and not re.search(r'(HH|SC|CCZ)-?\d+', line)
                and line != result.get("case_name")
                and not re.search(r'\b(J|JA|CJ)\b$', line)):
            result["summary"] = line.strip()
            break
    return result

def classify_case_with_ai(text: str, filename: str) -> dict:
    categories = [
        "Constitutional Law", "Administrative Law & Review", "Civil Procedure",
        "Appeals & Review", "Contract Law", "Property Law",
        "Family Law & Matrimonial", "Customary Law & Succession",
        "Company & Commercial Law", "Employment & Labour Law", "Delict",
        "Criminal Law & Procedure", "Revenue & Tax Law",
        "Insolvency & Sequestration", "Intellectual Property",
        "Environmental & Mining Law", "Human Rights"
    ]
    text_lower = text.lower()
    keyword_map = {
        "Revenue & Tax Law": ["zimra", "zimbabwe revenue authority", "income tax act", "value added tax", "vat"],
        "Constitutional Law": ["constitutional court", "declaration of rights", "bill of rights", "constitutionality"],
        "Property Law": ["deeds registry", "deed of transfer", "immoveable property", "rei vindicatio", "eviction"],
        "Family Law & Matrimonial": ["divorce", "matrimonial causes", "custody", "maintenance", "lobola"],
        "Administrative Law & Review": ["judicial review", "administrative court", "minister of public service"],
        "Employment & Labour Law": ["labour court", "labour act", "unfair dismissal", "retrenchment", "nec"],
        "Criminal Law & Procedure": ["accused", "state v", "criminal procedure", "magistrate", "bail"],
        "Company & Commercial Law": ["companies act", "cobe act", "shareholders", "liquidation", "winding up"],
    }
    for category, keywords in keyword_map.items():
        if any(kw in text_lower for kw in keywords):
            return {"taxonomy_category": category, "summary": None, "case_type": None, "subject_chains": []}
    try:
        msg = client.messages.create(
            model="claude-haiku-4-5",
            max_tokens=300,
            messages=[{"role": "user", "content": f"""Classify this Zimbabwe case law excerpt.
Return ONLY JSON: {{"taxonomy_category": "...", "summary": "...", "case_type": "...", "subject_chains": []}}
Categories: {', '.join(categories)}
Text: {text[:1500]}
JSON:"""}]
        )
        raw = msg.content[0].text
        m = re.search(r'\{[\s\S]*\}', raw)
        return json.loads(m.group(0)) if m else {}
    except Exception:
        return {}

@app.get("/api/zlr")
async def list_zlr_entries(category: Optional[str] = None, limit: int = 50, request: Request = None):
    if request:
        user = await get_current_user(request)
        _check_permission(user, "matter:read")
    async with _db_pool.acquire() as conn:
        if category:
            rows = await conn.fetch(
                "SELECT * FROM zlr_entries WHERE firm_id=$1 AND taxonomy_category=$2 ORDER BY uploaded_at DESC LIMIT $3",
                FIRM_ID, category, limit
            )
        else:
            rows = await conn.fetch(
                "SELECT * FROM zlr_entries WHERE firm_id=$1 ORDER BY uploaded_at DESC LIMIT $2",
                FIRM_ID, limit
            )
    result = []
    for r in rows:
        d = dict(r)
        d["id"] = str(d["id"])
        d["firm_id"] = str(d["firm_id"])
        if d.get("uploaded_at"):
            d["uploaded_at"] = d["uploaded_at"].isoformat()
        if isinstance(d.get("subject_chains"), str):
            try:
                d["subject_chains"] = json.loads(d["subject_chains"])
            except Exception:
                d["subject_chains"] = []
        result.append(d)
    return result

@app.get("/api/zlr/categories")
async def zlr_categories(request: Request = None):
    if request:
        user = await get_current_user(request)
        _check_permission(user, "matter:read")
    async with _db_pool.acquire() as conn:
        rows = await conn.fetch("""
            SELECT taxonomy_category, COUNT(*) as count
            FROM zlr_entries WHERE firm_id=$1
            GROUP BY taxonomy_category ORDER BY count DESC
        """, FIRM_ID)
    return [{"category": r["taxonomy_category"], "count": r["count"]} for r in rows]

async def _process_zlr_background(item_id: str, content: bytes, filename: str, ext: str,
                                   source: str, volume_year: Optional[str], zimlii_url: Optional[str],
                                   scraper_meta: Optional[dict] = None):
    """Background task: parse, classify, chunk, and index a ZLR entry.

    `scraper_meta` carries whatever the feed service already extracted
    (case_name, citation, court, judge, judgment_date, summary) so that
    items pushed without a PDF (download failed, or none exists) don't lose
    that metadata — previously it was accepted as form fields but never
    passed through, so it was silently discarded on every push.
    """
    scraper_meta = scraper_meta or {}
    text = ""
    page_count = 1
    ocr_used = False
    ocr_confidence = None

    if content:
        try:
            # A couple of legacy ZimLII entries were confirmed to be RTF
            # content saved/served under a .pdf filename -- the extension
            # alone isn't trustworthy. Sniff the real content type for the
            # PDF case specifically (the one confirmed to occur) rather
            # than trusting `ext`, and route to the real RTF extractor
            # instead of extract_pdf_text() choking on non-PDF bytes.
            effective_ext = ext
            if ext == "pdf" and not content.startswith(b"%PDF-"):
                if content.lstrip()[:10].startswith(b"{\\rtf"):
                    effective_ext = "rtf"

            if effective_ext == "pdf":
                text, page_count, ocr_used, ocr_confidence = extract_pdf_text(content)
            elif effective_ext in ("docx", "doc"):
                text = extract_docx_text(content)
            elif effective_ext == "rtf":
                text = extract_rtf_text(content)
            elif effective_ext == "txt":
                text = content.decode("utf-8", errors="replace")
            elif effective_ext in ("jpg", "jpeg", "png", "webp"):
                text, ocr_confidence = ocr_image_bytes(content, effective_ext)
                ocr_used = True
            else:
                text = content.decode("utf-8", errors="replace")
        except Exception as e:
            print(f"[zlr] text extraction failed for {filename}: {e}")

    if not text and scraper_meta.get("summary"):
        # No PDF attached/downloadable — use the scraper's own summary as the
        # raw text so the entry still lands with usable content instead of
        # being marked 'error' and losing everything the scraper found.
        text = scraper_meta["summary"]

    if not text:
        # zlr_entries has no `status` column (confirmed by grep -- nothing
        # else in this file references one either); the row is simply left
        # as-is, with no raw_text/chunk_count ever populated, rather than
        # raising UndefinedColumnError on this early-exit path.
        return

    parsed = parse_zlr_headnote(text)

    # scraper_meta's citation (JSC's own structured citation/judgment-number
    # field, when present) is more reliable than parse_zlr_headnote()'s
    # regex-based extraction from raw text, which can only recognize the
    # old-style printed-report format ("YYYY (N) ZLR NNN") and has no way
    # to tell "this judgment's own citation" apart from a precedent it
    # happens to quote elsewhere in its reasoning. Prefer it outright
    # rather than only filling in when parsing found nothing -- the old
    # "fill only if empty" logic let a wrong-but-non-empty parsed value
    # silently block a correct scraper-supplied one. Falls further back to
    # judgment_number (the modern HH/SC/CCZ-style reference,
    # parse_zlr_headnote()'s one genuinely reliable identifier field) when
    # neither the scraper nor the old-style regex found anything.
    if scraper_meta.get("citation"):
        parsed["citation"] = scraper_meta["citation"]
    elif not parsed.get("citation") and parsed.get("judgment_number"):
        parsed["citation"] = parsed["judgment_number"]

    # Fill any remaining gaps in what parse_zlr_headnote found from raw
    # text with whatever the scraper already told us directly (most
    # reliable when there's no PDF/full text to parse from).
    for key in ("case_name", "court", "judge", "judgment_date", "summary"):
        if not parsed.get(key) and scraper_meta.get(key):
            parsed[key] = scraper_meta[key]
    if parsed.get("taxonomy_category") == "General" or not parsed.get("summary") or len(parsed.get("subject_chains", [])) == 0:
        ai_meta = await asyncio.to_thread(classify_case_with_ai, text, filename)
        if ai_meta:
            if ai_meta.get("taxonomy_category") and ai_meta["taxonomy_category"] != "General":
                parsed["taxonomy_category"] = ai_meta["taxonomy_category"]
            if ai_meta.get("summary") and not parsed.get("summary"):
                parsed["summary"] = ai_meta["summary"]
            if ai_meta.get("case_type") and not parsed.get("case_type"):
                parsed["case_type"] = ai_meta["case_type"]
            if ai_meta.get("subject_chains") and not parsed.get("subject_chains"):
                parsed["subject_chains"] = ai_meta["subject_chains"]

    jurisdiction = get_jurisdiction(source)
    authority_weight = get_authority_weight(source)
    subject_chains_json = json.dumps(parsed.get("subject_chains", []))

    legal_source_type = classify_zlr_entry(parsed.get("court"))
    authority_strength = authority_strength_for(legal_source_type)

    enriched_text = f"""CASE: {parsed.get('case_name') or ''}
CITATION: {parsed.get('citation') or ''}
JUDGMENT: {parsed.get('judgment_number') or ''}
COURT: {parsed.get('court') or ''}
JUDGE: {parsed.get('judge') or ''}
CATEGORY: {parsed.get('taxonomy_category') or ''}
SUBJECT: {' | '.join(parsed.get('subject_chains', []))}
SUMMARY: {parsed.get('summary') or ''}

FULL TEXT:
{text}"""

    new_chunks = chunk_text(enriched_text, page_count, item_id, "zlr")
    for c in new_chunks:
        c["chunk_source"] = "zlr"
        c["zlr_item_id"] = item_id
        c["citation"] = parsed.get("citation")
        c["case_name"] = parsed.get("case_name")
        c["taxonomy_category"] = parsed.get("taxonomy_category")

    if new_chunks:
        await asyncio.to_thread(index_chunks_in_chroma, new_chunks, "zlr")
        async with _db_pool.acquire() as conn:
            for c in new_chunks:
                await conn.execute("""
                    INSERT INTO chunks (id, firm_id, document_id, matter_id, chunk_source,
                                       text, chunk_index, page_number, zlr_item_id, citation,
                                       case_name, taxonomy_category, created_at)
                    VALUES ($1,$2,$3,'zlr','zlr',$4,$5,$6,$7,$8,$9,$10,NOW())
                    ON CONFLICT (id) DO NOTHING
                """,
                c["id"], FIRM_ID, _uuid_mod.UUID(item_id),
                c["text"], c["chunk_index"], c.get("page_number", 1),
                item_id, c.get("citation"), c.get("case_name"), c.get("taxonomy_category")
                )

    needs_review = ocr_used and (ocr_confidence is not None) and (ocr_confidence < 80)
    async with _db_pool.acquire() as conn:
        await conn.execute("""
            UPDATE zlr_entries SET
                case_name=$1, citation=$2, judgment_number=$3, court=$4, judge=$5,
                case_type=$6, hearing_date=$7, judgment_date=$8,
                subject_chains=$9::jsonb, taxonomy_category=$10, summary=$11,
                raw_text=$12, word_count=$13, chunk_count=$14, ocr_used=$15,
                jurisdiction=$16, authority_weight=$17,
                zimlii_url=COALESCE($18, zimlii_url),
                ocr_confidence=$21, needs_review=$22,
                legal_source_type=$23, authority_strength=$24
            WHERE id=$19 AND firm_id=$20
        """,
        parsed.get("case_name") or filename,
        parsed.get("citation"), parsed.get("judgment_number"),
        parsed.get("court"), parsed.get("judge"), parsed.get("case_type"),
        parsed.get("hearing_date"), parsed.get("judgment_date"),
        subject_chains_json, parsed.get("taxonomy_category", "General"),
        parsed.get("summary"), text, len(text.split()), len(new_chunks), ocr_used,
        jurisdiction, authority_weight, zimlii_url or parsed.get("zimlii_url"),
        _uuid_mod.UUID(item_id), FIRM_ID, ocr_confidence, needs_review,
        legal_source_type.value, authority_strength.value
        )
    if needs_review:
        print(f"[zlr] ⚠ {filename}: OCR confidence {ocr_confidence}% (below 80%) — flagged for manual review")

@app.post("/api/zlr/upload", status_code=202)
async def upload_zlr_document(
    background_tasks: BackgroundTasks,
    file: Optional[UploadFile] = File(None),
    source: str = Form("ZLR"),
    volume_year: Optional[str] = Form(None),
    zimlii_url: Optional[str] = Form(None),
    source_url: str = Form(""),
    case_name: str = Form(""),
    citation: str = Form(""),
    court: str = Form(""),
    judge: str = Form(""),
    judgment_date: str = Form(""),
    summary: str = Form(""),
    scraped_at: str = Form(""),
    request: Request = None,
):
    if request:
        feed_token_header = request.headers.get("X-Feed-Service-Token", "")
        if not (LEGAL_FEED_SERVICE_TOKEN and feed_token_header == LEGAL_FEED_SERVICE_TOKEN):
            user = await get_current_user(request)
            _check_permission(user, "legal:upload")

    # `file` is optional — the feed service pushes metadata-only when a PDF
    # couldn't be downloaded (e.g. ZimLII Cloudflare blocking the download).
    # Prefer the actual case name as the display name when there's no file —
    # same class of gap as legal-updates: without this, entries render with
    # a generic "<source>_entry.txt" heading instead of the case name.
    if file is not None:
        content = await file.read()
        filename = file.filename or "zlr_entry"
    else:
        content = b""
        filename = case_name.strip() or f"{source or 'zlr'}_entry.txt"
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else "bin"
    item_id = str(_uuid_mod.uuid4())
    zimlii_url = zimlii_url or source_url or None

    async with _db_pool.acquire() as conn:
        row = await conn.fetchrow("""
            INSERT INTO zlr_entries (id, firm_id, filename, source, volume_year, zimlii_url,
                                     jurisdiction, authority_weight, uploaded_at)
            VALUES ($1,$2,$3,$4,$5,$6,$7,$8,NOW())
            ON CONFLICT (firm_id, zimlii_url) WHERE zimlii_url IS NOT NULL DO NOTHING
            RETURNING *
        """,
        _uuid_mod.UUID(item_id), FIRM_ID, filename, source, volume_year, zimlii_url,
        get_jurisdiction(source), get_authority_weight(source)
        )

    if not row:
        return {"status": "duplicate", "zimlii_url": zimlii_url}

    scraper_meta = {
        "case_name": case_name, "citation": citation, "court": court,
        "judge": judge, "judgment_date": judgment_date, "summary": summary,
    }
    background_tasks.add_task(
        _process_zlr_background, item_id, content, filename, ext, source, volume_year, zimlii_url,
        scraper_meta
    )

    d = dict(row)
    d["id"] = str(d["id"])
    d["firm_id"] = str(d["firm_id"])
    if d.get("uploaded_at"):
        d["uploaded_at"] = d["uploaded_at"].isoformat()
    return {**d, "processing": True, "message": "ZLR entry received. Parsing and indexing are running in the background."}

@app.post("/api/zlr/bulk-import")
async def bulk_import_zlr(
    file: UploadFile = File(...),
    source: str = Form("ZLR"),
    volume_year: Optional[str] = Form(None),
    request: Request = None,
):
    if request:
        user = await get_current_user(request)
        _check_permission(user, "legal:upload")

    content = await file.read()
    filename = file.filename or "zlr_index"
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else "bin"

    try:
        if ext in ("docx", "doc"):
            text = extract_docx_text(content)
        elif ext == "pdf":
            text, _, _, _ = extract_pdf_text(content)
        elif ext in ("txt",):
            text = content.decode("utf-8", errors="replace")
        else:
            raise HTTPException(status_code=422, detail=f"Unsupported file type: {ext}")
    except Exception as e:
        raise HTTPException(status_code=422, detail=f"Could not extract text: {e}")

    if not text:
        raise HTTPException(status_code=422, detail="No text extracted from document")

    parsed_cases = await asyncio.to_thread(parse_zlr_subject_index, text, source, volume_year)
    if not parsed_cases:
        raise HTTPException(status_code=422, detail="No cases could be parsed from this document.")

    imported = 0
    all_chunks = []
    async with _db_pool.acquire() as conn:
        for case in parsed_cases:
            item_id = str(_uuid_mod.uuid4())
            case["id"] = item_id
            raw_text = f"""CASE: {case['case_name']}
JUDGMENT: {case['judgment_number']}
COURT: {case['court']}
JUDGE: {case.get('judge') or ''}
DATE: {case.get('judgment_date') or ''}
CATEGORY: {case['taxonomy_category']}
SUBJECT: {' | '.join(case['subject_chains'])}
SUMMARY: {case.get('summary') or ''}"""

            subject_chains_json = json.dumps(case.get("subject_chains", []))
            case_legal_source_type = classify_zlr_entry(case.get("court"))
            case_authority_strength = authority_strength_for(case_legal_source_type)
            await conn.execute("""
                INSERT INTO zlr_entries (id, firm_id, filename, source, volume_year,
                    jurisdiction, authority_weight, case_name, judgment_number, court, judge,
                    judgment_date, subject_chains, taxonomy_category, summary, raw_text,
                    word_count, chunk_count, uploaded_at, legal_source_type, authority_strength)
                VALUES ($1,$2,$3,$4,$5,$6,$7,$8,$9,$10,$11,$12,$13::jsonb,$14,$15,$16,$17,0,NOW(),$18,$19)
            """,
            _uuid_mod.UUID(item_id), FIRM_ID,
            f"{case['case_name']} [{case['judgment_number']}]",
            source, volume_year,
            case.get("jurisdiction", get_jurisdiction(source)),
            case.get("authority_weight", get_authority_weight(source)),
            case.get("case_name"), case.get("judgment_number"),
            case.get("court"), case.get("judge"),
            case.get("judgment_date"), subject_chains_json,
            case.get("taxonomy_category", "General"), case.get("summary"),
            raw_text, len((case.get("summary") or "").split()),
            case_legal_source_type.value, case_authority_strength.value
            )

            new_chunks = chunk_text(raw_text, 1, item_id, "zlr")
            for c in new_chunks:
                c["chunk_source"] = "zlr"
                c["zlr_item_id"] = item_id
                c["citation"] = case.get("citation")
                c["case_name"] = case.get("case_name")
                c["taxonomy_category"] = case.get("taxonomy_category")
            all_chunks.extend(new_chunks)

            await conn.execute(
                "UPDATE zlr_entries SET chunk_count=$1 WHERE id=$2",
                len(new_chunks), _uuid_mod.UUID(item_id)
            )
            imported += 1

    if all_chunks:
        await asyncio.to_thread(index_chunks_in_chroma, all_chunks, "zlr")
        async with _db_pool.acquire() as conn:
            for c in all_chunks:
                await conn.execute("""
                    INSERT INTO chunks (id, firm_id, document_id, matter_id, chunk_source,
                                       text, chunk_index, page_number, zlr_item_id, citation,
                                       case_name, taxonomy_category, created_at)
                    VALUES ($1,$2,$3,'zlr','zlr',$4,$5,$6,$7,$8,$9,$10,NOW())
                    ON CONFLICT (id) DO NOTHING
                """,
                c["id"], FIRM_ID, _uuid_mod.UUID(c["document_id"]),
                c["text"], c["chunk_index"], c.get("page_number", 1),
                c.get("zlr_item_id"), c.get("citation"), c.get("case_name"), c.get("taxonomy_category")
                )

    from collections import Counter
    categories = Counter(c["taxonomy_category"] for c in parsed_cases)
    return {"imported": imported, "total_parsed": len(parsed_cases), "categories": dict(categories), "source": source, "volume_year": volume_year}

@app.delete("/api/zlr/{item_id}")
async def delete_zlr_entry(item_id: str, request: Request):
    user = await get_current_user(request)
    _check_permission(user, "legal:delete")
    async with _db_pool.acquire() as conn:
        chunk_rows = await conn.fetch(
            "SELECT id FROM chunks WHERE document_id=$1 AND firm_id=$2",
            _uuid_mod.UUID(item_id), FIRM_ID
        )
        chunk_ids = [r["id"] for r in chunk_rows]
        result = await conn.execute(
            "DELETE FROM zlr_entries WHERE id=$1 AND firm_id=$2",
            _uuid_mod.UUID(item_id), FIRM_ID
        )
    if result == "DELETE 0":
        raise HTTPException(status_code=404, detail="Not found")
    if chunk_ids:
        await asyncio.to_thread(remove_chunks_from_chroma, chunk_ids, "zlr")
    return {"deleted": True}

@app.post("/api/zlr/search")
async def search_zlr(req: LegalUpdateSearchRequest, request: Request):
    user = await get_current_user(request)
    _check_permission(user, "search")
    _require_retrieval_ready()
    async with _db_pool.acquire() as conn:
        chunk_rows = await conn.fetch(
            "SELECT * FROM chunks WHERE firm_id=$1 AND chunk_source='zlr'",
            FIRM_ID
        )
    zlr_chunks = [dict(r) for r in chunk_rows]
    if not zlr_chunks:
        return {"results": [], "message": "No ZLR entries indexed yet."}
    results = await asyncio.to_thread(_zlr_semantic_search, zlr_chunks, req.query, req.source_type, req.limit)
    return {"results": results, "count": len(results)}

# Chunks from chunk_text() run ~2500 chars (500 words); a ZLR judgment's
# chunk almost always opens with case-caption/coram boilerplate (case name,
# court, judge, hearing dates) before reaching the actual facts or holding.
# A short excerpt taken from chunk start therefore tends to show only that
# boilerplate — enough for the case NAME to match the query, but not enough
# substantive content for a drafting model to recognize the case as
# genuinely on point. Widened well past the old 400-char cutoff, and
# centered on whichever sentence in the chunk most overlaps the query's own
# words, rather than always the start.
ZLR_EXCERPT_CHARS = 1400

def _best_excerpt_window(text: str, query: str, window_chars: int = ZLR_EXCERPT_CHARS) -> str:
    """
    Picks a window_chars-sized slice of text by sliding a probe across it
    and scoring each position by query-word overlap — a simple,
    dependency-free stand-in for "where in this chunk did the actual match
    come from," since ChromaDB gives us a similarity score for the whole
    chunk but not which part of it drove that score.

    The first ~300 chars are deliberately excluded from candidate
    positions. chunk_text() rejoins words with single spaces, so no
    newlines survive chunking — a judgment's caption/coram block (case
    name, court, judge, dates) then has no punctuation boundary separating
    it from the first real sentence, and it legitimately scores well
    against any query that names the same parties as the case caption
    (exactly the Kombayi case: "MINISTER OF LOCAL GOVERNMENT" in the
    caption vs. a query about "Ministry of Local Government" overlap on
    3+ words). Left unexcluded, the excerpt would anchor right back on the
    caption it was supposed to get past. A judgment's caption is reliably
    near chunk start (especially chunk 0), so skipping straight past it is
    a safe bet without needing to actually parse judgment structure.

    Falls back to text[:window_chars] if the text already fits, the query
    has no usable words, or nothing beyond the skipped lead-in scores
    above zero — same behavior as the old flat slice in that case, just at
    the new, wider length.
    """
    if len(text) <= window_chars:
        return text

    query_words = set(re.findall(r"[a-z]{3,}", query.lower()))
    if not query_words:
        return text[:window_chars]

    probe_chars = min(300, window_chars // 3)
    stride = 100
    skip_chars = min(300, max(0, len(text) - probe_chars))

    best_pos, best_score = None, 0
    for pos in range(skip_chars, len(text) - probe_chars + 1, stride):
        probe_words = set(re.findall(r"[a-z]{3,}", text[pos:pos + probe_chars].lower()))
        score = len(query_words & probe_words)
        if score > best_score:
            best_score, best_pos = score, pos

    if best_pos is None:
        return text[:window_chars]

    # A little lead-in before the match so the excerpt doesn't start
    # mid-sentence, without giving up much of the window to it.
    lead_in = min(200, window_chars // 4)
    window_start = max(0, min(best_pos - lead_in, len(text) - window_chars))
    if window_start > 0:
        # Snap forward to the next word boundary rather than starting
        # mid-word — purely cosmetic, doesn't change which content is in
        # the window by more than a few characters.
        next_space = text.find(" ", window_start)
        if 0 <= next_space < window_start + 50:
            window_start = next_space + 1
    return text[window_start:window_start + window_chars]

def _zlr_semantic_search(zlr_chunks: list, query: str, category_filter: Optional[str], limit: int) -> list:
    results = []
    try:
        _, _, zlr_col = get_chroma_collections()
        if zlr_col.count() > 0:
            query_vec = embed_texts([query])[0]
            if hasattr(query_vec[0], "__len__"): query_vec = query_vec[0]
            res = zlr_col.query(query_embeddings=[query_vec], n_results=min(limit * 3, zlr_col.count()))
            ids = res["ids"][0] if res["ids"] else []
            distances = res["distances"][0] if res["distances"] else []
            chunk_by_id = {c["id"]: c for c in zlr_chunks}
            seen_items = set()
            for cid, dist in zip(ids, distances):
                chunk = chunk_by_id.get(cid)
                if not chunk:
                    continue
                item_id = str(chunk["document_id"])
                if item_id in seen_items:
                    continue
                if category_filter and chunk.get("taxonomy_category") != category_filter:
                    continue
                seen_items.add(item_id)
                similarity = max(0.0, 1.0 - dist)
                results.append({
                    "item_id": item_id,
                    "similarity": round(similarity, 3),
                    "case_name": chunk.get("case_name"),
                    "citation": chunk.get("citation"),
                    "taxonomy_category": chunk.get("taxonomy_category"),
                    "relevant_excerpt": _best_excerpt_window(chunk["text"], query),
                    "legal_source_type": chunk.get("legal_source_type"),
                    "authority_strength": chunk.get("authority_strength"),
                })
                if len(results) >= limit:
                    break
    except Exception as e:
        print(f"[zlr_search] semantic search failed, using keyword fallback: {e}")
        query_words = set(query.lower().split())
        scored = []
        for chunk in zlr_chunks:
            if category_filter and chunk.get("taxonomy_category") != category_filter:
                continue
            score = len(query_words & set(chunk["text"].lower().split())) / max(len(query_words), 1)
            if score > 0:
                scored.append((score, chunk))
        scored.sort(key=lambda x: x[0], reverse=True)
        for score, chunk in scored[:limit]:
            results.append({
                "item_id": str(chunk["document_id"]),
                "similarity": round(score, 3),
                "case_name": chunk.get("case_name"),
                "citation": chunk.get("citation"),
                "taxonomy_category": chunk.get("taxonomy_category"),
                "relevant_excerpt": _best_excerpt_window(chunk["text"], query),
                "legal_source_type": chunk.get("legal_source_type"),
                "authority_strength": chunk.get("authority_strength"),
            })
    return results

def parse_zlr_subject_index(text: str, source: str, volume_year: Optional[str]) -> list:
    """Parse a ZLR 'Cases Decided' subject index into individual case records."""
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    cases = []
    current_subject_chains = []
    i = 0
    while i < len(lines):
        line = lines[i]
        judgment_match = re.search(r'((?:HH|SC|CCZ|LC|HB|HM|HMT)[-\s]?\d+[-/]\d+)', line, re.IGNORECASE)
        if judgment_match:
            judgment_number = judgment_match.group(1).strip()
            case_name = None
            court = None
            judge = None
            judgment_date = None
            if i > 0:
                prev = lines[i-1]
                if re.search(r'\bv\b', prev, re.IGNORECASE) and len(prev) > 10:
                    case_name = prev
            court_keywords = {
                "HH": "High Court, Harare", "HB": "High Court, Bulawayo",
                "HM": "High Court, Masvingo", "HMT": "High Court, Mutare",
                "SC": "Supreme Court", "CCZ": "Constitutional Court", "LC": "Labour Court",
            }
            prefix = judgment_match.group(1)[:2].upper()
            court = court_keywords.get(prefix, "High Court, Harare")
            for j in range(i, min(i+5, len(lines))):
                if re.search(r'\b(J|JA|CJ|DCJ|AJA)\b$', lines[j].strip()):
                    judge = lines[j].strip()
                    break
            for j in range(i, min(i+5, len(lines))):
                if re.search(r'\d+\s+(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}', lines[j]):
                    judgment_date = lines[j].strip()
                    break
            summary_parts = []
            j = i + 1
            while j < min(i + 10, len(lines)):
                next_line = lines[j]
                if re.search(r'((?:HH|SC|CCZ|LC|HB|HM|HMT)[-\s]?\d+[-/]\d+)', next_line, re.IGNORECASE):
                    break
                if next_line.startswith('See below') or next_line.startswith('See above'):
                    j += 1
                    continue
                if len(next_line) > 40:
                    summary_parts.append(next_line)
                j += 1
            taxonomy = classify_zlr_subject(current_subject_chains)
            cases.append({
                'case_name': case_name or f"Case {judgment_number}",
                'judgment_number': judgment_number,
                'court': court,
                'judge': judge,
                'judgment_date': judgment_date,
                'subject_chains': list(current_subject_chains),
                'taxonomy_category': taxonomy,
                'summary': ' '.join(summary_parts)[:600] if summary_parts else None,
                'citation': None,
                'source': source,
                'volume_year': volume_year,
                'jurisdiction': get_jurisdiction(source),
                'authority_weight': get_authority_weight(source),
            })
            current_subject_chains = []
        elif ' – ' in line or ' — ' in line:
            if not re.search(r'\d{4}.*ZLR', line):
                current_subject_chains.append(line)
        i += 1
    return cases

# ── Search ────────────────────────────────────────────────────────────────────

async def _run_plain_search_job(job_id: str, req: SearchRequest, user: dict):
    """
    Background half of POST /api/search -- same fire-and-poll shape and
    same reason as _run_document_search_job below: full multi-source
    retrieval plus Sonnet synthesis can run long enough on a genuinely
    broad query to exceed Cloudflare's ~100s edge timeout if held behind
    one synchronous request. Confirmed directly against a real broad
    3-issue query in production: 123.5s wall-clock, well past that
    ceiling, after the max_tokens fix let the model actually finish
    instead of cutting off early. The endpoint returns a job_id
    immediately; the frontend polls /api/search/status/{job_id}.
    """
    _search_jobs[job_id]["status"] = JobStatus.RUNNING
    try:
        # Load chunks from DB for keyword fallback
        async with _db_pool.acquire() as conn:
            firm_chunk_rows = await conn.fetch(
                """
                SELECT c.*, d.filename AS document_filename, d.document_type, d.court,
                       d.matter_type, d.legal_source_type, d.authority_strength,
                       d.document_status, d.provenance_document_type,
                       m.matter_number, m.name AS matter_name, m.client_id AS matter_client_id,
                       cl.full_name AS client_name
                FROM chunks c
                JOIN documents d ON d.id = c.document_id
                LEFT JOIN matters m ON m.id = d.matter_id
                LEFT JOIN clients cl ON cl.id = m.client_id
                WHERE c.firm_id=$1 AND c.chunk_source='firm'
                """,
                FIRM_ID
            )
            legal_chunk_rows = await conn.fetch(
                """
                SELECT c.*, lu.legal_source_type, lu.authority_strength
                FROM chunks c
                LEFT JOIN legal_updates lu ON lu.id = c.document_id
                WHERE c.firm_id=$1 AND c.chunk_source='legal'
                """,
                FIRM_ID
            )
            zlr_chunk_rows = await conn.fetch(
                """
                SELECT c.*, z.legal_source_type, z.authority_strength
                FROM chunks c
                LEFT JOIN zlr_entries z ON z.id = c.document_id
                WHERE c.firm_id=$1 AND c.chunk_source='zlr'
                """,
                FIRM_ID
            )

        firm_chunks = [dict(r) for r in firm_chunk_rows]
        legal_chunks = [dict(r) for r in legal_chunk_rows]
        zlr_chunks_list = [dict(r) for r in zlr_chunk_rows]

        results = await asyncio.to_thread(_semantic_search_firm, req, firm_chunks)
        legal_results = []
        if req.include_legal_updates:
            legal_results = await asyncio.to_thread(_semantic_search_legal, req, legal_chunks)

        zlr_results = []
        if zlr_chunks_list:
            raw_zlr = await asyncio.to_thread(_zlr_semantic_search, zlr_chunks_list, req.query, None, 3)
            for r in raw_zlr:
                zlr_results.append({
                    "result_source": "zlr",
                    "chunk_id": r.get("item_id"),
                    "text": r.get("relevant_excerpt", ""),
                    "similarity": r.get("similarity", 0),
                    "document_id": r.get("item_id"),
                    "filename": r.get("case_name") or r.get("citation") or "ZLR Entry",
                    "citation": r.get("citation"),
                    "taxonomy_category": r.get("taxonomy_category"),
                    "summary": r.get("summary"),
                    "legal_source_type": r.get("legal_source_type"),
                    "authority_strength": r.get("authority_strength"),
                })

        deadline_info = try_compute_deadline(req.query, legal_results, zlr_results)

        all_results = results + legal_results + zlr_results
        for r in all_results:
            r["display_label"] = display_label(r)
        if not all_results:
            _search_jobs[job_id]["result"] = {
                "answer": None, "results": [],
                "message": f'No relevant documents found for: "{req.query}"',
            }
            _search_jobs[job_id]["status"] = JobStatus.COMPLETE
            return

        # Authority-first reranking (backend/authority_ranker.py) — additive:
        # existing `results`/grounding fields are untouched, this is a second,
        # independent view of the same candidates for clients that want it.
        authority_ranking = rerank(all_results, req.query)

        grounding = compute_grounding(results, legal_results, zlr_results)

        research_map = None
        if not grounding["sources_sufficient"]:
            context_for_agent = format_context(results, legal_results, zlr_results)
            research_map = await asyncio.to_thread(run_legal_research_agent, req.query, context_for_agent)

        synthesis_context = format_context(results[:5], legal_results[:3], zlr_results[:3])

        firm = await get_firm_identity()
        answer = await asyncio.to_thread(
            synthesise_answer_sync, req.query, results[:5], legal_results[:3], zlr_results[:3],
            deadline_info=deadline_info, research_map=research_map,
            firm_name=firm["name"], firm_city=firm["city"],
        )

        answer, qc_log = verify_citations(answer, synthesis_context)
        answer, inline_qc_log = verify_inline_case_citations(answer, synthesis_context)
        qc_log = qc_log + inline_qc_log

        answer, confidence_qc_log = enforce_confidence_consistency(answer)
        qc_log = qc_log + confidence_qc_log

        answer = apply_confidence_safeguard(answer, grounding)

        research_agent_status = (
            "success" if research_map and "error" not in research_map
            else "failed" if research_map
            else "not_triggered"
        )

        async with _db_pool.acquire() as conn:
            await conn.execute(
                """
                INSERT INTO audit_logs (firm_id, user_id, actor_name, actor_role, action, target_type, target_id, details)
                VALUES ($1, $2, $3, $4, 'SEARCH', 'QUERY', NULL, $5)
                """,
                FIRM_ID,
                _uuid_mod.UUID(str(user["id"])) if user.get("id") else None,
                user.get("display_name", "Unknown"),
                user.get("role", "unknown"),
                json.dumps({
                    "query": req.query,
                    "max_similarity_score": grounding["max_similarity_score"],
                    "sources_sufficient": grounding["sources_sufficient"],
                    "source_tier_breakdown": grounding["source_tier_breakdown"],
                    "research_agent_status": research_agent_status,
                    "research_agent_gaps": research_map.get("gaps", []) if research_map else [],
                    "qc_downgrades": qc_log,
                    "authority_confidence": authority_ranking["confidence"],
                    "authority_excluded_count": authority_ranking["excluded_count"],
                }),
            )

        _search_jobs[job_id]["result"] = {
            "answer": answer, "results": all_results, **grounding,
            "authority_ranking": authority_ranking,
        }
        _search_jobs[job_id]["status"] = JobStatus.COMPLETE
    except Exception as e:
        print(f"[search_job:{job_id}] FAILED: {e}")
        _search_jobs[job_id]["error"] = str(e)
        _search_jobs[job_id]["status"] = JobStatus.FAILED

@app.post("/api/search", status_code=202)
async def search_documents(req: SearchRequest, request: Request):
    """
    Search Vault's main query path — full multi-source retrieval (firm
    precedent, indexed legislation/case law, ZLR judgments) plus Sonnet
    synthesis.

    Runs as a fire-and-poll background job rather than a single
    synchronous request/response, same pattern as /api/search/document
    below and for the same reason: the pipeline can run long enough on a
    genuinely broad query to exceed Cloudflare's ~100s edge timeout.
    Returns a job_id immediately; poll /api/search/status/{job_id} for
    progress and the eventual result.
    """
    user = await get_current_user(request)
    _check_permission(user, "search")
    _require_retrieval_ready()

    now = datetime.utcnow()
    for jid, job in list(_search_jobs.items()):
        if now - datetime.fromisoformat(job["created_at"]) > _SEARCH_JOB_MAX_AGE:
            del _search_jobs[jid]

    job_id = str(_uuid_mod.uuid4())

    # Capture only the minimum authenticated user context required by the
    # pipeline, same reasoning as job_user in search_with_document below —
    # the Request object is tied to this request/response cycle and must
    # not be relied on after it ends.
    job_user = {
        "id": user.get("id"), "firm_id": user.get("firm_id"),
        "display_name": user.get("display_name"), "role": user.get("role"),
    }

    _search_jobs[job_id] = {
        "status": JobStatus.PENDING,
        "result": None,
        "error": None,
        "firm_id": str(user.get("firm_id") or FIRM_ID),
        "created_at": now.isoformat(),
    }

    asyncio.create_task(_run_plain_search_job(job_id, req, job_user))

    return {"job_id": job_id, "status": "pending"}

@app.get("/api/search/status/{job_id}")
async def get_plain_search_job_status(job_id: str, request: Request):
    user = await get_current_user(request)
    _check_permission(user, "search")

    job = _search_jobs.get(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    if job["firm_id"] != str(user.get("firm_id") or FIRM_ID):
        raise HTTPException(status_code=403, detail="Not authorized")

    return {
        "job_id": job_id,
        "status": job["status"],
        "result": job["result"],
        "error": job["error"],
    }

# Cap on attached-document text sent to the model — generous enough for a
# full lease agreement, contract, or affidavit (roughly 40k chars is well
# within Sonnet's context window even alongside retrieved chunks), while
# still bounding cost/latency for anything unusually long.
MAX_ATTACHED_DOC_CHARS = 40_000

# ── Search: fire-and-poll job pattern ───────────────────────────────────────────
# Shared by both Search Vault paths — plain query (/api/search) and
# document-attached query (/api/search/document). Either pipeline (multi-
# source retrieval, optional OCR, Sonnet synthesis) can run long enough on
# a genuinely broad query to exceed Cloudflare's ~100s edge timeout when
# held behind one synchronous request — confirmed directly in production
# for the plain-query path (123.5s on a real broad 3-issue query, after
# the max_tokens fix let it actually finish instead of cutting off early).
# Each endpoint returns a job_id immediately and runs its pipeline in a
# background asyncio task; the frontend polls that job's own status
# endpoint (/api/search/status/{job_id} or /api/search/document/status/{job_id})
# for the result. In-process dict is an MVP job store, shared across both
# paths by job_id — fine for a single instance, not a durable queue.
class JobStatus(str, Enum):
    PENDING = "pending"
    RUNNING = "running"
    COMPLETE = "complete"
    FAILED = "failed"

_search_jobs: dict[str, dict] = {}
_SEARCH_JOB_MAX_AGE = timedelta(minutes=30)

def _extract_attached_document_text(content: bytes, filename: str) -> tuple:
    """
    Extracts text from one ad-hoc attached document, exactly the same way
    for every file in a multi-file attach — one file's extraction failure
    (corrupt PDF, unreadable photo) fails the whole query rather than
    silently dropping that document from consideration, same fail-fast
    behavior as the original single-file path. Returns (doc_text, ocr_confidence).
    """
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    ocr_confidence = None
    try:
        if ext == "pdf":
            doc_text, _, _, ocr_confidence = extract_pdf_text(content)
        elif ext in ("docx", "doc"):
            doc_text = extract_docx_text(content)
        elif ext in ("jpg", "jpeg", "png", "webp"):
            doc_text, ocr_confidence = ocr_image_bytes(content, ext)
            if not doc_text:
                raise ValueError(
                    f'Could not read text from "{filename}". Make sure the photo is clear, '
                    "well-lit, and the document fills most of the frame."
                )
        else:
            doc_text = content.decode("utf-8", errors="replace")
    except ValueError:
        raise
    except Exception as e:
        raise ValueError(f'Could not read "{filename}": {e}')

    if not doc_text or not doc_text.strip():
        raise ValueError(f'No readable text found in "{filename}".')

    return doc_text, ocr_confidence

def _combine_attached_documents(docs: list) -> tuple:
    """
    docs: list of {"filename", "text"} dicts (already extracted/truncated).
    A single document is passed through unlabeled — exactly the prompt
    shape the original single-file path always sent, so that behavior is
    unchanged. Two or more get a clear "=== DOCUMENT: filename ===" header
    each, so the model can attribute a fact to the specific document it
    came from instead of an undifferentiated blob. Returns
    (combined_text, combined_name) for synthesise_answer_sync.
    """
    if len(docs) == 1:
        return docs[0]["text"], docs[0]["filename"]
    combined_text = "\n\n".join(f"=== DOCUMENT: {d['filename']} ===\n{d['text']}" for d in docs)
    combined_name = ", ".join(d["filename"] for d in docs)
    return combined_text, combined_name

async def _run_document_search_job(
    job_id: str,
    files: list,
    query: str,
    user: dict,
    matter_id: Optional[str],
    include_legal_updates: bool,
    limit: int,
):
    _search_jobs[job_id]["status"] = JobStatus.RUNNING
    print(f"[search_job:{job_id}] STARTED")
    try:
        docs = []
        for f in files:
            doc_text, ocr_confidence = _extract_attached_document_text(f["content"], f["filename"])
            truncated = len(doc_text) > MAX_ATTACHED_DOC_CHARS
            if truncated:
                doc_text = doc_text[:MAX_ATTACHED_DOC_CHARS]
            docs.append({
                "filename": f["filename"], "text": doc_text, "truncated": truncated,
                "char_count": len(doc_text), "ocr_confidence": ocr_confidence,
            })

        print(f"[search_job:{job_id}] OCR_COMPLETE")

        doc_text, filename = _combine_attached_documents(docs)

        req = SearchRequest(
            query=query, matter_id=matter_id, limit=limit,
            include_legal_updates=include_legal_updates,
        )

        async with _db_pool.acquire() as conn:
            firm_chunk_rows = await conn.fetch(
                """
                SELECT c.*, d.filename AS document_filename, d.document_type, d.court,
                       d.legal_source_type, d.authority_strength,
                       d.document_status, d.provenance_document_type,
                       m.matter_number, m.name AS matter_name, m.client_id AS matter_client_id,
                       cl.full_name AS client_name
                FROM chunks c
                JOIN documents d ON d.id = c.document_id
                LEFT JOIN matters m ON m.id = d.matter_id
                LEFT JOIN clients cl ON cl.id = m.client_id
                WHERE c.firm_id=$1 AND c.chunk_source='firm'
                """,
                FIRM_ID
            )
            legal_chunk_rows = await conn.fetch(
                """
                SELECT c.*, lu.legal_source_type, lu.authority_strength
                FROM chunks c
                LEFT JOIN legal_updates lu ON lu.id = c.document_id
                WHERE c.firm_id=$1 AND c.chunk_source='legal'
                """,
                FIRM_ID
            )
            zlr_chunk_rows = await conn.fetch(
                """
                SELECT c.*, z.legal_source_type, z.authority_strength
                FROM chunks c
                LEFT JOIN zlr_entries z ON z.id = c.document_id
                WHERE c.firm_id=$1 AND c.chunk_source='zlr'
                """,
                FIRM_ID
            )
        firm_chunks = [dict(r) for r in firm_chunk_rows]
        legal_chunks = [dict(r) for r in legal_chunk_rows]
        zlr_chunks_list = [dict(r) for r in zlr_chunk_rows]

        results = await asyncio.to_thread(_semantic_search_firm, req, firm_chunks)
        legal_results = []
        if include_legal_updates:
            legal_results = await asyncio.to_thread(_semantic_search_legal, req, legal_chunks)
        zlr_results = []
        if zlr_chunks_list:
            raw_zlr = await asyncio.to_thread(_zlr_semantic_search, zlr_chunks_list, query, None, 3)
            for r in raw_zlr:
                zlr_results.append({
                    "result_source": "zlr", "chunk_id": r.get("item_id"),
                    "text": r.get("relevant_excerpt", ""), "similarity": r.get("similarity", 0),
                    "document_id": r.get("item_id"),
                    "filename": r.get("case_name") or r.get("citation") or "ZLR Entry",
                    "citation": r.get("citation"), "taxonomy_category": r.get("taxonomy_category"),
                    "summary": r.get("summary"),
                    "legal_source_type": r.get("legal_source_type"),
                    "authority_strength": r.get("authority_strength"),
                })

        all_results = results + legal_results + zlr_results
        print(f"[search_job:{job_id}] RETRIEVAL_COMPLETE")

        # This endpoint's existing pipeline does not call run_legal_research_agent —
        # that only exists on /api/search today. Logged as skipped rather than
        # silently omitted so the lifecycle trace is honest about what ran.
        print(f"[search_job:{job_id}] RESEARCH_AGENT_SKIPPED")

        firm = await get_firm_identity()
        answer = await asyncio.to_thread(
            synthesise_answer_sync, query, results[:5], legal_results[:3], zlr_results[:3], doc_text, filename,
            firm_name=firm["name"], firm_city=firm["city"],
        )
        print(f"[search_job:{job_id}] SYNTHESIS_COMPLETE")

        grounding = compute_grounding(results, legal_results, zlr_results, has_attached_doc=True)
        answer = apply_confidence_safeguard(answer, grounding)

        result = {
            "answer": answer,
            "results": all_results,
            "attached_documents": [
                {
                    "filename": d["filename"], "truncated": d["truncated"], "char_count": d["char_count"],
                    "ocr_confidence": d["ocr_confidence"],
                    "low_confidence": d["ocr_confidence"] is not None and d["ocr_confidence"] < 80,
                }
                for d in docs
            ],
            **grounding,
        }

        # TEMP diagnostic logging — remove once the stale-frontend vs.
        # result-shape-mismatch question is settled.
        print(f"[search_job:{job_id}] RESULT_SHAPE keys={list(result.keys())} "
              f"results_type={type(all_results).__name__} results_len={len(all_results)}")

        _search_jobs[job_id]["result"] = result
        _search_jobs[job_id]["status"] = JobStatus.COMPLETE
        print(f"[search_job:{job_id}] COMPLETE")
    except Exception as e:
        print(f"[search_job:{job_id}] FAILED: {e}")
        _search_jobs[job_id]["error"] = str(e)
        _search_jobs[job_id]["status"] = JobStatus.FAILED

@app.post("/api/search/document", status_code=202)
async def search_with_document(
    request: Request,
    files: List[UploadFile] = File(...),
    query: str = Form(...),
    matter_id: Optional[str] = Form(None),
    include_legal_updates: bool = Form(True),
    limit: int = Form(8),
):
    """
    Search Vault, extended: upload one or more documents ad-hoc (lease,
    contract, affidavit, etc.) and ask a question that spans them — e.g.
    applying a question of law to facts drawn from several uploaded
    documents at once. Each document is analyzed directly and kept clearly
    labeled by filename in what's sent to the model, so the answer can
    attribute a fact to the specific document it came from rather than
    treating them as one undifferentiated blob (see
    _combine_attached_documents). Nothing is permanently stored or
    indexed, same as a single-file attach — this is a one-off query, not
    matter documents — and the answer is grounded in both the attached
    documents and the firm's existing indexed knowledge (firm precedent,
    legal updates, ZLR judgments), same as a normal Search Vault query.

    Runs as a fire-and-poll background job rather than a single synchronous
    request/response — the pipeline can run long enough to trip Cloudflare's
    524 edge timeout. This returns a job_id immediately; poll
    /api/search/document/status/{job_id} for progress and the eventual result.
    """
    user = await get_current_user(request)
    _check_permission(user, "search")
    _require_retrieval_ready()

    if not files:
        raise HTTPException(status_code=422, detail="At least one file is required.")

    now = datetime.utcnow()
    for jid, job in list(_search_jobs.items()):
        if now - datetime.fromisoformat(job["created_at"]) > _SEARCH_JOB_MAX_AGE:
            del _search_jobs[jid]

    # Read every file's content upfront — UploadFile objects are tied to
    # this request/response cycle and must not be relied on after it ends
    # (same reasoning as job_user below).
    job_files = [{"filename": f.filename or "document", "content": await f.read()} for f in files]

    job_id = str(_uuid_mod.uuid4())

    # Capture only the minimum authenticated user context required by the
    # pipeline — not the Request or UploadFile objects, which are tied to
    # this request/response cycle and must not be relied on after it ends.
    job_user = {
        "id": user.get("id"),
        "firm_id": user.get("firm_id"),
        "display_name": user.get("display_name"),
        "role": user.get("role"),
    }

    _search_jobs[job_id] = {
        "status": JobStatus.PENDING,
        "result": None,
        "error": None,
        "firm_id": str(user.get("firm_id") or FIRM_ID),
        "created_at": now.isoformat(),
    }

    asyncio.create_task(
        _run_document_search_job(
            job_id, job_files, query, job_user, matter_id, include_legal_updates, limit,
        )
    )

    return {"job_id": job_id, "status": "pending"}

@app.get("/api/search/document/status/{job_id}")
async def get_search_job_status(job_id: str, request: Request):
    user = await get_current_user(request)
    _check_permission(user, "search")

    job = _search_jobs.get(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    if job["firm_id"] != str(user.get("firm_id") or FIRM_ID):
        raise HTTPException(status_code=403, detail="Not authorized")

    # TEMP diagnostic logging — remove once the stale-frontend vs.
    # result-shape-mismatch question is settled.
    result_keys = list(job["result"].keys()) if job["result"] else None
    print(f"[search_job:{job_id}] STATUS_POLL_RESPONSE status={job['status']} "
          f"has_result={job['result'] is not None} result_keys={result_keys}")

    return {
        "job_id": job_id,
        "status": job["status"],
        "result": job["result"],
        "error": job["error"],
    }

# ── Contract Review ────────────────────────────────────────────────────────────
# Two-stage design, same principle as the grounding check elsewhere in this
# file: stage 1 (Sonnet) identifies findings; stage 2 independently verifies
# each finding that claims specific text exists in the contract, by actually
# checking that text against the real document — before anything is shown
# to the user. This is a direct defence against exactly the failure mode in
# Pulserate Investments (Pvt) Ltd v Andrew Zuze and Others [SC202/25], where
# AI-generated fictitious citations in heads of argument were rejected by
# the Supreme Court. Verification only covers claims about text that IS
# present ("this clause says X") — it can't fully verify claims about
# absence ("this is missing a termination clause"), since proving a
# negative isn't the same kind of check. Those are flagged separately as
# unverified-by-quote rather than silently treated the same as a verified one.

CONTRACT_REVIEW_SYSTEM = """You are a contract review assistant for {FIRM_NAME}, {FIRM_CITY}, reviewing
documents under Zimbabwean law. Analyse the contract and identify genuine
issues only — do not manufacture findings to pad out the list. Use the
submit_contract_review tool to report your findings.

Categories:
- missing_clause: a standard clause this type of agreement would normally have, that isn't present
- risky_term: a clause that IS present but exposes the client to unusual risk
- non_standard: unusual wording/structure that deviates from normal market practice
- ambiguity: language that is genuinely unclear or could be read multiple ways
- compliance: potential conflict with Zimbabwean statutory requirements (e.g. Deeds Registries Act
  formalities, Labour Act provisions, Companies and Other Business Entities Act requirements)

Every "quote" must be copied EXACTLY from the contract text provided — do not paraphrase or
reconstruct from memory. If you cannot find the exact text to quote, omit the quote field and
rely on the category/description alone (this applies to missing_clause findings by definition,
since there's no text to quote for something that isn't there).

The "findings" field must be a native JSON array of objects, exactly as submit_contract_review's
schema defines it — never a JSON-encoded string containing an array. When referencing contract
text inside "description" (as opposed to the "quote" field, which must be verbatim), prefer
single quotes or paraphrasing over embedding double-quoted excerpts — this keeps the output
easier to encode correctly and avoids wasting output budget on escape sequences."""

# Structured tool schema instead of asking the model to hand-format free-text
# JSON — contract text routinely contains quote marks (defined terms like
# "Employee") and embedded line breaks within clauses, both of which broke
# manual JSON parsing in production (a real "Expecting ',' delimiter" error
# from a genuine employment contract upload). Anthropic's tool-use handles
# the encoding reliably; this sidesteps the whole class of escaping bugs
# rather than trying to patch the free-text JSON parsing after the fact.
CONTRACT_REVIEW_TOOL = {
    "name": "submit_contract_review",
    "description": "Submit the structured contract review findings.",
    "input_schema": {
        "type": "object",
        "properties": {
            "overall_summary": {
                "type": "string",
                "description": "2-3 sentence plain-language summary of the contract's overall risk profile",
            },
            "findings": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "category": {
                            "type": "string",
                            "enum": ["missing_clause", "risky_term", "non_standard", "ambiguity", "compliance"],
                        },
                        "severity": {"type": "string", "enum": ["high", "medium", "low"]},
                        "title": {"type": "string", "description": "Short label, e.g. 'No termination for convenience clause'"},
                        "description": {"type": "string", "description": "1-3 sentences explaining the issue and why it matters"},
                        "quote": {
                            "type": "string",
                            "description": "Exact verbatim text from the contract this finding is based on. "
                                           "Omit entirely for missing_clause findings (nothing to quote for an absence).",
                        },
                    },
                    "required": ["category", "severity", "title", "description"],
                },
            },
        },
        "required": ["overall_summary", "findings"],
    },
}

# Typographic substitutions a word processor's autocorrect applies that an
# LLM transcribing a quote won't reproduce (it defaults to plain ASCII) —
# without this, a single autocorrected quote mark in a Schedule-of-
# definitions-style clause (very common in contracts: "Employee" means...)
# turns an exact quote into a mismatch. Confirmed empirically: a realistic
# curly-quote defined-term clause scored 0.91 similarity against the old
# fuzzy fallback's 0.92 threshold — a real finding, silently dropped.
_TYPOGRAPHIC_NORMALIZE = str.maketrans({
    "“": '"', "”": '"', "‘": "'", "’": "'",
    "–": "-", "—": "-",
})

def _normalize_for_match(s: str) -> str:
    s = (s or "").translate(_TYPOGRAPHIC_NORMALIZE)
    return re.sub(r"\s+", " ", s).strip().lower()

def _verify_quote_in_text(quote: str, doc_text: str) -> bool:
    """
    Checks a finding's quoted text is actually present in the source
    document — the core safeguard against a finding being fabricated
    rather than genuinely drawn from the contract. Normalizes whitespace
    and typographic quote/dash variants before checking, then falls back
    to a fuzzy check anchored on the middle of the quote for minor
    formatting differences, before giving up.

    The fuzzy fallback anchors on a chunk from the *middle* of the quote
    (deliberately avoiding the first/last few characters, which is exactly
    where a corrected quote mark tends to sit) and locates it exactly in
    the document, then compares only the small window around each anchor
    hit. This replaces an earlier version that ran difflib over the whole
    document per candidate window at a fixed step size — too slow at
    real document length (measured ~14s for 40 findings against a 50k-char
    doc with autojunk off) and, with autojunk on, unreliable on ordinary
    prose (measured finding only 13/250 matching chars on a genuinely
    present 250-char excerpt, because common characters get treated as
    "junk" and excluded from matching blocks).
    """
    if not quote or not quote.strip():
        return False
    norm_quote = _normalize_for_match(quote)
    norm_doc = _normalize_for_match(doc_text)
    if norm_quote in norm_doc:
        return True
    if len(norm_quote) < 15:
        return False  # too short to fuzzy-match reliably — must be exact
    anchor_len = min(24, len(norm_quote) // 2 or len(norm_quote))
    anchor_start = max(0, (len(norm_quote) - anchor_len) // 2)
    anchor = norm_quote[anchor_start:anchor_start + anchor_len]
    margin = 20
    search_from = 0
    checked = 0
    while checked < 25:
        pos = norm_doc.find(anchor, search_from)
        if pos == -1:
            break
        checked += 1
        search_from = pos + 1
        window_start = max(0, pos - anchor_start - margin)
        window_end = min(len(norm_doc), pos + (len(norm_quote) - anchor_start) + margin)
        window = norm_doc[window_start:window_end]
        ratio = difflib.SequenceMatcher(None, norm_quote, window).ratio()
        if ratio >= 0.90:
            return True
    return False

async def _run_contract_review_job(job_id: str, content: bytes, filename: str, focus_areas: Optional[str]):
    """
    Background job for /api/contract-review -- see that endpoint's docstring
    for why this runs as fire-and-poll rather than a single synchronous
    request/response. Contains the full original synchronous logic
    unchanged (extraction, stage-1 Sonnet generation, stage-2 per-finding
    verification); only the control flow at the boundaries changed --
    HTTPException raises became ValueError (caught uniformly below and
    written into the job's error field, same pattern as
    _run_document_search_job/_run_document_generation_job), and the final
    return became a write into _search_jobs[job_id]["result"].
    """
    _search_jobs[job_id]["status"] = JobStatus.RUNNING
    print(f"[contract_review_job:{job_id}] STARTED")
    try:
        firm = await get_firm_identity()

        ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
        ocr_confidence = None
        try:
            if ext == "pdf":
                doc_text, _, _, ocr_confidence = extract_pdf_text(content)
            elif ext in ("docx", "doc"):
                doc_text = extract_docx_text(content)
            elif ext in ("jpg", "jpeg", "png", "webp"):
                doc_text, ocr_confidence = ocr_image_bytes(content, ext)
                if not doc_text:
                    raise ValueError(
                        "Could not read text from this image. Make sure the photo is clear, "
                        "well-lit, and the document fills most of the frame."
                    )
            else:
                doc_text = content.decode("utf-8", errors="replace")
        except ValueError:
            raise
        except Exception as e:
            raise ValueError(f"Could not read document: {e}")

        if not doc_text or not doc_text.strip():
            raise ValueError("No readable text found in the uploaded document.")

        print(f"[contract_review_job:{job_id}] EXTRACTION_COMPLETE")

        truncated = len(doc_text) > MAX_ATTACHED_DOC_CHARS
        review_text = doc_text[:MAX_ATTACHED_DOC_CHARS]

        focus_line = f"\n\nPay particular attention to: {focus_areas}" if focus_areas else ""
        prompt = f"""Contract to review ({filename}):
---
{review_text}
---
{focus_line}

Review this contract now and call submit_contract_review with your findings."""

        # asyncio.to_thread is required here, not optional -- confirmed by a
        # real production incident: a request landed right as a deploy was
        # rolling over, the (then-)direct client.messages.create() call
        # blocked this worker's single event loop for the whole ~56s the
        # generation was in flight, the worker couldn't even respond to
        # health checks during that window (observed as repeated 499s on
        # /health/alerts immediately before), and when the old container
        # was torn down mid-block the request died with a bare 502 and no
        # application-level log line at all -- the coroutine never got a
        # chance to run past the blocking call. Every other Anthropic call
        # of comparable length in this file already goes through
        # asyncio.to_thread (see _call_document_generation_model's
        # docstring for the same reasoning); this one just hadn't been
        # brought in line with that convention yet.
        msg = await asyncio.to_thread(
            client.messages.create,
            model="claude-sonnet-4-5",
            # A flat 4096 was cutting off mid-generation on longer/denser
            # contracts (confirmed via stop_reason logged below on a real
            # production failure) -- a 9-page employment contract with
            # several substantial findings, each carrying a title,
            # multi-sentence description, and a verbatim quote, genuinely
            # needs more headroom than a short/simple contract. Scaled by
            # source length rather than picking another flat guess, same
            # pattern already used for the search/grounding endpoint
            # elsewhere in this file.
            max_tokens=min(6000 + len(review_text) // 5, 16000) if review_text else 4096,
            system=CONTRACT_REVIEW_SYSTEM.format(FIRM_NAME=firm["name"], FIRM_CITY=firm["city"]),
            tools=[CONTRACT_REVIEW_TOOL],
            tool_choice={"type": "tool", "name": "submit_contract_review"},
            messages=[{"role": "user", "content": prompt}]
        )
        tool_use = next((b for b in msg.content if b.type == "tool_use"), None)
        review = tool_use.input if tool_use else {"overall_summary": "", "findings": []}
        if msg.stop_reason == "max_tokens":
            print(f"[contract_review_job:{job_id}] generation hit max_tokens (usage={msg.usage}) -- "
                  f"output was likely cut off mid-structure")
        print(f"[contract_review_job:{job_id}] GENERATION_COMPLETE")

        # Stage 2 — verify each finding's quote against the real document text.
        # unverified_absence findings (no quote — a "missing clause" claim)
        # can't be checked this way; they're marked distinctly so the UI can
        # show them with appropriately lower confidence, not silently equal
        # to a verified quote-backed finding.
        findings_raw = review.get("findings", [])
        if isinstance(findings_raw, str):
            # Observed in production on a real employment contract, repeatedly:
            # the model flattens `findings` into a JSON-encoded string instead
            # of emitting a native array in the tool-use input. The first two
            # occurrences produced a complete, cleanly closed string that
            # json.loads() recovers correctly (see the isinstance(dict) guard
            # below for the per-item shape check). A later occurrence on the
            # same document instead cut off mid-string with no closing
            # bracket -- that's max_tokens truncation compounding on top of
            # the same string-flattening behaviour (the escaping overhead from
            # stringifying the array eats extra output budget, which is part
            # of why max_tokens was scaled up above and the system prompt now
            # explicitly forbids this). json.loads() correctly refuses to
            # fabricate a partial list from truncated JSON in that case, so
            # this recovery step only ever succeeds on genuinely complete
            # content -- every recovered item still has to pass the
            # isinstance(dict) check below like any other finding, so this
            # doesn't weaken the guard against genuinely malformed input.
            try:
                parsed = json.loads(findings_raw)
            except (json.JSONDecodeError, ValueError):
                parsed = None
            if isinstance(parsed, list):
                print(f"[contract_review_job:{job_id}] findings field was a JSON-encoded string; "
                      f"recovered {len(parsed)} item(s) via json.loads()")
                findings_raw = parsed

        if not isinstance(findings_raw, list):
            # Defensive guard against a wrong-typed top-level `findings` value
            # that recovery couldn't fix (e.g. genuinely not JSON, or JSON
            # truncated mid-string by max_tokens). Observed in practice, before
            # any of this existed, as a wildly inflated dropped_unverified_count
            # (11047 on a 9-page contract) with zero findings shown — because
            # `for f in "some string"` iterates one character at a time and
            # every "item" then fails the isinstance(f, dict) check below, so
            # the per-item guard (which exists for a single malformed *entry*
            # inside an otherwise-valid list) silently absorbed the whole thing
            # as thousands of individually "dropped" findings. Fail loudly
            # instead of returning what would look like a clean "no issues
            # flagged" result — the entire point of two-stage verification is
            # to never show a misleading result, and a silently-empty findings
            # list after a malformed generation is exactly that.
            print(f"[contract_review_job:{job_id}] findings field was not recoverable as a list "
                  f"(type={type(findings_raw).__name__}, stop_reason={msg.stop_reason}): "
                  f"{findings_raw!r:.3000}")
            truncated_hint = (
                " The model's response was cut off before completing — this contract may need "
                "a shorter focus_areas request, or is hitting a generation length limit."
                if msg.stop_reason == "max_tokens" else ""
            )
            raise ValueError(
                "The contract review didn't come back in the expected format. "
                "This is a generation error, not a clean review — please try again."
                + truncated_hint
            )

        verified_findings = []
        dropped_count = 0
        for f in findings_raw:
            if not isinstance(f, dict):
                # Defensive guard — the tool schema requires each finding to be
                # an object, and this hit production once already (a finding
                # came back as a plain string instead), crashing the whole
                # request. Skip anything malformed rather than fail the whole
                # job over one bad entry.
                print(f"[contract_review_job:{job_id}] Skipping malformed finding (not an object): {f!r}")
                dropped_count += 1
                continue
            quote = f.get("quote")
            if not quote:
                f["verification"] = "unverifiable_absence_claim"
                verified_findings.append(f)
                continue
            if _verify_quote_in_text(quote, doc_text):
                f["verification"] = "verified"
                verified_findings.append(f)
            else:
                # The model claimed this text exists in the contract, but it
                # doesn't — this is exactly the fabrication risk the two-stage
                # design exists to catch. Drop it rather than show something
                # that could be a hallucinated citation-equivalent.
                dropped_count += 1
                print(f"[contract_review_job:{job_id}] Dropped unverified finding '{f.get('title')}' — quoted text not found in document")

        result = {
            "overall_summary": review.get("overall_summary", ""),
            "findings": verified_findings,
            "dropped_unverified_count": dropped_count,
            "document": {
                "filename": filename, "truncated": truncated,
                "ocr_confidence": ocr_confidence,
                "low_confidence": ocr_confidence is not None and ocr_confidence < 80,
            },
        }

        _search_jobs[job_id]["result"] = result
        _search_jobs[job_id]["status"] = JobStatus.COMPLETE
        print(f"[contract_review_job:{job_id}] COMPLETE")
    except Exception as e:
        print(f"[contract_review_job:{job_id}] FAILED: {e}")
        _search_jobs[job_id]["error"] = str(e)
        _search_jobs[job_id]["status"] = JobStatus.FAILED

@app.post("/api/contract-review", status_code=202)
async def review_contract(
    request: Request,
    file: UploadFile = File(...),
    focus_areas: Optional[str] = Form(None),
):
    """
    Upload a contract and get a structured review: missing clauses, risky
    terms, non-standard wording, ambiguities, and Zimbabwe-specific
    compliance concerns — each finding independently verified against the
    actual document text before being returned. Not stored, same as
    Search Vault's document upload — this is a one-off analysis tool.

    Runs as a fire-and-poll background job rather than a single synchronous
    request/response, same pattern and same job store (_search_jobs) as
    /api/generate-document and /api/search/document -- extraction + Sonnet
    generation (up to 16000 tokens on long/dense contracts, see
    _run_contract_review_job's max_tokens comment) + per-finding
    verification, held behind one synchronous request, is exactly the
    pipeline shape that was already found to regularly exceed Cloudflare's
    ~100s edge timeout on /api/generate-document before that got converted.
    Confirmed via investigation (2026-08-26) that this endpoint had never
    received the same fix despite an already-documented blocking-event-loop
    incident on this exact endpoint (see _run_contract_review_job) and a
    structurally identical pipeline to /api/generate-document's own
    pre-fix incident. Returns a job_id immediately; poll
    /api/contract-review/status/{job_id} for progress and the eventual result.
    """
    user = await get_current_user(request)
    _check_permission(user, "draft:document")

    # Read the file upfront -- UploadFile is tied to this request/response
    # cycle and must not be relied on after it ends, same reasoning as
    # job_files in search_with_document above.
    content = await file.read()
    filename = file.filename or "contract"

    now = datetime.utcnow()
    for jid, job in list(_search_jobs.items()):
        if now - datetime.fromisoformat(job["created_at"]) > _SEARCH_JOB_MAX_AGE:
            del _search_jobs[jid]

    job_id = str(_uuid_mod.uuid4())

    _search_jobs[job_id] = {
        "status": JobStatus.PENDING,
        "result": None,
        "error": None,
        "firm_id": str(user.get("firm_id") or FIRM_ID),
        "created_at": now.isoformat(),
    }

    asyncio.create_task(_run_contract_review_job(job_id, content, filename, focus_areas))

    return {"job_id": job_id, "status": "pending"}

@app.get("/api/contract-review/status/{job_id}")
async def get_contract_review_job_status(job_id: str, request: Request):
    user = await get_current_user(request)
    _check_permission(user, "draft:document")

    job = _search_jobs.get(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    if job["firm_id"] != str(user.get("firm_id") or FIRM_ID):
        raise HTTPException(status_code=403, detail="Not authorized")

    return {
        "job_id": job_id,
        "status": job["status"],
        "result": job["result"],
        "error": job["error"],
    }

def _semantic_search_firm(req, chunks: list) -> list:
    results = []
    try:
        firm_col, _, _ = get_chroma_collections()
        if firm_col.count() > 0:
            query_vec = embed_texts([req.query])[0]
            if hasattr(query_vec[0], "__len__"): query_vec = query_vec[0]
            # Multi-tenancy hardening (Part 3) -- defense-in-depth on top of
            # the existing Postgres-side scoping (chunk_by_id.get(cid) below
            # already drops results for another firm's chunk_id, since
            # `chunks` is queried firm-scoped by the caller). Explicit here
            # too so a cross-firm leak can't happen silently if that
            # Postgres pre-filtering is ever weakened or bypassed.
            # firm_id is always present; Chroma's where clause allows only
            # one top-level operator, so a second condition (matter_id) must
            # be combined via $and rather than added as a second dict key.
            where = {"firm_id": str(FIRM_ID)}
            if req.matter_id:
                where = {"$and": [where, {"matter_id": req.matter_id}]}
            n_fetch = max(req.limit * 4, 20)
            query_kwargs = {"query_embeddings": [query_vec], "n_results": n_fetch, "where": where}
            res = firm_col.query(**query_kwargs)
            ids = res["ids"][0] if res["ids"] else []
            distances = res["distances"][0] if res["distances"] else []
            chunk_by_id = {c["id"]: c for c in chunks}
            for cid, dist in zip(ids, distances):
                chunk = chunk_by_id.get(cid)
                if not chunk:
                    continue
                similarity = max(0.0, 1.0 - dist)
                results.append({
                    "result_source": "firm",
                    "chunk_id": chunk["id"],
                    "text": chunk["text"],
                    "similarity": round(similarity, 3),
                    "document_id": str(chunk["document_id"]),
                    "matter_id": chunk.get("matter_id"),
                    "page_number": chunk.get("page_number"),
                    "chunk_index": chunk.get("chunk_index"),
                    "filename": chunk.get("document_filename") or "Unknown Document",
                    "document_type": chunk.get("document_type"),
                    "court": chunk.get("court"),
                    "matter_type": chunk.get("matter_type"),
                    "legal_source_type": chunk.get("legal_source_type"),
                    "authority_strength": chunk.get("authority_strength"),
                    "document_status": chunk.get("document_status"),
                    "provenance_document_type": chunk.get("provenance_document_type"),
                    "matter_number": chunk.get("matter_number"),
                    "matter_name": chunk.get("matter_name"),
                    "client_id": str(chunk["matter_client_id"]) if chunk.get("matter_client_id") else None,
                    "client_name": chunk.get("client_name"),
                })
                if len(results) >= req.limit:
                    break
    except Exception as e:
        print(f"[search] semantic search failed, falling back to keyword: {e}")
        query_words = set(req.query.lower().split())
        scored = []
        for chunk in chunks:
            score = len(query_words & set(chunk["text"].lower().split())) / max(len(query_words), 1)
            if score > 0:
                scored.append((score, chunk))
        scored.sort(key=lambda x: x[0], reverse=True)
        for score, chunk in scored[:req.limit]:
            results.append({
                "result_source": "firm",
                "chunk_id": chunk["id"],
                "text": chunk["text"],
                "similarity": round(score, 3),
                "document_id": str(chunk["document_id"]),
                "matter_id": chunk.get("matter_id"),
                "filename": chunk.get("document_filename") or "Unknown Document",
                "document_type": chunk.get("document_type"),
                "court": chunk.get("court"),
                "matter_type": chunk.get("matter_type"),
                "legal_source_type": chunk.get("legal_source_type"),
                "authority_strength": chunk.get("authority_strength"),
                "document_status": chunk.get("document_status"),
                "provenance_document_type": chunk.get("provenance_document_type"),
                "matter_number": chunk.get("matter_number"),
                "matter_name": chunk.get("matter_name"),
                "client_id": str(chunk["matter_client_id"]) if chunk.get("matter_client_id") else None,
                "client_name": chunk.get("client_name"),
            })
    return results

def _semantic_search_legal(req, chunks: list) -> list:
    results = []
    try:
        _, legal_col, _ = get_chroma_collections()
        if legal_col.count() > 0:
            query_vec = embed_texts([req.query])[0]
            if hasattr(query_vec[0], "__len__"): query_vec = query_vec[0]
            res = legal_col.query(query_embeddings=[query_vec], n_results=min(req.limit * 2, legal_col.count()))
            ids = res["ids"][0] if res["ids"] else []
            distances = res["distances"][0] if res["distances"] else []
            chunk_by_id = {c["id"]: c for c in chunks}
            for cid, dist in zip(ids, distances):
                chunk = chunk_by_id.get(cid)
                if not chunk:
                    continue
                similarity = max(0.0, 1.0 - dist)
                results.append({
                    "result_source": "legal",
                    "chunk_id": chunk["id"],
                    "text": chunk["text"],
                    "similarity": round(similarity, 3),
                    "document_id": str(chunk["document_id"]),
                    "source_type": chunk.get("source_type"),
                    "source_name": chunk.get("source_name"),
                    "reference": chunk.get("reference"),
                    "legal_source_type": chunk.get("legal_source_type"),
                    "authority_strength": chunk.get("authority_strength"),
                })
                if len(results) >= req.limit:
                    break
    except Exception as e:
        print(f"[search] legal semantic search failed: {e}")
    return results

def synthesise_answer_sync(query: str, results: list, legal_results: list, zlr_results: list,
                            attached_doc_text: Optional[str] = None,
                            attached_doc_name: Optional[str] = None,
                            deadline_info: Optional[dict] = None,
                            research_map: Optional[dict] = None,
                            firm_name: Optional[str] = None,
                            firm_city: Optional[str] = None) -> str:
    # Sync function (runs via asyncio.to_thread from the async job runners
    # below) — can't await a DB call itself, so callers resolve
    # get_firm_identity() beforehand and pass the live values in. Falls
    # back to the frozen constants only if a caller doesn't pass them.
    firm_name = firm_name or FIRM_NAME
    firm_city = firm_city or FIRM_CITY
    if not results and not legal_results and not zlr_results and not attached_doc_text:
        return None
    context = format_context(results, legal_results, zlr_results)

    attached_block = ""
    if attached_doc_text:
        attached_block = f"""

ATTACHED DOCUMENT ({attached_doc_name or 'uploaded document'}) — this is the primary
subject of the query. Analyze it directly and specifically, quoting or
referencing its actual clauses/wording where relevant:
---
{attached_doc_text}
---"""

    deadline_block = ""
    if deadline_info:
        deadline_block = f"""
PRE-CALCULATED DEADLINE (computed deterministically — state this exactly, do not recalculate):
Event date: {deadline_info['event_date']}
Notice period required: {deadline_info['notice_period_days']} days (per {deadline_info['source_reference']})
Calculated deadline: {deadline_info['deadline']}
Status: {deadline_info['status']}
"""
    else:
        deadline_block = "\nNo notice-period figure with an exact day count was found in the retrieved sources — do not state or calculate any specific deadline; say plainly that this cannot be determined from the retrieved sources.\n"

    research_map_block = ""
    if research_map and research_map.get("gaps"):
        gap_lines = "\n".join(
            f"- {g['issue']}: missing {g['missing_authority']} ({g['reason']})"
            for g in research_map["gaps"]
        )
        research_map_block = f"""
RESEARCH GAP MAP (this is a completeness analysis of the retrieved material, NOT legal authority — never cite it as a source; use it only to state precisely what the retrieved sources do not establish):
{gap_lines}
"""

    if attached_doc_text:
        instructions = """Answer the question about the attached document directly and specifically:
- Ground your analysis in the document's actual wording — reference specific clauses, dates, or terms where relevant
- Where firm precedent or legal sources below are relevant, cross-reference them explicitly (e.g. "this clause is consistent with/departs from [reference]")
- If the document appears to have a legal defect, gap, or unusual provision, flag it clearly
- If firm precedents or legislation/case law don't materially bear on this question, say so briefly rather than forcing a connection
- If a firm precedent source below is labeled DRAFT, REVIEW, or SUPERSEDED, explicitly caveat any reliance on it (e.g. "based on a draft document, not yet finalized") — never present it as settled firm precedent""" + TEXTURE_RULES + FACT_EXTRACTION_RULES + LAWYER_JUDGMENT_RULES + STATUTORY_MECHANISM_PRECISION + IRAC_STRUCTURE_RULES
    else:
        instructions = """Answer directly and practically:
- If firm precedents are present, identify patterns and note them by document ID
- If legislation or case law is present, summarise the relevant legal position and cite by reference
- Flag variations over time
- For drafting queries, suggest specific language from the firm precedents
- If a firm precedent source below is labeled DRAFT, REVIEW, or SUPERSEDED, explicitly caveat any reliance on it (e.g. "based on a draft document, not yet finalized") — never present it as settled firm precedent""" + TEXTURE_RULES + FACT_EXTRACTION_RULES + LAWYER_JUDGMENT_RULES + STATUTORY_MECHANISM_PRECISION + IRAC_STRUCTURE_RULES

    try:
        msg = client.messages.create(
            model="claude-sonnet-4-5",
            # 900, then 3000, both still cut off mid-sentence in production for
            # attached-document analysis — this level of thorough, multi-section
            # analysis (headed sections, tables, extensive verbatim quoting)
            # genuinely needs significant headroom. The plain-query branch (no
            # attached document) used to be a flat 4000 regardless of how much
            # was actually being asked — fine for a single narrow issue, but a
            # genuinely broad multi-issue query hits IRAC_STRUCTURE_RULES' full
            # 7-section structure once per issue and blows straight through it.
            # Both branches now scale the same way, off however much source
            # material/context this particular call actually has to work with,
            # rather than one branch guessing a fixed number.
            max_tokens=min(8000 + len(attached_doc_text) // 5, 24000) if attached_doc_text
                       else min(8000 + len(context) // 5, 24000),
            messages=[{"role": "user", "content": f"""You are a legal research assistant for {firm_name}, {firm_city}.
Today's date: {datetime.utcnow().strftime('%Y-%m-%d')}

Query: {query}
{attached_block}
{deadline_block}
{research_map_block}

Sources:
{context if context else '(no additional firm/legal sources retrieved)'}

{instructions}

Professional, direct{
    ", using clear headed sections for a thorough document review. Clearly distinguish the attached document's own content from firm precedent and from public legal sources."
    if attached_doc_text else
    ", following the issue-by-issue structure above in full for every distinct legal issue the query raises — do not compress multiple issues together or shorten/omit sections to fit a target length."
}. Finish every section you start — do not begin a point and leave it incomplete."""}]
        )
        answer_text = msg.content[0].text
        if msg.stop_reason == "max_tokens":
            # This has hit production twice already at lower limits (900,
            # then 3000) — rather than assume a fixed number will never be
            # too small again, detect it directly and say so, instead of
            # silently handing over an analysis that stops mid-sentence as
            # if it were complete.
            answer_text += (
                "\n\n---\n**⚠ This analysis was cut off before completing "
                "— it ran out of space rather than reaching a natural end. "
                "Treat the final section as incomplete and re-run the query "
                "for the rest, or ask a narrower follow-up question.**"
            )
        return answer_text
    except Exception as e:
        print(f"[synthesise_answer_sync] synthesis failed: {e}")
        total = len(results) + len(legal_results or [])
        return f"Found {total} relevant excerpt(s). Review the sources below."

# ── Affidavit Generator ───────────────────────────────────────────────────────

AFFIDAVIT_SYSTEM = """You are a legal drafting assistant for {FIRM_NAME}, {FIRM_CITY}.
Draft affidavits in proper Zimbabwe High Court form per SI 202/2021.
- Full court caption with case number, party names and designations
- Opening: deponent full name, ID, capacity, competency declaration
- Numbered paragraphs, first person, chronological facts
- Prayer paragraph with specific relief
- Commissioner of oaths block at end
- Use [_____] for unknown specifics"""

DOCUMENT_SYSTEM_BASE = """You are a legal drafting assistant for {FIRM_NAME}, {FIRM_CITY}, drafting for the
Zimbabwean legal system. Produce a complete, properly formatted document —
not a template or outline. Use [_____] for any specific detail (dates,
amounts, ID numbers) not supplied. Number paragraphs/clauses where that is
standard practice for this document type. Do not add commentary before or
after the document itself — output the document only.

This document is being drafted for {FIRM_NAME} specifically. Any firm
name, letterhead, or legal practitioners' signature block in this document
must read exactly "{FIRM_NAME}" — never a different firm name. If a firm
precedent or retrieved authority provided below was drafted by or
mentions a different firm, match only its language, structure, and
formatting — never copy a firm identity, letterhead, or signature block
from it verbatim."""

# Per-type drafting guidance — the specifics that make a Zimbabwean legal
# document actually usable rather than a generic template. Litigation types
# get a full court caption; everything else gets letterhead-style framing.
DOC_TYPE_GUIDANCE = {
    "summons_matrimonial": """MATRIMONIAL SUMMONS (High Court, Matrimonial Causes Act [Chapter 5:13]):
- Full caption: court, case number, Plaintiff/Defendant designations
- Summons proper: command to enter appearance to defend within the prescribed period
- Declaration: marriage particulars (date, place, type — civil/customary), breakdown allegation,
  children of the marriage (names, ages, custody sought), matrimonial property, maintenance claim
- Prayer: decree of divorce, custody, maintenance, property distribution, costs
- Certificate/notice of appearance to defend format if applicable""",

    "summons_civil": """CIVIL SUMMONS (High Court Rules, 2021):
- Full caption: court, case number, Plaintiff/Defendant designations
- Summons proper: command to enter appearance to defend within the prescribed period
- Declaration: numbered paragraphs — jurisdiction, cause of action (contract/delict), material facts,
  quantum/damages claimed with basis for calculation
- Prayer: relief sought, interest, costs on the scale claimed""",

    "court_application": """COURT APPLICATION (Notice of Motion + Draft Order, High Court Rules 2021):
- Notice of Motion: caption, "TAKE NOTICE THAT" formula, relief sought, respondent's right to oppose
  and time period, address for service
- Draft Order: precise operative wording of the order sought, ready for a judge to grant as-is
- Founding affidavit reference (draft the Notice of Motion and Draft Order; note that a founding
  affidavit should accompany this but is a separate document)""",

    "urgent_chamber": """URGENT CHAMBER APPLICATION (High Court Rules 2021, r59):
- Certificate of Urgency: legal practitioner's certificate stating why the matter cannot wait
  for the ordinary roll, irreparable harm if not heard urgently
- Notice of Motion (urgent form): caption, relief sought, interim relief pending return date
- Draft Order: interim relief + return date + final relief sought on return date
- Emphasise the urgency test (self-created urgency must be addressed if relevant)""",

    "notice_of_appeal": """NOTICE OF APPEAL (High Court/Supreme Court, Rules of the relevant court):
- Caption showing court a quo, case number below, and the appellate court above
- Grounds of appeal: numbered, each ground a distinct, specific error of law or fact
  (avoid vague/generic grounds — Zimbabwean appellate courts require specificity)
- Relief sought on appeal
- Notice of set down / heads of argument filing timeline reference where relevant""",

    "letter_of_demand": """LETTER OF DEMAND (formal pre-litigation demand):
- Firm letterhead style (no court caption) — addressed directly to the debtor/wrongdoer
- Clear statement of the claim/breach and its factual basis
- Precise amount demanded (or specific performance required) with basis for the figure
- Firm deadline for compliance (state exact date, not just "within X days")
- Clear statement of consequences of non-compliance (legal proceedings, interest, costs)
- Professional but firm tone — this is often the last step before litigation""",

    "review": """APPLICATION FOR REVIEW (High Court Rules 2021, judicial review grounds):
- Notice of Motion + Draft Order in review form
- Grounds of review: gross irregularity, failure to apply mind, procedural unfairness,
  irrationality/unreasonableness, ultra vires — be specific about which ground(s) apply and why
- Relief: setting aside the decision, remittal for fresh determination, costs
- Distinguish clearly from an appeal (review concerns the process, not the merits)""",

    "heads_of_argument": """HEADS OF ARGUMENT (structured legal submission):
- Introduction: brief statement of the issue(s) before the court
- Factual background: concise, only facts relevant to the legal argument
- Issues for determination: numbered
- Argument: structured by issue, statutory/case authority cited for each proposition,
  applied to the facts of this matter, addressing the strongest counter-argument
- Conclusion and relief sought""",

    "legal_opinion": """LEGAL OPINION (formal written opinion for a client):
- Addressed to the client, headed "RE: [matter]" with a clear scope-of-opinion statement
- Executive summary / short answer up front
- Factual background as instructed
- Legal analysis: statutory and case law basis, applied to these specific facts,
  including genuine risks/uncertainties — do not overstate certainty
- Conclusion and recommended course of action
- Standard opinion caveats (based on facts as instructed, subject to further information)""",

    "client_letter": """CLIENT LETTER (formal correspondence):
- Firm letterhead style, clear subject line
- Plain-language explanation (client may not be legally trained) while remaining precise
- Clear statement of purpose, next steps, and any action required of the client with a deadline
- Professional, warm but not informal tone""",

    "agreement": """AGREEMENT / CONTRACT (general commercial agreement):
- Parties clause with full legal names and addresses/registration details
- Recitals (background/purpose)
- Definitions clause for defined terms used throughout
- Operative clauses: obligations of each party, payment terms, duration/termination,
  breach and remedies, dispute resolution, governing law (Zimbabwe)
- Signature blocks for all parties, witnesses if required""",

    "joint_venture": """JOINT VENTURE / SHAREHOLDERS AGREEMENT (Companies and Other Business Entities Act [Chapter 24:31]):
- Parties and recitals (purpose of the joint venture)
- Shareholding structure, capital contributions, valuation basis
- Governance: board composition, reserved matters requiring unanimous/special consent,
  deadlock resolution mechanism
- Transfer restrictions (pre-emption rights, tag-along/drag-along if relevant)
- Exit mechanisms, dispute resolution, governing law""",

    "agreement_of_sale": """AGREEMENT OF SALE — IMMOVEABLE PROPERTY (Deeds Registries Act [Chapter 20:05]):
- Full description of the property matching the title deed (stand number, township, registration
  details) — flag clearly that this must be verified against the actual title deed
- Purchase price, payment terms (deposit, balance, timing)
- Conditions precedent (bond approval, subdivision consent, etc. if relevant)
- Transfer obligations: who bears transfer costs, rates clearance, timeline to transfer
- Occupation/possession date, risk and benefit passing, breach and cancellation clauses""",

    "sale_of_vehicle_or_equipment": """SALE OF VEHICLE OR EQUIPMENT — MOVABLE PROPERTY (not the Deeds Registries Act — a different regime entirely):

Applies across a wide range — from a single vehicle to heavy industrial/mining equipment (lathes,
milling machines, conveyor systems, mining plant). Adapt which considerations below actually apply
based on what is being sold; not everything applies to every transaction.

Common to all:
- Establish whether the Seller is selling as principal (already the owner) or as an agent/dealer selling
  on behalf of a third party disposing of the item — a common arrangement for vehicle dealerships in
  particular. If selling as agent: the contract should say so, and should not simply name the dealer as
  "the Seller" without disclosing this, since a covenant to "convey ownership" or "execute all documents
  necessary to finalise transfer" may be a promise the dealer cannot keep alone — for a vehicle, CVR
  registration transfer requires the registered owner's own signature/authority, not merely whoever
  collected the purchase price. Where the Seller is acting as agent, either the actual registered owner
  should be a party (or provide a power of attorney), or the dealer should expressly warrant it has that
  owner's authority to sell and that the owner will cooperate with the registration transfer
- Full description of the item — for a vehicle: make, model, registration number, engine number, and
  chassis/VIN number, matching the registration book (logbook) exactly; for equipment: make, model,
  and serial number, matching the equipment's own identification plate/documentation. For multiple
  items (e.g. a set of factory machinery), a schedule listing each item individually is preferable to a
  single lump description. Flag clearly that these details must be verified against the actual
  registration book or identification markings, not just what the parties have stated — note also that
  for a vehicle, the chassis number and VIN are normally the same value for modern vehicles; if a
  contract lists them as two different numbers, that is worth querying rather than assuming correct
- Seller's warranty of title: that the Seller is the lawful and sole owner, and that the item is free of any
  lien, encumbrance, hire-purchase agreement, equipment finance/leasing arrangement, or other
  third-party claim
- Purchase price and payment terms (cash on signature is typical for smaller items; for higher-value
  industrial equipment, consider staged payment tied to delivery/installation/commissioning milestones
  rather than full payment upfront)
- Sold "voetstoots" (as-is) — no warranty as to mechanical/operational condition or fitness for
  purpose, unless specific representations are made and recorded. For equipment represented as
  operational/functional, consider whether a testing or commissioning period with a defined
  acceptance mechanism is more appropriate than a bare as-is sale
- Breach and cancellation clauses proportionate to the value of the transaction

Additional, for a vehicle specifically:
- Confirmation that a Central Vehicle Registry (CVR) search/clearance and police clearance
  (confirming it is not reported stolen) have been or will be obtained before transfer
- Transfer of registration into the Purchaser's name at CVR/Vehicle Registration Department — who is
  responsible for effecting this, the timeline, and who bears the transfer fee
- Roadworthy certificate: note whether one exists or is required (mandatory for public service
  vehicles/commercial use in some cases; not always required for a private sale — flag as something
  to confirm rather than assume)

Additional, for larger/industrial/mining equipment specifically:
- Decommissioning, rigging, and removal from the Seller's premises — who arranges and pays for
  this, and what access to the site the Purchaser (or their contractors) requires
- Delivery, transport, reinstallation, and commissioning at the new site — who is responsible, and
  whether risk passes on removal from the Seller's site, on delivery, or only on successful
  commissioning (these can be materially different points in time for heavy equipment)
- Whether existing technical documentation, service/maintenance manuals, spare parts inventory,
  and any manufacturer warranty registration transfer with the equipment
- Safety and regulatory compliance — certification relevant to the specific equipment (e.g. pressure
  vessel/boiler inspection certificates for factory machinery, Ministry of Mines/Environmental
  Management Act compliance for mining equipment) — flag as something to confirm applies and is
  current, not something to assume
- If the equipment is being sold in situ or as part of broader business assets, consider whether VAT
  applies (this is more likely to arise here than in a private individual vehicle sale) and who bears
  responsibility for any site/environmental conditions left behind after removal""",

    "acknowledgement_of_debt": """ACKNOWLEDGEMENT OF DEBT (liquid document, consent to judgment):
- Clear acknowledgement of the specific amount owed and its basis
- Repayment terms (schedule or lump sum, interest rate if applicable)
- Consent to judgment clause (if instructed) — acceleration on default
- Security/surety details if applicable
- Signature by debtor, ideally with a competent witness — note this document's evidentiary
  value depends on proper execution""",

    "power_of_attorney_transfer": """POWER OF ATTORNEY TO PASS TRANSFER (Deeds Registries Act [Chapter 20:05], conveyancing):
- Principal's full details, clear appointment of the conveyancer/attorney
- Specific property description matching the title deed
- Precise scope of authority (to sign all documents necessary to pass transfer of the specific
  property to the specific purchaser)
- Note this must be executed per Deeds Registry formalities (commissioning requirements)""",

    "declaration_transferor": """DECLARATION BY TRANSFEROR (Deeds Registry seller's declaration):
- Transferor's confirmation of the sale, property details matching title deed
- Standard declarations required by the Deeds Registry: no other unregistered sale/prior
  disposal, marital status and consent of spouse where relevant, citizenship/status declarations
- Signature and commissioning block per Deeds Registry practice""",

    "declaration_transferee": """DECLARATION BY TRANSFEREE (Deeds Registry buyer's declaration):
- Transferee's confirmation of the purchase, property details matching title deed
- Standard declarations: citizenship/residency status (relevant to land acquisition rules),
  marital status, source of funds where applicable
- Signature and commissioning block per Deeds Registry practice""",

    "special_power_of_attorney": """SPECIAL POWER OF ATTORNEY (limited, transaction-specific):
- Principal's details, clear and narrow statement of the specific act(s) authorised
  (avoid broad/general authority language — specificity is the point of a "special" POA)
- Explicit limitation to the named transaction/purpose only
- Duration/expiry if applicable, revocation clause
- Signature and commissioning block""",

    "sale_of_business": """SALE OF BUSINESS AGREEMENT (going concern — assets, goodwill, employees):
- Parties and description of the business being sold (as a going concern)
- Assets included/excluded, goodwill, stock valuation basis
- Employees: transfer of employment terms, any retrenchment provisions
- Purchase price, payment/completion mechanics, warranties (title, no undisclosed liabilities)
- Restraint of trade clause (reasonable in scope, area, and duration to be enforceable)""",

    "memorandum_of_understanding": """MEMORANDUM OF UNDERSTANDING (pre-agreement):
- Parties and purpose/intent of the arrangement
- Clear statement of whether this MOU is binding or non-binding (state explicitly — this is
  the single most important clause in an MOU and ambiguity here causes real disputes)
- Key terms contemplated for the eventual full agreement
- Exclusivity/confidentiality during the MOU period if relevant, term and termination""",

    "freeform": """Draft the document based entirely on the facts and instructions provided. Infer the
appropriate structure and formality from context. If the document type is ambiguous,
choose the most standard Zimbabwean legal form that fits the stated purpose.""",
}

@app.post("/api/generate-affidavit")
async def generate_affidavit(req: AffidavitRequest, request: Request):
    user = await get_current_user(request)
    _check_permission(user, "draft:affidavit")
    firm = await get_firm_identity()
    precedent_block = ""
    if req.precedent_context:
        fname = req.precedent_context.get("filename", "precedent")
        mname = req.precedent_context.get("matter_name", "")
        text = str(req.precedent_context.get("text", ""))[:2000]
        precedent_block = f"\n\nFIRM PRECEDENT ({fname} \u2014 {mname}):\n---\n{text}\n---"
    prompt = f"""Draft a Zimbabwe High Court affidavit.

Matter type: {req.matter_type or 'General'}
Court: {req.court}
Deponent: {req.deponent_name or '[DEPONENT NAME]'}
ID Number: {req.deponent_id or '[ID NUMBER]'}
Capacity: {req.deponent_capacity or 'the Applicant'}
Parties: {req.parties or '[PARTIES]'}
Matter summary: {req.matter_summary}
Key facts: {req.key_facts or 'As per matter summary above'}
Relief sought: {req.relief or '[RELIEF TO BE SPECIFIED]'}
{precedent_block}

Draft the complete affidavit in proper Zimbabwe High Court form. Number all paragraphs. Include the commissioner of oaths block at the end."""

    try:
        msg = client.messages.create(
            model="claude-sonnet-4-5",
            max_tokens=4096,
            system=AFFIDAVIT_SYSTEM.format(FIRM_NAME=firm["name"], FIRM_CITY=firm["city"]),
            messages=[{"role": "user", "content": prompt}]
        )
        return {"affidavit": msg.content[0].text}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Affidavit generation failed: {e}")

# Upload-time classification uses its own taxonomy (affidavit, lease_agreement,
# correspondence, contract, etc. — see classify_document_sync) which doesn't
# line up 1:1 with the 20 drafting types (e.g. a demand letter gets classified
# as generic "correspondence" at upload, never literally "letter_of_demand").
# This maps each draft type to the upload classification(s) most likely to
# actually be the same kind of document, for the exact-match half of
# find_precedents below.
DRAFT_TYPE_TO_UPLOAD_TYPES = {
    "summons_matrimonial": ["summons", "declaration"],
    "summons_civil": ["summons", "declaration"],
    "court_application": ["notice_of_motion", "founding_affidavit"],
    "urgent_chamber": ["notice_of_motion", "founding_affidavit"],
    "notice_of_appeal": ["notice_of_motion", "court_order"],
    "letter_of_demand": ["correspondence"],
    "review": ["notice_of_motion", "founding_affidavit"],
    "heads_of_argument": ["heads_of_argument"],
    "legal_opinion": ["opinion"],
    "client_letter": ["correspondence"],
    "agreement": ["contract"],
    "joint_venture": ["contract", "deed_of_settlement"],
    "agreement_of_sale": ["contract", "lease_agreement"],
    "sale_of_vehicle_or_equipment": ["contract"],
    "acknowledgement_of_debt": ["deed_of_settlement", "contract"],
    "power_of_attorney_transfer": ["power_of_attorney"],
    "declaration_transferor": ["declaration", "power_of_attorney"],
    "declaration_transferee": ["declaration", "power_of_attorney"],
    "special_power_of_attorney": ["power_of_attorney"],
    "sale_of_business": ["contract"],
    "memorandum_of_understanding": ["contract", "correspondence"],
    "freeform": [],
}

@app.get("/api/documents/find-precedents")
async def find_precedents(doc_type: str, facts: str = "", request: Request = None):
    """
    Finds real, existing firm documents to use as drafting precedent —
    replaces the old flow, which only let you manually pick a matter and
    then silently grabbed whatever document in it happened to be most
    recently uploaded, regardless of whether it was even the same kind of
    document. Combines two signals:
      1. Exact/mapped document_type match (from upload-time classification)
      2. Semantic similarity search across all firm chunks, using the
         document type label + the facts just entered as the query — this
         is what catches good precedents even when classification labels
         don't line up (e.g. a real demand letter filed as "correspondence").
    Type-matches are ranked first (most reliable "same kind of document"),
    semantic matches fill in after.
    """
    user = await get_current_user(request)
    _check_permission(user, "matter:read")

    candidates = {}
    mapped_types = DRAFT_TYPE_TO_UPLOAD_TYPES.get(doc_type, [])

    async with _db_pool.acquire() as conn:
        if mapped_types:
            rows = await conn.fetch("""
                SELECT d.id, d.filename, d.document_type, m.name as matter_name
                FROM documents d JOIN matters m ON m.id = d.matter_id
                WHERE d.firm_id=$1 AND d.document_type = ANY($2::text[]) AND d.status='complete'
                ORDER BY d.uploaded_at DESC LIMIT 5
            """, FIRM_ID, mapped_types)
            for r in rows:
                candidates[str(r["id"])] = {
                    "id": str(r["id"]), "filename": r["filename"], "matter_name": r["matter_name"],
                    "document_type": r["document_type"],
                    "match_reason": f"same document type ({r['document_type'].replace('_',' ')})",
                    "score": 1.0,
                }

        firm_chunk_rows = await conn.fetch(
            "SELECT * FROM chunks WHERE firm_id=$1 AND chunk_source='firm'", FIRM_ID
        )

    firm_chunks = [dict(r) for r in firm_chunk_rows]
    label = DOC_TYPE_LABELS_BACKEND.get(doc_type, doc_type)
    query_text = f"{label}. {facts[:500]}" if facts else label
    fake_req = SearchRequest(query=query_text, limit=8)

    try:
        semantic_results = await asyncio.to_thread(_semantic_search_firm, fake_req, firm_chunks)
    except Exception:
        semantic_results = []

    doc_best_score = {}
    for r in semantic_results:
        did = r["document_id"]
        if did not in doc_best_score or r["similarity"] > doc_best_score[did]:
            doc_best_score[did] = r["similarity"]

    new_doc_ids = [did for did in doc_best_score if did not in candidates]
    if new_doc_ids:
        async with _db_pool.acquire() as conn:
            rows = await conn.fetch("""
                SELECT d.id, d.filename, d.document_type, m.name as matter_name
                FROM documents d JOIN matters m ON m.id = d.matter_id
                WHERE d.id = ANY($1::uuid[]) AND d.firm_id=$2 AND d.status='complete'
            """, [_uuid_mod.UUID(did) for did in new_doc_ids], FIRM_ID)
        for r in rows:
            did = str(r["id"])
            candidates[did] = {
                "id": did, "filename": r["filename"], "matter_name": r["matter_name"],
                "document_type": r["document_type"],
                "match_reason": f"similar content ({round(doc_best_score[did] * 100)}% match)",
                "score": doc_best_score[did],
            }

    ranked = sorted(candidates.values(), key=lambda c: c["score"], reverse=True)
    return {"candidates": ranked[:6]}

# Mirrors the frontend's DOC_TYPE_LABELS exactly — used to give the model a
# readable document-type name in the prompt rather than the raw type key.
DOC_TYPE_LABELS_BACKEND = {
    "summons_matrimonial": "Matrimonial Summons",
    "summons_civil": "Civil Summons",
    "court_application": "Court Application (Notice of Motion & Draft Order)",
    "urgent_chamber": "Urgent Chamber Application",
    "notice_of_appeal": "Notice of Appeal",
    "letter_of_demand": "Letter of Demand",
    "review": "Application for Review",
    "heads_of_argument": "Heads of Argument",
    "legal_opinion": "Legal Opinion",
    "client_letter": "Client Letter",
    "agreement": "Agreement / Contract",
    "joint_venture": "Joint Venture / Shareholders Agreement",
    "agreement_of_sale": "Agreement of Sale (Immoveable Property)",
    "sale_of_vehicle_or_equipment": "Agreement of Sale (Vehicle or Equipment)",
    "acknowledgement_of_debt": "Acknowledgement of Debt",
    "power_of_attorney_transfer": "Power of Attorney to Pass Transfer",
    "declaration_transferor": "Declaration by Transferor",
    "declaration_transferee": "Declaration by Transferee",
    "special_power_of_attorney": "Special Power of Attorney",
    "sale_of_business": "Sale of Business Agreement",
    "memorandum_of_understanding": "Memorandum of Understanding",
    "freeform": "legal document",
}

ADVOCACY_RULES = """
LITIGATION DRAFTING VOICE — this document is an advocate's submission, not a
legal textbook. The reader is a judge who must be persuaded, not a student
being taught the law. Every substantive paragraph should follow this rhythm,
not "state law → quote statute → move on":

1. STATE THE PROPOSITION FIRST, not the statute. Open with what you are
   arguing, in your own words, before any citation. Wrong: "Section 68(1)
   provides that administrative conduct must be procedurally fair." Right:
   "The Respondent's conduct falls squarely within the prohibition in s 68
   because the impugned decision was taken without notice, consultation, or
   reasons — conduct that cannot be described as lawful or procedurally
   fair."

2. APPLY TO THE FACTS IMMEDIATELY. Do not recite a legal principle in the
   abstract and leave the application implicit or for a later paragraph.
   The proposition and its application to this client's facts belong in the
   same breath.

3. WEAVE AUTHORITY INTO THE ARGUMENT, don't just cite it. Prefer "This
   principle has been repeatedly affirmed by this Court — see X v Y, where
   [holding]..." over a bare citation dropped after a legal statement.
   Authorities support a proposition you have already made; they do not
   replace making it.

4. ANTICIPATE THE OPPONENT, THEN REFUTE. For any point where a contrary
   argument is realistically available, name it plainly, then dismantle it
   — don't wait to be asked. This is what makes Heads persuasive rather than
   descriptive.

5. VARY PARAGRAPH LENGTH TO MATCH WEIGHT, not uniform blocks. A minor,
   uncontested point can be one sentence. A contested, central issue may
   need several paragraphs of sustained argument. Mechanically equal
   paragraph lengths read as generated, not argued — real Heads of Argument
   expand exactly where the case is won or lost and compress everywhere
   else.

6. USE ADVOCATE'S DICTION where it fits naturally — "it is respectfully
   submitted that," "with respect," "the inevitable consequence is,"
   "surely," "this Court has already held" — but only where it earns its
   place; don't insert stock phrases as decoration. The goal is genuine
   persuasive register, not a checklist of phrases.

7. THE GOAL OF EVERY PARAGRAPH IS TO ADVANCE WHY THE CLIENT MUST WIN, not
   to inform the reader about the law. If a paragraph could be deleted
   without weakening the argument, it does not belong.
"""

LITIGATION_DOC_TYPES = {
    "summons_matrimonial", "summons_civil", "court_application",
    "urgent_chamber", "notice_of_appeal", "review", "heads_of_argument",
}

def _call_document_generation_model(prompt: str, max_tokens: int, firm_name: str, firm_city: str):
    """
    Synchronous Anthropic call factored out so it can run via
    asyncio.to_thread() from inside the background job below — calling
    the SDK directly inside an async function would block the event loop
    for the whole generation, stalling every other concurrent request
    (including other jobs' status polls) for however long this document
    takes, defeating the point of moving this to a background job.

    firm_name/firm_city are resolved by the async caller (get_firm_identity())
    before this runs, not read from the frozen FIRM_NAME/FIRM_CITY
    constants here — this function can't await a DB call itself.
    """
    return client.messages.create(
        model="claude-sonnet-4-5",
        max_tokens=max_tokens,
        system=DOCUMENT_SYSTEM_BASE.format(FIRM_NAME=firm_name, FIRM_CITY=firm_city),
        messages=[{"role": "user", "content": prompt}]
    )


async def _run_document_generation_job(job_id: str, req: "DocumentRequest", user: dict):
    """
    Runs generate_document's full pipeline (retrieval, prompt construction,
    Sonnet generation, citation verification) in the background — same
    reasoning as _run_document_search_job above: retrieval + generation +
    verification held behind one synchronous request regularly exceeds
    Cloudflare's edge timeout once retrieval was added ahead of generation.
    """
    _search_jobs[job_id]["status"] = JobStatus.RUNNING
    print(f"[generate_job:{job_id}] STARTED")
    try:
        guidance = DOC_TYPE_GUIDANCE.get(req.doc_type, DOC_TYPE_GUIDANCE["freeform"])
        is_litigation = req.doc_type in LITIGATION_DOC_TYPES

        precedent_block = ""
        if req.precedent_context:
            fname = req.precedent_context.get("filename", "precedent")
            mname = req.precedent_context.get("matter_name", "")
            text = str(req.precedent_context.get("text", ""))[:4000]
            precedent_block = (
                f"\n\nFIRM PRECEDENT ({fname} — {mname}) — match this document's language, "
                f"structure, and drafting style where appropriate:\n---\n{text}\n---"
            )

        # Corpus-wide retrieval — the same search /api/search runs, so drafting
        # draws on genuinely relevant case law/legislation/firm precedent from
        # the whole Vault, not just the single manually-attached precedent
        # above. Additive alongside precedent_block, not a replacement for it.
        retrieval_query = req.facts + (f" {req.instructions}" if req.instructions else "")
        retrieval_req = SearchRequest(query=retrieval_query, limit=8)

        async with _db_pool.acquire() as conn:
            firm_chunk_rows = await conn.fetch(
                """
                SELECT c.*, d.filename AS document_filename, d.document_type, d.court,
                       d.matter_type, d.legal_source_type, d.authority_strength
                FROM chunks c
                JOIN documents d ON d.id = c.document_id
                WHERE c.firm_id=$1 AND c.chunk_source='firm'
                """,
                FIRM_ID
            )
            legal_chunk_rows = await conn.fetch(
                """
                SELECT c.*, lu.legal_source_type, lu.authority_strength
                FROM chunks c
                LEFT JOIN legal_updates lu ON lu.id = c.document_id
                WHERE c.firm_id=$1 AND c.chunk_source='legal'
                """,
                FIRM_ID
            )
            zlr_chunk_rows = await conn.fetch(
                """
                SELECT c.*, z.legal_source_type, z.authority_strength
                FROM chunks c
                LEFT JOIN zlr_entries z ON z.id = c.document_id
                WHERE c.firm_id=$1 AND c.chunk_source='zlr'
                """,
                FIRM_ID
            )

        retrieved_results = await asyncio.to_thread(
            _semantic_search_firm, retrieval_req, [dict(r) for r in firm_chunk_rows]
        )
        retrieved_legal_results = await asyncio.to_thread(
            _semantic_search_legal, retrieval_req, [dict(r) for r in legal_chunk_rows]
        )
        retrieved_zlr_results = []
        zlr_chunks_list = [dict(r) for r in zlr_chunk_rows]
        if zlr_chunks_list:
            raw_zlr = await asyncio.to_thread(_zlr_semantic_search, zlr_chunks_list, retrieval_query, None, 3)
            for r in raw_zlr:
                retrieved_zlr_results.append({
                    "result_source": "zlr",
                    "chunk_id": r.get("item_id"),
                    "text": r.get("relevant_excerpt", ""),
                    "similarity": r.get("similarity", 0),
                    "document_id": r.get("item_id"),
                    "filename": r.get("case_name") or r.get("citation") or "ZLR Entry",
                    "citation": r.get("citation"),
                    "taxonomy_category": r.get("taxonomy_category"),
                    "summary": r.get("summary"),
                    "legal_source_type": r.get("legal_source_type"),
                    "authority_strength": r.get("authority_strength"),
                })

        retrieved_context = format_context(retrieved_results, retrieved_legal_results, retrieved_zlr_results)
        print(f"[generate_job:{job_id}] RETRIEVAL_RESULTS firm={len(retrieved_results)} legal={len(retrieved_legal_results)} zlr={len(retrieved_zlr_results)} context_chars={len(retrieved_context)}")
        if retrieved_zlr_results:
            print(f"[generate_job:{job_id}] ZLR_MATCHES: " + ", ".join(r.get('filename', 'unknown') for r in retrieved_zlr_results))
        if retrieved_legal_results:
            # _semantic_search_legal's result dicts have no "filename" key at
            # all (that's a "firm"/ZLR-result field) — legal_updates results
            # carry the name under "reference"/"source_name" instead, same
            # as format_context()'s own lookup a few lines below and the
            # Legal Updates list in the frontend (item.reference || item.filename).
            print(f"[generate_job:{job_id}] LEGAL_MATCHES: " + ", ".join(
                (r.get('reference') or r.get('source_name') or 'unknown') for r in retrieved_legal_results[:5]
            ))
        # Framed as citation-affirmative rather than caution-first: an earlier
        # version read "cite only what is genuinely on point; do not force a
        # connection... do not cite any case not shown here" — two negatives
        # to one hedge, and testing showed the model erring toward citing
        # nothing even when a retrieved case (by name/citation) was clearly
        # on point, treating citation as high-risk by default rather than
        # evaluating each source's actual relevance. The "not shown below"
        # guardrail against fabricating uncited case names is preserved.
        retrieved_block = (
            f"\n\nAVAILABLE AUTHORITY FROM THE FIRM'S VAULT (retrieved as potentially relevant to "
            f"these facts). When a retrieved case or source below is genuinely relevant to the "
            f"facts or legal issue, cite it directly by name — this is exactly the kind of "
            f"authority this document should draw on. Only omit citation for a retrieved source "
            f"that is not actually on point, and never cite a case that is not shown below:"
            f"\n---\n{retrieved_context}\n---"
        ) if retrieved_context else ""
        print(f"[generate_job:{job_id}] RETRIEVAL_COMPLETE")

        party_block = ""
        if is_litigation:
            party_block = f"""Court: {req.court or 'High Court of Zimbabwe'}
Case/Matter Number: {req.case_number or '[CASE NUMBER]'}
Plaintiff/Applicant: {req.plaintiff or '[PLAINTIFF/APPLICANT]'}
Defendant/Respondent: {req.defendant or '[DEFENDANT/RESPONDENT]'}"""
        else:
            party_block = f"""First Party: {req.plaintiff or '[FIRST PARTY]'}
Second Party: {req.defendant or '[SECOND PARTY]'}"""
            if req.case_number:
                party_block += f"\nReference/Matter Number: {req.case_number}"

        # Advocate's-voice instructions only apply to litigation documents
        # (Heads of Argument, applications, appeals, etc.) — a contract or
        # letter drafted in a persuasive/adversarial register would be wrong.
        advocacy_block = ADVOCACY_RULES if is_litigation else ""

        prompt = f"""Draft a {DOC_TYPE_LABELS_BACKEND.get(req.doc_type, 'legal document')}.

{party_block}

Facts and background:
{req.facts}

{f"Additional instructions: {req.instructions}" if req.instructions else ""}
{precedent_block}
{retrieved_block}

DOCUMENT-SPECIFIC REQUIREMENTS:
{guidance}
{advocacy_block}

Draft the complete document now."""

        # Scaled the same way as synthesise_answer_sync's research path —
        # litigation documents need real headroom. Capped well under
        # claude-sonnet-4-5's real 64,000-token ceiling.
        max_tokens = min(
            6000 + len(req.facts + (req.instructions or '') + precedent_block + retrieved_block) // 5,
            20000,
        )
        firm = await get_firm_identity()
        msg = await asyncio.to_thread(_call_document_generation_model, prompt, max_tokens, firm["name"], firm["city"])
        print(f"[generate_job:{job_id}] SYNTHESIS_COMPLETE")

        document_text = msg.content[0].text
        if msg.stop_reason == "max_tokens":
            # Same truncation safeguard as synthesise_answer_sync — without
            # this, a cut-off litigation document reads as complete and
            # could be filed with no indication the ending is missing.
            document_text += (
                "\n\n---\n**⚠ This document was cut off before completing "
                "— it ran out of space rather than reaching a natural end. "
                "Treat the final section as incomplete and re-run the "
                "request for the rest, or ask a narrower follow-up question.**"
            )

        # Same inline citation check as the research/synthesis path, against
        # the exact retrieved context this document was drafted from — but
        # softer wording than research's "UNVERIFIED": the corpus is not
        # comprehensive, so a citation absent from it is unconfirmed, not
        # proven fabricated.
        document_text, inline_qc_log = verify_inline_case_citations(
            document_text, retrieved_context,
            annotation_suffix="[⚠ Not found in retrieved sources — verify independently before filing]",
        )
        if inline_qc_log:
            print(f"[generate_job:{job_id}] {len(inline_qc_log)} citation(s) not found in retrieved context: "
                  f"{[q['case_name'] for q in inline_qc_log]}")
        print(f"[generate_job:{job_id}] CITATION_QC_COMPLETE")

        _search_jobs[job_id]["result"] = {"document": document_text, "doc_type": req.doc_type}
        _search_jobs[job_id]["status"] = JobStatus.COMPLETE
        print(f"[generate_job:{job_id}] COMPLETE")
    except Exception as e:
        print(f"[generate_job:{job_id}] FAILED: {e}")
        _search_jobs[job_id]["error"] = str(e)
        _search_jobs[job_id]["status"] = JobStatus.FAILED


@app.post("/api/generate-document", status_code=202)
async def generate_document(req: DocumentRequest, request: Request):
    """
    General-purpose drafting endpoint backing the "Draft Document" feature —
    covers 20 document types (litigation, conveyancing, commercial
    agreements, correspondence) via a shared prompt structure with
    per-type guidance, the same pattern as generate_affidavit but
    generalized.

    Runs as a fire-and-poll background job, same pattern and same job
    store (_search_jobs) as /api/search/document — retrieval + Sonnet
    generation + citation verification regularly exceeded Cloudflare's
    edge timeout once retrieval was added ahead of generation, held
    behind one synchronous request. Returns a job_id immediately; poll
    /api/generate-document/status/{job_id} for the result.
    """
    user = await get_current_user(request)
    _check_permission(user, "draft:document")
    _require_retrieval_ready()

    now = datetime.utcnow()
    for jid, job in list(_search_jobs.items()):
        if now - datetime.fromisoformat(job["created_at"]) > _SEARCH_JOB_MAX_AGE:
            del _search_jobs[jid]

    job_id = str(_uuid_mod.uuid4())
    job_user = {
        "id": user.get("id"),
        "firm_id": user.get("firm_id"),
        "display_name": user.get("display_name"),
        "role": user.get("role"),
    }

    _search_jobs[job_id] = {
        "status": JobStatus.PENDING,
        "result": None,
        "error": None,
        "firm_id": str(user.get("firm_id") or FIRM_ID),
        "created_at": now.isoformat(),
    }

    asyncio.create_task(_run_document_generation_job(job_id, req, job_user))

    return {"job_id": job_id, "status": "pending"}


@app.get("/api/generate-document/status/{job_id}")
async def get_document_generation_job_status(job_id: str, request: Request):
    user = await get_current_user(request)
    _check_permission(user, "draft:document")

    job = _search_jobs.get(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    if job["firm_id"] != str(user.get("firm_id") or FIRM_ID):
        raise HTTPException(status_code=403, detail="Not authorized")

    return {
        "job_id": job_id,
        "status": job["status"],
        "result": job["result"],
        "error": job["error"],
    }


# ── DOCX Export ───────────────────────────────────────────────────────────────

@app.post("/api/export-docx")
async def export_docx(req: ExportRequest, request: Request):
    """
    Builds affidavit_text into a .docx via backend/docx_export.py — no
    longer shells out to Node/npm per request (the previous implementation
    generated and ran a throwaway docx.js script, installing the npm
    `docx` package into a temp dir on every single export). Same visual
    result (Times New Roman, justified body, centered/bold caption lines,
    1" margins), built with python-docx, which was already a declared
    dependency (used elsewhere for reading uploaded .docx files).
    """
    user = await get_current_user(request)
    _check_permission(user, "draft:document")
    paragraph_blocks = paragraphs_from_plain_text(req.affidavit_text)
    if not paragraph_blocks:
        raise HTTPException(status_code=422, detail="No content to export.")
    docx_bytes = await asyncio.to_thread(build_docx_bytes, paragraph_blocks)
    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="affidavit_{req.document_id}.docx"'}
    )

@app.post("/api/export-document-docx")
async def export_document_docx(req: ExportDocumentDocxRequest, request: Request):
    """
    Exports the drafting editor's CURRENT (possibly hand-edited) HTML
    content as a .docx — by export time the content is rich HTML (headings,
    bold, centered captions from the editor's own formatting/toolbar), not
    the original plain-text AI output, so this parses content_html rather
    than re-deriving from the AI's first draft.
    """
    user = await get_current_user(request)
    _check_permission(user, "draft:document")
    paragraph_blocks = paragraphs_from_html(req.content_html)
    if not paragraph_blocks:
        raise HTTPException(status_code=422, detail="No content to export.")
    docx_bytes = await asyncio.to_thread(build_docx_bytes, paragraph_blocks)
    safe_filename = re.sub(r'[^\w\-. ]', '_', req.filename or "Document")
    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{safe_filename}.docx"'}
    )

# ── Calendar ──────────────────────────────────────────────────────────────────
#
# Visibility rule: a calendar event is visible only to its creator and to
# invited attendees who have accepted — not to the firm at large, and not
# to a role (e.g. partner) just by virtue of that role. See list_calendar
# and export_calendar_ics below, and _calendar_visibility_clause().
#
# attendees is a JSONB array of {email, name, user_id, status}. user_id is
# only stamped when the invited email matches an existing firm user's
# account email (_resolve_attendee_users) — an attendee who isn't a firm
# user (opposing counsel, a client, etc.) stays user_id=None/status=None
# and is purely an ICS-invite recipient, same as before this feature.
# status starts "pending" and only ever changes via
# POST /api/calendar/{id}/respond — there's no other way for it to become
# "accepted": the ICS Accept/Decline buttons an attendee's own mail client
# renders are cosmetic RSVP headers that go back to their mail server, not
# to this app.

async def _resolve_attendee_users(conn, firm_id, attendees: list) -> list:
    """
    Stamps user_id + status='pending' onto each attendee dict whose email
    matches an existing firm user's account email (case-insensitive).
    Non-matching attendees (external parties) are returned unchanged aside
    from explicit user_id=None/status=None keys, so every attendee has a
    consistent shape regardless of match.
    """
    if not attendees:
        return []
    emails = [a["email"].lower() for a in attendees if a.get("email")]
    rows = await conn.fetch(
        "SELECT id, email FROM users WHERE firm_id=$1 AND lower(email) = ANY($2::text[])",
        firm_id, emails,
    )
    by_email = {r["email"].lower(): str(r["id"]) for r in rows if r["email"]}
    resolved = []
    for a in attendees:
        matched_id = by_email.get((a.get("email") or "").lower())
        resolved.append({
            "email": a.get("email"), "name": a.get("name"),
            "user_id": matched_id, "status": "pending" if matched_id else None,
        })
    return resolved

def _calendar_visibility_clause(user: dict) -> tuple:
    """
    Returns (sql_fragment, params) implementing the visibility rule above:
    created_by = you, OR an attendee entry with your user_id and
    status='accepted'. A user with no real id (the synthetic
    AUTH_ENABLED=False dev user — see get_current_user) has no identity to
    scope by, so it falls back to the firm-wide view, matching how that
    synthetic user is already treated as a stand-in admin elsewhere.
    """
    user_id = user.get("id") if user else None
    if not user_id:
        return "firm_id=$1", [FIRM_ID]
    return (
        "firm_id=$1 AND (created_by=$2 OR EXISTS ("
        "SELECT 1 FROM jsonb_array_elements(attendees) att "
        "WHERE att->>'user_id' = $2::text AND att->>'status' = 'accepted'"
        "))",
        [FIRM_ID, _uuid_mod.UUID(str(user_id))],
    )

@app.get("/api/calendar")
async def list_calendar(request: Request):
    user = await get_current_user(request)
    _check_permission(user, "calendar:read")
    where, params = _calendar_visibility_clause(user)
    async with _db_pool.acquire() as conn:
        rows = await conn.fetch(
            f"SELECT * FROM calendar_events WHERE {where} ORDER BY date ASC, time ASC NULLS LAST",
            *params
        )
    return [_row_to_event(r) for r in rows]

@app.get("/api/calendar/pending-invites")
async def list_pending_calendar_invites(request: Request):
    """Events the current user is invited to but hasn't responded to yet —
    intentionally excluded from list_calendar above (not a confirmed event
    on your calendar until you accept), but still need to be discoverable
    somewhere so an invite isn't just silently unreachable."""
    user = await get_current_user(request)
    _check_permission(user, "calendar:read")
    user_id = user.get("id") if user else None
    if not user_id:
        return []
    async with _db_pool.acquire() as conn:
        rows = await conn.fetch(
            """
            SELECT * FROM calendar_events
            WHERE firm_id=$1 AND EXISTS (
                SELECT 1 FROM jsonb_array_elements(attendees) att
                WHERE att->>'user_id' = $2::text AND att->>'status' = 'pending'
            )
            ORDER BY date ASC, time ASC NULLS LAST
            """,
            FIRM_ID, str(user_id),
        )
    return [_row_to_event(r) for r in rows]

@app.post("/api/calendar/{event_id}/respond")
async def respond_to_calendar_invite(event_id: str, resp: CalendarInviteResponseRequest, request: Request):
    """Accept or decline your own invite on an event — the only thing that
    can ever move an attendee's status off 'pending' (see module note
    above). Matches by the current user's own account email against the
    event's attendee entries, so it also self-heals an attendee entry that
    predates this feature (no user_id stamped yet)."""
    user = await get_current_user(request)
    _check_permission(user, "calendar:read")
    if resp.status not in ("accepted", "declined"):
        raise HTTPException(status_code=422, detail="status must be 'accepted' or 'declined'")
    user_email = (user.get("email") or "").lower() if user else ""
    if not user_email:
        raise HTTPException(status_code=400, detail="Your account has no email on file to match against invites.")

    async with _db_pool.acquire() as conn:
        row = await conn.fetchrow(
            "SELECT * FROM calendar_events WHERE id=$1 AND firm_id=$2",
            _uuid_mod.UUID(event_id), FIRM_ID
        )
        if not row:
            raise HTTPException(status_code=404, detail="Event not found")

        event = _row_to_event(row)
        attendees = event.get("attendees") or []
        matched = False
        for a in attendees:
            if (a.get("email") or "").lower() == user_email:
                a["status"] = resp.status
                a["user_id"] = str(user["id"])
                matched = True
        if not matched:
            raise HTTPException(status_code=403, detail="You are not an attendee on this event.")

        updated_row = await conn.fetchrow(
            "UPDATE calendar_events SET attendees=$1::jsonb WHERE id=$2 RETURNING *",
            json.dumps(attendees), _uuid_mod.UUID(event_id)
        )
    return _row_to_event(updated_row)

@app.post("/api/calendar")
async def add_calendar_event(event: CalendarEvent, background_tasks: BackgroundTasks, request: Request):
    user = await get_current_user(request)
    _check_permission(user, "calendar:create")
    try:
        datetime.strptime(event.date, "%Y-%m-%d")
    except ValueError:
        raise HTTPException(status_code=422, detail=f"Invalid date format: {event.date}. Use YYYY-MM-DD.")

    matter_id_uuid = None
    if event.matter_id:
        try:
            matter_id_uuid = _uuid_mod.UUID(event.matter_id)
        except Exception:
            pass

    time_val = None
    if event.time:
        try:
            time_val = datetime.strptime(event.time, "%H:%M").time()
        except ValueError:
            pass

    attendees_list = [a.dict() for a in event.attendees] if event.attendees else []

    async with _db_pool.acquire() as conn:
        attendees_list = await _resolve_attendee_users(conn, FIRM_ID, attendees_list)
        row = await conn.fetchrow("""
            INSERT INTO calendar_events (firm_id, matter_id, title, date, time, event_type, court, matter_name, notes, attendees, created_by)
            VALUES ($1,$2,$3,$4,$5,$6,$7,$8,$9,$10::jsonb,$11) RETURNING *
        """,
        FIRM_ID, matter_id_uuid, event.title,
        datetime.strptime(event.date, "%Y-%m-%d").date(),
        time_val, event.event_type, event.court, event.matter_name, event.notes,
        json.dumps(attendees_list),
        _uuid_mod.UUID(str(user["id"])) if user and user.get("id") else None
        )
    result = _row_to_event(row)

    if attendees_list and is_email_configured():
        organizer_name = (user or {}).get("display_name") or FIRM_NAME
        background_tasks.add_task(send_event_invites, result, attendees_list, organizer_name, event.invite_message)

    return result

@app.post("/api/calendar/{event_id}/invite")
async def invite_to_calendar_event(event_id: str, req: CalendarInviteRequest, background_tasks: BackgroundTasks, request: Request):
    """
    Add one or more attendees to an existing event and send them a calendar
    invite. Covers the "just got off the phone, want to loop in another
    lawyer" case — no need to recreate the whole event.
    """
    user = await get_current_user(request)
    _check_permission(user, "calendar:create")

    async with _db_pool.acquire() as conn:
        row = await conn.fetchrow(
            "SELECT * FROM calendar_events WHERE id=$1 AND firm_id=$2",
            _uuid_mod.UUID(event_id), FIRM_ID
        )
        if not row:
            raise HTTPException(status_code=404, detail="Event not found")

        existing_event = _row_to_event(row)
        existing_attendees = existing_event.get("attendees") or []
        existing_emails = {a["email"].lower() for a in existing_attendees}

        new_attendees = [a.dict() for a in req.attendees if a.email.lower() not in existing_emails]
        if not new_attendees:
            return {"added": [], "message": "All provided attendees are already on this event"}
        new_attendees = await _resolve_attendee_users(conn, FIRM_ID, new_attendees)

        merged_attendees = existing_attendees + new_attendees
        updated_row = await conn.fetchrow(
            "UPDATE calendar_events SET attendees=$1::jsonb WHERE id=$2 RETURNING *",
            json.dumps(merged_attendees), _uuid_mod.UUID(event_id)
        )

    result = _row_to_event(updated_row)

    if is_email_configured():
        organizer_name = (user or {}).get("display_name") or FIRM_NAME
        background_tasks.add_task(send_event_invites, result, new_attendees, organizer_name, req.invite_message)

    return {"added": [a["email"] for a in new_attendees], "event": result}

@app.patch("/api/calendar/{event_id}")
async def update_calendar_event(event_id: str, update: CalendarEventUpdate,
                                 background_tasks: BackgroundTasks, request: Request):
    """
    Update an event — the postponement/rescheduling path. Only the fields
    provided are changed. If the event has attendees, they get an updated
    calendar invite (same UID, incremented SEQUENCE, so it replaces the
    existing entry on their calendar rather than creating a duplicate).
    """
    user = await get_current_user(request)
    _check_permission(user, "calendar:create")

    async with _db_pool.acquire() as conn:
        existing_row = await conn.fetchrow(
            "SELECT * FROM calendar_events WHERE id=$1 AND firm_id=$2",
            _uuid_mod.UUID(event_id), FIRM_ID
        )
        if not existing_row:
            raise HTTPException(status_code=404, detail="Event not found")

        existing = _row_to_event(existing_row)

        new_date = update.date
        if new_date:
            try:
                datetime.strptime(new_date, "%Y-%m-%d")
            except ValueError:
                raise HTTPException(status_code=422, detail=f"Invalid date format: {new_date}. Use YYYY-MM-DD.")

        new_time_val = None
        if update.time is not None:
            if update.time == "":
                new_time_val = None
            else:
                try:
                    new_time_val = datetime.strptime(update.time, "%H:%M").time()
                except ValueError:
                    raise HTTPException(status_code=422, detail=f"Invalid time format: {update.time}. Use HH:MM.")

        set_parts, values = [], []
        i = 1
        if update.title is not None:
            set_parts.append(f"title=${i}"); values.append(update.title); i += 1
        if new_date:
            set_parts.append(f"date=${i}"); values.append(datetime.strptime(new_date, "%Y-%m-%d").date()); i += 1
        if update.time is not None:
            set_parts.append(f"time=${i}"); values.append(new_time_val); i += 1
        if update.court is not None:
            set_parts.append(f"court=${i}"); values.append(update.court); i += 1
        if update.notes is not None:
            set_parts.append(f"notes=${i}"); values.append(update.notes); i += 1
        if update.event_type is not None:
            set_parts.append(f"event_type=${i}"); values.append(update.event_type); i += 1

        if not set_parts:
            raise HTTPException(status_code=400, detail="No fields to update")
        set_parts.append("sequence=sequence+1")
        values.append(_uuid_mod.UUID(event_id))
        updated_row = await conn.fetchrow(
            f"UPDATE calendar_events SET {', '.join(set_parts)} WHERE id=${i} RETURNING *",
            *values
        )

    result = _row_to_event(updated_row)
    attendees = result.get("attendees") or []

    if attendees and is_email_configured():
        organizer_name = (user or {}).get("display_name") or FIRM_NAME
        background_tasks.add_task(
            send_event_invites, result, attendees, organizer_name,
            update.update_message, "updated", result.get("sequence", 1)
        )

    return result

@app.delete("/api/calendar/{event_id}")
async def delete_calendar_event(event_id: str, background_tasks: BackgroundTasks, request: Request):
    user = await get_current_user(request)
    _check_permission(user, "calendar:delete")

    async with _db_pool.acquire() as conn:
        row = await conn.fetchrow(
            "SELECT * FROM calendar_events WHERE id=$1 AND firm_id=$2",
            _uuid_mod.UUID(event_id), FIRM_ID
        )
        if not row:
            raise HTTPException(status_code=404, detail="Event not found")

        event = _row_to_event(row)
        result = await conn.execute(
            "DELETE FROM calendar_events WHERE id=$1 AND firm_id=$2",
            _uuid_mod.UUID(event_id), FIRM_ID
        )

    if result == "DELETE 0":
        raise HTTPException(status_code=404, detail="Event not found")

    attendees = event.get("attendees") or []
    if attendees and is_email_configured():
        organizer_name = (user or {}).get("display_name") or FIRM_NAME
        # Cancellation notice — the row is already gone from our DB, but the
        # attendee's own calendar app still has it until we tell it to
        # remove/cancel the entry via METHOD:CANCEL.
        background_tasks.add_task(
            send_event_invites, event, attendees, organizer_name,
            None, "cancelled", event.get("sequence", 0) + 1
        )

    return {"deleted": True, "notified": [a["email"] for a in attendees]}

@app.get("/api/calendar/export-ics")
async def export_calendar_ics(request: Request):
    user = await get_current_user(request)
    _check_permission(user, "calendar:read")
    where, params = _calendar_visibility_clause(user)
    async with _db_pool.acquire() as conn:
        rows = await conn.fetch(
            f"SELECT * FROM calendar_events WHERE {where} ORDER BY date ASC",
            *params
        )
    events = [_row_to_event(r) for r in rows]
    lines = [
        "BEGIN:VCALENDAR",
        "VERSION:2.0",
        f"PRODID:-//Mutemo Desk//{FIRM_NAME}//EN",
        "CALSCALE:GREGORIAN",
        "METHOD:PUBLISH",
    ]
    for ev in events:
        uid = f"{ev['id']}@mutemodesk"
        dtstart = ev["date"].replace("-", "")
        if ev.get("time"):
            t = ev["time"].replace(":", "")
            dtstart = f"{dtstart}T{t}00"
        summary = ev["title"].replace(",", "\\,").replace(";", "\\;")
        desc = ""
        if ev.get("matter_name"):
            desc += f"Matter: {ev['matter_name']}\\n"
        if ev.get("court"):
            desc += f"Court: {ev['court']}\\n"
        if ev.get("notes"):
            desc += f"Notes: {ev['notes']}"
        lines += [
            "BEGIN:VEVENT",
            f"UID:{uid}",
            f"DTSTART:{dtstart}",
            f"SUMMARY:{summary}",
            f"DESCRIPTION:{desc}",
            "END:VEVENT",
        ]
    lines.append("END:VCALENDAR")
    ics_content = "\r\n".join(lines) + "\r\n"
    from fastapi.responses import Response as FastAPIResponse
    return FastAPIResponse(
        content=ics_content,
        media_type="text/calendar",
        headers={"Content-Disposition": f'attachment; filename="mutemo_calendar.ics"'}
    )

# ── Date Extraction from Documents ────────────────────────────────────────────

@app.post("/api/extract-dates")
async def extract_dates_from_document(
    file: UploadFile = File(...),
    matter_id: Optional[str] = Form(None),
    matter_name: Optional[str] = Form(None),
    request: Request = None,
):
    if request:
        user = await get_current_user(request)
        _check_permission(user, "calendar:create")

    content = await file.read()
    filename = file.filename or "document"
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else "bin"

    text = ""
    try:
        if ext == "pdf":
            text, _, _ = await asyncio.to_thread(extract_pdf_text, content)
        elif ext in ("docx", "doc"):
            text = await asyncio.to_thread(extract_docx_text, content)
        else:
            text = content.decode("utf-8", errors="replace")
    except Exception as e:
        raise HTTPException(status_code=422, detail=f"Could not extract text: {e}")

    if not text:
        raise HTTPException(status_code=422, detail="No text could be extracted from this document.")

    today = datetime.utcnow().date().isoformat()

    def extract_dates_sync():
        msg = client.messages.create(
            model="claude-sonnet-4-5",
            max_tokens=1500,
            messages=[{"role": "user", "content": f"""Extract all legal deadlines, hearing dates, filing dates, and appointments from this document.
Today is {today}. Focus on specific, actionable dates.

Return ONLY valid JSON:
{{
  "dates": [
    {{
      "title": "brief description",
      "date": "YYYY-MM-DD",
      "time": "HH:MM or null",
      "event_type": "deadline|hearing|meeting|filing|other",
      "party": "which party this applies to, or null",
      "notes": "any additional context"
    }}
  ],
  "document_summary": "one sentence summary of the document"
}}

Document text (first 8000 chars):
{text[:8000]}

JSON:"""}]
        )
        raw = msg.content[0].text.strip()
        raw = re.sub(r'^```json\s*|\s*```$', '', raw, flags=re.MULTILINE).strip()
        return json.loads(raw)

    try:
        parsed = await asyncio.to_thread(extract_dates_sync)
        dates = parsed.get("dates", [])
        summary = parsed.get("document_summary", "")
        for d in dates:
            d["matter_id"] = matter_id
            d["matter_name"] = matter_name
            d["source_document"] = filename
        return {"dates": dates, "count": len(dates), "document_summary": summary, "filename": filename}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Date extraction failed: {e}")


@app.post("/api/extract-dates-by-id")
async def extract_dates_by_document_id(
    request: Request,
    document_id: str = Form(...),
):
    user = await get_current_user(request)
    _check_permission(user, "calendar:create")

    # Get document record
    async with _db_pool.acquire() as conn:
        doc = await conn.fetchrow(
            "SELECT * FROM documents WHERE id=$1 AND firm_id=$2",
            _uuid_mod.UUID(document_id), FIRM_ID
        )
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    # Get chunks for this document
    async with _db_pool.acquire() as conn:
        chunk_rows = await conn.fetch(
            "SELECT text FROM chunks WHERE document_id=$1 ORDER BY chunk_index ASC",
            _uuid_mod.UUID(document_id)
        )

    if not chunk_rows:
        raise HTTPException(status_code=422, detail="No text content found for this document. It may still be processing.")

    text = " ".join(r["text"] for r in chunk_rows)
    filename = doc["filename"]
    matter_id = str(doc["matter_id"]) if doc["matter_id"] else None

    # Get matter name
    matter_name = None
    if matter_id:
        async with _db_pool.acquire() as conn:
            matter = await conn.fetchrow("SELECT name FROM matters WHERE id=$1", doc["matter_id"])
            if matter:
                matter_name = matter["name"]

    today = datetime.utcnow().date().isoformat()

    def extract_dates_sync():
        msg = client.messages.create(
            model="claude-sonnet-4-5",
            max_tokens=1500,
            messages=[{"role": "user", "content": f"""Extract all legal deadlines, hearing dates, filing dates, and appointments from this document.
Today is {today}. Focus on specific, actionable dates.

Return ONLY valid JSON:
{{
  "dates": [
    {{
      "title": "brief description",
      "date": "YYYY-MM-DD",
      "time": "HH:MM or null",
      "event_type": "deadline|hearing|meeting|filing|other",
      "party": "which party this applies to, or null",
      "notes": "any additional context"
    }}
  ],
  "document_summary": "one sentence summary of the document"
}}

Document text:
{text[:8000]}

JSON:"""}]
        )
        raw = msg.content[0].text.strip()
        raw = re.sub(r'^```json\s*|\s*```$', '', raw, flags=re.MULTILINE).strip()
        return json.loads(raw)

    try:
        parsed = await asyncio.to_thread(extract_dates_sync)
        dates = parsed.get("dates", [])
        summary = parsed.get("document_summary", "")
        for d in dates:
            d["matter_id"] = matter_id
            d["matter_name"] = matter_name
            d["source_document"] = filename
        return {"dates": dates, "count": len(dates), "document_summary": summary, "filename": filename}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Date extraction failed: {e}")

# ── Email / Reminders ─────────────────────────────────────────────────────────

EVENT_TYPE_LABELS = {
    "hearing": "Court Hearing",
    "deadline": "Deadline / Dies",
    "filing": "Filing Date",
    "round_table": "Round Table Conference",
    "ptc": "Pre-Trial Conference",
    "mediation": "Mediation",
    "meeting": "Client Meeting",
    "consultation": "Consultation",
    "signing": "Document Signing",
    "call": "Client Call",
    "speaking": "Speaking Engagement",
    "radio_tv": "Radio / TV Programme",
    "cpd": "CPD / Training",
    "lsz": "Law Society Event",
    "board": "Board Meeting",
    "staff_meeting": "Staff Meeting",
    "leave": "Leave",
    "billing_review": "Billing Review",
    "pro_bono": "Pro Bono",
    "networking": "Networking",
    "chamber_application": "Chamber Application",
    "judgment_awaiting": "Judgment Awaiting",
    "social": "Social Event",
    "church": "Church / Ministry",
    "personal": "Personal",
    "manual": "Appointment",
    "other": "Event",
}

def _escape_html(s: str) -> str:
    return (s or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

_CASE_REF_SEPARATOR = " — "

def extract_case_reference(text: Optional[str]) -> Optional[str]:
    """
    Pulls a leading court/case reference out of matter free text, following
    the onboarding template's "Reference/Case No. — description" convention
    (see bulk_onboard_from_excel above). There's no dedicated case-number
    column on matters — number/internal_ref/external_ref are all generic
    free-text refs used for other things — so this parses it out of the
    text instead of adding one.

    Requires a digit in the candidate: not every matter's text actually
    follows the convention (some pre-date it, or use the leading segment as
    a party/description label instead of a reference), so a purely
    alphabetic leading segment is treated as "no reference found" rather
    than displayed as if it were one.
    """
    if not text or _CASE_REF_SEPARATOR not in text:
        return None
    head = text.split(_CASE_REF_SEPARATOR, 1)[0].strip()
    if not head or not any(ch.isdigit() for ch in head):
        return None
    return head

def _matter_identity_prefix(e: dict) -> str:
    """
    Builds the "{matter_number} ({case_number}) — {client_name}" lead-in for
    a reminder line, e.g. "NGM-007-02 (HC 300/26) — Huang Li Qiang". Any
    piece that isn't available is omitted gracefully (never a literal
    "None" or a dangling separator) — returns "" if nothing is available at
    all, letting the caller fall back to the event's own display text.
    """
    parts = []
    if e.get("matter_number"):
        bit = e["matter_number"]
        if e.get("case_number"):
            bit += f" ({e['case_number']})"
        parts.append(bit)
    elif e.get("case_number"):
        parts.append(f"({e['case_number']})")
    if e.get("resolved_client_name"):
        parts.append(e["resolved_client_name"])
    return " — ".join(parts)

def build_reminder_email_body(events: list, review_matters: Optional[list] = None) -> tuple:
    """
    Returns (plain_text_body, html_body). Events must have a 'days_until'
    field; review_matters (matter review safety net, 2026-08-30) must too.

    review_matters is rendered as its own clearly-labeled section, kept
    deliberately separate from the events sections above \u2014 these are a
    soft "please look at this" nudge (a matter that's gone quiet), not a
    hard court/filing deadline, and mixing the two would blur a
    distinction lawyers actually rely on.
    """
    review_matters = review_matters or []
    if not events and not review_matters:
        text = "Good morning. You have no court dates, deadlines, or filings scheduled in the next 30 days.\n\n\u2014 Mutemo Desk"
        html = "<p>Good morning. You have no court dates, deadlines, or filings scheduled in the next 30 days.</p><p style='color:#6b6b64'>\u2014 Mutemo Desk</p>"
        return text, html

    today_items    = [e for e in events if e.get("days_until", 99) == 0]
    tomorrow_items = [e for e in events if e.get("days_until", 99) == 1]
    week_items     = [e for e in events if 1 < (e.get("days_until") or 0) <= 7]
    later_items    = [e for e in events if (e.get("days_until") or 0) > 7]

    def fmt_text(e):
        bits = [EVENT_TYPE_LABELS.get(e.get("event_type"), "Event") + ":", e.get("title", "")]
        if e.get("time"):  bits.append(f"at {e['time']}")
        if e.get("court"): bits.append(f"\u2014 {e['court']}")
        body = " ".join(bits)
        prefix = _matter_identity_prefix(e)
        if prefix:
            return f"{prefix}: {body}"
        if e.get("matter_name"):  # no linked matter/client \u2014 old trailing display, unchanged
            return f"{body} ({e['matter_name']})"
        return body

    def fmt_review_text(m):
        prefix = _matter_identity_prefix(m) or m.get("name", "")
        du = m.get("days_until", 0)
        status = f"overdue by {abs(du)}d" if du < 0 else ("due today" if du == 0 else f"due in {du}d")
        last_reviewed = f", last reviewed {m['last_reviewed_date']}" if m.get("last_reviewed_date") else ""
        return f"{prefix} \u2014 next review {m.get('next_review_date')} ({status}{last_reviewed})"

    text_lines = ["Good morning. Here is your Mutemo Desk reminder summary:\n"]
    if today_items:
        text_lines.append("TODAY:")
        for e in today_items: text_lines.append(f"  \u2022 {fmt_text(e)}")
        text_lines.append("")
    if tomorrow_items:
        text_lines.append("TOMORROW:")
        for e in tomorrow_items: text_lines.append(f"  \u2022 {fmt_text(e)}")
        text_lines.append("")
    if week_items:
        text_lines.append("LATER THIS WEEK:")
        for e in week_items: text_lines.append(f"  {e['date']}  \u2014  {fmt_text(e)}")
        text_lines.append("")
    if later_items:
        text_lines.append("COMING UP:")
        for e in later_items: text_lines.append(f"  {e['date']}  \u2014  {fmt_text(e)}")
        text_lines.append("")
    if review_matters:
        text_lines.append("MATTERS FOR REVIEW (no update in a while \u2014 not a hard deadline):")
        for m in review_matters: text_lines.append(f"  \u2022 {fmt_review_text(m)}")
        text_lines.append("")
    text_lines.append("A calendar file (.ics) is attached \u2014 open it to add these to your phone or computer calendar.")
    text_lines.append("\n\u2014 Mutemo Desk")
    text = "\n".join(text_lines)

    def fmt_html(e):
        type_chip = EVENT_TYPE_LABELS.get(e.get("event_type"), "Event")
        meta = []
        if e.get("time"):  meta.append(e["time"])
        if e.get("court"): meta.append(_escape_html(e["court"]))
        prefix = _matter_identity_prefix(e)
        if prefix:
            title_html = f'{_escape_html(prefix)}: {_escape_html(e.get("title",""))}'
        else:
            if e.get("matter_name"):  # no linked matter/client \u2014 old meta-line display, unchanged
                meta.append(_escape_html(e["matter_name"]))
            title_html = _escape_html(e.get("title", ""))
        meta_str = " \u00b7 ".join(meta)
        return (
            f'<div style="padding:8px 0;border-bottom:1px solid #e8e4da">'
            f'<span style="font-size:11px;font-weight:700;color:#b8922a;text-transform:uppercase;letter-spacing:0.5px">{type_chip}</span><br/>'
            f'<strong>{title_html}</strong><br/>'
            f'<span style="font-size:13px;color:#6b6b64">{meta_str}</span>'
            f'</div>'
        )

    html_sections = []
    if today_items:
        html_sections.append('<h3 style="color:#b83232;margin:16px 0 8px">Today</h3>' + "".join(fmt_html(e) for e in today_items))
    if tomorrow_items:
        html_sections.append('<h3 style="color:#b8922a;margin:16px 0 8px">Tomorrow</h3>' + "".join(fmt_html(e) for e in tomorrow_items))
    if week_items:
        html_sections.append(
            '<h3 style="color:#1b4d2e;margin:16px 0 8px">Later This Week</h3>' +
            "".join(
                f'<div style="padding:8px 0;border-bottom:1px solid #e8e4da">'
                f'<span style="font-size:12px;color:#6b6b64">{e["date"]}</span><br/>{fmt_html(e)}</div>'
                for e in week_items
            )
        )
    if later_items:
        html_sections.append(
            '<h3 style="color:#1b4d2e;margin:16px 0 8px">Coming Up</h3>' +
            "".join(
                f'<div style="padding:8px 0;border-bottom:1px solid #e8e4da">'
                f'<span style="font-size:12px;color:#6b6b64">{e["date"]}</span><br/>{fmt_html(e)}</div>'
                for e in later_items
            )
        )
    if review_matters:
        # Deliberately styled distinctly from the deadline sections above
        # (muted blue-grey, not the red/gold urgency palette) — a soft
        # nudge, not a hard deadline. Kept as its own clearly-labeled
        # section rather than merged into the list above.
        def fmt_review_html(m):
            prefix = _matter_identity_prefix(m) or _escape_html(m.get("name", ""))
            du = m.get("days_until", 0)
            status = f"overdue by {abs(du)}d" if du < 0 else ("due today" if du == 0 else f"due in {du}d")
            last_reviewed = f", last reviewed {m['last_reviewed_date']}" if m.get("last_reviewed_date") else ""
            return (
                f'<div style="padding:8px 0;border-bottom:1px solid #e8e4da">'
                f'<strong>{_escape_html(prefix)}</strong><br/>'
                f'<span style="font-size:13px;color:#6b6b64">Next review: {m.get("next_review_date")} ({status}{last_reviewed})</span>'
                f'</div>'
            )
        html_sections.append(
            '<h3 style="color:#4a5a6b;margin:16px 0 8px">\U0001F5C2 Matters for Review</h3>'
            '<p style="font-size:12px;color:#6b6b64;margin:0 0 8px">No update in a while — a nudge, not a hard deadline.</p>' +
            "".join(fmt_review_html(m) for m in review_matters)
        )

    html = f"""<div style="font-family:Georgia,serif;color:#1a1a18;max-width:560px">
        <div style="background:#1b4d2e;color:white;padding:16px 20px;border-radius:6px 6px 0 0">
            <strong style="font-size:18px">&#9878; Mutemo Desk</strong><br/>
            <span style="font-size:13px;opacity:0.8">Daily Calendar Reminder &mdash; {FIRM_NAME}</span>
        </div>
        <div style="padding:16px 20px;border:1px solid #d8d3c8;border-top:none;border-radius:0 0 6px 6px">
            <p>Good morning. Here is your reminder summary for the next 30 days.</p>
            {''.join(html_sections)}
            <p style="margin-top:16px;font-size:13px;color:#6b6b64">A calendar file (.ics) is attached &mdash; open it to add these events to your phone or computer calendar.</p>
        </div>
    </div>"""

    return text, html

async def get_firm_contact_email() -> Optional[str]:
    """
    The email address the firm already set up on the Calendar tab for daily
    reminders (reminder_settings.recipient_email). Reused as the calendar
    invite ORGANIZER/CC/Reply-To so Accept/Decline responses and any
    "reply" from an attendee actually reach the firm, not a generic system
    inbox nobody reads.
    """
    if not _db_pool:
        return None
    async with _db_pool.acquire() as conn:
        row = await conn.fetchrow(
            "SELECT recipient_email FROM reminder_settings WHERE firm_id=$1", FIRM_ID
        )
    email = row["recipient_email"] if row else None
    return email.strip() if email and email.strip() else None

def build_ics(events: list) -> str:
    """Build an ICS calendar string from a list of event dicts."""
    lines = [
        "BEGIN:VCALENDAR", "VERSION:2.0",
        f"PRODID:-//Mutemo Desk//{FIRM_NAME}//EN",
        "CALSCALE:GREGORIAN", "METHOD:PUBLISH",
    ]
    for ev in events:
        uid = f"{ev.get('id', ev.get('title','evt'))}@mutemodesk"
        dtstart = str(ev["date"]).replace("-", "")
        if ev.get("time"):
            t = str(ev["time"]).replace(":", "")[:4]
            dtstart = f"{dtstart}T{t}00"
        summary = str(ev.get("title", "")).replace(",", "\\,").replace(";", "\\;")
        lines += ["BEGIN:VEVENT", f"UID:{uid}", f"DTSTART:{dtstart}", f"SUMMARY:{summary}", "END:VEVENT"]
    lines.append("END:VCALENDAR")
    return "\r\n".join(lines) + "\r\n"

def _ics_escape(s: str) -> str:
    return (s or "").replace("\\", "\\\\").replace(",", "\\,").replace(";", "\\;").replace("\n", "\\n")

def build_invite_ics(event: dict, attendees: list, organizer_name: str, organizer_email: str,
                      method: str = "REQUEST", sequence: int = 0) -> str:
    """
    Build a calendar invite (METHOD:REQUEST for new/updated events,
    METHOD:CANCEL to cancel one) with an ORGANIZER and one ATTENDEE line per
    invitee. This is what makes Gmail/Outlook render Accept/Decline buttons —
    build_ics() above uses METHOD:PUBLISH instead, which is right for a
    personal reminder digest but doesn't get treated as an invite by mail
    clients.

    UID must stay identical across create/update/cancel for the same event
    so calendar clients treat them as the same entry rather than duplicates.
    SEQUENCE must increase on every update/cancel — clients use UID+SEQUENCE
    together to know a later message supersedes an earlier one.
    """
    uid = f"{event.get('id', 'evt')}@mutemodesk"
    now_stamp = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")

    dtstart = str(event["date"]).replace("-", "")
    has_time = bool(event.get("time"))
    if has_time:
        t = str(event["time"]).replace(":", "")[:4]
        dtstart = f"{dtstart}T{t}00"
        # Default 1-hour duration — the app doesn't track an explicit end time.
        start_dt = datetime.strptime(f"{event['date']} {event['time']}", "%Y-%m-%d %H:%M")
        end_dt = start_dt + timedelta(hours=1)
        dtend = end_dt.strftime("%Y%m%dT%H%M00")
    else:
        dtend = None

    summary = _ics_escape(event.get("title", ""))
    desc_parts = []
    if event.get("matter_name"):
        desc_parts.append(f"Matter: {event['matter_name']}")
    if event.get("notes"):
        desc_parts.append(f"Notes: {event['notes']}")
    description = _ics_escape("\\n".join(desc_parts))
    location = _ics_escape(event.get("court", "") or "")
    status = "CANCELLED" if method == "CANCEL" else "CONFIRMED"

    lines = [
        "BEGIN:VCALENDAR",
        "VERSION:2.0",
        f"PRODID:-//Mutemo Desk//{FIRM_NAME}//EN",
        "CALSCALE:GREGORIAN",
        f"METHOD:{method}",
        "BEGIN:VEVENT",
        f"UID:{uid}",
        f"DTSTAMP:{now_stamp}",
        f"DTSTART:{dtstart}",
    ]
    if dtend:
        lines.append(f"DTEND:{dtend}")
    lines += [
        f"SUMMARY:{summary}",
        f"DESCRIPTION:{description}",
        f"LOCATION:{location}",
        f"SEQUENCE:{sequence}",
        f"STATUS:{status}",
        f'ORGANIZER;CN="{_ics_escape(organizer_name)}":MAILTO:{organizer_email}',
    ]
    for a in attendees:
        name = a.get("name") or a.get("email")
        partstat = "NEEDS-ACTION" if method != "CANCEL" else "DECLINED"
        lines.append(
            f'ATTENDEE;CN="{_ics_escape(name)}";ROLE=REQ-PARTICIPANT;'
            f'PARTSTAT={partstat};RSVP=TRUE:MAILTO:{a["email"]}'
        )
    lines += ["END:VEVENT", "END:VCALENDAR"]
    return "\r\n".join(lines) + "\r\n"

def _send_via_resend_sync_with_method(to: str, subject: str, html_body: str, text_body: str,
                                       ics_content: str, method: str = "REQUEST",
                                       cc: Optional[str] = None, reply_to: Optional[str] = None) -> None:
    """
    Same as _send_via_resend_sync but sets the calendar attachment's
    content_type explicitly to text/calendar with the given method.

    This is required, not optional: without it, Resend infers the
    attachment's MIME type from the filename alone, which does not include
    the `method=REQUEST` parameter that Outlook specifically requires to
    render Accept/Decline buttons, and that Gmail also wants for full RSVP
    support rather than just silently adding the event.

    cc/reply_to are the firm's own contact email (reminder_settings.
    recipient_email) — CC'd so there's a visible record it went out, and
    Reply-To so a plain "Reply" in the recipient's mail client reaches the
    firm rather than a generic system inbox.
    """
    import base64
    api_key = os.environ.get("RESEND_API_KEY", "")
    if not api_key:
        raise RuntimeError("RESEND_API_KEY not configured")
    from_addr = os.environ.get("RESEND_FROM", f"reminders@{os.environ.get('RESEND_FROM_DOMAIN', 'tofamba.com')}")
    payload = {
        "from": f"Mutemo Desk <{from_addr}>",
        "to": [to],
        "subject": subject,
        "html": html_body,
        "text": text_body,
        "attachments": [{
            "filename": "invite.ics",
            "content": base64.b64encode(ics_content.encode("utf-8")).decode("utf-8"),
            "content_type": f'text/calendar; charset=utf-8; method={method}',
        }],
    }
    if cc:
        payload["cc"] = [cc]
    if reply_to:
        payload["reply_to"] = reply_to
    import httpx
    with httpx.Client(timeout=15) as http:
        resp = http.post(
            "https://api.resend.com/emails",
            json=payload,
            headers={"Authorization": f"Bearer {api_key}"},
        )
        if resp.status_code not in (200, 201):
            raise RuntimeError(f"Resend API error {resp.status_code}: {resp.text}")

async def send_event_invites(event: dict, attendees: list, organizer_name: str,
                              invite_message: Optional[str] = None,
                              kind: str = "new", sequence: int = 0) -> dict:
    """
    Send a calendar invite/update/cancellation email (with .ics attachment)
    to each attendee on an event. Sends one email per attendee — each ICS
    carries the full attendee list, so recipients' calendar apps still show
    the full guest list even though the email itself is addressed
    individually.

    kind: "new" (initial invite), "updated" (date/time/location changed —
    the .ics UID stays the same so it updates the existing calendar entry
    rather than creating a duplicate), or "cancelled" (event removed —
    sends METHOD:CANCEL so it's struck through / removed on the attendee's
    calendar, not just left dangling).

    invite_message is an optional free-text note from whoever is sending the
    invite, giving the recipient context — shown under the "From the Desk
    of ..." heading.
    Returns {"sent": [...], "failed": [...]}.
    """
    firm_contact_email = await get_firm_contact_email()
    organizer_email = firm_contact_email or os.environ.get(
        "RESEND_FROM", f"reminders@{os.environ.get('RESEND_FROM_DOMAIN', 'tofamba.com')}"
    )
    ics_method = "CANCEL" if kind == "cancelled" else "REQUEST"
    ics_content = build_invite_ics(event, attendees, organizer_name, organizer_email,
                                   method=ics_method, sequence=sequence)

    date_str = event.get("date", "")
    time_str = f" at {event['time']}" if event.get("time") else ""
    location_line = f"<p><strong>Location:</strong> {event['court']}</p>" if event.get("court") else ""
    notes_line = f"<p><strong>Notes:</strong> {event['notes']}</p>" if event.get("notes") else ""
    message_html = (
        f'<div style="background:#f5f3ee;border-left:3px solid #8a7a5c;padding:10px 14px;margin:12px 0;font-size:14px">'
        f'{_escape_html(invite_message)}</div>'
        if invite_message else ""
    )
    message_text = f"\n\"{invite_message}\"\n" if invite_message else ""
    contact_html = (
        f'<p style="font-size:13px">Should you need to discuss this further, please feel free to '
        f'contact me on {_escape_html(firm_contact_email)}.</p>'
        if firm_contact_email and kind != "cancelled" else ""
    )
    contact_text = (
        f"\nShould you need to discuss this further, please feel free to contact me on {firm_contact_email}.\n"
        if firm_contact_email and kind != "cancelled" else ""
    )

    if kind == "cancelled":
        lead_line = "The following event has been <strong>cancelled</strong>:"
        lead_text = "The following event has been CANCELLED:"
        subject = f"Cancelled: {event.get('title', 'Event')} — {date_str}{time_str}"
    elif kind == "updated":
        lead_line = "The following event has been <strong>updated</strong> — please check the new details:"
        lead_text = "The following event has been UPDATED — please check the new details:"
        subject = f"Updated: {event.get('title', 'Event')} — {date_str}{time_str}"
    else:
        lead_line = "You've been invited to the following:"
        lead_text = "You've been invited to:"
        subject = f"Invitation: {event.get('title', 'Event')} — {date_str}{time_str}"

    html_body = f"""
        <p style="font-size:15px;font-weight:600;margin-bottom:2px">From the Desk of {_escape_html(organizer_name)}</p>
        <p style="color:#6b6b64;font-size:13px;margin-top:0">{lead_line}</p>
        <p><strong>{_escape_html(event.get('title', ''))}</strong></p>
        <p><strong>Date:</strong> {date_str}{time_str}</p>
        {location_line}
        {notes_line}
        {message_html}
        {contact_html}
        <p style="color:#6b6b64;font-size:13px">Sent via Mutemo Desk.{' Open the attached invite to add this to your calendar.' if kind != 'cancelled' else ' This will remove or mark the event as cancelled on your calendar if you added it previously.'}</p>
    """
    text_body = (
        f"From the Desk of {organizer_name}\n\n"
        f"{lead_text} {event.get('title', '')}\n"
        f"Date: {date_str}{time_str}\n"
        + (f"Location: {event['court']}\n" if event.get("court") else "")
        + (f"Notes: {event['notes']}\n" if event.get("notes") else "")
        + message_text
        + contact_text
    )

    sent, failed = [], []
    for a in attendees:
        try:
            await asyncio.to_thread(
                _send_via_resend_sync_with_method, a["email"], subject, html_body, text_body,
                ics_content, ics_method, firm_contact_email, firm_contact_email
            )
            sent.append(a["email"])
        except Exception as e:
            print(f"[calendar-invite] failed to send to {a.get('email')}: {e}")
            failed.append(a["email"])
    return {"sent": sent, "failed": failed}

def is_email_configured() -> bool:
    return bool(os.environ.get("RESEND_API_KEY") or os.environ.get("SMTP_HOST"))

def _send_via_resend_sync(to: str, subject: str, html_body: str, text_body: str, ics_content: str = None) -> None:
    """Synchronous Resend send (called via asyncio.to_thread)."""
    import base64
    api_key = os.environ.get("RESEND_API_KEY", "")
    if not api_key:
        raise RuntimeError("RESEND_API_KEY not configured")
    from_addr = os.environ.get("RESEND_FROM", f"reminders@{os.environ.get('RESEND_FROM_DOMAIN', 'tofamba.com')}")
    payload = {
        "from": f"Mutemo Desk <{from_addr}>",
        "to": [to],
        "subject": subject,
        "html": html_body,
        "text": text_body,
    }
    if ics_content:
        payload["attachments"] = [{
            "filename": "mutemo-events.ics",
            "content": base64.b64encode(ics_content.encode("utf-8")).decode("utf-8"),
        }]
    import httpx
    with httpx.Client(timeout=15) as http:
        resp = http.post(
            "https://api.resend.com/emails",
            json=payload,
            headers={"Authorization": f"Bearer {api_key}"},
        )
        if resp.status_code not in (200, 201):
            raise RuntimeError(f"Resend API error {resp.status_code}: {resp.text}")

async def send_reminder_email(recipient: str, events: list, test: bool = False, review_matters: Optional[list] = None) -> bool:
    """Send rich HTML daily calendar reminder via Resend with ICS attachment."""
    review_matters = review_matters or []
    text_body, html_body = build_reminder_email_body(events, review_matters)
    if test:
        text_body = "[TEST EMAIL]\n\n" + text_body
        html_body = '<p style="background:#fdf6e8;padding:8px;border-radius:4px;font-size:13px"><strong>This is a test email.</strong></p>' + html_body
    subject_prefix = "[TEST] " if test else ""
    if any(e.get("days_until") == 0 for e in events):
        subject = f"{subject_prefix}\u2696 Mutemo Desk \u2014 Court date TODAY + upcoming"
    elif events:
        subject = f"{subject_prefix}\u2696 Mutemo Desk \u2014 Daily reminder ({len(events)} upcoming)"
    elif review_matters:
        # No hard deadlines, but there's still something worth a look \u2014
        # don't say "nothing upcoming" when there genuinely is something.
        subject = f"{subject_prefix}\u2696 Mutemo Desk \u2014 Daily reminder ({len(review_matters)} for review)"
    else:
        subject = f"{subject_prefix}\u2696 Mutemo Desk \u2014 Daily reminder (nothing upcoming)"
    ics_content = build_ics(events) if events else None
    try:
        await asyncio.to_thread(_send_via_resend_sync, recipient, subject, html_body, text_body, ics_content)
        return True
    except Exception as e:
        print(f"[email] send failed: {e}")
        return False

@app.get("/api/reminders/settings")
async def get_reminder_settings(request: Request):
    user = await get_current_user(request)
    _check_permission(user, "admin:settings")
    async with _db_pool.acquire() as conn:
        row = await conn.fetchrow("SELECT * FROM reminder_settings WHERE firm_id=$1", FIRM_ID)
    if not row:
        return {"enabled": False, "recipient_email": "", "send_hour_utc": 5}
    d = dict(row)
    d["firm_id"] = str(d["firm_id"])
    if d.get("last_run_date"):
        d["last_run_date"] = str(d["last_run_date"])
    return d

@app.post("/api/reminders/settings")
async def update_reminder_settings(settings: ReminderSettings, request: Request):
    user = await get_current_user(request)
    _check_permission(user, "admin:settings")
    async with _db_pool.acquire() as conn:
        await conn.execute("""
            INSERT INTO reminder_settings (firm_id, enabled, recipient_email, send_hour_utc)
            VALUES ($1,$2,$3,$4)
            ON CONFLICT (firm_id) DO UPDATE SET
                enabled=$2, recipient_email=$3, send_hour_utc=$4
        """, FIRM_ID, settings.enabled, settings.recipient_email, settings.send_hour_utc)
    return {"saved": True}

@app.get("/api/digest/settings")
async def get_digest_settings(request: Request):
    user = await get_current_user(request)
    _check_permission(user, "admin:settings")
    async with _db_pool.acquire() as conn:
        row = await conn.fetchrow("SELECT * FROM reminder_settings WHERE firm_id=$1", FIRM_ID)
    if not row:
        return {"enabled": False, "recipient_email": "", "send_hour_utc": 6}
    return {
        "enabled": row["digest_enabled"],
        "recipient_email": row["digest_recipient_email"] or "",
        "send_hour_utc": row["digest_send_hour_utc"],
        "last_run_date": str(row["digest_last_run_date"]) if row["digest_last_run_date"] else None,
    }

@app.post("/api/digest/settings")
async def update_digest_settings(settings: DigestSettings, request: Request):
    user = await get_current_user(request)
    _check_permission(user, "admin:settings")
    async with _db_pool.acquire() as conn:
        await conn.execute("""
            INSERT INTO reminder_settings (firm_id, digest_enabled, digest_recipient_email, digest_send_hour_utc)
            VALUES ($1,$2,$3,$4)
            ON CONFLICT (firm_id) DO UPDATE SET
                digest_enabled=$2, digest_recipient_email=$3, digest_send_hour_utc=$4
        """, FIRM_ID, settings.enabled, settings.recipient_email, settings.send_hour_utc)
    return {"saved": True}

@app.post("/api/digest/test")
async def test_digest(request: Request):
    user = await get_current_user(request)
    _check_permission(user, "admin:settings")
    if not is_email_configured():
        raise HTTPException(status_code=503, detail="Email is not configured on this server.")
    async with _db_pool.acquire() as conn:
        row = await conn.fetchrow("SELECT * FROM reminder_settings WHERE firm_id=$1", FIRM_ID)
    if not row or not row["digest_recipient_email"]:
        raise HTTPException(status_code=400, detail="No digest recipient email configured.")

    since = datetime.utcnow() - timedelta(hours=24)
    async with _db_pool.acquire() as conn:
        news_rows = await conn.fetch(
            "SELECT * FROM legal_updates WHERE firm_id=$1 AND source_type='news' AND uploaded_at >= $2 ORDER BY uploaded_at DESC",
            FIRM_ID, since
        )
        legislation_rows = await conn.fetch(
            "SELECT * FROM legal_updates WHERE firm_id=$1 AND source_type='legislation' AND uploaded_at >= $2 ORDER BY uploaded_at DESC",
            FIRM_ID, since
        )
        judgment_rows = await conn.fetch(
            "SELECT * FROM zlr_entries WHERE firm_id=$1 AND uploaded_at >= $2 ORDER BY uploaded_at DESC",
            FIRM_ID, since
        )
    news_items = [dict(r) for r in news_rows]
    legislation_items = [dict(r) for r in legislation_rows]
    judgment_items = [dict(r) for r in judgment_rows]

    text_body, html_body = build_digest_email_body(news_items, legislation_items, judgment_items)
    text_body = "[TEST EMAIL]\n\n" + text_body
    html_body = '<p style="background:#fdf6e8;padding:8px;border-radius:4px;font-size:13px"><strong>This is a test email.</strong></p>' + html_body
    total = len(news_items) + len(legislation_items) + len(judgment_items)
    subject = f"[TEST] \u2696 Mutemo Desk \u2014 Daily vault digest ({total} item{'s' if total != 1 else ''})"

    try:
        await asyncio.to_thread(_send_via_resend_sync, row["digest_recipient_email"], subject, html_body, text_body)
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Failed to send: {e}")
    return {"sent": True, "recipient": row["digest_recipient_email"], "items_included": total}

@app.post("/api/reminders/test")
async def test_reminder(request: Request):
    user = await get_current_user(request)
    _check_permission(user, "admin:settings")
    if not is_email_configured():
        raise HTTPException(status_code=503, detail="Email is not configured on this server.")
    async with _db_pool.acquire() as conn:
        row = await conn.fetchrow("SELECT * FROM reminder_settings WHERE firm_id=$1", FIRM_ID)
    if not row or not row["recipient_email"]:
        raise HTTPException(status_code=400, detail="No recipient email configured.")
    # Send a test with a dummy upcoming event so the HTML template renders
    today = datetime.utcnow().date()
    async with _db_pool.acquire() as conn:
        rows = await conn.fetch("""
            SELECT * FROM calendar_events
            WHERE firm_id=$1 AND date >= $2
            ORDER BY date ASC, time ASC NULLS LAST
            LIMIT 10
        """, FIRM_ID, today)
    test_events = [_row_to_event(r) for r in rows]
    for e in test_events:
        try:
            event_date = datetime.strptime(e["date"], "%Y-%m-%d").date()
            e["days_until"] = (event_date - today).days
        except Exception:
            e["days_until"] = 99
    if not test_events:
        test_events = [{
            "id": "test", "title": "No events scheduled yet",
            "date": today.isoformat(), "time": None, "event_type": "other",
            "court": None, "matter_name": FIRM_NAME, "notes": None, "days_until": 0,
        }]
    # Real matters for review (not dummy data) — same query the actual
    # scheduler uses, so this test send genuinely verifies the section
    # renders correctly for whatever real data exists right now, rather
    # than a synthetic stand-in.
    review_matters = await _get_review_matters_for_digest(today)
    sent = await send_reminder_email(row["recipient_email"], test_events, test=True, review_matters=review_matters)
    if not sent:
        raise HTTPException(status_code=500, detail="Failed to send test email.")
    return {
        "sent": True, "recipient": row["recipient_email"],
        "event_count": len(test_events), "review_matter_count": len(review_matters),
    }

# ── Reminder Scheduler ────────────────────────────────────────────────────────

def build_digest_email_body(news_items: list, legislation_items: list, judgment_items: list) -> tuple:
    """Returns (plain_text_body, html_body) for the daily legal-updates digest."""
    total = len(news_items) + len(legislation_items) + len(judgment_items)
    if total == 0:
        text = "Good morning. No new news, legislation, or judgments were added in yesterday's scrape.\n\n\u2014 Mutemo Desk"
        html = "<p>Good morning. No new news, legislation, or judgments were added in yesterday's scrape.</p><p style='color:#6b6b64'>\u2014 Mutemo Desk</p>"
        return text, html

    text_lines = [f"Good morning. Here's what was added to the vault today ({total} item(s)):\n"]
    html_parts = ["<p>Good morning. Here's what was added to the vault today:</p>"]

    def section(label: str, items: list, fmt_text_fn, fmt_html_fn):
        if not items:
            return
        text_lines.append(f"{label.upper()}:")
        for it in items:
            text_lines.append(f"  \u2022 {fmt_text_fn(it)}")
        text_lines.append("")
        html_parts.append(f"<p style='font-weight:600;margin-bottom:4px'>{label}</p><ul style='margin-top:0'>")
        for it in items:
            html_parts.append(f"<li>{fmt_html_fn(it)}</li>")
        html_parts.append("</ul>")

    section(
        "\U0001F4F0 News", news_items,
        lambda it: f"{it.get('reference') or it.get('filename', 'Untitled')} \u2014 {it.get('source_url', '')}",
        lambda it: (
            f'<a href="{_escape_html(it.get("source_url",""))}" target="_blank">'
            f'{_escape_html(it.get("reference") or it.get("filename","Untitled"))}</a>'
            if it.get("source_url") else _escape_html(it.get("reference") or it.get("filename", "Untitled"))
        ),
    )
    section(
        "\U0001F4DC New Legislation", legislation_items,
        lambda it: f"{it.get('reference') or it.get('filename', 'Untitled')} \u2014 {it.get('source_url', '')}",
        lambda it: (
            f'<a href="{_escape_html(it.get("source_url",""))}" target="_blank">'
            f'{_escape_html(it.get("reference") or it.get("filename","Untitled"))}</a>'
            if it.get("source_url") else _escape_html(it.get("reference") or it.get("filename", "Untitled"))
        ),
    )
    section(
        "\u2696 New Judgments", judgment_items,
        lambda it: f"{it.get('case_name') or it.get('filename', 'Untitled')} ({it.get('citation','')}) \u2014 {it.get('zimlii_url', '')}",
        lambda it: (
            f'<a href="{_escape_html(it.get("zimlii_url",""))}" target="_blank">'
            f'{_escape_html(it.get("case_name") or it.get("filename","Untitled"))}</a> '
            f'({_escape_html(it.get("citation",""))})'
            if it.get("zimlii_url") else
            f'{_escape_html(it.get("case_name") or it.get("filename","Untitled"))} ({_escape_html(it.get("citation",""))})'
        ),
    )

    text_lines.append("\u2014 Mutemo Desk")
    html_parts.append("<p style='color:#6b6b64'>\u2014 Mutemo Desk</p>")
    return "\n".join(text_lines), "".join(html_parts)


async def _maybe_send_digest():
    """
    Daily digest of new news/legislation/judgments — mirrors _maybe_send_reminder's
    pattern exactly. Default send hour (6 UTC) sits after ZimLII (04:00),
    Veritas (04:30), and News (05:30) all complete on weekdays, so a single
    daily digest naturally covers that whole morning's scrape.
    """
    if not _db_pool:
        return
    async with _db_pool.acquire() as conn:
        settings = await conn.fetchrow("SELECT * FROM reminder_settings WHERE firm_id=$1", FIRM_ID)
    if not settings or not settings["digest_enabled"] or not settings["digest_recipient_email"]:
        return

    now_utc = datetime.utcnow()
    if now_utc.hour != settings["digest_send_hour_utc"]:
        return
    today = now_utc.date()
    if today.weekday() >= 5:
        return
    if settings.get("digest_last_run_date") == today:
        return

    # "Since last digest" — falls back to the last 24 hours on first run
    since = datetime.combine(settings["digest_last_run_date"], datetime.min.time()) if settings.get("digest_last_run_date") else now_utc - timedelta(hours=24)

    async with _db_pool.acquire() as conn:
        news_rows = await conn.fetch(
            "SELECT * FROM legal_updates WHERE firm_id=$1 AND source_type='news' AND uploaded_at >= $2 ORDER BY uploaded_at DESC",
            FIRM_ID, since
        )
        legislation_rows = await conn.fetch(
            "SELECT * FROM legal_updates WHERE firm_id=$1 AND source_type='legislation' AND uploaded_at >= $2 ORDER BY uploaded_at DESC",
            FIRM_ID, since
        )
        judgment_rows = await conn.fetch(
            "SELECT * FROM zlr_entries WHERE firm_id=$1 AND uploaded_at >= $2 ORDER BY uploaded_at DESC",
            FIRM_ID, since
        )

    news_items = [dict(r) for r in news_rows]
    legislation_items = [dict(r) for r in legislation_rows]
    judgment_items = [dict(r) for r in judgment_rows]

    text_body, html_body = build_digest_email_body(news_items, legislation_items, judgment_items)
    total = len(news_items) + len(legislation_items) + len(judgment_items)
    subject = f"\u2696 Mutemo Desk \u2014 Daily vault digest ({total} new item{'s' if total != 1 else ''})" if total else "\u2696 Mutemo Desk \u2014 Daily vault digest (nothing new today)"

    try:
        await asyncio.to_thread(_send_via_resend_sync, settings["digest_recipient_email"], subject, html_body, text_body)
        print(f"[digest] Sent daily digest to {settings['digest_recipient_email']} ({total} items)")
    except Exception as e:
        print(f"[digest] Failed to send: {e}")
        return  # don't mark as sent if it failed — retry next hour

    async with _db_pool.acquire() as conn:
        await conn.execute(
            "UPDATE reminder_settings SET digest_last_run_date=$1 WHERE firm_id=$2",
            today, FIRM_ID
        )


async def reminder_scheduler_loop():
    """Runs every hour. Sends daily digest of upcoming deadlines if enabled."""
    await asyncio.sleep(30)  # brief startup delay
    while True:
        try:
            await _maybe_send_reminder()
        except Exception as e:
            print(f"[reminder] scheduler error: {e}")
        try:
            await _maybe_send_digest()
        except Exception as e:
            print(f"[digest] scheduler error: {e}")
        await asyncio.sleep(3600)

async def _get_review_matters_for_digest(today: date) -> list:
    """
    Matter review safety net (2026-08-30): matters whose next_review_date
    has arrived or passed, shaped for build_reminder_email_body's own
    "Matters for Review" section (see there — deliberately kept separate
    from the deadline events list, not merged into it). Shared between
    the real scheduler (_maybe_send_reminder) and the manual test-send
    endpoint (/api/reminders/test) so a real end-to-end check on staging
    doesn't need to wait for the hourly scheduler tick.

    No lower bound on the date range — overdue matters must always show,
    however overdue — only an upper bound of REVIEW_DIGEST_LOOKAHEAD_DAYS
    out. Same status/is_sentinel exclusions as the deadline query in
    _maybe_send_reminder — no reason to nudge about a closed or
    sentinel/demo matter.
    """
    if not _db_pool:
        return []
    async with _db_pool.acquire() as conn:
        review_rows = await conn.fetch("""
            SELECT m.id, m.name, m.next_review_date, m.last_reviewed_date, m.matter_number,
                   m.client_name AS matter_client_name, c.full_name AS client_full_name
            FROM matters m
            LEFT JOIN clients c ON c.id = m.client_id
            WHERE m.firm_id=$1 AND m.next_review_date IS NOT NULL
              AND m.next_review_date <= $2
              AND m.status != 'Closed' AND NOT m.is_sentinel
            ORDER BY m.next_review_date ASC
        """, FIRM_ID, today + timedelta(days=REVIEW_DIGEST_LOOKAHEAD_DAYS))
    review_matters = []
    for r in review_rows:
        review_matters.append({
            "name": r["name"],
            "matter_number": r["matter_number"],
            "case_number": extract_case_reference(r["name"]),
            "resolved_client_name": r["client_full_name"] or r["matter_client_name"],
            "next_review_date": str(r["next_review_date"]),
            "last_reviewed_date": str(r["last_reviewed_date"]) if r["last_reviewed_date"] else None,
            "days_until": (r["next_review_date"] - today).days,
        })
    return review_matters

async def _maybe_send_reminder():
    if not _db_pool:
        return
    async with _db_pool.acquire() as conn:
        settings = await conn.fetchrow("SELECT * FROM reminder_settings WHERE firm_id=$1", FIRM_ID)
    if not settings or not settings["enabled"] or not settings["recipient_email"]:
        return

    now_utc = datetime.utcnow()
    if now_utc.hour != settings["send_hour_utc"]:
        return
    today = now_utc.date()
    # Skip weekends — reminders only on Mon–Fri
    if today.weekday() >= 5:
        return
    if settings.get("last_run_date") == today:
        return

    # Collect upcoming events (next 30 days, including today). LEFT JOINs to
    # matters/clients (via calendar_events.matter_id, when a calendar event
    # is actually linked to one) resolve matter_number/case_number/client
    # name for the new reminder-line format — a plain event with no linked
    # matter (e.g. a staff meeting) simply gets NULLs here and falls back to
    # its own free-text matter_name in build_reminder_email_body, unchanged.
    async with _db_pool.acquire() as conn:
        rows = await conn.fetch("""
            SELECT ce.*, m.matter_number AS _joined_matter_number, m.name AS _joined_matter_full_name,
                   m.client_name AS _joined_matter_client_name, c.full_name AS _joined_client_full_name
            FROM calendar_events ce
            LEFT JOIN matters m ON m.id = ce.matter_id
            LEFT JOIN clients c ON c.id = m.client_id
            WHERE ce.firm_id=$1 AND ce.date >= $2 AND ce.date <= $3
            ORDER BY ce.date ASC, ce.time ASC NULLS LAST
        """, FIRM_ID, today, today + timedelta(days=30))

    events = []
    for r in rows:
        e = _row_to_event(r)
        e["matter_number"] = r["_joined_matter_number"]
        e["case_number"] = extract_case_reference(r["_joined_matter_full_name"])
        e["resolved_client_name"] = r["_joined_client_full_name"] or r["_joined_matter_client_name"]
        events.append(e)

    # Matter-level critical deadlines, converted into the same event shape
    # so they reuse the existing rendering logic in build_reminder_email_body
    # rather than needing a parallel display path. These are the deadlines
    # set directly on a matter (e.g. a jurisdictional time limit, an appeal
    # window) — distinct from calendar events, and easy to miss if they only
    # lived on the matter list without also surfacing here.
    async with _db_pool.acquire() as conn:
        deadline_rows = await conn.fetch("""
            SELECT m.id, m.name, m.next_deadline, m.next_deadline_note, m.matter_number,
                   m.client_name AS matter_client_name, c.full_name AS client_full_name
            FROM matters m
            LEFT JOIN clients c ON c.id = m.client_id
            WHERE m.firm_id=$1 AND m.next_deadline >= $2 AND m.next_deadline <= $3
              AND m.status != 'Closed' AND NOT m.is_sentinel
            ORDER BY m.next_deadline ASC
        """, FIRM_ID, today, today + timedelta(days=30))
    for r in deadline_rows:
        events.append({
            "event_type": "deadline",
            "title": r["next_deadline_note"] or f"Matter deadline: {r['name']}",
            "date": str(r["next_deadline"]),
            "matter_name": r["name"],
            "matter_number": r["matter_number"],
            "case_number": extract_case_reference(r["name"]),
            "resolved_client_name": r["client_full_name"] or r["matter_client_name"],
            "time": None,
            "court": None,
        })
    events.sort(key=lambda e: (e.get("date") or "", e.get("time") or ""))

    review_matters = await _get_review_matters_for_digest(today)

    if not events and not review_matters:
        async with _db_pool.acquire() as conn:
            await conn.execute(
                "UPDATE reminder_settings SET last_run_date=$1 WHERE firm_id=$2",
                today, FIRM_ID
            )
        return

    # Enrich events with days_until for the HTML email builder
    for e in events:
        try:
            event_date = datetime.strptime(str(e["date"])[:10], "%Y-%m-%d").date()
            e["days_until"] = (event_date - today).days
        except Exception:
            e["days_until"] = 0  # treat as today if date parse fails

    sent = await send_reminder_email(settings["recipient_email"], events, review_matters=review_matters)
    if sent:
        async with _db_pool.acquire() as conn:
            await conn.execute(
                "UPDATE reminder_settings SET last_run_date=$1 WHERE firm_id=$2",
                today, FIRM_ID
            )
        print(f"[reminder] digest sent to {settings['recipient_email']}: {len(events)} events, {len(review_matters)} for review")

# ── Inactivity Alerts ─────────────────────────────────────────────────────────

@app.post("/api/reminders/inactivity-check")
async def inactivity_check(request: Request):
    user = await get_current_user(request)
    _check_permission(user, "admin:settings")
    async with _db_pool.acquire() as conn:
        settings = await conn.fetchrow("SELECT * FROM reminder_settings WHERE firm_id=$1", FIRM_ID)
    if not settings or not settings["recipient_email"]:
        raise HTTPException(status_code=400, detail="No recipient email configured.")

    threshold = datetime.utcnow() - timedelta(days=14)
    async with _db_pool.acquire() as conn:
        rows = await conn.fetch("""
            SELECT id, name, internal_ref, last_activity
            FROM matters
            WHERE firm_id=$1 AND status='Active' AND NOT is_sentinel
              AND (last_activity IS NULL OR last_activity < $2)
            ORDER BY last_activity ASC NULLS FIRST
        """, FIRM_ID, threshold)

    inactive = [dict(r) for r in rows]
    for m in inactive:
        m["id"] = str(m["id"])
        if m.get("last_activity"):
            m["last_activity"] = m["last_activity"].isoformat()

    if not inactive:
        return {"inactive_count": 0, "message": "All active matters have recent activity."}

    lines = [f"Mutemo Desk — Inactivity Alert for {FIRM_NAME}", ""]
    lines.append(f"The following {len(inactive)} active matter(s) have had no activity in 14+ days:")
    lines.append("")
    for m in inactive:
        ref = m.get("internal_ref") or m.get("id", "")[:8]
        last = m.get("last_activity", "Never")[:10] if m.get("last_activity") else "Never"
        lines.append(f"  • [{ref}] {m['name']} — last activity: {last}")
    body = "\n".join(lines)
    sent = await send_reminder_email(
        settings["recipient_email"],
        f"Mutemo Desk — {len(inactive)} inactive matter(s)",
        body
    )
    return {"inactive_count": len(inactive), "matters": inactive, "email_sent": sent}

# ── Document status polling ───────────────────────────────────────────────────

@app.get("/api/documents/{doc_id}/status")
async def get_document_status(doc_id: str, request: Request):
    """Poll this endpoint after upload to check if background processing is complete."""
    user = await get_current_user(request)
    _check_permission(user, "matter:read")
    async with _db_pool.acquire() as conn:
        row = await conn.fetchrow(
            "SELECT id, filename, status, chunk_count, word_count, error_message FROM documents WHERE id=$1 AND firm_id=$2",
            _uuid_mod.UUID(doc_id), FIRM_ID
        )
    if not row:
        raise HTTPException(status_code=404, detail="Document not found")
    d = dict(row)
    d["id"] = str(d["id"])
    return d


@app.get("/api/documents/{doc_id}/text")
async def get_document_text(doc_id: str, request: Request):
    """
    Returns the actual stored text of a document, reconstructed from its
    chunks — used for loading a real precedent into the Draft Document
    feature. Previously the frontend only had a metadata placeholder
    ("Document from matter: X. Type: Y.") with no real content at all,
    which defeated the point of "precedent-aware" drafting.
    """
    user = await get_current_user(request)
    _check_permission(user, "matter:read")
    async with _db_pool.acquire() as conn:
        doc_row = await conn.fetchrow(
            "SELECT filename FROM documents WHERE id=$1 AND firm_id=$2",
            _uuid_mod.UUID(doc_id), FIRM_ID
        )
        if not doc_row:
            raise HTTPException(status_code=404, detail="Document not found")
        chunk_rows = await conn.fetch(
            "SELECT text FROM chunks WHERE document_id=$1 AND chunk_source='firm' ORDER BY chunk_index",
            _uuid_mod.UUID(doc_id)
        )
    if not chunk_rows:
        raise HTTPException(status_code=404, detail="No indexed text available for this document.")
    full_text = "\n\n".join(r["text"] for r in chunk_rows)
    # Cap at a reasonable length for prompt context — long enough to capture
    # real structure/language/clauses, short enough to not blow out the
    # drafting prompt alongside facts/instructions.
    MAX_PRECEDENT_CHARS = 4000
    truncated = len(full_text) > MAX_PRECEDENT_CHARS
    return {
        "filename": doc_row["filename"],
        "text": full_text[:MAX_PRECEDENT_CHARS],
        "truncated": truncated,
    }

@app.get("/api/documents/{doc_id}/view-url")
async def get_document_view_url(doc_id: str, request: Request):
    """Generate a presigned R2 URL for viewing/downloading a document.
    Bypasses Cloudflare Access — URL is valid for 1 hour."""
    user = await get_current_user(request)
    _check_permission(user, "matter:read")

    async with _db_pool.acquire() as conn:
        row = await conn.fetchrow(
            "SELECT filename, r2_key FROM documents WHERE id=$1 AND firm_id=$2",
            _uuid_mod.UUID(doc_id), FIRM_ID
        )
    if not row:
        raise HTTPException(status_code=404, detail="Document not found")
    if not row["r2_key"]:
        raise HTTPException(status_code=404, detail="File not available — document was uploaded before R2 storage was enabled. Please re-upload.")
    if not R2_ENABLED or not _r2_client:
        raise HTTPException(status_code=503, detail="File storage not configured.")

    try:
        presigned_url = await asyncio.to_thread(
            _r2_client.generate_presigned_url,
            "get_object",
            Params={"Bucket": R2_BUCKET, "Key": row["r2_key"]},
            ExpiresIn=3600  # 1 hour
        )
        return {"url": presigned_url, "filename": row["filename"]}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Could not generate file URL: {e}")


@app.delete("/api/documents/{doc_id}")
async def delete_document(doc_id: str, request: Request):
    """Delete a document — removes from DB, chunks, ChromaDB, and R2."""
    user = await get_current_user(request)
    _check_permission(user, "document:delete")

    async with _db_pool.acquire() as conn:
        row = await conn.fetchrow(
            "SELECT filename, matter_id, r2_key FROM documents WHERE id=$1 AND firm_id=$2",
            _uuid_mod.UUID(doc_id), FIRM_ID
        )
        if not row:
            raise HTTPException(status_code=404, detail="Document not found")

        matter_id = row["matter_id"]
        r2_key = row["r2_key"]

        await conn.execute(
            "DELETE FROM chunks WHERE document_id=$1 AND firm_id=$2",
            _uuid_mod.UUID(doc_id), FIRM_ID
        )
        await conn.execute(
            "DELETE FROM documents WHERE id=$1 AND firm_id=$2",
            _uuid_mod.UUID(doc_id), FIRM_ID
        )
        if matter_id:
            await conn.execute(
                "UPDATE matters SET document_count = GREATEST(0, document_count - 1), last_activity=NOW() WHERE id=$1 AND firm_id=$2",
                matter_id, FIRM_ID
            )

    if r2_key and R2_ENABLED and _r2_client:
        try:
            await asyncio.to_thread(_r2_client.delete_object, Bucket=R2_BUCKET, Key=r2_key)
        except Exception as e:
            print(f"[r2] delete failed: {e}")

    try:
        firm_col, _, _ = get_chroma_collections()
        firm_col.delete(where={"document_id": doc_id})
    except Exception as e:
        print(f"[chroma] delete failed: {e}")

    return {"deleted": True, "doc_id": doc_id}


# ── Firm settings ─────────────────────────────────────────────────────────────

@app.get("/api/settings")
async def get_settings(request: Request):
    user = await get_current_user(request)
    _check_permission(user, "admin:settings")
    async with _db_pool.acquire() as conn:
        firm = await conn.fetchrow("SELECT * FROM firms WHERE id=$1", FIRM_ID)
    if not firm:
        return {"firm_name": FIRM_NAME, "firm_city": FIRM_CITY}
    d = dict(firm)
    d["id"] = str(d["id"])
    if d.get("created_at"):
        d["created_at"] = d["created_at"].isoformat()
    return d

@app.patch("/api/settings")
async def update_settings(body: dict, request: Request):
    user = await get_current_user(request)
    _check_permission(user, "admin:settings")
    allowed = {"name", "short_name", "city", "country", "features"}
    updates = {k: v for k, v in body.items() if k in allowed}
    if not updates:
        raise HTTPException(status_code=400, detail="No valid fields to update")

    if "features" in updates:
        if not isinstance(updates["features"], list) or not all(
            isinstance(f, str) for f in updates["features"]
        ):
            raise HTTPException(status_code=400, detail="features must be a list of strings")
        updates["features"] = json.dumps(updates["features"])

    set_clauses = ", ".join(
        f"{k}=${i+2}::jsonb" if k == "features" else f"{k}=${i+2}"
        for i, k in enumerate(updates.keys())
    )
    values = list(updates.values())
    async with _db_pool.acquire() as conn:
        await conn.execute(
            f"UPDATE firms SET {set_clauses} WHERE id=$1",
            FIRM_ID, *values
        )
    return {"saved": True}

# ── PWA assets ────────────────────────────────────────────────────────────────
# Explicit routes, registered before the frontend catch-all below -- without
# these, /manifest.json and /sw.js would fall through to serve_frontend()
# and silently return index.html (wrong content, wrong content-type)
# instead of 404ing or serving the real file, since FastAPI/Starlette
# match routes in registration order and the catch-all matches everything.
# sw.js specifically must be served from the site root (not /assets/sw.js)
# -- a service worker's scope is limited to its own directory and below,
# so root-scoped install (controlling the whole app) requires a root path.

@app.get("/manifest.json")
async def serve_manifest():
    path = os.path.join(frontend_path, "manifest.json")
    if os.path.exists(path):
        return FileResponse(path, media_type="application/manifest+json")
    return JSONResponse(status_code=404, content={"detail": "manifest.json not found"})

@app.get("/sw.js")
async def serve_service_worker():
    path = os.path.join(frontend_path, "sw.js")
    if os.path.exists(path):
        return FileResponse(path, media_type="application/javascript")
    return JSONResponse(status_code=404, content={"detail": "sw.js not found"})

@app.get("/icons/{filename}")
async def serve_icon(filename: str):
    # filename comes straight from the URL path; restrict to a plain
    # basename with no path separators so this can't be used to read
    # arbitrary files elsewhere on disk (e.g. "../../backend/main.py").
    if "/" in filename or "\\" in filename or ".." in filename:
        return JSONResponse(status_code=404, content={"detail": "Not found"})
    path = os.path.join(frontend_path, "icons", filename)
    if os.path.exists(path):
        return FileResponse(path, media_type="image/png")
    return JSONResponse(status_code=404, content={"detail": "Not found"})

# ── Frontend catch-all ────────────────────────────────────────────────────────

@app.get("/{full_path:path}")
async def serve_frontend(full_path: str):
    index = os.path.join(frontend_path, "index.html")
    if os.path.exists(index):
        return FileResponse(index)
    return JSONResponse(status_code=404, content={"detail": "Frontend not found"})
