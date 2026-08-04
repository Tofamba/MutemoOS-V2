"""
Shared DOCX paragraph generation for the document-export endpoints in
backend/main.py (/api/export-docx for affidavits, /api/export-document-docx
for drafted documents from the rich-text editor).

Both endpoints previously built (or, for the new one, would have built)
one python-docx/docx.js Paragraph per input LINE, including blank lines —
so a run of N blank lines between two paragraphs became N+1 separately
margined paragraphs stacked on top of each other, compounding visible
spacing. The fix, shared by both callers here: a run of one or more blank
lines is a single paragraph boundary, not one per blank line.

Parsing (plain text -> blocks, HTML -> blocks) is kept separate from and
independent of the python-docx library — both `paragraphs_from_*` functions
return plain data (lists of (segments, alignment) tuples, alignment as the
strings "center"/"justify"), so they're fully unit-testable without
python-docx installed. Only build_docx_bytes()/write_docx_paragraph() at
the bottom actually touch the docx library.
"""

import re
from html.parser import HTMLParser
from typing import Optional

# Lines starting with any of these are treated as a caption heading —
# centered and bold — matching the affidavit export's existing convention.
_HEADING_PREFIXES = ("IN THE", "CASE NO", "BETWEEN")


def paragraphs_from_plain_text(text: str) -> list:
    """
    Splits plain text into DOCX paragraph blocks: a run of one or more
    blank lines is one paragraph boundary. Returns
    [(segments, alignment), ...] where alignment is "center" or "justify"
    and segments is a list of {"text": str, "bold": bool, "italic": bool,
    "underline": bool} dicts, with {"break": True} entries marking a
    forced line break within a paragraph (a single, non-blank-run newline
    inside one block — e.g. an address split across lines).
    """
    blocks = re.split(r'\n\s*\n+', text or '')
    paragraphs = []
    for block in blocks:
        block = block.strip('\n')
        if not block.strip():
            continue
        lines = block.split('\n')
        first_line = lines[0].strip()
        is_heading = first_line.startswith(_HEADING_PREFIXES)
        alignment = "center" if is_heading else "justify"
        segments = []
        for i, line in enumerate(lines):
            if i > 0:
                segments.append({"break": True})
            segments.append({"text": line, "bold": is_heading, "italic": False, "underline": False})
        paragraphs.append((segments, alignment))
    return paragraphs


class _DocxHtmlParser(HTMLParser):
    """
    Minimal HTML -> paragraph-segment parser scoped to what the drafting
    editor's contenteditable/execCommand toolbar and generateDocument()
    actually produce: <p>/<div> paragraphs, <br>, <strong>/<b>, <em>/<i>,
    <u>, and text-align:center via a style attribute. Not a general-purpose
    HTML-to-DOCX converter.
    """
    BLOCK_TAGS = {"p", "div", "h1", "h2", "h3", "h4", "h5", "h6", "li"}

    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.paragraphs = []
        self._segments = []
        self._bold_depth = 0
        self._italic_depth = 0
        self._underline_depth = 0
        self._alignment = "justify"

    def _flush_paragraph(self):
        if any(s.get("text", "").strip() for s in self._segments):
            self.paragraphs.append((self._segments, self._alignment))
        self._segments = []
        self._alignment = "justify"

    def handle_starttag(self, tag, attrs):
        attrs_dict = dict(attrs)
        if tag in self.BLOCK_TAGS:
            self._flush_paragraph()
            style = (attrs_dict.get("style") or "").replace(" ", "")
            if "text-align:center" in style:
                self._alignment = "center"
        elif tag == "br":
            self._segments.append({"break": True})
        elif tag in ("strong", "b"):
            self._bold_depth += 1
        elif tag in ("em", "i"):
            self._italic_depth += 1
        elif tag == "u":
            self._underline_depth += 1

    def handle_endtag(self, tag):
        if tag in self.BLOCK_TAGS:
            self._flush_paragraph()
        elif tag in ("strong", "b"):
            self._bold_depth = max(0, self._bold_depth - 1)
        elif tag in ("em", "i"):
            self._italic_depth = max(0, self._italic_depth - 1)
        elif tag == "u":
            self._underline_depth = max(0, self._underline_depth - 1)

    def handle_data(self, data):
        if not data:
            return
        self._segments.append({
            "text": data,
            "bold": self._bold_depth > 0,
            "italic": self._italic_depth > 0,
            "underline": self._underline_depth > 0,
        })

    def close(self):
        super().close()
        self._flush_paragraph()


def paragraphs_from_html(html_content: str) -> list:
    """Same return shape as paragraphs_from_plain_text(), parsed from editor HTML."""
    parser = _DocxHtmlParser()
    parser.feed(html_content or "")
    parser.close()
    return parser.paragraphs


def write_docx_paragraph(document, segments: list, alignment: str = "justify",
                          font_name: str = "Times New Roman", font_size_pt: int = 12):
    """
    Adds ONE paragraph to a python-docx Document from a segment list —
    the single place both export endpoints go through, so the blank-line
    fix (and any future paragraph-formatting fix) only has to be correct
    once, not once per endpoint.
    """
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if alignment == "center" else WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)  # was spacing:{after:120} twips in the old docx.js script = 6pt
    for seg in segments:
        if seg.get("break"):
            p.add_run().add_break()
            continue
        text = seg.get("text", "")
        if not text:
            continue
        run = p.add_run(text)
        run.font.name = font_name
        run.font.size = Pt(font_size_pt)
        run.bold = bool(seg.get("bold"))
        run.italic = bool(seg.get("italic"))
        run.underline = bool(seg.get("underline"))
    return p


def build_docx_bytes(paragraph_blocks: list) -> bytes:
    """Builds a complete .docx file from paragraph blocks, returns its raw bytes."""
    import io
    from docx import Document
    from docx.shared import Inches

    document = Document()
    section = document.sections[0]
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)
    for segments, alignment in paragraph_blocks:
        write_docx_paragraph(document, segments, alignment)
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()
